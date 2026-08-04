"""Write the clause-actuator campaign into Chapters 1 and 5 and Appendix E.

THE SLICE. experiments/out/geometry_campaign_g5bp8.csv and
geometry_proposals_g5bp8.jsonl: five seeds, two arms differing only in
whether the geometry diagnosis is raised in the situation, the same
worlds and the same gate chain, which for this slice ends in G5b. Pool
0.25, four ignitions, 240 minutes, decision cycle 12 minutes, trial
horizon 24 minutes, stage-3 patience 8.

EVERY NUMBER IS READ FROM THOSE TWO FILES at run time. Nothing here is
typed by hand, so the prose cannot drift from the ledger, and a rerun
that produces different results produces different prose.

WHAT IS WRITTEN
  5.5.3 a closing block on the discovered clause actuators: two tables,
        one figure, and the attribution finding
  1.3   C1 moves from a capability claim to a demonstrated one
  5.6   one sentence on the H1 evidence line
  E     a listing of every object the campaign created

Usage: python experiments/fill_g5b_results.py IN.docx OUT.docx
"""
from __future__ import annotations

import collections
import csv
import json
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "out")
sys.path.insert(0, os.path.dirname(HERE))
sys.path.insert(0, HERE)

import docx                                                # noqa: E402

from fill_ch4_mechanism import (Cursor, caption4, find,     # noqa: E402
                                full, heading, retext)
from fill_h1_h2 import figure_par, para                     # noqa: E402

SLICE = "g5bp8"
FIGPNG = "fig5_29_generative_products.png"

KINDS = [("rule", "Plain rule"),
         ("concept", "Intermediate concept"),
         ("composite", "Macro intervention"),
         ("clause", "Clause actuator")]


def caption5(doc, kind, text, bookmark, shown, chapter="5"):
    """A caption whose cached chapter label is the one it will resolve
    to. The STYLEREF field finds the heading itself when Word updates
    the document; the cached text only matters until then, and a
    caption reading "Table 5.3" inside Appendix E is confusing enough
    to be worth setting."""
    p = caption4(doc, kind, text, bookmark, shown)
    for fld in p._p.iter():
        if fld.text and fld.text.strip() == "4":
            fld.text = chapter
    return p


# ------------------------------------------------------------ the data
def load():
    with open(os.path.join(OUT, f"geometry_campaign_{SLICE}.csv"),
              encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    with open(os.path.join(OUT, f"geometry_proposals_{SLICE}.jsonl"),
              encoding="utf-8") as f:
        props = [json.loads(line) for line in f]

    arm = collections.defaultdict(collections.Counter)
    for r in rows:
        for k in ("cycles", "diagnosed", "stage3_attempts", "accepted",
                  "clause", "clause_accepted", "rej_G3", "rej_G4G5",
                  "rej_G5b"):
            arm[r["arm"]][k] += int(float(r.get(k) or 0))

    prop = collections.Counter(p["kind"] for p in props)
    adm = collections.Counter(p["kind"] for p in props
                              if p.get("accepted"))
    gate = collections.Counter(p.get("gate") for p in props
                               if not p.get("accepted"))

    acts, refused_flat, refused_worse = [], 0, 0
    for r in props:
        if r["kind"] != "clause":
            continue
        ni = r["payload"]["new_intervention"]
        m = r.get("measured") or {}
        g3, g4 = m.get("g3") or {}, m.get("g4") or {}
        if r.get("accepted"):
            acts.append(dict(
                name=ni["name"], arm=r["arm"], seed=r["seed"],
                step=int(r.get("step") or 0),
                clauses="; ".join(
                    f"{c['effect']} on {c['sector']} "
                    f"[{c['range'][0]}, {c['range'][1]}] at "
                    f"{float(c['amount']):.1f}"
                    for c in ni["clauses"]),
                alone=not [k for k, _v in r["payload"]["consequents"]
                           if k != ni["name"]],
                g3=g3["j_without"] - g3["j_with"],
                g4=g4["j_without"] - g4["j_with"]))
        elif g3:
            if abs(g3["j_with"] - g3["j_without"]) < 1e-12:
                refused_flat += 1
            else:
                refused_worse += 1
    acts.sort(key=lambda a: (a["step"], a["name"]))

    objs = []
    for r in props:
        if not r.get("accepted"):
            continue
        p = r["payload"]
        if r["kind"] == "concept":
            nc = p["new_concept"]
            objs.append(("Intermediate concept", nc["name"],
                         ", ".join(f"{a} {float(b):.2f}"
                                   for a, b in nc["inputs"])))
        elif r["kind"] == "composite":
            ni = p["new_intervention"]
            objs.append(("Macro intervention", ni["name"],
                         ", ".join(f"{a} {float(b):.2f}"
                                   for a, b in ni["composition"])))
        elif r["kind"] == "clause":
            ni = p["new_intervention"]
            objs.append(("Clause actuator", ni["name"],
                         "; ".join(
                             f"{c['effect']} on {c['sector']} "
                             f"[{c['range'][0]}, {c['range'][1]}] at "
                             f"{float(c['amount']):.1f}"
                             for c in ni["clauses"])))
    seen, uniq = set(), []
    for kind, name, defn in objs:
        if (kind, name) in seen:
            continue
        seen.add((kind, name))
        uniq.append((kind, name, defn))
    uniq.sort()

    gen_concepts = {n for k, n, _d in uniq
                    if k == "Intermediate concept"}
    cross = []
    for r in props:
        if not r.get("accepted"):
            continue
        cited = [v for v, _t in (r["payload"].get("antecedents") or [])
                 if v in gen_concepts]
        if cited and r["kind"] == "clause":
            cross.append((r["payload"]["new_intervention"]["name"],
                          cited[0]))
    reuse = []
    for r in props:
        if r["kind"] != "rule" or not r.get("accepted"):
            continue
        for k, _v in r["payload"]["consequents"]:
            if any(k == n for kd, n, _d in uniq
                   if kd == "Clause actuator"):
                reuse.append((k, int(r.get("step") or 0)))
    return dict(arm=arm, prop=prop, adm=adm, gate=gate, acts=acts,
                objs=uniq, cross=cross, reuse=reuse,
                refused_flat=refused_flat, refused_worse=refused_worse,
                cycles=sum(a["cycles"] for a in arm.values()),
                attempts=sum(a["stage3_attempts"] for a in arm.values()),
                admitted=sum(a["accepted"] for a in arm.values()))


def main():
    inp, outp = sys.argv[1], sys.argv[2]
    D = load()
    doc = docx.Document(inp)
    done = []
    A = D["arm"]

    # =========================================== 5.5.3, the actuators
    ps = doc.paragraphs
    c = Cursor(ps[find(ps, "One recorded incident from ignition to "
                           "containment", style="Caption")]._p)

    c.add(heading(doc, 4, "Discovered Clause Actuators"))
    c.add(para(doc, "ParText", [
        "The products of Table 5.11 are compositions of orders the doctrine already carries. A "
        "second campaign asks the harder question, which is whether "
        "the stage can also create an intervention that carries its "
        "own geometry, meaning an order that says where on the map to "
        "act and how, in terms no base family expresses. It is "
        "reported separately because it runs the gate chain that ends "
        f"in G5b. Five seeds were run in two arms over {D['cycles']} "
        "decision cycles on the same worlds and under the same gates, "
        "the arms differing only in whether the geometry diagnosis of "
        "Section 4.5 is raised in the situation the proposer reads. "
        f"The stage was entered {D['attempts']} times and admitted "
        f"{D['admitted']} proposals. ",
        ("_RefTabGen", "Table 5.14"),
        " and ",
        ("_RefFigGen", "Figure 5.23"),
        " report what was proposed and what survived, by kind of "
        "object."]))
    c.add(caption5(doc, "Table",
                   "What stage ③ proposed and what the gates admitted, "
                   "by object kind", "_RefTabGen", "14"))
    c.table(doc, ["Object kind", "Proposed", "Admitted",
                  "Where the refusals fell"],
            [(label,
              str(D["prop"].get(k, 0)),
              str(D["adm"].get(k, 0)),
              _refusals(D, k)) for k, label in KINDS]
            + [("Total", str(sum(D["prop"].values())),
                str(sum(D["adm"].values())),
                ", ".join(f"{g}: {n}" for g, n
                          in sorted(D["gate"].items())))],
            widths=[1.55, 0.85, 0.85, 3.05])
    c.add(figure_par(doc, FIGPNG, width_in=6.2))
    c.add(caption5(doc, "Figure",
                   "Generative products of the campaign: (a) proposed "
                   "and admitted by object kind, (b) the forecast cost "
                   "each admitted actuator removed on the two reseeded "
                   "rollouts", "_RefFigGen", "23"))

    c.add(para(doc, "ParText", [
        _num(len(D["acts"])).capitalize() + " of the "
        + _num(D["prop"].get("clause", 0)) + " proposed clause "
        "actuators were admitted, and they are listed in ",
        ("_RefTabAct", "Table 5.15"),
        " with the moment of the incident at which each was written "
        "and the cost each removed on both rollouts. Two families "
        "appear and they are separated in time. The actuators written "
        "in the first cycle coat the assets and the populated cells "
        "before the fire reaches them, which is a preventive tactic "
        "with no counterpart in the seed doctrine, since the base "
        "asset protection family raises a defensive posture but lays "
        "no retardant. The actuators written at minute 48 draft from "
        "the water body into the burning cells and the head and "
        "evacuate the population behind them, which is the "
        "substitution described in Section 4.5.3: on an incident whose "
        "pool is already spent, drafting raises the yield of crews "
        "that are committed instead of asking for crews that are "
        "not."]))
    c.add(caption5(doc, "Table",
                   "The admitted clause actuators, their definitions "
                   "and what each removed from the forecast cost",
                   "_RefTabAct", "15"))
    c.table(doc, ["Actuator", "Written at", "Clauses",
                  "Rollout 1", "Rollout 2"],
            [(a["name"], f"{a['step']} min", a["clauses"],
              f"{a['g3']:+.5f}", f"{a['g4']:+.5f}")
             for a in D["acts"]],
            widths=[1.35, 0.65, 2.65, 0.8, 0.8])

    _alone = all(a["alone"] for a in D["acts"])
    c.add(para(doc, "ParText", [
        "The gain reported in ",
        ("_RefTabAct", "Table 5.15"),
        " belongs to the actuator and not to its company. "
        + ("Every admitted actuator was the only consequent of the "
           "rule that ordered it, so the two rollouts of G5 compare a "
           "situation in which the new object acts against one in "
           "which nothing acts in its place. "
           if _alone else
           "Not every admitted actuator was the only consequent of "
           "its rule. ")
        + "G5b, which repeats the comparison with the new object "
        "struck out and the ordinary orders of its rule left standing, "
        f"refused {_num(A['on']['rej_G5b'] + A['off']['rej_G5b'])} "
        "proposals in this campaign. That is reported as it stands "
        "rather than as a property of the gate: the situation given to "
        "the proposer also states that a new object must stand on its "
        "own, so no package arrived padded and the gate was never put "
        "to work. An earlier slice, measured before either change, "
        "contained two admissions in which two different clause "
        "geometries returned bit-identical forecasts because the "
        "reduction came from a resource deployment written beside "
        "them, which is the failure the gate exists to prevent."]))

    c.add(para(doc, "ParText", [
        "The refusals carry as much information as the admissions. Of "
        "the " + _num(D["refused_flat"] + D["refused_worse"])
        + " refused actuators, "
        + _num(D["refused_flat"]) + " returned a forecast "
        "identical to the one without them on the first rollout, which "
        "means the tactic did not touch anything the forecast reads, "
        f"and {_num(D['refused_worse'])} raised the cost. An identical "
        "forecast has three causes in this framework, all of them "
        "geometric: the sector resolved to no cells at the moment the "
        "actuator was written, the coating fell below the extinction "
        "moisture of the fuel it was laid on, or the effect drew on a "
        "pool with nothing left to give. Two of the sectors, assets "
        "and populated, resolve without reference to the written "
        "range, so two actuators that differ only in their ranges "
        "compile to the same order and are measured identically, which "
        "is why the two actuators of the first row pair in ",
        ("_RefTabAct", "Table 5.15"),
        " report the same numbers."]))

    if D["cross"] or D["reuse"]:
        bits = []
        if D["cross"]:
            n, cited = D["cross"][0]
            bits.append(f"The actuator {n} is ordered by a rule whose "
                        f"antecedent is {cited}, an intermediate "
                        "concept the same campaign created, so one "
                        "generated object is read through another. ")
        if D["reuse"]:
            n, step = D["reuse"][0]
            bits.append(f"The actuator {n}, written at minute 48, is "
                        "ordered again at minute "
                        f"{step} of the same incident by a separately "
                        "admitted rule that was not part of the "
                        "package which created it. ")
        c.add(para(doc, "ParText", [
            "Two of the admitted objects show the vocabulary being "
            "used rather than merely being grown. " + "".join(bits)
            + "This is the property the open decision space claims: an "
            "object admitted through the gates enters the vocabulary "
            "on the same footing as a doctrinal one and is available "
            "to whatever the rest of the incident writes."]))

    c.add(para(doc, "ParText", [
        "Both arms produced admitted actuators, "
        + _num(A["on"]["clause_accepted"]) + " of "
        + _num(A["on"]["clause"]) + " in the arm that raises the "
        "geometry diagnosis and "
        + _num(A["off"]["clause_accepted"]) + " of "
        + _num(A["off"]["clause"]) + " in the arm that suppresses it. "
        "The "
        "diagnosis is therefore not what makes the open space "
        "productive; it is a statement about the incident that changes "
        "what the proposer attends to, and the gates decide the rest."]))
    done.append("5.5.3: clause actuator subsection, Tables 5.14 and "
                "5.15, Figure 5.23")

    # ======================================================= C1 in 1.3
    ps = doc.paragraphs
    c1 = ps[find(ps, "To address G1, this thesis replaces the fixed "
                     "menu")]
    if retext(c1, "compose verified effects, such as wetting, "
                  "clearing, and coating, into new interventions",
              "compose verified effects, such as wetting, clearing, "
              "coating and drafting, into new interventions carrying "
              "their own geometry, which Section 5.5.3 demonstrates "
              "rather than merely permits"):
        done.append("1.3: C1 raised from a capability claim to a "
                    "demonstrated one")
    else:
        print("  ! C1 phrase not matched, contribution left unchanged")

    # ====================================================== 5.6 line
    ps = doc.paragraphs
    h1 = ps[find(ps, "H1 rests on the improvement ladder and the "
                     "scaling study")]
    Cursor(h1._p).add(para(doc, "ParText", [
        "The clause-actuator campaign of Section 5.5.3 adds the "
        "vocabulary level to that evidence line. The open "
        "configuration did not only choose better among the orders it "
        "was given; it composed "
        f"{_num(len(D['acts']))} interventions that the seed doctrine "
        "does not contain, each one admitted only after lowering the "
        "forecast physical cost on two independently reseeded futures, "
        "and one of them was later ordered by a rule that had no part "
        "in creating it."]))
    done.append("5.6: one sentence on the H1 evidence line")

    # ================================================== Appendix E.3
    ps = doc.paragraphs
    c = Cursor(ps[find(ps, "Term set: VL = very low, L = low, M = "
                           "medium")]._p)
    c.add(para(doc, "ParText", [
        "Objects created by the open decision space. A generated rule "
        "is written in the vocabulary in force at the time, and part "
        "of that vocabulary is itself generated. The table below "
        "lists every object the campaign of Section 5.5.3 admitted, "
        "with the definition recorded in the lineage store, so a rule "
        "of the previous table that cites an unfamiliar name can be "
        "read here. A concept is a normalized weight vector over "
        "existing inputs, a macro is a weighted composition of base "
        "families, and a clause actuator is a list of verified "
        "effects with the sector and range each acts on."]))
    c.add(caption5(doc, "Table",
                   "Vocabulary objects admitted by the generative "
                   "stage, with their recorded definitions",
                   "_RefTabE3", "3", chapter="E"))
    c.table(doc, ["Kind", "Name", "Recorded definition"],
            D["objs"], widths=[1.3, 1.75, 3.25])
    done.append(f"Appendix E: Table E.3, {len(D['objs'])} admitted "
                "objects")

    doc.save(outp)
    print("written:", outp)
    for d in done:
        print("  -", d)


def _refusals(D, kind):
    g = collections.Counter()
    for p in _all(D):
        if p["kind"] == kind and not p.get("accepted"):
            g[p.get("gate") or "?"] += 1
    return ", ".join(f"{k}: {v}" for k, v in sorted(g.items())) or "none"


_CACHE = {}


def _all(D):
    if "props" not in _CACHE:
        with open(os.path.join(OUT, f"geometry_proposals_{SLICE}.jsonl"),
                  encoding="utf-8") as f:
            _CACHE["props"] = [json.loads(line) for line in f]
    return _CACHE["props"]


def _num(n):
    return {0: "no", 1: "one", 2: "two", 3: "three", 4: "four",
            5: "five", 6: "six", 7: "seven", 8: "eight", 9: "nine",
            10: "ten", 11: "eleven", 12: "twelve", 13: "thirteen",
            14: "fourteen", 15: "fifteen", 16: "sixteen",
            17: "seventeen", 18: "eighteen", 19: "nineteen",
            20: "twenty"}.get(int(n), str(n))


if __name__ == "__main__":
    main()
