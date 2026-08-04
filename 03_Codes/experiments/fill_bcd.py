"""Write the traceability, counter and ledger corrections into the thesis.

Three items, each from a logged source and none of them invented.

  B  one decision reproduced end to end, from the observation the sensors
     delivered to the order the region received, as a table in 5.5.3
  C  what the campaign actually fielded, counted over the recorded cycle
     logs, as one sentence at the end of the ladder subsection
  D  two ledger corrections: the macro whose name claims a fire its
     definition does not carry, and the appendix caption that counts
     rules the table does not show

WHAT IS NOT WRITTEN. The clause-actuator slice produced ten proposals and
admitted none, so Chapter 1 keeps its capability wording and Section
5.5.3 gains no actuator paragraph. The constraint-violation count the
reviewer asked for is not recoverable from the recorded runs: the pool
block that carries demand against funded budget was added to the cycle
record after those runs were made, so the logs are silent rather than
clean, and the sentence says so instead of reporting a zero.

Usage: python experiments/fill_bcd.py IN.docx OUT.docx
"""
from __future__ import annotations

import os
import sys

import docx
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(HERE))
sys.path.insert(0, HERE)

from docx_track import (AUTHOR, DATE, _el, _ins_row, _nid,   # noqa: E402
                        ins_run)
from fill_h1_h2 import caption_par, para, table_after       # noqa: E402

# ---------------------------------------------------------------- TASK B
#: logs/DSS_20260726_175523/cycles.jsonl, step 5, t = 5 min, Agent_1.
#: Chosen because the acceptance test decides something there: the
#: candidate is far inside the bound rather than marginal, the quality
#: gate passes without derating, and twelve rules carry the decision.
TRACE_SOURCE = "logs/DSS_20260726_175523/cycles.jsonl, step 5, Agent_1"
TRACE = [
    ("Observation",
     "fire intensity 0.97 (confidence 0.97), ignition proximity 1.00 "
     "(0.97), asset exposure 1.00 (1.00), fuel load 0.70 (0.12), "
     "suppression availability 1.00 (1.00)"),
    ("Concept gate",
     "fire threat level 0.12, asset exposure risk 0.97, suppression "
     "feasibility 1.00, intervention urgency 0.12, evacuation pressure "
     "0.97"),
    ("Concept activation, effective",
     "fire threat level 0.36, asset exposure risk 0.85, suppression "
     "feasibility 0.79, intervention urgency 0.38, evacuation pressure "
     "0.62"),
    ("Rules fired, firing strength",
     "G5 0.50, R20 0.19, G2 0.19, G4 0.19, G1 0.17, G3 0.17, G6 0.17, "
     "G7 0.17, G8 0.17, G9 0.17, R26 0.11, R7 0.07"),
    ("Candidate orders",
     "water drafting 0.84, containment line 0.83, retardant drop 0.82, "
     "suppression effort 0.80, asset protection 0.75, resource "
     "deployment 0.70, evacuation 0.66"),
    ("Satisficing test",
     "forecast cost 0.0490 against a no-action forecast of 0.2568 and a "
     "bound of 0.2500, accepted"),
    ("Quality gate",
     "decision quality 0.776 against a gate of 0.60, no derating"),
    ("Fielded order",
     "the candidate as written, at an attention share of 1.00, the "
     "region being attended"),
]

# ---------------------------------------------------------------- TASK C
#: counted over logs/DSS_*/cycles.jsonl, 160 runs, 7130 decision cycles
COUNTERS = dict(runs=160, cycles=7130, candidate=157247, fielded=151745,
                withheld=5502, failsafe=5338, veto=113)

# ---------------------------------------------------------------- TASK D
OLD_MACRO = "downwind_backburn"
NEW_MACRO = "downwind_containment_shield"


def full(p):
    return "".join((t.text or "") for t in p._p.iter(qn("w:t")))


def find(ps, needle, start=0):
    for i in range(start, len(ps)):
        if needle in full(ps[i]):
            return i
    raise SystemExit(f"anchor not found: {needle!r}")


def retext(par, old, new):
    """Replace a phrase inside a paragraph as a tracked edit."""
    for r in list(par._p.iter(qn("w:r"))):
        if r.getparent().tag == qn("w:del"):
            continue
        ts = r.findall(qn("w:t"))
        txt = "".join(t.text or "" for t in ts)
        if old not in txt:
            continue
        parent = r.getparent()
        d = _el("w:del", **{"w:id": _nid(), "w:author": AUTHOR,
                            "w:date": DATE})
        parent.replace(r, d)
        d.append(r)
        for t in ts:
            t.tag = qn("w:delText")
        d.addnext(ins_run(txt.replace(old, new)))
        return True
    return False


def main():
    inp, outp = sys.argv[1], sys.argv[2]
    doc = docx.Document(inp)
    changed = []

    # ---- B: one decision, end to end -------------------------------
    ps = doc.paragraphs
    i = find(ps, "The complete admitted set is listed in Appendix E")
    cur = ps[i]

    cur = _after(cur, para(doc, "ParText", [
        "The rule base and the gates are also observable one decision at "
        "a time. ",
        ("_RefTabTrace", "5.13"),
        " reproduces a single decision of a recorded incident from end "
        "to end, five minutes after ignition, for the region the "
        "coordinator was attending. The row order is the order the "
        "layer works in, so each line is produced by the line above it: "
        "the sensors deliver ten readings with a confidence apiece, the "
        "confidences gate the concepts, the gated concepts fire the "
        "rules, the fired rules compose the candidate orders, and the "
        "two acceptance tests decide whether those orders are fielded "
        "as written."]))

    cap = caption_par(doc, "Table", "One decision, end to end",
                      "_RefTabTrace", "13")
    cur = _after(cur, cap)
    table_after(doc, cap, ["Step of the loop", "What the record holds"],
                TRACE, widths=[1.9, 4.4])
    changed.append("Task B: end-to-end decision table, " + TRACE_SOURCE)

    # ---- C: what the campaign fielded ------------------------------
    ps = doc.paragraphs
    j = find(ps, "The two evolving-fuzzy funnels appear on the left")
    c = COUNTERS
    _after(ps[j], para(doc, "ParText", [
        f"Over the {c['runs']} recorded runs and {c['cycles']:,} "
        f"decision cycles the rule base produced {c['candidate']:,} "
        f"candidate interventions and {c['fielded']:,} of them reached "
        f"the field, so {c['withheld']:,} were withheld between the "
        "rule output and the order, the graduated fail-safe engaged in "
        f"{c['failsafe']:,} region cycles, and the no-harm veto held "
        f"the offensive channels in {c['veto']} cycles. The number of "
        "orders the physics could not honour after execution is not "
        "reported here, because the pool record that would carry it was "
        "added to the cycle log after these runs were made; the logs "
        "are silent on that quantity rather than clear of it."
        .replace(",", " ")]))
    changed.append("Task C: fielded counters over logs/DSS_*/cycles.jsonl")

    # ---- D: the macro whose name claims a fire it does not carry ----
    hits = 0
    for p in doc.paragraphs:
        if OLD_MACRO in full(p) and retext(p, OLD_MACRO, NEW_MACRO):
            hits += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if OLD_MACRO in full(p) and retext(p, OLD_MACRO,
                                                       NEW_MACRO):
                        hits += 1
    if hits:
        ps = doc.paragraphs
        k = find(ps, "Table 5.11 lists every vocabulary object")
        _after(ps[k], para(doc, "ParText", [
            "One product is renamed here. The logged definition of the "
            f"macro formerly called {OLD_MACRO} is a containment line "
            "at full intensity with suppression effort behind it, and "
            "it carries no igniting channel and no igniting clause, so "
            "the name claimed a counter-fire the definition does not "
            "perform. It is listed as "
            f"{NEW_MACRO}. The macro that does carry fire is "
            "counterfire_strip, which composes the tactical burn with "
            "a containment line."]))
    changed.append(f"Task D: {OLD_MACRO} renamed in {hits} places")

    # ---- D: the appendix caption that counts what it does not show --
    for p in doc.paragraphs:
        if "distinct generative rules" in full(p):
            if retext(p, "The forty distinct generative rules admitted "
                         "over the recorded runs",
                      "Twenty of the forty-two generative rules "
                      "admitted over the recorded runs"):
                changed.append("Task D: Table E.2 caption reconciled "
                               "with its twenty rows and the forty-two "
                               "in the store")
            break

    doc.save(outp)
    print("written:", outp)
    for c_ in changed:
        print("  -", c_)


def _after(anchor, par):
    anchor._p.addnext(par._p)
    return par


if __name__ == "__main__":
    main()
