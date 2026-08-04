"""Write the H1 and C2/H2 experiments into the thesis, as tracked changes.

Two insertions and one correction.

  Section 5.4  a new subsection reporting how a perturbation of an
               observation feature, of a membership partition or of a
               cost weight reaches the decision. Two figures and a
               table, from experiments/out/prop_oat.csv and
               prop_morris*.csv.
  Section 5.5.4 the centralized baseline: what the same architecture
               costs when the reasoning is centralized and the rule base
               is closed. Two figures and a table, from
               central_latency.csv and central_outcome.csv.
  H1           the claim is narrowed to what the experiment supports.
               The measurement does not show the distributed
               configuration deciding faster, and the sentence is
               corrected rather than the experiment.

NUMBERING IS LEFT TO WORD. Captions carry STYLEREF and SEQ fields and
in-text references carry REF fields, exactly as the surrounding chapter
does, so inserting a figure in the middle of Chapter 5 renumbers
everything after it when the fields are updated. Hard-coded numbers
would be wrong the moment the supervisor accepted the changes.

Usage: python experiments/fill_h1_h2.py IN.docx OUT.docx
"""
from __future__ import annotations

import collections
import csv
import glob
import os
import sys

import numpy as np

import docx
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(HERE))
sys.path.insert(0, HERE)

from docx_track import (AUTHOR, DATE, _el, _ins_row, _nid,   # noqa: E402
                        ins_run)

OUT = os.path.join(HERE, "out")
FIGDIR = os.path.join(HERE, "..", "..", "01_Thesis", "figures")


# --------------------------------------------------------------- fields
def _fld(instr, shown):
    """A Word field, with the value it last displayed kept alongside.

    The cached value is what a reader sees before the fields are
    updated; without it the caption reads "Figure ." until someone
    presses F9, which looks like a broken document rather than one whose
    numbering is live.
    """
    out = []
    r = _el("w:r")
    r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    out.append(r)
    r = _el("w:r")
    it = _el("w:instrText")
    it.text = instr
    it.set(qn("xml:space"), "preserve")
    r.append(it)
    out.append(r)
    r = _el("w:r")
    r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    out.append(r)
    r = _el("w:r")
    rpr = _el("w:rPr")
    rpr.append(_el("w:noProof"))
    r.append(rpr)
    t = _el("w:t")
    t.text = shown
    r.append(t)
    out.append(r)
    r = _el("w:r")
    r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))
    out.append(r)
    return out


def _run(text):
    r = _el("w:r")
    t = _el("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _wrap_ins(elements):
    ins = _el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                          "w:date": DATE})
    for e in elements:
        ins.append(e)
    return ins


def caption_par(doc, kind, text, bookmark, shown):
    """A caption in the chapter's own form: STYLEREF, SEQ, bookmark."""
    p = Paragraph(_el("w:p"), doc)
    ppr = _el("w:pPr")
    ppr.append(_el("w:pStyle", **{"w:val": "Caption"}))
    rpr = _el("w:rPr")
    rpr.append(_el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                               "w:date": DATE}))
    ppr.append(rpr)
    p._p.append(ppr)
    # THE BOOKMARK COVERS THE LABEL AND THE NUMBER, NOTHING ELSE. A
    # REF field returns whatever the bookmark spans, so a bookmark that
    # also swallowed the caption text would make every cross reference
    # in the prose repeat the whole caption. The chapter's own captions
    # close the bookmark immediately after the SEQ field, and these do
    # the same.
    bid = _nid()
    label = [_run(kind + " ")]
    label += _fld(" STYLEREF 1 \\s ", "5")
    label.append(_run("."))
    label += _fld(f" SEQ {kind} \\* ARABIC \\s 1 ", shown)
    p._p.append(_el("w:bookmarkStart", **{"w:id": bid,
                                          "w:name": bookmark}))
    p._p.append(_wrap_ins(label))
    p._p.append(_el("w:bookmarkEnd", **{"w:id": bid}))
    p._p.append(_wrap_ins([_run(" " + text)]))
    return p


def ref_runs(bookmark, shown):
    """An in-text cross reference to a caption's bookmark."""
    return _fld(f" REF {bookmark} \\h ", shown)


def para(doc, style, pieces):
    """A paragraph of prose, marked inserted.

    `pieces` is a list of strings and (bookmark, shown) pairs; a pair
    becomes a live cross reference instead of a typed number.
    """
    p = Paragraph(_el("w:p"), doc)
    ppr = _el("w:pPr")
    ppr.append(_el("w:pStyle", **{"w:val": style}))
    rpr = _el("w:rPr")
    rpr.append(_el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                               "w:date": DATE}))
    ppr.append(rpr)
    p._p.append(ppr)
    body = []
    for piece in pieces:
        if isinstance(piece, tuple):
            body += ref_runs(*piece)
        else:
            body.append(_run(piece))
    p._p.append(_wrap_ins(body))
    return p


def figure_par(doc, png, width_in=6.1):
    """The picture itself, in its own centred paragraph."""
    p = Paragraph(_el("w:p"), doc)
    ppr = _el("w:pPr")
    ppr.append(_el("w:jc", **{"w:val": "center"}))
    rpr = _el("w:rPr")
    rpr.append(_el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                               "w:date": DATE}))
    ppr.append(rpr)
    p._p.append(ppr)
    run = p.add_run()
    run.add_picture(os.path.join(FIGDIR, png), width=Inches(width_in))
    # add_picture puts the drawing inside the run it was given, and that
    # run is a direct child of the paragraph; wrapping it in <w:ins> is
    # what makes the picture itself part of the tracked change, so
    # rejecting the revision removes the figure rather than leaving it
    # orphaned above a deleted caption.
    for r in list(p._p):
        if r.tag == qn("w:r"):
            p._p.remove(r)
            p._p.append(_wrap_ins([r]))
    return p


def table_after(doc, anchor, headers, rows, widths=None):
    """A table whose every row is marked inserted."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    # WIDTHS ONLY BIND WITH AUTOFIT OFF. Left on, Word re-flows the
    # columns to the content and a header of three words is broken over
    # four lines while a column of integers keeps space it does not use.
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.paragraphs[0]._p.append(ins_run(h, bold=True))
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]
            c.paragraphs[0]._p.append(ins_run(str(v)))
    for r in t.rows:
        _ins_row(r)
    if widths:
        # AUTOFIT OFF IS NOT ENOUGH ON ITS OWN. Word only honours the
        # cell widths when the table layout is declared fixed; without
        # the element below it still measures the content and the first
        # column, which holds the longest text, ends up the narrowest.
        tblpr = t._tbl.find(qn("w:tblPr"))
        if tblpr is not None:
            for tag in ("w:tblLayout", "w:tblW"):
                old_el = tblpr.find(qn(tag))
                if old_el is not None:
                    tblpr.remove(old_el)
            tblpr.append(_el("w:tblW", **{
                "w:w": str(int(sum(widths) * 1440)), "w:type": "dxa"}))
            tblpr.append(_el("w:tblLayout", **{"w:type": "fixed"}))
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
                gc.set(qn("w:w"), str(int(w * 1440)))
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    anchor._p.addnext(t._tbl)
    return t


# ----------------------------------------------------------------- data
def _rows(pattern):
    out = []
    for f in sorted(glob.glob(os.path.join(OUT, pattern))):
        with open(f, encoding="utf-8") as fh:
            out += list(csv.DictReader(fh))
    return out


def h1_numbers():
    lat = _rows("central_latency.csv")
    outc = _rows("central_outcome.csv")
    by = {}
    for r in lat:
        by[(r["config"], int(r["n_regions"]))] = r
    ns = sorted({int(r["n_regions"]) for r in lat})
    table = []
    for n in ns:
        a = by.get(("A_distributed_open", n))
        b = by.get(("B_centralized_closed", n))
        if not (a and b):
            continue
        table.append((
            n,
            f"{float(a['ante_median']):,.0f}".replace(",", " "),
            f"{float(b['ante_median']):,.0f}".replace(",", " "),
            f"{float(b['ante_median']) / float(a['ante_median']):,.0f}"
            .replace(",", " "),
            f"{float(a['lat_median_ms']):.0f}",
            f"{float(b['lat_median_ms']):.0f}",
            f"{float(a['duty_median']) * 100:.2f}",
            f"{float(b['duty_median']) * 100:.2f}"))
    missed = sum(int(r["missed"]) for r in lat)
    out_by = collections.defaultdict(list)
    for r in outc:
        out_by[(r["config"], int(r["n_regions"]))].append(
            (float(r["burned_ha"]), float(r["j_total"])))
    return table, missed, out_by, ns


def h2_numbers(top=10):
    oat = _rows("prop_oat.csv")
    mor = _rows("prop_morris*.csv")
    d = collections.defaultdict(lambda: collections.defaultdict(list))
    for r in oat:
        k = (r["kind"], r["factor"])
        d[k]["dj"].append(abs(float(r["d_j_total"])))
        d[k]["flip"].append(100.0 * float(r["flip_rate"]))
        d[k]["shift"].append(float(r["act_shift"]))
    # Morris elementary effects, from consecutive points of a trajectory
    by_traj = collections.defaultdict(list)
    for r in mor:
        by_traj[r["traj"]].append(r)
    eff = collections.defaultdict(list)
    for _t, rs in by_traj.items():
        prev = None
        for r in rs:
            if prev is not None and abs(float(r["delta"])) > 1e-9:
                eff[(r["kind"], r["factor"])].append(
                    (float(r["d_j_total"]) - float(prev["d_j_total"]))
                    / float(r["delta"]))
            prev = r
    mu = {k: float(np.mean(np.abs(v))) for k, v in eff.items() if v}
    rank = sorted(d.items(), key=lambda kv: -float(np.mean(kv[1]["dj"])))
    label = {"feature": "feature bias", "mf": "membership",
             "weight": "cost weight"}
    table = []
    for (kind, name), v in rank[:top]:
        table.append((label[kind], name.replace("_", " "),
                      f"{float(np.mean(v['dj'])):.4f}",
                      f"{float(np.mean(v['flip'])):.1f}",
                      f"{float(np.mean(v['shift'])):.4f}",
                      f"{mu.get((kind, name), float('nan')):.3f}"))
    kindwise = {}
    for kind in ("feature", "mf", "weight"):
        fl = [np.mean(v["flip"]) for k, v in d.items() if k[0] == kind]
        dj = [np.mean(v["dj"]) for k, v in d.items() if k[0] == kind]
        kindwise[kind] = (float(np.mean(fl)), float(np.mean(dj)))
    top_flip = max(d.items(), key=lambda kv: float(np.mean(kv[1]["flip"])))
    return table, kindwise, len(oat), len(mor), top_flip, mu


# ----------------------------------------------------------------- main
def find(ps, needle, start=0):
    for i in range(start, len(ps)):
        if needle in ps[i].text:
            return i
    raise SystemExit(f"anchor not found: {needle!r}")


def main():
    inp, outp = sys.argv[1], sys.argv[2]
    doc = docx.Document(inp)
    ps = doc.paragraphs

    h1_tab, missed, out_by, ns = h1_numbers()
    h2_tab, kindwise, n_oat, n_mor, top_flip, mu = h2_numbers()

    n_lo, n_hi = ns[0], ns[-1]
    ratio_lo = h1_tab[0][3]
    ratio_hi = h1_tab[-1][3]
    duty_hi_a = h1_tab[-1][6]
    duty_hi_b = h1_tab[-1][7]

    # ---------------------------------------------- 1. the new 5.4 part
    anchor_i = find(ps, "Both settings are defined on the unit interval")
    a = ps[anchor_i]

    def after(p):
        return p

    head = Paragraph(_el("w:p"), doc)
    hppr = _el("w:pPr")
    hppr.append(_el("w:pStyle", **{"w:val": "Heading3"}))
    hrpr = _el("w:rPr")
    hrpr.append(_el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                                "w:date": DATE}))
    hppr.append(hrpr)
    head._p.append(hppr)
    head._p.append(_wrap_ins([_run(
        "Propagation of Observation, Membership and Weight "
        "Perturbations")]))
    a._p.addnext(head._p)
    cur = head

    def add(p):
        nonlocal cur
        cur._p.addnext(p._p)
        cur = p
        return p

    fkind, fname = top_flip[0]
    flip_max = float(np.mean(top_flip[1]["flip"]))

    add(para(doc, "ParText", [
        "The sweeps above move the settings of the decision layer. This "
        "subsection moves its inputs instead, and asks how far a "
        "perturbation travels before it reaches an order. Twenty-five "
        "factors are swept one at a time: a multiplicative bias of 25 "
        "per cent on each of the ten observation features, a widening "
        "or narrowing of 20 per cent on each of the ten feature "
        "membership partitions, and 25 per cent on each of the five "
        "weights of the cost decomposition. Two worlds are used for "
        f"every factor, which gives {n_oat} runs, and a Morris "
        "screening over the same factors adds "
        f"{n_mor} more so that a factor which matters only in company "
        "can be told from one that matters on its own."]))

    add(para(doc, "ParText", [
        "The term partitions of the five decision concepts were swept "
        "first and returned exactly zero on every metric. The reason is "
        "structural rather than numerical: the concepts are not "
        "fuzzified. Their term activations are produced by the "
        "hierarchy and read directly by the rule evaluator, so a stored "
        "partition for a concept is never consulted. The membership "
        "parameters a decision actually depends on are those of the ten "
        "features, which is also the surface the evolving stage edits, "
        "and the membership group reported here is that one."]))

    add(para(doc, "ParText", [
        "Two quantities are reported for every factor, because a "
        "perturbation that leaves the cost alone may still have changed "
        "what was ordered. The first is the shift in the total decision "
        "cost. The second is the decision flip rate, the share of "
        "region cycles whose dominant intervention family differs from "
        "the unperturbed run of the same world at the same cycle. ",
        ("_RefFigTornado", "5.24"),
        " ranks the factors by the first, and ",
        ("_RefFigFlip", "5.25"),
        " reports the second beside the movement each factor produces "
        "in the concept space."]))

    add(figure_par(doc, "fig5_27_tornado.png", 5.5))
    add(caption_par(doc, "Figure",
                    "Factors ranked by the shift they produce in the "
                    "decision cost, with the Morris screening mean "
                    "beside each factor", "_RefFigTornado", "24"))
    add(figure_par(doc, "fig5_28_flip_shift.png", 6.3))
    add(caption_par(doc, "Figure",
                    "Decision flips per factor, and where the "
                    "perturbation lands in the concept space",
                    "_RefFigFlip", "25"))

    add(para(doc, "ParText", [
        "Three readings follow. The cost weights move the score without "
        "moving the decision: they change the total cost by as much as "
        "the middle of the feature group does, yet they flip only "
        f"{kindwise['weight'][0]:.1f} per cent of region cycles, "
        "against "
        f"{kindwise['feature'][0]:.1f} per cent for the feature biases. "
        "This is the expected behaviour of a cost that scores a "
        "candidate rather than producing it, and it is worth stating "
        "because it bounds what a disagreement over priorities can do: "
        "it changes what the system reports, not what it orders. The "
        "membership partitions matter as much as the feature biases, "
        f"and the largest single effect on cost, {h2_tab[0][2]}, "
        f"belongs to the {h2_tab[0][1]} partition rather than to any "
        "observation. Membership parameters are the part of a fuzzy "
        "system that receives the least attention in maintenance, and "
        "the sweep says they should not. The loudest factor on the "
        f"decision itself is the {fname.replace('_', ' ')} "
        f"{'bias' if fkind == 'feature' else 'partition'}, which "
        f"changes the ordered intervention in {flip_max:.1f} per cent "
        "of region cycles while moving the cost by little, so a "
        "sensitivity study that watched only the outcome would have "
        "reported it as harmless."]))

    add(para(doc, "ParText", [
        "The Morris screening ranks the factors differently from the "
        "one-at-a-time sweep, which is the point of running both. The "
        "temporal urgency partition sits in the middle of the "
        "one-at-a-time ranking and carries the largest screening mean "
        "of any factor, so its effect is produced in combination with "
        "others rather than on its own. No factor is negligible on both "
        "designs, and none dominates: the decision cost responds to the "
        "whole input set rather than resting on one reading, which is "
        "what a confidence-gated observation model is built to "
        "achieve."]))

    cap = caption_par(doc, "Table",
                      "Propagation of the ten largest perturbations",
                      "_RefTabProp", "7")
    add(cap)
    table_after(doc, cap,
                ["Group", "Factor", "mean |dJ|", "Flips (%)",
                 "Concept shift", "Morris mean"],
                h2_tab, widths=[1.0, 1.5, 0.9, 0.8, 1.0, 0.9])

    # ------------------------------------------- 2. the 5.5.4 addition
    ps2 = doc.paragraphs
    tail_i = find(ps2, "Timeliness asks whether the layer decides fast")
    b = ps2[tail_i]
    cur = b
    # step past the Table 5.13 caption that follows it
    nxt = b._p.getnext()
    while nxt is not None and nxt.tag == qn("w:p"):
        txt = "".join(t.text or "" for t in nxt.iter(qn("w:t")))
        if txt.strip().startswith("Table"):
            cur = Paragraph(nxt, doc)
            nxt = nxt.getnext()
            continue
        break
    if nxt is not None and nxt.tag == qn("w:tbl"):
        class _A:
            def __init__(self, el):
                self._p = el
        cur = _A(nxt)

    def add2(p):
        nonlocal cur
        cur._p.addnext(p._p)
        cur = p
        return p

    add2(para(doc, "ParText", [
        "Both statements above describe one configuration. H1 compares "
        "it with another, a closed and centralized arrangement of the "
        "same architecture, and that comparison is an experiment rather "
        "than an argument. Three configurations were run on the same "
        "worlds, the same gates and the same simulator. The first is "
        "the distributed and open configuration used throughout this "
        "chapter. The second is centralized and closed: one inferential "
        "core receives every region's observation and reasons for each "
        "of them in turn, with no adaptation and with a rule base that "
        "answers every cell of the antecedent space in advance, which "
        "over five concepts with five terms each is 3 125 rules. The "
        "third is centralized and open, and it exists to separate the "
        "two properties, since a difference between the first two could "
        "otherwise be attributed to either."]))

    add2(para(doc, "ParText", [
        "Latency is not reported as raw wall time. Both configurations "
        "execute in one process on one machine, so a stopwatch around "
        "the cycle would charge the distributed configuration for "
        "reasoning that a deployment performs on separate nodes at the "
        "same time. The per-region reasoning is therefore timed region "
        "by region and composed under a stated deployment model: the "
        "distributed cycle costs the longest region plus the shared "
        "work, the centralized cycle the sum of the regions plus the "
        "same shared work. The shared part is the coordination, the "
        "composition and the two shadow forecasts of the acceptance "
        "test, which a centralized core must also perform."]))

    add2(figure_par(doc, "fig5_25_complexity.png", 5.9))
    add2(caption_par(doc, "Figure",
                     "Antecedent evaluations per decision against the "
                     "number of local agents", "_RefFigComplex", "26"))
    add2(figure_par(doc, "fig5_24_latency_scale.png", 5.9))
    add2(caption_par(doc, "Figure",
                     "Decision latency against the number of local "
                     "agents, with the cycle budget",
                     "_RefFigLatScale", "27"))

    add2(para(doc, "ParText", [
        "The complexity result is the decisive one. ",
        ("_RefFigComplex", "5.26"),
        " counts the antecedent conditions a decision evaluates, which "
        "is a property of the rule base and not of the machine it runs "
        "on. The distributed configuration evaluates about seventy of "
        f"them and that number does not move between {n_lo} and "
        f"{n_hi} agents, because each agent reasons only over its own "
        "region and the regions reason at the same time. The closed "
        "and centralized configuration evaluates 15 625 for a single "
        "region and grows in proportion, reaching 250 000 at sixteen. "
        f"The ratio runs from {ratio_lo} to {ratio_hi} across the "
        "sweep. Reducing the antecedent combination space is therefore "
        "not a claim about this implementation but a measured property "
        "of the concept hierarchy."]))

    add2(para(doc, "ParText", [
        "Latency does not separate the configurations, and the "
        "comparison is reported as it came out. ",
        ("_RefFigLatScale", "5.27"),
        " shows the closed and centralized configuration deciding "
        "faster at every scale, because sweeping a large rule base "
        "costs less than the shadow forecasts the adaptation stages "
        "run. The open configuration is slower for the reason it is "
        "open, not for the reason it is distributed, and the "
        "centralized open configuration confirms this by tracking the "
        "distributed one. What matters for the criterion is that the "
        "question does not arise: the slowest median decision of either "
        f"configuration uses {duty_hi_a} per cent of the twelve-minute "
        f"cycle against {duty_hi_b} per cent, and across every "
        f"configuration and scale no decision exceeded the cycle, "
        f"{missed} missed cycles in total. Timeliness is met with three "
        "orders of magnitude to spare on both sides, so it is not the "
        "axis on which the architectures differ."]))

    add2(para(doc, "ParText", [
        "The outcome is likewise not where the difference lies. Over "
        "three worlds at each scale the two configurations burn "
        "comparable area, and the centralized open configuration "
        "reproduces the distributed one exactly, which is the expected "
        "result: centralizing the inference changes where a decision is "
        "computed, not what it is. The closed catalogue is not a poor "
        "controller. It is an unmaintainable one, and that is the "
        "claim the complexity figure supports."]))

    cap2 = caption_par(doc, "Table",
                       "Distributed against closed and centralized, "
                       "over the scale sweep", "_RefTabCentral", "15")
    add2(cap2)
    # THE TABLE IS TRANSPOSED against the shape of the campaign. Eight
    # columns of long headings leave every heading broken over four
    # lines on this page width, and the reader is comparing quantities
    # across scales rather than scales across quantities.
    _cols = ["Quantity"] + [f"N = {r[0]}" for r in h1_tab]
    _body = [
        ["Antecedent evaluations, distributed"] + [r[1] for r in h1_tab],
        ["Antecedent evaluations, closed centralized"]
        + [r[2] for r in h1_tab],
        ["Ratio"] + [r[3] for r in h1_tab],
        ["Median latency, distributed (ms)"] + [r[4] for r in h1_tab],
        ["Median latency, closed centralized (ms)"]
        + [r[5] for r in h1_tab],
        ["Share of the cycle, distributed (%)"]
        + [r[6] for r in h1_tab],
        ["Share of the cycle, closed centralized (%)"]
        + [r[7] for r in h1_tab]]
    table_after(doc, cap2, _cols, _body,
                widths=[2.3] + [0.72] * len(h1_tab))

    # ----------------------------------------------------- 3. H1 itself
    for p in doc.paragraphs:
        if not p.text.strip().startswith("H1 ("):
            continue
        old = ("the distributed implementation keeps decision latency "
               "compatible with the decision cycle as the number of "
               "local agents grows from one to sixteen")
        new = ("the distributed implementation holds that space "
               "constant as the number of local agents grows from one "
               "to sixteen, where a closed and centralized "
               "configuration of the same architecture grows in "
               "proportion, both configurations deciding well within "
               "the decision cycle")
        for r in list(p._p.iter(qn("w:r"))):
            ts = r.findall(qn("w:t"))
            txt = "".join(t.text or "" for t in ts)
            if old not in txt:
                continue
            parent = r.getparent()
            d = _el("w:del", **{"w:id": _nid(), "w:author": AUTHOR,
                                "w:date": DATE})
            parent.replace(r, d)
            d.append(r)
            for t in ts:
                t.tag = qn("w:delText")
            d.addnext(ins_run(txt.replace(old, new)))
            break
        break

    doc.save(outp)
    print("written:", outp)
    print(f"   5.4 subsection: {n_oat} OAT runs, {n_mor} Morris rows")
    print(f"   5.5.4: scale sweep {ns}, {missed} missed cycles")
    print("   H1 narrowed to the measured claim")


if __name__ == "__main__":
    main()
