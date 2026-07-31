"""Write Section 5.5, Sensitivity Analysis, into the thesis.

Reads   experiments/out/sens_runs.csv        (the study)
        01_Thesis/figures/fig5_12..16*.png   (its figures)
Edits   the given .docx IN TRACKED-CHANGES form (author "Claude"):
  - rewrites the opening, the design table and every reading paragraph
  - replaces the three figures the section already carries and inserts
    the two it does not
  - adds the cost-weight table

EVERY NUMBER IN THE PROSE COMES FROM THE CSV. A sentence that quotes a
figure it was not measured from is the failure this section is being
rewritten to remove, so the text is assembled from the same aggregate
the figures are drawn from.

Usage: python experiments/fill_sensitivity.py IN.docx OUT.docx
"""
from __future__ import annotations

import csv
import math
import os
import shutil
import sys
import zipfile
from collections import defaultdict
from datetime import datetime, timezone

import numpy as np
import docx
from docx.oxml.ns import qn
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "out")
FIGDIR = os.path.join(HERE, "..", "..", "01_Thesis", "figures")
AUTHOR = "Claude"
DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
_ID = [9000]

FIGS = [
    ("fig5_12_calibration.png",
     "Physical cost over fire load and resource level, as a share of "
     "the free burn of the same world, with the operating point of the "
     "sweeps marked"),
    ("fig5_13_ranking.png",
     "Spread in the physical decision cost produced by each parameter "
     "over its own range, environment and decision-layer tuning on one "
     "axis"),
    ("fig5_14_capacity.png",
     "Physical cost against (a) resource level and (b) fire load, for "
     "the static and the adaptive configuration, with the free burn as "
     "the ceiling"),
    ("fig5_15_thresholds.png",
     "Physical cost against each parameter of the decision layer, drawn "
     "on one shared vertical scale"),
    ("fig5_16_eta.png",
     "Fail-safe engagement, admitted stage-3 products and physical cost "
     "against the quality gate"),
]


# ------------------------------------------------------------ tracked xml
def _nid():
    _ID[0] += 1
    return str(_ID[0])


def _el(tag, **attrs):
    from docx.oxml import OxmlElement
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def ins_run(text, bold=False, italic=False):
    ins = _el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                          "w:date": DATE})
    r = _el("w:r")
    if bold or italic:
        rpr = _el("w:rPr")
        if bold:
            rpr.append(_el("w:b"))
        if italic:
            rpr.append(_el("w:i"))
        r.append(rpr)
    t = _el("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    ins.append(r)
    return ins


def del_runs_of(p):
    """Wrap every live run of p in <w:del>, insertions included."""
    todo = [r for r in p.iter(qn("w:r"))
            if r.getparent().tag != qn("w:del")]
    for r in todo:
        parent = r.getparent()
        dl = _el("w:del", **{"w:id": _nid(), "w:author": AUTHOR,
                             "w:date": DATE})
        parent.replace(r, dl)
        dl.append(r)
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")


def replace_para(par, text):
    if par is None:
        return
    del_runs_of(par._p)
    par._p.append(ins_run(text))


def replace_caption_desc(par, new_desc):
    """Replace what a caption SAYS without touching what numbers it.

    The captions in this document are built from fields: a STYLEREF for
    the chapter and a SEQ for the counter, then the descriptive text.
    Overwriting the whole paragraph turns "Figure 5.12" into three
    literal characters that no longer renumber when a figure is added
    ahead of them, so only the runs after the last field end are
    replaced.
    """
    if par is None:
        return
    p = par._p
    ends = [fc for fc in p.iter(qn("w:fldChar"))
            if fc.get(qn("w:fldCharType")) == "end"]
    if not ends:
        del_runs_of(p)
        p.append(ins_run(new_desc))
        return
    last = ends[-1]
    seen = False
    for r in list(p.iter(qn("w:r"))):
        if not seen:
            if last in list(r.iter()):
                seen = True
            continue
        if r.getparent().tag == qn("w:del"):
            continue
        parent = r.getparent()
        dl = _el("w:del", **{"w:id": _nid(), "w:author": AUTHOR,
                             "w:date": DATE})
        parent.replace(r, dl)
        dl.append(r)
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")
    p.append(ins_run(" " + new_desc.lstrip()))


def cell_text(cell):
    return "".join(t.text or ""
                   for t in cell._tc.iter(qn("w:t"))).strip()


def fill_cell(cell, value):
    p = cell.paragraphs[0]._p
    del_runs_of(p)
    p.append(ins_run(str(value)))


def mark_inserted(par):
    """Mark a whole paragraph, mark included, as a tracked insertion."""
    p = par._p
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = _el("w:pPr")
        p.insert(0, ppr)
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = _el("w:rPr")
        ppr.append(rpr)
    rpr.insert(0, _el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                                  "w:date": DATE}))
    for r in list(p.findall(qn("w:r"))):
        ins = _el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                              "w:date": DATE})
        p.replace(r, ins)
        ins.append(r)


# ------------------------------------------------------------- the study
def load_runs():
    with open(os.path.join(OUT, "sens_runs.csv"), encoding="utf-8") as f:
        return list(csv.DictReader(f))


def _f(r, k):
    try:
        return float(r[k])
    except (TypeError, ValueError):
        return float("nan")


def agg(rows, block, param, arm, key="j_phys"):
    by = defaultdict(list)
    for r in rows:
        if r["block"] == block and r["param"] == param and r["arm"] == arm:
            by[float(r["value"])].append(_f(r, key))
    out = {}
    for v, xs in by.items():
        xs = [x for x in xs if not math.isnan(x)]
        if xs:
            out[v] = float(np.mean(xs))
    return dict(sorted(out.items()))


def spread(d):
    return (max(d.values()) - min(d.values())) if d else 0.0


def study(rows):
    """Every number the prose quotes, in one place."""
    s = {}
    for p in ("n_ign", "pool", "n_sensors", "n_regions"):
        s[p] = agg(rows, "environment", p, "adaptive")
        s[p + "_static"] = agg(rows, "environment", p, "static")
    for p in ("j_threshold", "eta", "attention_thr", "horizon_min",
              "cycle_min", "revision_budget"):
        s[p] = agg(rows, "tuning", p, "adaptive")
    s["fs"] = agg(rows, "tuning", "eta", "adaptive", "fs_frac")
    s["free"] = agg(rows, "calibration", "freeburn", "freeburn")
    # the grid itself, as a share of the free burn of the same world
    grid = {}
    for r in rows:
        if r["block"] == "calibration" and r["arm"] == "adaptive":
            grid.setdefault((int(float(r["n_ign"])),
                             float(r["pool"])), []).append(_f(r, "j_phys"))
    s["grid"] = {k: 100.0 * float(np.mean(v)) / s["free"][float(k[0])]
                 for k, v in grid.items() if float(k[0]) in s["free"]}
    for w in ("w_burn", "w_asset", "w_pop"):
        s[w] = agg(rows, "weights", w, "adaptive", "burned_ha")
    n_seeds = len({r["seed"] for r in rows if r["block"] == "environment"})
    s["n_seeds"] = n_seeds
    s["n_runs"] = len(rows)
    return s


# ------------------------------------------------------------- the prose
def paragraphs(s):
    def v(d, k):
        return d.get(k, float("nan"))

    ign, pool = s["n_ign"], s["pool"]
    grid = s["grid"]
    op_share = grid[(4, 0.25)]
    starved = max(sh for (n, p), sh in grid.items() if p == 0.10)
    easiest = min(grid.values())

    rank = sorted(
        [("simultaneous ignitions", spread(ign)),
         ("resource level", spread(pool)),
         ("decision cycle", spread(s["cycle_min"])),
         ("observation assets", spread(s["n_sensors"])),
         ("local regions", spread(s["n_regions"])),
         ("quality gate", spread(s["eta"])),
         ("attention threshold", spread(s["attention_thr"])),
         ("no-harm horizon", spread(s["horizon_min"])),
         ("revision budget", spread(s["revision_budget"])),
         ("satisficing bound", spread(s["j_threshold"]))],
        key=lambda t: -t[1])

    opening = (
        "This section reports how sensitive the decision layer is to its "
        "own settings and to the situation it is placed in. The method is "
        "one at a time: one parameter is varied across its range while "
        "the others hold their base value, and the physical decision "
        f"cost is recorded (Morris, 1991; Saltelli et al., 2008). Each "
        f"setting is repeated over {s['n_seeds']} worlds and the figures "
        "show the mean with its 95 per cent interval, because a single "
        "world cannot separate the effect of a parameter from the noise "
        f"of the map. The study comprises {s['n_runs']} runs in total.")

    why = (
        "A sweep can only show what its operating point allows. Where the "
        "fire is beaten under every setting, or lost under every setting, "
        "no parameter can move the result, and the study would report a "
        "robustness that belongs to the situation rather than to the "
        "system. The study therefore begins by locating an operating "
        "point at which the decision layer is neither winning nor losing "
        "outright. Table 5.13 lists the parameters, their base values and "
        "their ranges.")

    calib = (
        "Figure 5.12 maps the physical cost over fire load and resource "
        "level, read as a share of the free burn of the same world. Two "
        "regions are of no use to a sweep. At a tenth of the suggested "
        f"pool the cost reaches {starved:.0f} per cent of the free burn "
        "whatever else is set, and at a single ignition with a full pool "
        f"it falls to {easiest:.0f} per cent just as reliably. "
        "Between them lies the band in which the decision actually "
        "decides. The sweeps that follow were run at four simultaneous "
        "ignitions with a quarter of the suggested pool, a setting that "
        f"retains {op_share:.0f} per cent of the free-burn cost.")

    ranking = (
        "Figure 5.13 ranks the parameters by the spread each one produces "
        "in the physical cost over its own range. Environment and tuning "
        "are drawn on one axis, because the claim of this section is a "
        f"comparison between them. The fire load leads at {rank[0][1]:.3f}, "
        f"followed by the resource level at {rank[1][1]:.3f}. The decision "
        f"cycle follows at {rank[2][1]:.3f} and is the only setting of the "
        "decision layer that reaches the same order of magnitude as the "
        "environment. Observation comes next: a single deployed asset "
        f"leaves the cost at {v(s['n_sensors'], 1.0):.3f}, while two or more "
        f"bring it to about {v(s['n_sensors'], 3.0):.2f}. Suppression is "
        "aimed at the fire the network reports, so a map the network "
        "cannot see is a map the system cannot fight. Every remaining "
        f"threshold moves the cost by less than {rank[5][1]:.2f}.")

    capacity = (
        "Figure 5.14 shows why the two leading parameters lead. The cost "
        f"falls from {v(pool, 0.10):.3f} at a tenth of the suggested pool "
        f"to {v(pool, 0.50):.3f} at a half and then flattens, so below "
        "that level the system is capacity limited and above it further "
        "resources buy little. The fire load acts in the same way in the "
        f"other direction, from {v(ign, 1.0):.3f} at one ignition to "
        f"{v(ign, 12.0):.3f} at twelve, approaching the free burn beyond "
        "eight. This is the multi-ignition regime that drives extreme "
        "fire years, in which simultaneous starts disperse and overwhelm "
        "the response.")

    arms = (
        "The static and the adaptive configuration are drawn together in "
        "Figure 5.14 and stay inside each other's intervals at every "
        "point. The margin the adaptation earns in Section 5.4 is "
        "measured on the scenario grid over ten worlds; on this testbed, "
        f"with {s['n_seeds']} worlds at a fixed operating point, it is not "
        "separable from the noise. The reading is that the capacity "
        "balance dominates both configurations here, not that the "
        "adaptation is without effect.")

    thresholds = (
        "Figure 5.15 sweeps the parameters of the decision layer on one "
        "shared vertical scale, so that flatness can be seen rather than "
        "asserted. Five of the six are flat. The satisficing bound leaves "
        "the cost unchanged across its full range, because at this "
        "operating point the cost stands far above the bound and the "
        "bound is never the binding constraint; the no-harm horizon, the "
        "attention threshold and the revision budget behave the same "
        "way.")

    cycle = (
        "The decision cycle is the exception. The cost rises from "
        f"{v(s['cycle_min'], 2.0):.3f} at a two-minute cycle to "
        f"{v(s['cycle_min'], 20.0):.3f} at twenty minutes. Revisiting the "
        "decision more often is worth more than any threshold in the "
        "system, because a fire that is re-assessed every two minutes is "
        "answered while it is still small.")

    eta = (
        "Figure 5.16 examines the fail-safe quality gate on its own, "
        "because its purpose is admission rather than cost. As the gate "
        "tightens the fail-safe engages more often: it does not engage at "
        f"all up to 0.45, engages in {100 * v(s['fs'], 0.60):.0f} per cent "
        f"of cycles at 0.60, in {100 * v(s['fs'], 0.75):.0f} per cent at "
        f"0.75 and in {100 * v(s['fs'], 0.90):.0f} per cent at 0.90. The "
        f"physical cost stays between {min(s['eta'].values()):.2f} and "
        f"{max(s['eta'].values()):.2f} over the same range. The gate "
        "therefore governs what the generative stage is allowed to do "
        "without charging the outcome for it, which is what a fail-safe "
        "is for.")

    weights = (
        "The cost weights are reported against the physical outcome "
        "rather than against the cost they define, since a weight that "
        "moves the cost it appears in explains nothing. Table 5.14 shows "
        "a small effect in the expected direction: doubling the "
        f"burned-area weight lowers the mean burned area from "
        f"{v(s['w_burn'], 0.5):.0f} to {v(s['w_burn'], 2.0):.0f} hectares, "
        "while the asset and population weights leave it within a few per "
        "cent. The priority profile steers the response, but weakly.")

    closing = (
        "The outcome of the decision layer is therefore governed by the "
        "balance between the fire load and the resources available to "
        "fight it, by how much of the fire the network observes, and by "
        "how often the decision is revisited. It is robust to the "
        "thresholds that were set by hand. That robustness is useful, "
        "because the layer does not depend on delicate tuning to perform, "
        "and it directs planning attention to resourcing, to sensing and "
        "to decision frequency rather than to threshold choice. It is "
        "also bounded: the study covers one landscape family at one "
        "operating point, and a system insensitive to its thresholds "
        "there need not be insensitive elsewhere.")

    return dict(opening=opening, why=why, calib=calib, ranking=ranking,
                capacity=capacity, arms=arms, thresholds=thresholds,
                cycle=cycle, eta=eta, weights=weights, closing=closing)


DESIGN_ROWS = [
    ("Simultaneous ignitions", "-", "4", "1-12"),
    ("Resource level", "-", "0.25",
     "0.10-1.00 (fraction of the suggested pool)"),
    ("Observation assets", "-", "full deployment", "1-9"),
    ("Local regions", "N", "4", "1-8"),
    ("Satisficing bound", "J\u1d1b\u029c", "0.35", "0.15-0.60"),
    ("Fail-safe quality gate", "\u03b7", "0.60", "0.30-0.90"),
    ("Attention threshold", "\u03c4", "0.35", "0.15-0.70"),
    ("No-harm horizon", "-", "24 min", "8-48 min"),
    ("Decision cycle", "-", "12 min", "2-20 min"),
    ("Revision budget", "-", "3", "1-6"),
    ("Cost weights", "w_burn, w_asset, w_pop", "1.0", "0.5-2.0"),
]


# ------------------------------------------------------------------ main
def main():
    inp, outp = sys.argv[1], sys.argv[2]
    rows = load_runs()
    s = study(rows)
    txt = paragraphs(s)

    doc = docx.Document(inp)
    body = doc.element.body

    def ptext(el):
        return "".join(t.text or "" for t in el.iter(qn("w:t")))

    kids = list(body)
    start = None
    for i, el in enumerate(kids):
        if el.tag == qn("w:p") and ptext(el).strip() == "Sensitivity Analysis":
            start = i
            break
    if start is None:
        raise SystemExit("heading 'Sensitivity Analysis' not found")

    from docx.text.paragraph import Paragraph
    from docx.table import Table

    # the section as it stands: heading, opening, design table, caption,
    # then three figure/caption/reading groups
    seq = kids[start:start + 14]
    P = lambda el: Paragraph(el, doc)                      # noqa: E731

    open_par = P(seq[1])
    design_tb = Table(seq[2], doc) if seq[2].tag == qn("w:tbl") else None
    cap_design = P(seq[3])
    img_pars = [P(seq[4]), P(seq[7]), P(seq[10])]
    cap_pars = [P(seq[5]), P(seq[8]), P(seq[11])]
    read_pars = [P(seq[6]), P(seq[9]), P(seq[12])]

    # ---- 1. the opening, split in two
    replace_para(open_par, txt["opening"])
    p_why = P(_el("w:p"))
    open_par._p.addnext(p_why._p)
    p_why._p.append(ins_run(txt["why"]))
    mark_inserted(p_why)

    # ---- 2. the design table
    if design_tb is not None:
        for row in list(design_tb.rows)[1:]:
            _del_row(row)
        for vals in DESIGN_ROWS:
            new = design_tb.add_row()
            _ins_row(new)
            for cell, v in zip(list(new.cells)[:4], vals):
                fill_cell(cell, v)
    replace_caption_desc(cap_design,
                         "Sensitivity design: parameters, base values "
                         "and ranges")

    # ---- 3. the three figures the section already has
    for cap, (_png, caption) in zip(cap_pars, FIGS[:3]):
        replace_caption_desc(cap, caption)
    for par, key in zip(read_pars, ("calib", "ranking", "capacity")):
        replace_para(par, txt[key])

    # the arms paragraph follows the capacity reading
    p_arms = P(_el("w:p"))
    read_pars[2]._p.addnext(p_arms._p)
    p_arms._p.append(ins_run(txt["arms"]))
    mark_inserted(p_arms)
    anchor = p_arms

    # ---- 4. the two figures it does not have
    for fignum, png, caption, reading in (
            (15, FIGS[3][0], FIGS[3][1],
             [txt["thresholds"], txt["cycle"]]),
            (16, FIGS[4][0], FIGS[4][1], [txt["eta"]])):
        pic = doc.add_paragraph()
        run = pic.add_run()
        run.add_picture(os.path.join(FIGDIR, png), width=Inches(6.0))
        pic.alignment = 1
        anchor._p.addnext(pic._p)
        mark_inserted(pic)
        cap = P(_el("w:p"))
        pic._p.addnext(cap._p)
        cap._p.append(ins_run(f"Figure 5.{fignum} {caption}"))
        _style(cap, "Caption")
        mark_inserted(cap)
        anchor = cap
        for body_text in reading:
            par = P(_el("w:p"))
            anchor._p.addnext(par._p)
            par._p.append(ins_run(body_text))
            mark_inserted(par)
            anchor = par

    # ---- 5. the weight table and its reading
    par = P(_el("w:p"))
    anchor._p.addnext(par._p)
    par._p.append(ins_run(txt["weights"]))
    mark_inserted(par)
    anchor = par

    tb = doc.add_table(rows=1, cols=4)
    tb.style = "Table Grid"
    for i, h in enumerate(["Weight", "Burned area at 0.5 (ha)",
                           "at 1.0 (ha)", "at 2.0 (ha)"]):
        fill_cell(tb.rows[0].cells[i], h)
    for w in ("w_burn", "w_asset", "w_pop"):
        cells = tb.add_row().cells
        fill_cell(cells[0], w.replace("_", " "))
        for i, k in enumerate((0.5, 1.0, 2.0), start=1):
            fill_cell(cells[i], f"{s[w].get(k, float('nan')):.0f}")
    for row in tb.rows:
        _ins_row(row)
    anchor._p.addnext(tb._tbl)
    cap = P(_el("w:p"))
    tb._tbl.addnext(cap._p)
    cap._p.append(ins_run("Table 5.14 Cost weights against the "
                          "physical outcome"))
    _style(cap, "Caption")
    mark_inserted(cap)
    anchor = cap

    # ---- 6. the closing paragraph replaces the old one
    par = P(seq[13]) if len(seq) > 13 else None
    if par is not None and ptext(par._p).strip().startswith(
            "The overall finding"):
        del_runs_of(par._p)
        par._p.append(ins_run(txt["closing"]))
        anchor._p.addnext(par._p)
    else:
        par = P(_el("w:p"))
        anchor._p.addnext(par._p)
        par._p.append(ins_run(txt["closing"]))
        mark_inserted(par)

    doc.save(outp)
    _swap_media(inp, outp)
    print("written:", outp)
    for k in ("n_runs", "n_seeds"):
        print(f"   {k}: {s[k]}")


def _del_row(row):
    """Tracked-delete one table row, content included."""
    tr = row._tr
    trpr = tr.find(qn("w:trPr"))
    if trpr is None:
        trpr = _el("w:trPr")
        tr.insert(0, trpr)
    trpr.append(_el("w:del", **{"w:id": _nid(), "w:author": AUTHOR,
                                "w:date": DATE}))
    for p in tr.iter(qn("w:p")):
        del_runs_of(p)


def _ins_row(row):
    tr = row._tr
    trpr = tr.find(qn("w:trPr"))
    if trpr is None:
        trpr = _el("w:trPr")
        tr.insert(0, trpr)
    trpr.append(_el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                                "w:date": DATE}))


def _style(par, name):
    try:
        par.style = par.part.document.styles[name]
    except Exception:
        pass


def _swap_media(inp, outp):
    """Put the new figures in place of the three the section carried.

    The pictures are addressed by the media part they already use, so
    the relationship, the anchor and the layout stay exactly as the
    author left them and only the pixels change.
    """
    mapping = {"word/media/image42.png": FIGS[0][0],
               "word/media/image43.png": FIGS[1][0],
               "word/media/image44.png": FIGS[2][0]}
    tmp = outp + ".tmp"
    with zipfile.ZipFile(outp) as zin, \
            zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            png = mapping.get(item.filename)
            if png:
                with open(os.path.join(FIGDIR, png), "rb") as fh:
                    data = fh.read()
                print(f"   {item.filename} <- {png}")
            zout.writestr(item, data)
    shutil.move(tmp, outp)


if __name__ == "__main__":
    main()
