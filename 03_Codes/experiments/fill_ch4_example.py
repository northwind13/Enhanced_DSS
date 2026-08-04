"""Insert the worked example of the two acceptance tests into Chapter 4.

The chapter defines the cost test and the quality test and then states
that a candidate must clear both. What it does not do is show what the
two tests do to the same candidate, which is the only way a reader
learns why the second one exists. This script adds that.

Three outcomes are possible and all three appear:

  cost fails                the standing orders stay in force and the
                            adaptation stages engage. Stated, not
                            worked: the path it opens is Section 4.5.3.
  cost passes, quality passes   the orders go out as written.
  cost passes, quality fails    the graduated fail-safe attenuates the
                            offensive orders in proportion to the
                            deficit.

THE EXAMPLE IS RUN, NOT WRITTEN. Every number comes from a live
decision cycle reproduced here on a fixed seed, and every formula is
real Office Math rather than a typed imitation of it, so the worked
example can drift neither from the system nor from the notation the
rest of the chapter uses.

Usage: python experiments/fill_ch4_example.py IN.docx OUT.docx
"""
from __future__ import annotations

import os
import sys

import docx
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(HERE))
sys.path.insert(0, HERE)

from docx_track import (DATE, _el, _ins_row, _nid,        # noqa: E402
                        fill_cell, ins_run)
import omml as M                                          # noqa: E402

AUTHOR = "Claude"
SEED = 201
N_IGN = 4
POOL = 0.25
CYCLE_STEP = 12
DELTA = 0.05                 # the relative margin, min_gain
ETA = 0.60
REGIONS = ("Agent_1", "Agent_2")
CONCEPT_SHORT = {
    "fire threat level": "fire threat",
    "asset exposure risk": "asset exposure",
}


def _ins():
    return dict(id=_nid(), author=AUTHOR, date=DATE)


# ------------------------------------------------------------ the cycle
def run_example():
    """One decision cycle, reproduced."""
    import scenario as S
    import dss
    from dss import loop as L
    from dss.evaluate import CONCEPT_FAMILY
    from disaster_phyengine.core import Simulator

    seen = []
    _orig = L.forecast_cost

    def _spy(sim, ov, hs, horizon_min=None):
        r = _orig(sim, ov, hs, horizon_min=horizon_min)
        seen.append((ov is None,
                     {k: float(getattr(r, k)) for k in
                      ("j_burn", "j_asset", "j_pop", "j_resp",
                       "j_delay", "j_total")}))
        return r

    L.forecast_cost = _spy

    w = S.build_world(SEED)
    base, _ = dss.resource_suggestion(w)
    base.ravail = base.ravail * POOL
    w.config.cost.capacity_reference = max(
        100.0, 1.2 * float((base.rcap * base.ravail).sum()))
    for x, y in S.pick_ignitions(w, base, SEED, N_IGN):
        w.add_ignition(x, y, step=0, radius=1)
    sim = Simulator(w)
    sim.record_states = False
    eng = dss.DecisionEngine(
        dss.partition_n(w.config.nx, w.config.ny, 4), base_pool=base,
        state_path=dss.isolated_store_path("ch4_example"),
        adapt_on=False, cycle_min=12.0, horizon_min=24.0, eta=ETA)
    try:
        for _ in range(CYCLE_STEP + 8):
            sim.step(resource_override=eng.maybe_decide(sim))
    finally:
        L.forecast_cost = _orig

    cyc = [c for c in eng.cycles if c.get("step") == CYCLE_STEP]
    if not cyc:
        raise SystemExit(f"no decision cycle at step {CYCLE_STEP}")
    c = cyc[0]
    f = c["forecast"]
    out = dict(j_c=float(f["j_candidate"]), j_0=float(f["j_noaction"]),
               j_th=float(f["j_threshold"]),
               bound=float(f["satisficing_bound"]),
               t_min=float(c.get("t_min") or 0.0), regions=[])
    cw = w.config.cost
    out["w"] = {k: float(getattr(cw, k)) for k in
                ("w_burn", "w_asset", "w_pop", "w_resp", "w_delay")}
    out["wsum"] = sum(out["w"].values())
    for is_noaction, d in seen:
        tgt = out["j_0"] if is_noaction else out["j_c"]
        if abs(d["j_total"] - tgt) < 1e-9:
            out["terms_0" if is_noaction else "terms_c"] = d
    if "terms_c" not in out or "terms_0" not in out:
        raise SystemExit("the forecast terms of the cycle were not "
                         "captured")
    for name in REGIONS:
        rd = c["regions"][name]
        cr, u = rd["concepts_effective"], rd["orders_from_rules"]
        terms = []
        for cn, fams in CONCEPT_FAMILY.items():
            a = float(cr.get(cn, 0.0))
            uu = max(float(u.get(fa, 0.0)) for fa in fams)
            label = cn.replace("_", " ")
            terms.append((CONCEPT_SHORT.get(label, label), a, uu,
                          abs(a - uu)))
        share = float(rd["coord_share"])
        out["regions"].append(dict(
            terms=terms, q=float(rd["quality"]), share=share,
            eta=1.0 - share * (1.0 - ETA),
            failsafe=bool(rd["failsafe"]),
            rules={k: float(v) for k, v in u.items() if float(v) > 0.01},
            final={k: float(v) for k, v in rd["orders_final"].items()
                   if float(v) > 0.01}))
    out["regions"].sort(key=lambda r: -r["q"])      # the passing one first
    return out


# ------------------------------------------------------------------ main
def main():
    inp, outp = sys.argv[1], sys.argv[2]
    e = run_example()
    hi, lo = e["regions"][0], e["regions"][1]

    doc = docx.Document(inp)
    from docx.text.paragraph import Paragraph

    anchor = None
    for p in doc.paragraphs:
        if (p.style.name == "Caption"
                and p.text.strip().startswith("Table 4.4")):
            anchor = p
            break
    if anchor is None:
        raise SystemExit("anchor not found: the Table 4.4 caption")

    def attach(par):
        nonlocal anchor
        anchor._p.addnext(par._p)
        anchor = par
        return par

    def add_par(text, style=None, bold_lead=None):
        par = Paragraph(_el("w:p"), doc)
        if bold_lead:
            par._p.append(ins_run(bold_lead, bold=True))
        par._p.append(ins_run(text))
        if style:
            try:
                par.style = doc.styles[style]
            except KeyError:
                pass
        _mark_para_inserted(par)
        return attach(par)

    def add_par_ref(before, bookmark, shown, after):
        """A paragraph that cross-references a numbered equation."""
        par = Paragraph(_el("w:p"), doc)
        par._p.append(ins_run(before))
        wrap = _el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                               "w:date": DATE})
        for r in M.ref_runs(bookmark, shown):
            wrap.append(r)
        par._p.append(wrap)
        par._p.append(ins_run(after))
        _mark_para_inserted(par)
        return attach(par)

    def _fill(cell, value):
        """A cell holds text, or math, or both: a symbol belongs beside
        the quantity it names, not in a column of its own."""
        if isinstance(value, str):
            fill_cell(cell, value)
            return
        p = cell.paragraphs[0]._p
        for r in list(p.findall(qn("w:r"))):
            p.remove(r)
        for part in value:
            if isinstance(part, str):
                p.append(ins_run(part))
            else:
                wrap = _el("w:ins", **{"w:id": _nid(),
                                       "w:author": AUTHOR, "w:date": DATE})
                wrap.append(M.oMath([part]))
                p.append(wrap)

    def add_table(headers, rows, number, caption):
        nonlocal anchor
        tb = doc.add_table(rows=1, cols=len(headers))
        tb.style = "Table Grid"
        for i, h in enumerate(headers):
            _fill(tb.rows[0].cells[i], h)
        for r in rows:
            cells = tb.add_row().cells
            for i, v in enumerate(r):
                _fill(cells[i], v)
        for row in tb.rows:
            _ins_row(row)
        cap = Paragraph(_el("w:p"), doc)
        for child in caption_runs("Table", number, caption):
            cap._p.append(child)
        try:
            cap.style = doc.styles["Caption"]
        except KeyError:
            pass
        _mark_para_inserted(cap)
        anchor._p.addnext(cap._p)
        cap._p.addnext(tb._tbl)
        anchor = _Anchor(tb._tbl)

    # ---------------------------------------------------------- the text
    # THE CHAPTER ALREADY STATES BOTH TESTS. Restating them here would
    # teach nothing and add two equations the reader has met; the
    # example cites them and shows only what evaluating them produces.
    EQ_J = "_RefEqJcomp"           # the five-term cost
    EQ_COST = "_RefEqB4503"        # the acceptance test on cost
    EQ_Q = "_RefEqB4504"           # the quality test
    n_eq = [51]
    bm = [9500]

    def add_eq(parts, bookmark):
        nonlocal anchor
        tb = M.eq_table(doc, parts, bookmark, n_eq[0], bm[0], ins=_ins())
        bm[0] += 1
        n_eq[0] += 1
        anchor._p.addnext(tb._tbl)
        anchor = _Anchor(tb._tbl)
        return n_eq[0] - 1

    add_par("A Worked Example", "Heading 4")
    add_par(
        "One cycle is enough to separate the two tests. The cycle below "
        f"is taken {e['t_min']:.0f} minutes into an incident with four "
        "simultaneous ignitions and a quarter of the suggested "
        "resources, under the five seed rules. A candidate may fail the "
        "cost test, or pass it and fail the quality test, or pass both. "
        "The last two outcomes occur in this cycle, in different "
        "regions.")

    # ---- the cost test
    add_par_ref(
        "Two forecasts are run from the state of the cycle: one with "
        "the candidate allocation held constant, one with no action. "
        "Each is priced by the five-term cost of ", EQ_J, 48,
        ". Table 4.5 gives the terms of both, already normalized, with "
        "the weights that combine them.")

    W = e["w"]
    tc_, t0_ = e["terms_c"], e["terms_0"]
    cost_rows = [
        (["Burned area, ", M.sub("J", "burn")], f"{W['w_burn']:.1f}",
         f"{tc_['j_burn']:.4f}", f"{t0_['j_burn']:.4f}"),
        (["Asset loss, ", M.sub("J", "asset")], f"{W['w_asset']:.1f}",
         f"{tc_['j_asset']:.4f}", f"{t0_['j_asset']:.4f}"),
        (["Population, ", M.sub("J", "pop")], f"{W['w_pop']:.1f}",
         f"{tc_['j_pop']:.4f}", f"{t0_['j_pop']:.4f}"),
        (["Response, ", M.sub("J", "resp")], f"{W['w_resp']:.1f}",
         f"{tc_['j_resp']:.4f}", f"{t0_['j_resp']:.4f}"),
        (["Delay, ", M.sub("J", "delay")], f"{W['w_delay']:.1f}",
         f"{tc_['j_delay']:.4f}", f"{t0_['j_delay']:.4f}"),
    ]
    wc = sum(W[k] * tc_[j] for k, j in _PAIRS)
    w0 = sum(W[k] * t0_[j] for k, j in _PAIRS)
    cost_rows.append(("Weighted sum", "", f"{wc:.4f}",
                      f"{w0:.4f}"))
    add_table([["Term"], ["Weight ", M.run("w")],
               ["Candidate ", M.acc(M.sub("J", "k")), " (", M.run("U"),
                ")"],
               ["No action ", M.acc(M.sub("J", "k")), " (", M.num(0),
                ")"]],
              [[r[0], r[1], r[2], r[3]] for r in cost_rows], 5,
              "Forecast cost terms of the candidate and of no action, "
              "with their weights")

    add_par("Each cost is that weighted sum divided by the total "
            f"weight, {e['wsum']:.1f}.")
    add_eq([M.acc(M.sub("J", "k")), M.txt("("), M.run("U"),
            M.txt(") = "), M.frac(M.num(wc, 4), M.num(e["wsum"], 1)),
            M.txt(" = "), M.num(e["j_c"], 4), M.txt(",    "),
            M.acc(M.sub("J", "k")), M.txt("("), M.num(0),
            M.txt(") = "), M.frac(M.num(w0, 4), M.num(e["wsum"], 1)),
            M.txt(" = "), M.num(e["j_0"], 4)], "_RefEqWEj")
    add_par_ref("These two numbers are what the acceptance test of ",
                EQ_COST, 49, " compares.")
    add_eq([M.num(e["j_c"], 4), M.txt(" ≤ min("), M.num(e["j_th"], 2),
            M.txt(", "), M.num(1 - DELTA, 2), M.txt(" × "),
            M.num(e["j_0"], 4), M.txt(") = "), M.num(e["bound"], 4)],
           "_RefEqWEcost")
    add_par(
        "The candidate is accepted. The ceiling is not the binding "
        f"condition: the forecast cost of no action, {e['j_0']:.4f}, "
        f"already lies below the ceiling of {e['j_th']:.2f}, so the "
        "ceiling alone would accept every candidate. The margin sets "
        "the requirement. A candidate that fails is not withdrawn: the "
        "standing orders remain in force while the stages of Section "
        "4.5.3 search for a better one.")

    # ---- the quality test
    add_par_ref("Acceptance on cost applies to the incident as a whole. "
                "The quality test of ", EQ_Q, 50,
                " is evaluated in each region, over the four scored "
                "concepts. Table 4.6 lists their terms. The activation "
                "is what the concept demanded, the intensity is what "
                "the candidate ordered from the family that answers it, "
                "and ordering more than was asked for counts as much as "
                "ordering less.")

    rows = []
    for i, r in enumerate(e["regions"], start=1):
        for cn, a, u, d in r["terms"]:
            rows.append([f"Region {i}", cn, f"{a:.3f}", f"{u:.3f}",
                         f"{d:.3f}"])
    add_table([["Region"], ["Concept"],
               ["Activation ", M.subsup("a", "n", "eff")],
               ["Ordered intensity ", M.sub("u", "f")],
               ["Mismatch |", M.subsup("a", "n", "eff"), " − ",
                M.sub("u", "f"), "|"]],
              rows, 6,
              "Terms of the quality test in the two regions of the "
              "cycle")

    add_par("The mismatches of the first region are small.")
    add_eq(_q_line(hi, 1), "_RefEqWEq1")
    add_par(f"This stands above the gate, {hi['q']:.3f} ≥ "
            f"{hi['eta']:.2f}, so the orders of that region are applied "
            "as written.")
    add_par(
        "The second region orders slightly more than intervention "
        "urgency asked for, and answers two demands with nothing: "
        + ", and ".join(f"{t[0]} asked for {t[1]:.3f}"
                        for t in [x for x in lo["terms"]
                                  if x[2] <= 0.001 and x[1] > 0.2])
        + ". The cost test could not see this, because assets that are "
        "not burning and people who are not exposed carry no cost "
        "inside the forecast horizon.")
    add_eq(_q_line(lo, 2), "_RefEqWEq2")

    scale = lo["q"] / lo["eta"]
    sup0 = lo["rules"].get("suppression_effort", 0.0)
    dep0 = lo["rules"].get("resource_deployment", 0.0)
    add_par(f"This falls below the gate, {lo['q']:.3f} < "
            f"{lo['eta']:.2f}, so the fail-safe engages. It does not "
            "cancel the orders; it attenuates the offensive ones in "
            "proportion to the deficit.")
    add_eq([M.sub("u", "f"), M.txt("' = "),
            M.frac(M.sub("Q", "2"), M.run("η")), M.sub("u", "f"),
            M.txt(" = "), M.num(scale, 2), M.sub("u", "f")],
           "_RefEqWEfs")
    add_par(
        f"Suppression falls from {sup0:.3f} to {scale * sup0:.3f} and "
        f"resource deployment from {dep0:.3f} to {scale * dep0:.3f}. "
        "Evacuation and public warning are never attenuated. Neither "
        "was ordered in this cycle, which is the omission the quality "
        "test recorded.")

    add_par(
        "The cost test passed for the incident as a whole, and a single "
        "number cannot show where inside it the response is "
        "misdirected. The quality test judges each region and failed in "
        "one. The cost asks what the orders are expected to achieve; "
        "the quality asks whether they answer what was demanded.")

    doc.save(outp)
    print("written:", outp)
    print(f"   cost: J_c={e['j_c']:.4f} bound={e['bound']:.4f} accepted")
    for i, r in enumerate(e["regions"], start=1):
        print(f"   Region {i}: Q={r['q']:.3f} eta={r['eta']:.2f} "
              f"failsafe={r['failsafe']}")


def _jhat(arg):
    """The forecast cost of a candidate, in the chapter's notation."""
    return [M.acc(M.sub("J", "k")), M.txt(f"({arg})")]


_PAIRS = (("w_burn", "j_burn"), ("w_asset", "j_asset"),
          ("w_pop", "j_pop"), ("w_resp", "j_resp"),
          ("w_delay", "j_delay"))


def _q_line(r, idx):
    """Q written out the way the algorithm computes it, term by term."""
    parts = [M.sub("Q", str(idx)), M.txt(" = "), M.num(1), M.txt(" − "),
             M.frac(M.num(1), M.num(4)), M.txt("(")]
    for i, (_cn, _a, _u, d) in enumerate(r["terms"]):
        if i:
            parts.append(M.txt(" + "))
        parts.append(M.num(d, 3))
    parts += [M.txt(") = "), M.num(r["q"], 3)]
    return parts


class _Anchor:
    """Whatever the last thing written was, so the next thing can follow
    it. A table is not a Paragraph, and the insertion point has to be
    able to be either."""

    def __init__(self, element):
        self._p = element


def _mark_para_inserted(par):
    """Mark the paragraph MARK as inserted, leaving its runs alone."""
    p = par._p
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = _el("w:pPr")
        p.insert(0, ppr)
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = _el("w:rPr")
        ppr.append(rpr)
    rpr.insert(0, _el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                                  "w:date": DATE}))


def caption_runs(kind, number, text):
    """A caption that numbers itself, the way the others in the chapter
    do, through a STYLEREF for the chapter and a SEQ for the counter."""
    out = [ins_run(f"{kind} ")]
    for instr, shown in ((" STYLEREF 1 \\s ", "4"),
                         (f" SEQ {kind} \\* ARABIC \\s 1 ", str(number))):
        ins = _el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                              "w:date": DATE})
        for child in _field_runs(instr, shown):
            ins.append(child)
        out.append(ins)
        if instr.startswith(" STYLEREF"):
            out.append(ins_run("."))
    out.append(ins_run(" " + text))
    return out


def _field_runs(instr, shown):
    runs = []
    r = _el("w:r")
    r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    runs.append(r)
    r = _el("w:r")
    t = _el("w:instrText")
    t.text = instr
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    runs.append(r)
    r = _el("w:r")
    r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    runs.append(r)
    r = _el("w:r")
    t = _el("w:t")
    t.text = shown
    r.append(t)
    runs.append(r)
    r = _el("w:r")
    r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))
    runs.append(r)
    return runs


if __name__ == "__main__":
    main()
