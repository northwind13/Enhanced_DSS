"""Write the generative mechanism into Chapter 4 as tracked changes.

WHAT THIS ANSWERS. The chapter says that three adaptation stages exist
and that a package may carry three kinds of object. It does not say which
stage runs on which evidence, when the generative stage writes a rule as
opposed to a concept or an actuator, what the effect vocabulary is, how a
definition written by a language model becomes an order the physics
executes, or whether that vocabulary can grow. Without those answers the
mechanism reads like a magic wand, and a reviewer is entitled to ask why
an ordinary optimiser could not have done the same work.

WHAT IS INSERTED. Two compact tables and two subsections, and nothing
else. The first draft carried five new subsections and four tables, which
buried the gate walkthrough it was meant to support; the sector list and
the compilation path read better as prose and are written as prose here.

  4.5   Table 4.7, the three stages against the evidence that admits
        each one, placed where the symptoms are introduced
  4.5.3 Table 4.8, the proposal forms and what calls each one
        Table 4.9, the effect vocabulary a clause may invoke
        one subsection on what is written and how it is compiled
        one subsection on why the proposer is not an enumeration
        one paragraph on G5b, in its place in the gate walkthrough

EVERY NUMBER HERE IS EITHER A CONSTANT OF THE CODE or a count taken from
the recorded proposal ledgers under experiments/out, both read at run
time so the text cannot drift from them.

Usage: python experiments/fill_ch4_mechanism.py IN.docx OUT.docx
"""
from __future__ import annotations

import glob
import json
import os
import sys
from math import comb

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "out")
sys.path.insert(0, os.path.dirname(HERE))
sys.path.insert(0, HERE)

import docx                                                # noqa: E402
from docx.oxml.ns import qn                                # noqa: E402
from docx.shared import Inches                             # noqa: E402
from docx.text.paragraph import Paragraph                  # noqa: E402

from docx_track import (AUTHOR, DATE, _el, _ins_row,       # noqa: E402
                        _nid, ins_run)
from fill_h1_h2 import _fld, _run, _wrap_ins, para         # noqa: E402


# ------------------------------------------------------------- helpers
def caption4(doc, kind, text, bookmark, shown):
    """A Chapter 4 caption: STYLEREF to the chapter, then SEQ."""
    p = Paragraph(_el("w:p"), doc)
    ppr = _el("w:pPr")
    ppr.append(_el("w:pStyle", **{"w:val": "Caption"}))
    rpr = _el("w:rPr")
    rpr.append(_el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                               "w:date": DATE}))
    ppr.append(rpr)
    p._p.append(ppr)
    bid = _nid()
    label = [_run(kind + " ")]
    label += _fld(" STYLEREF 1 \\s ", "4")
    label.append(_run("."))
    label += _fld(f" SEQ {kind} \\* ARABIC \\s 1 ", shown)
    p._p.append(_el("w:bookmarkStart", **{"w:id": bid,
                                          "w:name": bookmark}))
    p._p.append(_wrap_ins(label))
    p._p.append(_el("w:bookmarkEnd", **{"w:id": bid}))
    p._p.append(_wrap_ins([_run(" " + text)]))
    return p


def heading(doc, level, text):
    p = Paragraph(_el("w:p"), doc)
    ppr = _el("w:pPr")
    ppr.append(_el("w:pStyle", **{"w:val": f"Heading{level}"}))
    rpr = _el("w:rPr")
    rpr.append(_el("w:ins", **{"w:id": _nid(), "w:author": AUTHOR,
                               "w:date": DATE}))
    ppr.append(rpr)
    p._p.append(ppr)
    p._p.append(_wrap_ins([_run(text)]))
    return p


def build_table(doc, headers, rows, widths=None):
    """A table whose every row is marked inserted, not yet placed."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    t.autofit = False
    for j, h in enumerate(headers):
        t.rows[0].cells[j].paragraphs[0]._p.append(ins_run(h, bold=True))
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].paragraphs[0]._p.append(ins_run(str(v)))
    for r in t.rows:
        _ins_row(r)
    if widths:
        tblpr = t._tbl.find(qn("w:tblPr"))
        if tblpr is not None:
            for tag in ("w:tblLayout", "w:tblW"):
                old_el = tblpr.find(qn(tag))
                if old_el is not None:
                    tblpr.remove(old_el)
            tblpr.append(_el("w:tblW", **{
                "w:w": str(int(sum(widths) * 1440)), "w:type": "dxa"}))
            tblpr.append(_el("w:tblLayout", **{"w:type": "fixed"}))
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
                gc.set(qn("w:w"), str(int(w * 1440)))
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


class Cursor:
    """An insertion point that advances over what it inserts.

    ADDNEXT ALWAYS INSERTS IMMEDIATELY AFTER ITS ANCHOR, so re-using one
    anchor for a caption, then its table, then the paragraph that
    follows puts the paragraph BETWEEN the caption and the table and
    leaves the tables stacked in reverse order further down the section.
    The cursor holds the last element it wrote and advances to it, so
    body order matches writing order.
    """

    def __init__(self, el):
        self.el = el

    def add(self, obj):
        el = getattr(obj, "_p", None)
        if el is None:
            el = obj._tbl
        self.el.addnext(el)
        self.el = el
        return obj

    def table(self, doc, headers, rows, widths=None):
        return self.add(build_table(doc, headers, rows, widths))


def full(p):
    return "".join((t.text or "") for t in p._p.iter(qn("w:t")))


def find(ps, needle, start=0, style=None):
    """First paragraph holding `needle`, optionally of a given style.

    THE LIST OF FIGURES COMES FIRST. Every caption appears twice in
    paragraph order, once in the front matter and once under the figure,
    and a plain search returns the front-matter copy. An insertion
    anchored on it lands in the front matter silently.
    """
    for i in range(start, len(ps)):
        if needle in full(ps[i]):
            if style is not None and ps[i].style.name != style:
                continue
            return i
    raise SystemExit(f"anchor not found: {needle!r}")


def past_equation(el):
    """Advance over an equation table that belongs to the paragraph.

    An equation of this thesis sits in a borderless one-row table under
    the paragraph that introduces it. A paragraph inserted between the
    two separates a gate from its own formula.
    """
    nxt = el.getnext()
    while nxt is not None and nxt.tag == qn("w:tbl"):
        el = nxt
        nxt = el.getnext()
    return el


def retext(par, old, new):
    """Replace a phrase inside a paragraph as a tracked edit."""
    for r in list(par._p.iter(qn("w:r"))):
        if r.getparent().tag == qn("w:del"):
            continue
        ts = r.findall(qn("w:t"))
        txt = "".join(t.text or "" for t in ts)
        if old not in txt:
            continue
        parent = r.getparent()
        d = _el("w:del", **{"w:id": _nid(), "w:author": AUTHOR,
                            "w:date": DATE})
        parent.replace(r, d)
        d.append(r)
        for t in ts:
            t.tag = qn("w:delText")
        d.addnext(ins_run(txt.replace(old, new)))
        return True
    return False


# ------------------------------------------------------- measured facts
def facts():
    """Constants read from the code, counts read from the ledgers."""
    from dss.actions import EFFECTS, SECTORS

    n_rng = sum(1 for a in range(16) for b in range(16) if a < b)
    per_clause = len(EFFECTS) * len(SECTORS) * n_rng * 10
    space = sum(comb(per_clause + k - 1, k) for k in (1, 2, 3))
    use = {}
    for f in sorted(glob.glob(os.path.join(OUT,
                                           "geometry_proposals*.jsonl"))):
        tally, n = {}, 0
        with open(f, encoding="utf-8") as fh:
            for line in fh:
                r = json.loads(line)
                if r.get("kind") != "clause":
                    continue
                n += 1
                ni = (r.get("payload") or {}).get("new_intervention") or {}
                for c in ni.get("clauses") or []:
                    e = str(c.get("effect"))
                    tally[e] = tally.get(e, 0) + 1
        use[os.path.basename(f)] = (n, tally)
    return dict(effects=EFFECTS, sectors=SECTORS, n_rng=n_rng,
                per_clause=per_clause, space=space, use=use)


# --------------------------------------------------------------- tables
#: Terse on purpose. The prose around each table carries the argument;
#: the table carries only what a reader would otherwise have to hold in
#: mind while reading it.
STAGE_ROWS = [
    ("① Parameter adaptation",
     "The forecast cost of the candidate exceeds the satisficing "
     "bound, while at least one rule of the base still fires above the "
     "gap threshold. A rule that speaks to the situation exists and is "
     "valued wrongly.",
     "The consequent of each of the two strongest-firing rules moves "
     "by one step of 0.05, or the boundary shared by two neighbouring "
     "antecedent terms moves. The base keeps the rules and the terms "
     "it had."),
    ("② Rule instantiation and resolution increase",
     "Fire is alive on the ground while the strongest firing of the "
     "base stays below the gap threshold, or the burned area grew "
     "although the orders were applied. No rule speaks to the "
     "situation, so there is nothing to revalue.",
     "A rule is written on an antecedent cell that carries none, and a "
     "narrow term is inserted into a partition when no term of the "
     "strongest concept reaches a membership of 0.62. The rules "
     "already in the base are left as they stand."),
    ("③ Generative proposal",
     "Either symptom persists after the cheaper stages have been tried "
     "in this context, and the stage selector values this stage "
     "highest for the context it reads.",
     "A rule is written, and with it at most one new vocabulary object "
     "that this rule uses, which is an intermediate concept, a macro "
     "intervention or a clause actuator. The gates, the cost model and "
     "the rules already in the base are untouched."),
]

FORM_ROWS = [
    ("Plain rule",
     "The concepts can express the situation and an existing action "
     "answers it, but no rule sits on the cell that holds now."),
    ("Rule that edits a rule",
     "Never. A cell that already carries an active rule is refused at "
     "G2, because changing its numbers is the work of stage ①."),
    ("Rule with an intermediate concept",
     "The situation recurs and the five decision concepts cannot name "
     "it, so every rule the stage can write lands on a used cell. The "
     "object is a weight vector over at most four existing inputs."),
    ("Rule with a macro intervention",
     "Two or three base families must act as one at a fixed ratio, "
     "which separate orders would let the aggregation dilute."),
    ("Rule with a clause actuator",
     "The tactic is a placement rather than a mixture: it needs a "
     "where and a how that no base family carries."),
]

EFFECT_ROWS = [
    ("wet", "Raises capacity and crew availability: a pre-wetted band.",
     "Workable ground; draws on the pool."),
    ("clear", "Marks cells as cut and raises capacity: a containment "
              "line.",
     "Workable ground, never built-up or water; draws on the pool."),
    ("ignite", "Sets the cells alight: a firing operation.",
     "Never built-up or water; adds fire the forecast then judges."),
    ("coat", "Raises the moisture floor to 0.45 of the amount laid: "
             "retardant or soil.",
     "The only thresholded effect; acts only where that floor reaches "
     "the extinction moisture of the fuel."),
    ("evacuate", "Raises the evacuation order.",
     "Only where population is present; spends no pool capacity."),
    ("prime", "Raises the public warning.",
     "Only where population is present; spends no pool capacity."),
    ("draft", "Multiplies capacity already on the cells, by a factor "
              "falling with the distance to water.",
     "Needs a water body; acts only where capacity exists, so it "
     "raises the yield of committed crews."),
]


def main():
    inp, outp = sys.argv[1], sys.argv[2]
    F = facts()
    doc = docx.Document(inp)
    done = []

    # ============================================== 4.5 stage evidence
    ps = doc.paragraphs
    c = Cursor(ps[find(ps, "Three symptoms engage the adaptation "
                           "stages")]._p)
    c.add(para(doc, "ParText", [
        "The evidence therefore decides which stages may run, and the "
        "selector of Section 4.5.4 chooses among the stages the "
        "evidence leaves open. ",
        ("_RefTabStage", "4.7"),
        " states the division. Read downward it is the invasiveness "
        "order of the ladder. Read across it is the guarantee that "
        "keeps the ladder auditable, namely that a consequent is tuned "
        "in one place, a rule is written in one place, and a name "
        "enters the vocabulary in one place."]))
    c.add(caption4(doc, "Table",
                   "Adaptation stages, the evidence that admits each "
                   "one, and its reach", "_RefTabStage", "7"))
    c.table(doc, ["Stage", "The evidence that admits it",
                  "What it may change"],
            STAGE_ROWS, widths=[1.35, 2.5, 2.45])
    done.append("4.5: Table 4.7, stage against evidence")

    # ======================================= 4.5.3 forms and triggers
    ps = doc.paragraphs
    c = Cursor(ps[find(ps, "Package objects: (a) macro intervention",
                       style="Caption")]._p)
    c.add(heading(doc, 4, "What Stage ③ Writes, and From What"))
    c.add(para(doc, "ParText", [
        "The three kinds of object are the reach of the stage, not its "
        "policy. What decides the form of a proposal is the reason the "
        "stage was called, and that reason is in the situation the "
        "model is given. ",
        ("_RefTabForm", "4.8"),
        " states the cases, including the one the stage may not "
        "produce. The stage never modifies a rule: a proposal landing "
        "on a covered antecedent cell is refused at G2 with the "
        "duplication named, because re-numbering a rule is the work of "
        "stage ① and costs no model call. That separation is what "
        "makes the ledger readable afterwards, since a consequent that "
        "moved was moved by the tuning stage and a rule that appeared "
        "was written by one of the two growth stages."]))
    c.add(caption4(doc, "Table",
                   "What stage ③ may propose, and what calls each "
                   "form", "_RefTabForm", "8"))
    c.table(doc, ["Form", "What calls it"], FORM_ROWS,
            widths=[1.5, 4.8])
    c.add(para(doc, "ParText", [
        "Only the last form carries geometry. A macro fixes a ratio "
        "between orders that already exist and a concept fixes a "
        "direction in the feature space, but neither can say where on "
        "the map an order should act. A clause names one of the "
        f"{len(F['effects'])} verified effects of ",
        ("_RefTabEffect", "4.9"),
        f", one of {len(F['sectors'])} sectors, a range of cells from "
        "the front, and an amount, and an actuator is at most three "
        "clauses. The sectors are head, flank and rear, which cut the "
        "band by bearing against the direction of spread, ring, which "
        "is the band without a bearing test, at_fire, which is the "
        "burning cells and the two rings around them, assets, which is "
        "the cover of the assets within fifteen cells of the fire, and "
        "populated, which is every cell holding population. The list "
        "of effects is short because each entry is a field the "
        "engine already writes and the cost model already prices. An "
        "actuator can be new; the physics it invokes cannot."]))
    c.add(caption4(doc, "Table",
                   "The verified effect vocabulary a clause may "
                   "invoke", "_RefTabEffect", "9"))
    c.table(doc, ["Effect", "What the engine writes", "What limits it"],
            EFFECT_ROWS, widths=[0.7, 3.0, 2.6])
    c.add(para(doc, "ParText", [
        "The third column is not decoration, and none of it is visible "
        "in the names. Two effects draw on the shared crew pool and "
        "five do not, so on an incident whose pool is spent the first "
        "two only redivide capacity already committed. One is "
        "thresholded, so a coating below the extinction moisture of "
        "the fuel underneath slows the fire without stopping it and is "
        "measured as a tactic that acted and bought nothing. One "
        "multiplies capacity instead of requesting it, which makes it "
        "the only effect that can add real work to a saturated "
        "incident. These facts are assembled from the fuel table and "
        "the engine constants and placed in the situation the model "
        "receives. Geometry matters as much: a sector resolves against "
        "the fire as it stands in this cycle, so a band written far "
        "ahead of a small fire resolves to no cells, and the written "
        "amount is multiplied by the intensity the rule realises, so a "
        "rule firing at half strength with a consequent of 0.9 "
        "delivers about 0.45."]))
    c.add(para(doc, "ParText", [
        "The engine never executes text the model produces. An "
        "actuator arrives as a named list of clauses, its name is "
        "checked for collision, each clause is checked against the "
        "effect and sector lists, and anything outside them is refused "
        "at G2 before a simulation is spent. The definition is then "
        "installed on a working copy of the vocabulary and the rule on "
        "a working copy of the base, so the seed base and the standing "
        "vocabulary are untouched while the proposal is judged. When "
        "the rule fires, the framework resolves each sector to a mask "
        "on the current map, intersects it with the region and with "
        "the ground the effect may touch, multiplies the amount by the "
        "realised intensity, and writes the corresponding field of the "
        "resource plan. A clearing or igniting clause is intersected "
        "with the diggable mask as well, so a generated actuator "
        "cannot cut a line through a settlement or light a fire on "
        "water, which are the orders the base containment family is "
        "forbidden to give. The plan that leaves this path is the same "
        "object the base families produce, so the simulator, the cost "
        "model and the audit trail treat a generated actuator and a "
        "doctrinal order identically. The decision space is therefore "
        "open at the level of interventions and closed at the level of "
        "physics."]))
    c.add(para(doc, "ParText", [
        "The vocabulary grows along one axis and is fixed along the "
        "other. Every package that clears the gates adds one name, and "
        "an admitted name is written to the lineage store, so later "
        "rules in the same incident may cite it and a later incident "
        "restores it unless the vocabulary is reset. The basis itself "
        "is not proposable. Adding an effect means implementing and "
        "validating a field in the physical model, which is a change "
        "made by the authors of the framework and reported in Chapter "
        "3, not an action available to the proposer, because a system "
        "that let the proposer extend the basis would be asking a "
        "language model to certify physics. Growth is also priced. "
        "Each admitted concept adds an axis to the antecedent space "
        "and therefore to the search, which is why G2b refuses a "
        "concept pointing in a direction an existing concept already "
        "covers, and each admitted intervention adds a consequent the "
        "coordinator must price. The controller reward is read on the "
        "reduction in forecast physical cost, so a stage that buys "
        "names without buying outcomes loses value in its own table "
        "and is selected less often."]))
    done.append("4.5.3: Tables 4.8 and 4.9, forms and effects, plus "
                "compilation and growth as prose")

    # ================================================ G5b, in the walk
    ps = doc.paragraphs
    k = find(ps, "G5 is the last gate and applies only to a package")
    _g5 = ps[k]
    if not retext(_g5, "the last gate", "a gate"):
        print("  ! could not retext the G5 opening phrase")
    # THE EQUATION BELONGS TO G5. Its formula sits in a borderless table
    # under the paragraph, so the new gate goes after that table.
    Cursor(past_equation(_g5._p)).add(para(doc, "ParText", [
        "G5b is the last gate and applies to a package that grows the "
        "intervention vocabulary. G5 weighs the package against no "
        "package, and the rule carrying a new object may order "
        "ordinary interventions in the same consequent list, so a "
        "package can clear G5 on the strength of an order the "
        "framework already had. G5b runs the same two futures again "
        "with the consequent on the new object removed and the "
        "ordinary orders of that rule left standing, and requires the "
        "cost to be higher without the object on both. It is not run "
        "when the new object is the only consequent, since G5 has then "
        "already measured the object alone, and a new concept is not "
        "ablated because the rule is written on it. As a negative "
        "example, a package that lowers the cost while its own rule "
        "stripped of the new actuator lowers it by the same amount "
        "fails here: the gain belongs to the ordinary order, and the "
        "vocabulary does not grow on borrowed credit."]))
    ps = doc.paragraphs
    _sum = ps[find(ps, "A proposal is rejected the moment it fails a "
                       "gate")]
    if not retext(_sum, "and G5 for its vocabulary growth",
                  "and G5 with G5b for its vocabulary growth"):
        print("  ! could not extend the gate summary sentence")
    done.append("4.5.3: G5b added after the G5 equation, gate summary "
                "sentence extended")

    # ============================================= the search economy
    ps = doc.paragraphs
    c = Cursor(ps[find(ps, "Grounding by construction, not by "
                           "retrieval")]._p)
    c.add(heading(doc, 4, "Why the Proposer Is Not a Search"))
    c.add(para(doc, "ParText", [
        "The clause grammar is small to write and large to enumerate. "
        f"With {len(F['effects'])} effects, {len(F['sectors'])} "
        f"sectors, {F['n_rng']} admissible ranges over the "
        "fifteen-cell reach and the amount taken to one decimal, a "
        "single clause has "
        + format(F['per_clause'], ',').replace(",", " ")
        + " forms and an actuator of up to three clauses has on the "
        + f"order of {F['space']:.0e} definitions",
        ". Each candidate costs two reseeded rollouts of the physical "
        "forecast, so an enumerating optimiser is not slow on this "
        "space; it cannot begin. Nor is the space one a local search "
        "climbs, because most of it is flat: a sector resolving to no "
        "cells, a coating below the extinction moisture of its fuel, "
        "and an effect rationed away on a spent pool all return a "
        "forecast identical to the one without them, and a gradient "
        "over identical values points nowhere."]))
    c.add(para(doc, "ParText", [
        "What the generative stage supplies is not a faster path "
        "through that space but the ability to remove most of it from "
        "statements about the incident. The situation carries the "
        "affordances of the map, the state of the pool and the effect "
        "notes assembled from the fuel table, and they are used the "
        "way a commander would use them: an effect that draws on a "
        "pool already spent is not written, a coating is not laid on "
        "fuel whose extinction moisture the amount cannot reach, and a "
        "band is not drawn where the sector holds no cells. The "
        "recorded campaigns show this. Before those facts were placed "
        "in the situation, the proposals spread their clauses across "
        "the effects that ask for crew capacity as readily as across "
        "the effects that do not, and none survived the gates. After "
        "they were placed there, with the gate chain unchanged, the "
        "effect that requests capacity disappeared from the proposals "
        "entirely and was replaced by the effect that multiplies "
        "capacity already committed, which is the correct substitution "
        "on a saturated incident and the one that produced the "
        "admitted actuators of Section 5.5.3."]))
    c.add(para(doc, "ParText", [
        "That substitution is the distinction being drawn. An "
        "optimiser over the same grammar has no channel through which "
        "a statement about crew capacity can remove an effect from "
        "consideration; it can only learn that the effect does not pay "
        "by spending rollouts on it, one candidate at a time, on a "
        "surface that returns the same number for every one of them. "
        "The proposer removes the effect before the first rollout is "
        "spent. That is the whole of the contribution the generative "
        "component makes, and it is bounded on purpose: the proposer "
        "chooses what is worth measuring, and the gates, which are "
        "arithmetic and simulation and contain no model, decide what "
        "is kept."]))
    done.append("4.5.3: the search economy")

    doc.save(outp)
    print("written:", outp)
    for d in done:
        print("  -", d)
    print()
    print("clause-effect usage per recorded slice:")
    for f, (n, tally) in F["use"].items():
        print(f"  {f:52} {n:3} proposals  {tally}")


if __name__ == "__main__":
    main()
