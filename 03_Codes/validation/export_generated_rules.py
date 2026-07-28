"""Export every rule the DSS wrote or changed at run time, as a Word table.

THE SEED RULE BASE IS DELIBERATELY ABSENT. Those rules were written by
hand before any fire started and are listed in their own appendix; a
rule appears here only if the system produced it or moved it during a
run. Three things qualify, and the document keeps them apart because
they are different claims:

  written, generative (Stage 3)   proposed by the generative stage and
                                  admitted through the full gate chain
  written, resolution (Stage 2)   instantiated on an antecedent cell the
                                  standing base did not answer, built
                                  from the nearest standing rule
  tuned, evFIS (Stage 1)          an existing rule whose consequents the
                                  evolving-fuzzy stage walked, one
                                  accepted step at a time

The source is the learned store, which is the system's own record and
carries the provenance of every entry. Nothing here is typed by hand: a
hand-kept list of learned rules is a second store, and the two drift.

    python validation/export_generated_rules.py
    python validation/export_generated_rules.py --store ../other.json \
        --out generated_rules.docx
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from collections import defaultdict

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Twips

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(HERE))

DEFAULT_STORE = os.path.join(HERE, "..", "logs",
                             "dss_generated_state.json")
DEFAULT_OUT = os.path.join(HERE, "figures", "generated_rules.docx")

# ------------------------------------------------------------- vocabulary
# A rule reads as fuzzy shorthand; the table has to read as English, so
# every symbol in it is given a word here and nowhere else.
TERM = {"VL": "very low", "L": "low", "M": "moderate",
        "H": "high", "VH": "very high"}
VAR = {
    "fire_threat_level": "the fire threat",
    "suppression_feasibility": "suppression feasibility",
    "asset_exposure_risk": "asset exposure risk",
    "evacuation_pressure": "evacuation pressure",
    "intervention_urgency": "intervention urgency",
    "interface_exposure": "interface exposure",
}
ACT = {
    "suppression_effort": "press the suppression effort",
    "retardant_drop": "drop retardant",
    "resource_deployment": "deploy resources",
    "asset_protection": "protect the exposed assets",
    "evacuation": "evacuate",
    "water_drafting": "draft water",
    "containment_line": "cut containment line",
    "tactical_burn": "set a tactical burn",
}


def words(s):
    return str(s).replace("_", " ")


def term_words(t):
    s = str(t)
    if s.startswith(">="):
        return f"at least {TERM.get(s[2:], s[2:])}"
    if s.startswith("X") and s[1:].isdigit():
        return f"in the inserted band {s}"
    return TERM.get(s, s)


def act_words(name, value, made):
    base = ACT.get(name)
    if base is None:
        base = f"order {words(name)}" if name in made else words(name)
    v = float(value)
    band = ("at full strength" if v >= 0.85 else
            "hard" if v >= 0.65 else
            "moderately" if v >= 0.4 else "lightly")
    return f"{base} {band} ({v:.2f})"


def rule_text(ants, cons):
    a = " AND ".join(f"{words(v)} is {t}" for v, t in ants)
    c = ", ".join(f"{words(k)} {float(v):.2f}" for k, v in cons)
    return f"IF {a} THEN {c}"


def trigger_text(rec):
    t = rec.get("trigger") or {}
    bits = []
    if "step" in t:
        bits.append(f"step {t['step']}")
    if "minute" in t:
        bits.append(f"{t['minute']:g} min")
    if "deficit" in t:
        # the raw value carries sixteen digits of float; two is what the
        # number means
        bits.append(f"deficit {float(t['deficit']):.2f}")
    return ", ".join(bits)


# ------------------------------------------------------------------ input
def load(store_path):
    with open(store_path, encoding="utf-8") as fh:
        store = json.load(fh)
    made = {}
    for iv in store.get("genai_interventions", []):
        made[iv["name"]] = " + ".join(
            f"{words(k['channel'])} x{k['weight']}"
            for k in iv.get("composition", []))
    concept = {}
    for c in store.get("genai_concepts", []):
        concept[c["name"]] = " + ".join(
            f"{words(i['name'])} x{i['weight']}"
            for i in c.get("inputs", []))
    return store, made, concept


def seed_rules():
    """The hand-written base, needed only to print a tuned rule in full.

    A consequent update records what moved, not what the rule says, so
    the antecedents of a tuned base rule have to come from the base
    itself. This is also the one place the seed base is read at all.
    """
    try:
        from dss.rules import SEED_RULES
        return {r.name: (list(r.antecedents), list(r.consequents))
                for r in SEED_RULES}
    except Exception as exc:                       # pragma: no cover
        print(f"note: seed rules unavailable ({exc}); tuned base rules "
              "will be listed without their antecedents")
        return {}


def describe(ants, cons, made, concept, note=""):
    when = " and ".join(f"{VAR.get(v, words(v))} is {term_words(t)}"
                        for v, t in ants)
    then = ", ".join(act_words(k, v, made) for k, v in cons)
    s = f"When {when}, {then}." if when else f"{then}."
    for k, _v in cons:
        if k in made:
            s += (f' "{words(k)}" is a generated intervention '
                  f'({made[k]}).')
    for v, _t in ants:
        if v in concept:
            s += f' "{words(v)}" is a generated concept ({concept[v]}).'
    why = str(note).split("WHY:")
    if len(why) > 1:
        s += f" Recorded reason: {why[1].strip()}"
    elif "renamed from" in str(note):
        s += f" {str(note).split('|')[-1].strip()}."
    return s


# --------------------------------------------------------------- written
def written_rows(store, made, concept):
    rows = []
    for r in store.get("genai_rules", []):
        ants, cons = r.get("antecedents", []), r.get("consequents", [])
        rows.append(dict(
            id=r.get("name") or r["id"],
            rule=rule_text(ants, cons),
            desc=describe(ants, cons, made, concept, r.get("note", "")),
            origin="Generative (Stage 3)",
            made=f"{trigger_text(r)}\n{r['id']}",
            seq=r.get("seq", 0),
            key=json.dumps([ants, cons], sort_keys=True)))
    for m in store.get("evfis_rule_modifications", []):
        if m.get("modification_type") != "rule_add":
            continue
        r = (m.get("after") or {}).get("rule") or {}
        ants, cons = r.get("antecedents", []), r.get("consequents", [])
        rows.append(dict(
            id=r.get("name") or m["base_rule_id"],
            rule=rule_text(ants, cons),
            desc=(describe(ants, cons, made, concept, r.get("note", ""))
                  + f" Instantiated from the standing rule "
                    f"{m['base_rule_id']} of the {m.get('base_rule_set')} "
                    f"base, on an antecedent cell that base did not "
                    f"answer."),
            origin="Resolution (Stage 2)",
            made=f"{trigger_text(m)}\n{m['id']}",
            seq=m.get("seq", 0),
            key=json.dumps([ants, cons], sort_keys=True)))

    # THE SAME RULE MAY BE WRITTEN TWICE. Sessions run independently, so
    # a situation that recurs is answered again, sometimes under a name
    # that is already taken. Identical rules fold into one row that says
    # how often it was produced; the same name carrying different
    # content keeps both rows and is marked, because those are two rules.
    folded = {}
    for r in rows:
        if r["key"] in folded:
            folded[r["key"]]["times"] += 1
            continue
        folded[r["key"]] = dict(r, times=1)
    out = list(folded.values())
    n_name = defaultdict(int)
    for r in out:
        n_name[r["id"]] += 1
    seen = defaultdict(int)
    for r in out:
        if n_name[r["id"]] > 1:
            seen[r["id"]] += 1
            r["id"] = f"{r['id']}{'abcdefgh'[seen[r['id']] - 1]}"
            r["desc"] += (" Shares its name with another generated rule "
                          "but not its content: the two were written in "
                          "separate sessions.")
        if r["times"] > 1:
            r["desc"] += (f" Written {r['times']} times over the "
                          "recorded runs.")
    out.sort(key=lambda r: (0 if r["origin"].startswith("Gen") else 1,
                            r["seq"]))
    return out


# ----------------------------------------------------------------- tuned
def tuned_rows(store, made, concept, base, written):
    """One row per rule the evolving-fuzzy stage moved.

    The store holds one record per accepted step, and there are far more
    steps than rules: 157 of them land on 12 rules, and a table of 157
    rows would hide that. So the steps are folded per rule into the
    thing that is actually claimed - where the consequents started,
    where they ended, and how many accepted steps it took.
    """
    by_rule = defaultdict(list)
    for m in store.get("evfis_rule_modifications", []):
        if m.get("modification_type") == "consequent_update":
            by_rule[m["base_rule_id"]].append(m)

    known = {}
    for r in store.get("genai_rules", []):
        known[r.get("name")] = (r.get("antecedents", []),
                                r.get("consequents", []))
    for m in store.get("evfis_rule_modifications", []):
        if m.get("modification_type") == "rule_add":
            r = (m.get("after") or {}).get("rule") or {}
            known[r.get("name")] = (r.get("antecedents", []),
                                    r.get("consequents", []))
    known.update({k: v for k, v in base.items() if k not in known})

    rows = []
    for name, mods in by_rule.items():
        mods.sort(key=lambda m: m.get("seq", 0))
        first = dict((k, float(v)) for k, v in
                     (mods[0].get("before") or {}).get("consequents", []))
        last = dict((k, float(v)) for k, v in
                    (mods[-1].get("after") or {}).get("consequents", []))
        ants = known.get(name, ([], []))[0]
        moved, held = [], []
        for k in last:
            d = last[k] - first.get(k, last[k])
            if abs(d) >= 0.005:
                moved.append(f"{words(k)} {first.get(k, 0):.2f} -> "
                             f"{last[k]:.2f} ({d:+.2f})")
            else:
                held.append(words(k))
        desc = describe(ants, sorted(last.items()), made, concept)
        desc += (f" The evolving-fuzzy stage moved this rule over "
                 f"{len(mods)} accepted steps: "
                 + ("; ".join(moved) if moved
                    else "no consequent ended away from where it started")
                 + ".")
        if held:
            desc += f" Unchanged: {', '.join(held)}."
        desc += (" A step is only kept when the reseeded forecast shows "
                 "it costs less than the standing decision, so the walk "
                 "is directional rather than a search.")
        origin = ("a rule the system wrote" if name in
                  {w["id"].rstrip("abcdefgh") for w in written}
                  or name in known and name not in base
                  else "a seed rule")
        rows.append(dict(
            id=name,
            rule=rule_text(ants, sorted(last.items())) if ants
                 else "(antecedents unavailable)",
            desc=desc,
            steps=str(len(mods)),
            origin=f"Tuned (Stage 1), {origin}",
            made=f"{trigger_text(mods[-1])}\n{mods[0]['id']}"
                 f" .. {mods[-1]['id']}",
            seq=mods[0].get("seq", 0)))
    rows.sort(key=lambda r: -int(r["steps"]))
    return rows


# ------------------------------------------------------------- variables
def variable_rows(store):
    """Membership boundaries moved and linguistic terms inserted.

    Not rules, but they change what every rule using that variable
    means, so a list of rule changes that left them out would be
    incomplete.
    """
    rows = []
    for m in store.get("evfis_rule_modifications", []):
        t = m.get("modification_type")
        if t not in ("membership_shift", "term_insert"):
            continue
        a = m.get("after") or {}
        b = m.get("before") or {}
        var = (a.get("variable") or b.get("variable")
               or m.get("base_rule_id") or "(unnamed variable)")
        part = a.get("partition") or {}
        new = [k for k in part if k not in TERM]
        if t == "term_insert":
            what = (f"inserted the term {', '.join(new)} into the "
                    f"partition")
            why = ("the five standing terms were too coarse to separate "
                   "the situation the rule was facing, so a term was "
                   "added between them; the partition still sums to one "
                   "everywhere, so no existing rule changed meaning "
                   "outside the new band")
        else:
            what = "moved a shared boundary of the partition"
            why = ("the boundary between two terms sat where the "
                   "decision needed to distinguish; both sides move "
                   "together, so the partition stays intact")
        rows.append(dict(
            id=words(var),
            what=what,
            terms=", ".join(f"{k} [{', '.join(f'{x:.2f}' for x in v)}]"
                            for k, v in part.items()
                            if k in new) or "-",
            why=why,
            made=f"{trigger_text(m)}\n{m['id']}"))
    return rows


# ------------------------------------------------------------------ docx
def _cell(cell, text, bold=False, mono=False, size=8):
    cell.text = ""
    for i, line in enumerate(str(text).split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)
        if mono:
            run.font.name = "Consolas"


def _table(doc, headers, widths, rows, mono_cols=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True, size=9)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            _cell(cells[i], v, mono=(i in mono_cols))
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = Twips(widths[i])
    return t


def build(store_path, out_path):
    store, made, concept = load(store_path)
    base = seed_rules()
    written = written_rows(store, made, concept)
    tuned = tuned_rows(store, made, concept, base, written)
    variables = variable_rows(store)

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for side in ("left", "right", "top", "bottom"):
        setattr(sec, f"{side}_margin", Twips(720))

    doc.add_heading("DisasterAware - rules the system wrote and changed", 1)
    p = doc.add_paragraph()
    r = p.add_run(f"Source: {os.path.basename(store_path)} - "
                  f"{len(written)} written, {len(tuned)} tuned")
    r.italic = True
    r.font.size = Pt(9)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = intro.add_run(
        "Every rule below was produced or changed during a run. The seed "
        "rule base is not listed here: those rules were written by hand "
        "before any fire started and belong to their own appendix. What "
        "follows is what the adaptation stages did to that base. The "
        "resolution stage instantiates a rule when the situation in front "
        "of it falls on an antecedent cell the standing base does not "
        "answer, building it from the nearest standing rule. The "
        "generative stage proposes a rule outright, and nothing it "
        "proposes acts before it has passed the full gate chain. The "
        "evolving-fuzzy stage does not add rules; it walks the "
        "consequents of rules that already exist, one accepted step at a "
        "time, and a step is kept only when the reseeded forecast shows "
        "it costs less than the standing decision. Each row carries the "
        "step and minute at which the change was made and the record it "
        "came from, so any line here can be traced back to the run that "
        "produced it.")
    r.font.size = Pt(9)

    doc.add_heading("Summary", 2)
    n_gen = sum(1 for w in written if w["origin"].startswith("Gen"))
    _table(doc, ["What the system did", "Rules"], [9000, 2200], [
        ["Wrote a rule outright, generative stage (Stage 3)", str(n_gen)],
        ["Instantiated a rule on an unanswered antecedent cell, "
         "resolution stage (Stage 2)", str(len(written) - n_gen)],
        ["Tuned the consequents of an existing rule, evolving-fuzzy "
         "stage (Stage 1)", str(len(tuned))],
        ["Reshaped an input variable (boundary moved or term inserted)",
         str(len(variables))],
    ])

    doc.add_heading("Reading the rules", 2)
    _table(doc, ["Symbol", "Reading"], [2600, 8600], [
        ["VL / L / M / H / VH",
         "very low, low, moderate, high, very high: the five-term "
         "partition every input variable carries."],
        [">=H, >=M, >=L",
         "at least that term: the rule fires on the term and everything "
         "above it."],
        ["X1, X2, X3",
         "terms the resolution stage inserted into a variable when the "
         "standing five were too coarse to separate the situation it was "
         "facing. They sit between the standing terms and preserve the "
         "partition."],
        ["0.00 - 1.00",
         "the consequent value: how hard the named action is ordered, on "
         "the scale every rule shares."],
    ], mono_cols=(0,))

    doc.add_heading("Rules the system wrote", 2)
    _table(doc, ["ID", "Rule", "Description", "Origin", "Produced"],
           [800, 3900, 5300, 1500, 1700],
           [[w["id"], w["rule"], w["desc"], w["origin"], w["made"]]
            for w in written], mono_cols=(0, 1, 4))

    doc.add_heading("Rules the evolving-fuzzy stage tuned", 2)
    _table(doc, ["ID", "Rule after tuning", "Description", "Steps",
                 "Produced"],
           [800, 3700, 5300, 700, 2000],
           [[t["id"], t["rule"], t["desc"], t["steps"], t["made"]]
            for t in tuned], mono_cols=(0, 1, 4))

    if variables:
        doc.add_heading("Input variables the system reshaped", 2)
        _table(doc, ["Variable", "Change", "New terms", "Why", "Produced"],
               [1800, 2400, 2800, 4200, 2000],
               [[v["id"], v["what"], v["terms"], v["why"], v["made"]]
                for v in variables], mono_cols=(2, 4))

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"{len(written)} written + {len(tuned)} tuned "
          f"+ {len(variables)} variable changes -> {out_path}")
    print(f"   {n_gen}  generative (Stage 3)")
    print(f"   {len(written) - n_gen}  resolution (Stage 2)")


def main():
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("--store", default=DEFAULT_STORE)
    ap.add_argument("--out", default=DEFAULT_OUT)
    a = ap.parse_args()
    build(os.path.abspath(a.store), os.path.abspath(a.out))


if __name__ == "__main__":
    main()
