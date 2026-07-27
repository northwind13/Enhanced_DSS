"""DisasterAware interactive wildfire simulator dashboard.

Run with:  streamlit run app/streamlit_app.py

Uses page based navigation (one page runs per interaction) so the app stays
fast enough to animate the simulation and to build a DSS on top.
"""

import os
import sys
import time

import numpy as np
import streamlit as st

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import dataclasses

from disaster_phyengine import (Simulator, World, SimConfig, Asset, compute_costs,
                           scenarios, io_utils, maplib, terrain, viz,
                           FUEL_MODELS,
                           FUEL_NAME_TO_ID, SpreadParams, SuppressionParams,
                           IntensityParams, ValueWeights, CostParams)

# frozen copy of the default fuel table (A.1 + B.1) for the reset buttons
_REFERENCE_FUELS = {fid: dataclasses.replace(m) for fid, m in FUEL_MODELS.items()}


def _install_canvas_compat():
    try:
        import streamlit.elements.image as old_img
        if hasattr(old_img, "image_to_url"):
            return
        from streamlit.elements.lib import image_utils as IU
        from streamlit.elements.lib.layout_utils import LayoutConfig

        def _shim(image, width, clamp, channels, output_format, image_id):
            return IU.image_to_url(image, LayoutConfig(width=int(width)),
                                   clamp, channels, output_format, image_id)
        old_img.image_to_url = _shim
    except Exception:
        pass


_install_canvas_compat()
try:
    from streamlit_drawable_canvas import st_canvas
    HAS_CANVAS = True
except Exception:
    HAS_CANVAS = False


st.set_page_config(page_title="DisasterAware Simulator", layout="wide")

# a slimmer left panel so the map / simulation gets more room, and hide the
# default "Made with Streamlit" footer badge
st.markdown(
    "<style>section[data-testid='stSidebar']{width:265px!important;"
    "min-width:265px!important}"
    "footer{visibility:hidden;height:0}"
    "[data-testid='stStatusWidget']{visibility:hidden}"
    # The default toolbar is a full-width band roughly 3.5rem tall and the
    # content below it carries about 6rem of top padding, so the page opened
    # with an empty stripe above the first control. The band is shrunk rather
    # than removed: the ⋮ menu lives in it and is still needed.
    "header[data-testid='stHeader']{height:2.2rem;background:transparent}"
    "[data-testid='stToolbar']{height:2.2rem}"
    ".block-container,[data-testid='stAppViewBlockContainer']"
    "{padding-top:1.2rem!important}"
    # the sidebar carries its own top padding, on a different element than
    # the main area, so trimming only the main one left a band above the
    # title
    "section[data-testid='stSidebar'] .block-container,"
    "[data-testid='stSidebarUserContent'],"
    "section[data-testid='stSidebar'] > div:first-child"
    "{padding-top:0!important}"
    # COMPACT, CAREFULLY. Shrinking the flex gap and zeroing paragraph
    # margins made captions and buttons overlap: Streamlit lays these out on
    # a flex column and the gap IS the separation. Only the paddings inside
    # controls are trimmed, never the spacing between them.
    "section[data-testid='stSidebar'] .stButton>button"
    "{padding:0.15rem 0.5rem;font-size:0.88rem}"
    "section[data-testid='stSidebar'] [data-testid='stVerticalBlockBorder"
    "Wrapper']{padding:0.3rem 0.5rem}"
    # wide markdown tables (all-agents comparison) scroll sideways inside
    # their panel instead of overflowing the screen
    "[data-testid='stMarkdownContainer'] table{display:block;"
    "overflow-x:auto;max-width:100%}</style>",
    unsafe_allow_html=True)

FUEL_TYPES = ["grass", "shrub", "pine_litter", "hardwood"]
ASSET_KINDS = ["building", "critical", "population"]
ASSET_LABELS = {"building": "Building", "critical": "Critical facility",
                "population": "Population"}
FIREBREAK_TYPES = {"Water": 5, "Bare ground": 0}


# --------------------------------------------------------------------- state
def _resize_world(w: World, nx2: int, ny2: int,
                  keep_extent: bool = True) -> World:
    """Resample the world onto a new grid. All layers, roads, assets and
    ignitions come with it; slope and aspect are recomputed.

    `keep_extent` makes this a change of RESOLUTION rather than of area: the
    cell size is divided by the same factor the grid is multiplied by, so
    the landscape covers the same ground at a finer or coarser sampling.
    Without it, doubling nx doubled the map's physical width, which is why
    a resized map came out looking stretched and wrong: the fire then had
    twice as far to travel, the service radii covered half as much of the
    map, and the population per square kilometre stayed put while the
    square kilometres doubled.

    Asset footprints are scaled with the grid for the same reason. They are
    measured in CELLS, so a radius-6 town left alone on a doubled grid is
    physically half the town it was, while the built-up cells it painted
    were resampled and doubled: the block and the value written into it
    stopped describing the same place.
    """
    import dataclasses as _dc
    from disaster_phyengine.layers import (MeteoLayer, TopoLayer, FuelLayer,
                                      ValueLayer, ResourceLayer)
    ny1, nx1 = w.shape
    yi = np.minimum((np.arange(ny2) * ny1 / ny2).astype(int), ny1 - 1)
    xi = np.minimum((np.arange(nx2) * nx1 / nx2).astype(int), nx1 - 1)

    def R(a):
        return np.asarray(a)[yi][:, xi].copy()

    _cell2 = float(w.config.cell_size_m)
    if keep_extent:
        # one factor for both axes: a non-uniform change would stretch the
        # ground, so the mean keeps the cell square and the area honest
        _f = 0.5 * (nx2 / float(nx1) + ny2 / float(ny1))
        _cell2 = float(w.config.cell_size_m) / max(_f, 1e-9)
    cfg2 = SimConfig.from_dict({**w.config.to_dict(),
                                "nx": int(nx2), "ny": int(ny2),
                                "cell_size_m": _cell2})
    w2 = World.blank(cfg2)
    w2.meteo = MeteoLayer(**{k: R(getattr(w.meteo, k))
                             for k in ("temp", "rh", "wws", "wwd",
                                       "gust", "prec")})
    w2.topo = TopoLayer(elev=R(w.topo.elev), slope=R(w.topo.slope),
                        aspect=R(w.topo.aspect), access=R(w.topo.access))
    w2.fuel = FuelLayer(ftype=R(w.fuel.ftype).astype(int),
                        fload=R(w.fuel.fload), fmoist=R(w.fuel.fmoist),
                        fload0=R(w.fuel.fload0))
    w2.value = ValueLayer(vbld=R(w.value.vbld), vcrit=R(w.value.vcrit),
                          vpop=R(w.value.vpop), vevac=R(w.value.vevac))
    w2.resource = ResourceLayer(
        rcap=R(w.resource.rcap), ravail=R(w.resource.ravail),
        reff=R(w.resource.reff), rtime=R(w.resource.rtime),
        rair=(None if getattr(w.resource, "rair", None) is None
              else R(w.resource.rair)))
    if getattr(w, "roads", None) is not None:
        w2.roads = R(w.roads)
    sx, sy = nx2 / nx1, ny2 / ny1
    _rs = 0.5 * (sx + sy)          # footprints are in cells, so they scale
    w2.assets = [_dc.replace(a, x=int(a.x * sx), y=int(a.y * sy),
                             radius=max(0, int(round(
                                 getattr(a, "radius", 0) * _rs))))
                 for a in w.assets]
    w2.ignitions = [_dc.replace(e, x=int(e.x * sx), y=int(e.y * sy))
                    for e in w.ignitions]
    w2.recompute_slope_aspect()
    # the value layers are DERIVED from the assets, and the assets have just
    # moved and changed size, so they are rebuilt rather than resampled: a
    # resampled disc does not match the disc the new radius describes
    try:
        w2.rebuild_value_layers()
    except Exception:
        pass
    return w2


#: how far an export may be enlarged past the screen composition
EXPORT_MAX_MPX = 60.0


def _export_scale(nx: int, ny: int) -> int:
    """Pixels per cell for an export: THE SCREEN'S OWN.

    The first version rendered at ~27 px per cell to fill 4000 px, and the
    result was not the map enlarged: the marker glyphs and the lettering
    are drawn at sizes that do not all follow the cell size (the label font
    is capped, the sensor and depot icons are fixed pixel shapes), so at
    four times the scale the symbols and the words shrank into the terrain
    and the map lost the very labels it was exported for.

    The export is the composition on screen. Its RESOLUTION comes from a
    whole-image enlargement afterwards, which keeps every proportion.
    """
    return _fit_scale(nx)


def _export_maps(world, sim=None, factor: int = 1):
    """Two high-resolution PNGs of the CURRENT map.

    1. the map itself: terrain, land cover, roads, settlements and what is
       at risk on them - the world, with nothing operational on it;
    2. the same map with the SENSORS, the RESOURCE bases and their service
       radii, and the IGNITION points - what the response is set up to do
       about it.

    Two files rather than one because they answer different questions and
    a reader of the second cannot see the ground under it.
    """
    from io import BytesIO
    from PIL import Image
    _sc = _export_scale(world.config.nx, world.config.ny)
    _sens = st.session_state.get("dss_sensors_draw")
    _deps = st.session_state.get("dss_depots_draw")
    _common = dict(scale=_sc, show_fire=False, show_assets=True,
                   show_value=False, show_hillshade=True, show_roads=True,
                   show_labels=True, show_grid=False, show_perimeter=False,
                   show_wind=True)
    plain = viz.render_pil(world, sim=sim, show_ignitions=False, **_common)
    staged = viz.render_pil(world, sim=sim, show_ignitions=True,
                            sensors=_sens, depots=_deps, **_common)
    _f = max(1, int(factor))
    if (plain.width * _f) * (plain.height * _f) > EXPORT_MAX_MPX * 1e6:
        _f = max(1, int((EXPORT_MAX_MPX * 1e6
                         / max(1, plain.width * plain.height)) ** 0.5))
    out = []
    for img in (plain, staged):
        if _f > 1:
            # LANCZOS, not NEAREST: the enlargement is of the FINISHED
            # picture, so the terrain cells, the glyphs and the lettering
            # grow together and the export is what the screen shows, on
            # more pixels.
            img = img.resize((img.width * _f, img.height * _f),
                             Image.LANCZOS)
        b = BytesIO()
        img.save(b, format="PNG", dpi=(300, 300))
        out.append(b.getvalue())
    return out[0], out[1], _sc, (plain.width * _f, plain.height * _f), \
        (_sens, _deps)


def _export_panel(world, sim=None, key: str = "exp") -> None:
    """The export control, on both the simulation page and the editor."""
    with st.expander("Export map", expanded=False):
        _sc = _export_scale(world.config.nx, world.config.ny)
        _f = st.radio("Size", [1, 2, 3], horizontal=True, key=f"{key}_f",
                      format_func=lambda k: ("screen (1:1)" if k == 1
                                             else f"x{k} for print"),
                      help="The export is the map as it is on screen. The "
                           "multiplier enlarges the finished picture, so "
                           "the terrain, the symbols and the lettering keep "
                           "their proportions instead of the labels "
                           "shrinking into a giant map.")
        st.caption(f"{world.config.nx * _sc} x {world.config.ny * _sc} px "
                   f"at 1:1, 300 dpi. Two files: the map on its own, and "
                   "the map with the sensors, the resource bases and the "
                   "ignition points.")
        if st.button("Render", key=f"{key}_go", use_container_width=True,
                     type="primary"):
            with st.spinner("Rendering..."):
                _p, _s, _scl, _sz, _lists = _export_maps(world, sim,
                                                         factor=int(_f))
            st.session_state[f"{key}_png"] = (_p, _s, _scl, _sz)
            _sn, _dp = _lists
            if not _sn and not _dp:
                st.info("No sensors or resource units are staged yet, so "
                        "the second image differs only by the ignition "
                        "points. Stage them on Layer 1 first.")
        # THE KEY IS A FILE TOO, BUT NOT ON EVERY RERUN. The legend lived
        # in the page as HTML: readable on screen and impossible to put in
        # a document. Rendering the sheet here unconditionally cost 0.75 s
        # of every single interaction - a step, a slider, an animation
        # frame - and the app felt dead: pressing Step appeared to do
        # nothing. It is built when it is asked for, like the maps.
        if st.button("Render legend sheet", key=f"{key}_lego",
                     use_container_width=True):
            from io import BytesIO as _BIO
            with st.spinner("Rendering the legend..."):
                _lg = viz.legend_sheet(macros=_all_macros(None) or None,
                                       title="DisasterAware — map legend")
                _lb = _BIO()
                _lg.save(_lb, format="PNG", dpi=(300, 300))
            st.session_state[f"{key}_leg"] = (_lb.getvalue(), _lg.size)
        _lgot = st.session_state.get(f"{key}_leg")
        if _lgot:
            st.download_button(f"Legend ({_lgot[1][0]}x{_lgot[1][1]})",
                               _lgot[0], file_name="legend.png",
                               mime="image/png", use_container_width=True,
                               key=f"{key}_dleg")

        _got = st.session_state.get(f"{key}_png")
        if _got:
            _p, _s, _scl, _sz = _got
            c1, c2 = st.columns(2)
            c1.download_button(f"Map ({_sz[0]}x{_sz[1]})", _p,
                               file_name="map.png", mime="image/png",
                               use_container_width=True, key=f"{key}_d1")
            c2.download_button("Map + sensors + resources + ignitions", _s,
                               file_name="map_operational.png",
                               mime="image/png", use_container_width=True,
                               key=f"{key}_d2")


def _map_card():
    """Every map view sits in one bordered block of fixed width.

    The frames the animation writes were drawn at their natural pixel size
    while the paused view is a chart stretched to the column, so the map
    jumped in size on every play/pause and the page reflowed around it. The
    block gives all of them the same frame, and the images are stretched to
    it, so what changes between views is the content and not the layout.
    """
    try:
        return st.container(border=True)
    except TypeError:          # older streamlit: no border, same block
        return st.container()


def _new_simulator(world: World) -> None:
    st.session_state.world = world
    st.session_state.sim = Simulator(world)
    st.session_state.cost_series = []
    st.session_state["anim_stop"] = True   # applied before the toggle renders
    st.session_state.canvas_key = st.session_state.get("canvas_key", 0) + 1
    st.session_state.map_version = st.session_state.get("map_version", 0) + 1
    st.session_state.sim_applied = 0
    # a NEW MAP invalidates the whole DSS setup: sensors and depots
    # carry coordinates of the old terrain, so everything is
    # cleared and 'Apply decisions' drops to OFF (a fire RESET, in
    # contrast, keeps the infrastructure and clears only the
    # decisions)
    for _k in ("dss_sensors", "dss_sens_edit", "dss_network",
               "dss_net_sig", "dss_sensors_draw", "dss_res_items",
               "dss_res_base", "dss_res_base_v", "dss_res_sig",
               "dss_res_why", "res_edit", "dss_depots_draw",
               "dss_suggest_why", "dss_learned_override"):
        st.session_state.pop(_k, None)
    st.session_state["dss_apply"] = False
    try:
        _reset_dss_state(drop_engine=True)
    except Exception:
        pass


def _restore_rca_applied() -> None:
    """What the operator applied from a review STAYS applied: the
    settings come back as session defaults and the staged sensors /
    depots rejoin the staging lists (deduplicated), the same
    permanence the learned store already gives the rules."""
    if st.session_state.get("_rca_restored"):
        return
    st.session_state["_rca_restored"] = True
    import json as _js_rb
    _prb = os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "logs", "rca_applied.json")
    try:
        _d = _js_rb.load(open(_prb, encoding="utf-8"))
    except Exception:
        return
    for _k, _v in (_d.get("settings") or {}).items():
        st.session_state.setdefault(_k, _v)
    for _lst, _key in (("sensors", "dss_sensors"),
                       ("depots", "dss_res_items")):
        _add = _d.get(_lst) or []
        if not _add:
            continue
        _cur = list(st.session_state.get(_key, []) or [])
        _sigs = [(q.get("kind"), q.get("x"), q.get("y"))
                 for q in _cur]
        for _r in _add:
            if (_r.get("kind"), _r.get("x"), _r.get("y")) not in _sigs:
                _cur.append(_r)
        st.session_state[_key] = _cur


def _ensure_state() -> None:
    _restore_rca_applied()
    if "sim" not in st.session_state:
        # THE OPERATOR'S OWN MAP OPENS FIRST. The app always started on the
        # same procedural mountain landscape, so whatever had been built in
        # the editor had to be rebuilt or re-uploaded at the start of every
        # session. A map marked as the default in the library opens instead;
        # if none is marked, or the file will not open, the generated one
        # is still there to fall back on.
        _dw = None
        try:
            _dw = maplib.load_default()
        except Exception:
            _dw = None
        if _dw is not None:
            _new_simulator(_dw)
            st.session_state["_opened_map"] = maplib.default_name()
        else:
            _new_simulator(terrain.generate_landscape(
                SimConfig(nx=100, ny=70, cell_size_m=30.0), seed=42,
                preset="Mountain forest"))
            st.session_state.world.add_ignition(25, 35, step=0, radius=2)
    st.session_state.setdefault("tool", "Fuel")
    st.session_state.setdefault("cost_series", [])
    st.session_state.setdefault("anim_on", False)
    st.session_state.setdefault("canvas_key", 1)
    st.session_state.setdefault("map_version", 1)
    st.session_state.setdefault("edit_undo", [])


def _push_snapshot() -> None:
    """Save the editable layers so a map edit can be undone."""
    w = st.session_state.world
    snap = {
        "ftype": w.fuel.ftype.copy(), "fload": w.fuel.fload.copy(),
        "fload0": w.fuel.fload0.copy(), "fmoist": w.fuel.fmoist.copy(),
        "elev": w.topo.elev.copy(), "slope": w.topo.slope.copy(),
        "aspect": w.topo.aspect.copy(), "access": w.topo.access.copy(),
        "vbld": w.value.vbld.copy(), "vcrit": w.value.vcrit.copy(),
        "vpop": w.value.vpop.copy(), "vevac": w.value.vevac.copy(),
        "roads": None if w.roads is None else w.roads.copy(),
        "assets": list(w.assets), "ignitions": list(w.ignitions)}
    st.session_state.setdefault("edit_undo", []).append(snap)
    st.session_state.edit_undo = st.session_state.edit_undo[-12:]


def _restore_snapshot() -> None:
    stack = st.session_state.setdefault("edit_undo", [])
    if not stack:
        return
    s = stack.pop()
    w = st.session_state.world
    w.fuel.ftype[:] = s["ftype"]; w.fuel.fload[:] = s["fload"]
    w.fuel.fload0[:] = s["fload0"]; w.fuel.fmoist[:] = s["fmoist"]
    w.topo.elev[:] = s["elev"]; w.topo.slope[:] = s["slope"]
    w.topo.aspect[:] = s["aspect"]; w.topo.access[:] = s["access"]
    w.value.vbld[:] = s["vbld"]; w.value.vcrit[:] = s["vcrit"]
    w.value.vpop[:] = s["vpop"]; w.value.vevac[:] = s["vevac"]
    w.roads = None if s["roads"] is None else s["roads"].copy()
    w.assets[:] = s["assets"]; w.ignitions[:] = s["ignitions"]
    st.session_state.sim.reset(); st.session_state.cost_series = []
    _reset_dss_state()


def _shared_store_path() -> str:
    """Kept as the single name every panel asks for the store by; it now
    resolves to the generated-state file, so no view can end up reading the
    retired learned_rules.json while the engine reasons from the new one."""
    return _gstate_path()


def _active_store_path() -> str:
    """The store the DSS ACTUALLY reads: a loaded run's snapshot wins over
    the shared store."""
    return (getattr(st.session_state.get("dss_engine"),
                    "state_path", None)
            or st.session_state.get("dss_learned_override")
            or _shared_store_path())


def _gstate_path() -> str:
    """The generated-state store: the single source of truth for everything
    the adaptation stages produce."""
    return os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "logs", "dss_generated_state.json")


def _map_key(world=None) -> str | None:
    """A stable identity for the SCENE the DSS is learning on.

    The stage controller's value table is kept per map: what stage 2 is worth
    on a wooded ridge says nothing about what it is worth on a coastal town,
    so carrying the table across maps would teach it the wrong lesson. The
    key is content-based (grid size, fuel layout, assets), which means an
    EDIT to the map counts as a different map and resets the table. That is
    deliberate: an edited map is a different scene, and a stale value table
    is worse than an empty one.
    """
    w = world if world is not None else st.session_state.get("world")
    if w is None:
        return None
    try:
        import hashlib as _hl
        import numpy as _np
        h = _hl.blake2b(digest_size=8)
        _ft = _np.asarray(w.fuel.ftype)
        h.update(f"{_ft.shape}".encode())
        h.update(_np.ascontiguousarray(_ft, dtype=_np.int16).tobytes())
        for a in (getattr(w, "assets", None) or []):
            h.update(f"{getattr(a, 'name', '')}|{getattr(a, 'kind', '')}"
                     f"|{getattr(a, 'x', 0)}|{getattr(a, 'y', 0)}"
                     f"|{getattr(a, 'radius', 0)};".encode())
        return h.hexdigest()
    except Exception:
        return None


def _read_gstate(path: str | None = None):
    """The store as a plain dict, read HERE rather than through the dss
    package. Streamlit re-executes this script on every rerun but keeps
    imported sub-modules cached, so a view depending on a freshly added
    dss.* function stays broken until the whole process restarts."""
    import json as _js
    p = path or _gstate_path()
    try:
        with open(p, encoding="utf-8") as _f:
            return _js.load(_f)
    except Exception:
        return {}


def _store_vocab(path: str | None = None):
    """(concepts, macros) that stage ③ produced, keyed by name."""
    d = _read_gstate(path)
    cons = {}
    for c in d.get("genai_concepts") or []:
        if c.get("name"):
            cons[c["name"]] = dict(
                level=int(c.get("layer", c.get("level", 2))),
                inputs=[[i.get("name"), float(i.get("weight", 0.0))]
                        if isinstance(i, dict) else list(i)
                        for i in (c.get("inputs") or [])])
    macs = {}
    for m in d.get("genai_interventions") or []:
        if m.get("name"):
            macs[m["name"]] = dict(composition=[
                [i.get("channel"), float(i.get("weight", 0.0))]
                if isinstance(i, dict) else list(i)
                for i in (m.get("composition") or [])])
    return cons, macs


def _rebuild_run_stats(run_dir: str) -> dict:
    """The same per-run tally the engine keeps, recovered from a run's log.

    A loaded run has no live engine, and an analysis view that goes blank the
    moment a past run is opened is not much of an analysis. Two log layouts
    are read: decisions.jsonl (one row per REGION per cycle, current) and the
    older cycles.jsonl (one row per cycle)."""
    import json as _js
    st_ = dict(start_step=None, cycles=0, satisficing_failed=0, tried=0,
               accepted=0, rejected=0, blocked={}, per_stage={}, gates={},
               reasons={}, withheld=0, dJ_accepted=0.0, j_series=[])

    def _rows(name):
        try:
            return [ln for ln in open(os.path.join(run_dir, name),
                                      encoding="utf-8").read().splitlines()
                    if ln.strip()]
        except OSError:
            return []

    def _bump(d, k):
        d[str(k)[:70]] = d.get(str(k)[:70], 0) + 1

    # the log stores the RAW threshold; the engine tests against the
    # tightened satisficing bound min(J_TH, (1-min_gain) * J_noaction), so it
    # is recomputed here. Reading the raw value instead reported "the gate
    # never opened" for runs in which stages plainly ran.
    _mg = 0.05
    try:
        import json as _jm
        _mg = float((_jm.load(open(os.path.join(run_dir, "meta.json"),
                                   encoding="utf-8")).get("engine") or {})
                    .get("min_gain", 0.05))
    except Exception:
        pass

    def _bound(j_th, j0):
        return min(float(j_th or 0.0), (1.0 - _mg) * float(j0 or 0.0)) \
            if j0 else float(j_th or 0.0)

    lines = _rows("decisions.jsonl")
    if lines:
        # per-region rows: one CYCLE is one step, and the adaptation stage is
        # a property of the cycle, so each step is counted once
        seen = {}
        for ln in lines:
            try:
                o = _js.loads(ln)
            except Exception:
                continue
            step = int(o.get("step", 0))
            if step not in seen:
                seen[step] = o
            if o.get("failsafe"):
                seen[step]["_fs"] = True
        for step in sorted(seen):
            o = seen[step]
            jc = o.get("j_forecast")
            if not isinstance(jc, (int, float)):
                continue
            j0 = float(o.get("j_noaction") or 0.0)
            bd = _bound(o.get("j_threshold"), j0)
            if st_["start_step"] is None:
                st_["start_step"] = step
            st_["cycles"] += 1
            st_["j_series"].append((step, float(jc), j0, bd))
            if o.get("_fs") or o.get("failsafe"):
                st_["withheld"] += 1
            if jc > bd:
                st_["satisficing_failed"] += 1
            stg = o.get("stage_tried") or 0
            if not stg:
                continue
            st_["tried"] += 1
            ps = st_["per_stage"].setdefault(int(stg),
                                             dict(tried=0, accepted=0,
                                                  dJ=0.0))
            ps["tried"] += 1
            ok = bool(o.get("stage")) and int(o.get("stage") or 0) == int(stg)
            if ok:
                st_["accepted"] += 1
                ps["accepted"] += 1
            else:
                st_["rejected"] += 1
                _bump(st_["reasons"],
                      str(o.get("stage_detail") or "rejected").split(" | ")[0])
            if int(stg) == 3:
                g = o.get("gates")
                _bump(st_["gates"],
                      (g if isinstance(g, str) else
                       (g or {}).get("verdict") if isinstance(g, dict) else
                       ("admitted" if ok else "rejected")))
        return st_

    for ln in _rows("cycles.jsonl"):
        try:
            o = _js.loads(ln)
        except Exception:
            continue
        fc = o.get("forecast") or {}
        jc, j0 = fc.get("j_candidate"), fc.get("j_noaction")
        bd = fc.get("satisficing_bound", fc.get("j_threshold"))
        if not isinstance(jc, (int, float)):
            continue
        step = int(o.get("step", 0))
        if st_["start_step"] is None:
            st_["start_step"] = step
        st_["cycles"] += 1
        st_["j_series"].append((step, float(jc), float(j0 or 0.0),
                                float(bd or 0.0)))
        if o.get("no_harm_withheld"):
            st_["withheld"] += 1
        if isinstance(bd, (int, float)) and jc > bd:
            st_["satisficing_failed"] += 1
        ad = o.get("adaptation") or {}
        stg = ad.get("stage")
        if not stg:
            continue
        st_["tried"] += 1
        ps = st_["per_stage"].setdefault(int(stg),
                                         dict(tried=0, accepted=0, dJ=0.0))
        ps["tried"] += 1
        if ad.get("accepted"):
            st_["accepted"] += 1
            ps["accepted"] += 1
            ps["dJ"] += float(ad.get("dJ") or 0.0)
            st_["dJ_accepted"] += float(ad.get("dJ") or 0.0)
        else:
            st_["rejected"] += 1
            _bump(st_["reasons"],
                  str(ad.get("detail") or "rejected").split(" | ")[0])
        if int(stg) == 3:
            _bump(st_["gates"],
                  ((ad.get("info") or {}).get("gates") or {}).get("verdict")
                  or (ad.get("info") or {}).get("reason")
                  or ("admitted" if ad.get("accepted") else "rejected"))
    return st_


def _visible_interventions(engine=None) -> list:
    """The doctrine families plus only the DISCOVERED actuators.

    The actuator library (tactical_burn, water_drafting, ...) is
    factory physics that no seed rule orders. Until the generative
    stage writes a gate-passing rule that uses one, it must not
    appear in any intervention list or legend: seeing it there read
    as "generated content the wipe failed to delete"."""
    import dss as _dssv
    vis = list(_dssv.rules.DOCTRINE_INTERVENTIONS)
    used = set()
    for r in (getattr(engine, "rules", None) or []):
        if getattr(r, "active", True):
            for iv, _v in r.consequents:
                used.add(str(iv))
    for _mn, _md in (_all_macros(engine) or {}).items():
        for a, _b in (_md.get("composition") or []):
            used.add(str(a))
        for c in (_md.get("clauses") or []):
            used.add("clauses")
    for iv in _dssv.rules.DISCOVERABLE_INTERVENTIONS:
        if iv in used:
            vis.append(iv)
    return vis


def _all_macros(engine=None) -> dict:
    """Every generated intervention that should appear in a list or a legend.

    The engine holds them only while one is running and only when stage ③
    consumption is on. The STORE holds them always, so a macro created in an
    earlier simulation stayed in the file but vanished from every list, which
    read as knowledge lost rather than knowledge not currently used."""
    m = dict(getattr(engine, "macros", {}) or {})
    if m:
        return m
    try:
        _c, _m = _store_vocab()
        return dict(_m or {})
    except Exception:
        return {}


def _mirror(shared: str, wkey: str, default=None) -> None:
    """Point a mirrored control at the shared setting BEFORE it renders.

    The sidebar and the layer panels expose several of the same settings.
    Each widget keeps its own state, so two widgets writing one session key
    fight: the sidebar sets dss_apply=True, the panel toggle then returns its
    own stale False and writes it back, and the two displays disagree inside
    a single render. Assigning the widget's key from the shared value first,
    and copying it back on change, gives the setting ONE source of truth."""
    if shared not in st.session_state and default is not None:
        st.session_state[shared] = default
    if shared in st.session_state:
        st.session_state[wkey] = st.session_state[shared]


def _adopt(shared: str, wkey: str) -> None:
    """on_change partner of _mirror: the widget the user touched wins."""
    st.session_state[shared] = st.session_state[wkey]


def _genai_model_label() -> str:
    """Which model stage ③ will actually call.

    "(plan default)" names nothing: the point of the label is to answer
    "opus, sonnet or haiku?". When the request is left to the plan, the model
    the last live call was SERVED by is shown instead, because that is the
    only honest answer available without making another call."""
    ui = str(st.session_state.get("genai_model_ui", "(plan default)"))
    if ui and ui != "(plan default)":
        return ui
    served = str(st.session_state.get("genai_served_model", "") or "")
    if served:
        low = served.lower()
        for fam in ("opus", "sonnet", "haiku", "fable"):
            if fam in low:
                return f"plan default → {fam}"
        return f"plan default → {served}"
    return "plan default (untested)"


def _all_hierarchy(engine=None) -> dict:
    """The concept hierarchy a view should reason with: base plus whatever
    stage ③ has generated.

    Same reason as _all_macros. Without a live engine the generated concepts
    sit in the store, and a view that reads only the engine silently falls
    back to the base hierarchy, so a concept created in an earlier simulation
    stops being computed and stops being shown."""
    h = getattr(engine, "hierarchy", None)
    if h:
        return dict(h)
    # `dss` is imported panel-locally throughout this file, so it is not a
    # module-level name here
    from dss.concepts import HIERARCHY as _BASE_H
    out = dict(_BASE_H)
    try:
        _c, _m = _store_vocab()
        for name, spec in (_c or {}).items():
            out[name] = (int(spec.get("level", 2)),
                         [(a, float(b)) for a, b in spec.get("inputs", [])])
    except Exception:
        pass
    return out


_STAGE_NAME = {0: "0 — none tried", 1: "1 — evFIS tuning",
               2: "2 — resolution", 3: "3 — GenAI"}


_STEP_COL_HELP = (
    "- **cycle** — the row number in this run, 1, 2, 3 ...\n"
    "- **step / t_min** — the simulation step and the minute of fire time "
    "it stands for. Steps restart with every fire.\n"
    "- **agent** — the region whose situation the adaptation stage worked "
    "on this cycle. It is NOT the coordinator's hotspot. The coordinator "
    "ranks on operational priority, which decides where the capacity goes; "
    "the adaptation instead goes to the region the rule base is QUIETEST "
    "about, among the regions that actually have fire, because stage 2 and "
    "stage 3 exist to cover situations the base cannot answer. Sending them "
    "to the highest-priority region sent them to the best-covered one, and "
    "stage 2 was then turned away with 'cell already covered' in 150 of its "
    "162 attempts. Every agent still decides its own orders in the SAME "
    "cycle, on the Agents tab.\n"
    "- **why_agent** — the reason that region was picked, with each "
    "region's coverage (its strongest fired rule weight) so the choice can "
    "be checked.\n"
    "- **coordination** — the attention share the Global DSS gave each "
    "region this cycle, and which ones it chose to only monitor. The full "
    "ranking and its one-line statement are on the Global DSS tab.\n"
    "- **stage** — which adaptation stage the controller picked: 1 evFIS "
    "tuning, 2 resolution, 3 GenAI. **0 means no stage ran that cycle**, "
    "which is NOT the same as no decision: the concepts were still "
    "inferred, the rules still fired and the orders still went out (see "
    "**applied**), only the learning layer stayed out. The **reason** "
    "column says why: `rules as-is` is the satisficing test passing, so "
    "there was nothing to fix; `adaptation on cooldown` is the gate "
    "opening too soon after the previous attempt; and a stage cannot be "
    "picked at all when DSS active is off or both evFIS and GenAI are "
    "off. In practice the cooldown is the common one.\n"
    "- **target** — the rule the stage acted on (R7, A1, G9), or the "
    "variable for a membership change, or the antecedent cell it aimed at "
    "when it was turned away.\n"
    "- **change** — what it became: `R26 consequents 0.90→0.85`, a whole "
    "new rule body, an inserted term.\n"
    "- **produced** — whether a rule, a concept or an intervention was "
    "written to the store.\n"
    "- **G1 … G5** — `+` cleared, `-` stopped here, `·` never reached, "
    "`n/a` does not apply to this stage.\n"
    "- **bucket** — the situation class the stage controller learned on: "
    "`low`, `mid`, `high` is how far the forecast cost sits above the "
    "satisficing bound, and `+gap` marks a coverage void where no rule "
    "fires at all.\n"
    "- **failsafe** — `WITHHELD` means the no-harm guard pulled the "
    "offensive orders for that whole cycle because the candidate was "
    "forecast to end worse than doing nothing. Life-safety orders stand.\n"
    "- **rec_seq** — the store's replay order, filled only when the step "
    "actually wrote a record. A wipe reverts in reverse rec_seq, a restart "
    "replays forward.\n"
    "- **applied** — the orders that actually went out, per agent, after "
    "the quality gate and the fail-safe had their say.")

_GATE_HELP = (
    "A stage ③ proposal is admitted only if it clears every gate. The marks "
    "read: `+` passed, `-` stopped here, `·` not reached, `n/a` not "
    "applicable to this stage.\n\n"
    "- **G1 — form.** The reply parses as the required JSON and uses only "
    "the declared vocabulary: the five decision concepts, the five terms, "
    "the six base channels.\n"
    "- **G2 — constraints.** The antecedent cell is not already covered by "
    "an active rule, the intensities are in range, and any new object is "
    "well formed (a concept has 1 to 4 inputs and sits above them in the "
    "hierarchy; a composite names 1 to 3 base channels).\n"
    "- **G2b — redundancy.** A proposed CONCEPT must not be collinear with "
    "one that already exists (cosine < 0.95), otherwise the hierarchy grows "
    "without gaining anything.\n"
    "- **G3 — first simulation.** The rule must lower the physical forecast "
    "cost on seed 101: burned area, assets and population at a 45-minute "
    "basis. Response cost is deliberately excluded, so paying for the fleet "
    "is allowed but buying a worse fire is not.\n"
    "- **G4 — second simulation.** The same test on an independent seed "
    "202. A proposal has to survive both futures, not one lucky one.\n"
    "- **G5 — growth margin.** A proposal that also creates a concept or an "
    "intervention must beat no-action by a MARGIN on both seeds, not merely "
    "tie. An ordinary rule may be admitted for being harmless; a permanent "
    "addition to the vocabulary has to earn its place.")

STEP_COLS = ["cycle", "step", "t_min", "agent", "why_agent",
             "coordination", "stage",
             "target", "change", "produced", "G1", "G2", "G2b", "G3", "G4",
             "G5", "verdict", "reason", "dJ", "failsafe", "bucket",
             "rec_seq", "applied"]

# one row per AGENT per cycle: what each local DSS decided, not only the one
# the adaptation happened to work on
AGENT_COLS = ["cycle", "step", "t_min", "agent", "role", "priority", "share",
              "fired", "orders_from_rules", "orders_final", "Q", "failsafe"]

# one row per cycle for the coordinator itself
GLOBAL_COLS = ["cycle", "step", "t_min", "hotspot", "ranking", "shares",
               "monitored", "statement"]

_AGENT_COL_HELP = (
    "- **role** — `focus` is the region the coordinator ranked first, and "
    "the one the adaptation stage works on this cycle. `attended` regions "
    "get their full offensive tempo, `monitor` regions are deliberately "
    "held back so the capacity concentrates where it counts.\n"
    "- **priority** — the operational priority the coordinator ranked on.\n"
    "- **share** — the attention share it assigned. It multiplies the "
    "offensive orders and steers the budget concentration in the "
    "allocator.\n"
    "- **fired** — the rules that fired for THIS agent, with their "
    "weights (strongest first) and the interventions each rule orders "
    "(rule \u2192 orders). Agents see different situations, so they "
    "fire different rules.\n"
    "- **orders_from_rules → orders_final** — the agent's own decision "
    "before coordination, and what it became after the share was applied. "
    "The gap between the two IS the coordinator's intervention.\n"
    "- **Q** — the quality of that agent's decision.\n"
    "- **failsafe** — the graduated fail-safe engaged for that agent."
)

# THE LONG CELLS NEED ROOM. "reason", "change" and "produced" carry whole
# sentences; at the default width they were clipped mid-word and the reader
# lost exactly the part that says what to do about it.
_STEP_COL_CONFIG = {
    "reason": st.column_config.TextColumn("reason", width="large"),
    "coordination": st.column_config.TextColumn("coordination",
                                                width="medium"),
    "why_agent": st.column_config.TextColumn("why_agent", width="medium"),
    "fired": st.column_config.TextColumn("fired", width="medium"),
    "orders_from_rules": st.column_config.TextColumn(
        "orders_from_rules", width="medium"),
    "orders_final": st.column_config.TextColumn("orders_final",
                                                width="medium"),
    "change": st.column_config.TextColumn("change", width="medium"),
    "produced": st.column_config.TextColumn("produced", width="medium"),
    "target": st.column_config.TextColumn("target", width="small"),
}


def build_step_rows(_cycA, _rsA):
    """One row per decision cycle: what was decided, what the
    adaptation tried, which gate stopped it and what was applied.

    Extracted so the narrow panel preview, the modal and the full-width
    page all show the SAME table instead of three that can drift."""
    _rows = []
    _gd = _read_gstate()
    # SCOPE TO THIS RUN. Step numbers restart with every fire, so
    # matching records to cycles on the step alone dragged in
    # records from earlier runs that reached the same step. The
    # store sequence is global, so seq0 (its value when this run
    # started) is the cut.
    # WITHOUT seq0 NOTHING MAY BE ATTRIBUTED. The scope used to be skipped
    # whenever seq0 was missing, which silently turned the join into "match
    # on the step number alone". Step numbers restart with every fire, so an
    # old record from step 32 was then presented as this run's step 32.
    # Showing nothing is honest; showing another run's history as this one's
    # is not.
    _seq0 = _rsA.get("seq0")
    _byStep = {}
    if _seq0 is not None:
        for _sec in ("evfis_rule_modifications", "genai_rules",
                     "genai_concepts", "genai_interventions"):
            for _r in (_gd.get(_sec) or []):
                _s = (_r.get("trigger") or {}).get("step")
                if _s is None:
                    continue
                if int(_r.get("seq", 0)) < int(_seq0):
                    continue
                _byStep.setdefault(int(_s), []).append((_sec, _r))
    for _c in _cycA:
        _ad = _c.get("adaptation") or {}
        _inf = _ad.get("info") or {}
        _ctl = _c.get("stage_controller") or {}
        _st = int(_c.get("step", 0))
        _tried = int(_ad.get("tried") or 0)
        _ok = bool(_ad.get("accepted"))
        _g = _gate_marks(_ad)
        # WHICH RECORDS THIS CYCLE MAY CLAIM. Two conditions, both missing
        # before, which is how a REJECTED GenAI attempt came to display an
        # evFIS consequent tuning as its own result:
        #   the cycle was accepted   - a rejected attempt produced nothing,
        #                              so it may claim nothing
        #   the stage matches        - a record carries the stage that wrote
        #                              it, and stage 1 tunings are not stage
        #                              3 output
        _recs = ([(_sec, _r) for _sec, _r in _byStep.get(_st, [])
                  if int(_r.get("source_stage", 0) or 0) == _tried]
                 if (_ok and _tried) else [])
        # what this step produced, read from the store records
        # stamped with this step
        _prod = []
        for _sec, _r in _recs:
            if _sec == "genai_concepts":
                _prod.append(f"concept {_r.get('name')}")
            elif _sec == "genai_interventions":
                _prod.append(f"intervention {_r.get('name')}")
            elif _sec == "genai_rules":
                _prod.append(f"rule {_r.get('name')}")
            else:
                _mt = _r.get("modification_type")
                if _mt == "rule_add":
                    _prod.append(
                        "rule "
                        + str(((_r.get("after") or {})
                               .get("rule") or {}).get("name")))
                elif _mt == "term_insert":
                    _prod.append(f"term in {_r.get('variable')}")
                else:
                    _prod.append("parameter change")
        _seqs = [int(_r.get("seq", 0)) for _sec, _r in _recs]
        _gl = _c.get("global_dss") or {}
        # WHAT THE COORDINATOR DID. The agent column names only the hotspot,
        # so without this the other agents and the Global DSS were invisible:
        # the table read as if one region were the whole system.
        _shr = _gl.get("shares") or {}
        _att = set(_gl.get("attended") or [])
        _coord = " | ".join(
            f"{_n} {float(_v):.2f}" + ("" if _n in _att else " (monitor)")
            for _n, _v in sorted(_shr.items(), key=lambda kv: -float(kv[1]))
        ) or "—"
        _rows.append(dict(
            cycle=len(_rows) + 1,
            rec_seq=(min(_seqs) if _seqs else None),
            step=_st,
            # THE REGION THE ADAPTATION WORKED ON, which is no longer the
            # coordinator's hotspot: the coordinator ranks on priority, the
            # adaptation goes where the rule base is quietest.
            agent=(_c.get("adapt_region") or _gl.get("hotspot") or "—"),
            why_agent=str(_c.get("adapt_region_why") or ""),
            coordination=_coord,
            t_min=round(float(_c.get("t_min", 0.0)), 1),
            stage=_STAGE_NAME.get(_tried, "0 — none tried"),
            target=_adapt_target(_ad, _recs),
            change=_adapt_change(_ad, _recs),
            produced=", ".join(_prod) or "—",
            G1=_g["G1"], G2=_g["G2"], G2b=_g["G2b"],
            G3=_g["G3"], G4=_g["G4"], G5=_g["G5"],
            verdict=("ACCEPTED" if _ok else
                     ("rejected" if _tried else "—")),
            # FULL SENTENCE. Cutting at 120 characters chopped the useful
            # half off the longer reasons (the timeout one ends with the
            # advice on what to do about it), so the cell now carries the
            # whole text and the table is told to wrap it.
            reason=str(_ad.get("detail") or ""),
            dJ=round(float(_ad.get("dJ") or 0.0), 4),
            failsafe=("WITHHELD"
                      if _c.get("no_harm_withheld") else "ok"),
            bucket=_ctl.get("bucket") or "—",
            applied=_applied_orders(_c)))
    # NEWEST FIRST. The rows are built in cycle order, but the table is read
    # while the fire runs, so the step that just happened has to be the one
    # in view without scrolling to the bottom. Reversed here, at the single
    # source, so the panel preview, the modal and the page all agree.
    _rows.reverse()
    return _rows


def _fmt_orders(d) -> str:
    """The orders of one agent, shortest form that still identifies them."""
    return ", ".join(f"{k.split('_')[0]} {float(v):.2f}"
                     for k, v in (d or {}).items()
                     if float(v) > 0.02) or "none"


def build_agent_rows(_cycA):
    """One row per AGENT per cycle: what each local DSS actually decided.

    The step table carries one row per cycle and names only the hotspot, so
    the other agents and the coordinator never appeared anywhere. They do
    decide every cycle: each has its own concepts, its own rules fire, and
    the Global DSS then scales their offensive tempo by the attention share.
    Both halves of that are here, side by side, so the attenuation is
    visible as the difference between orders_from_rules and orders_final.
    """
    # rule name -> its ordered interventions, so the fired cell says
    # WHAT each rule ordered, not only that it fired
    _rc_map = {}
    try:
        # the step-view test harness runs this function without the
        # streamlit runtime, so the engine lookup must be optional
        _eng_fb = st.session_state.get("dss_engine")
    except Exception:
        _eng_fb = None
    if _eng_fb is not None:
        for _ru in (getattr(_eng_fb, "rules", None) or []):
            try:
                _rc_map[_ru.name] = ", ".join(
                    f"{str(_cch).split('_')[0]} {float(_cv):.2f}"
                    for _cch, _cv in (_ru.consequents or [])[:3])
            except Exception:
                pass
    _rows = []
    for _i, _c in enumerate(_cycA, start=1):
        _gl = _c.get("global_dss") or {}
        _prio = dict(_gl.get("ranking") or [])
        _att = set(_gl.get("attended") or [])
        _hot = _gl.get("hotspot")
        for _name, _r in (_c.get("regions") or {}).items():
            _fired = "; ".join(
                f"{_rn} {float(_w):.2f}"
                + (f" → {_rc_map[_rn]}"
                   if _rc_map.get(_rn) else "")
                for _rn, _w in (_r.get("fired") or [])[:4]
                if float(_w) > 0.05) or "none above 0.05"
            _rows.append(dict(
                cycle=_i,
                step=int(_c.get("step", 0)),
                t_min=round(float(_c.get("t_min", 0.0)), 1),
                agent=_name,
                role=("focus" if _name == _hot else
                      ("attended" if _name in _att else "monitor")),
                priority=(round(float(_prio[_name]), 3)
                          if _name in _prio else None),
                share=_r.get("coord_share"),
                fired=_fired,
                orders_from_rules=_fmt_orders(_r.get("orders_from_rules")),
                orders_final=_fmt_orders(_r.get("orders_final")),
                Q=_r.get("quality"),
                failsafe=("FAIL-SAFE" if _r.get("failsafe") else "ok")))
    _rows.reverse()          # same ordering rule as the step table
    return _rows


def build_global_rows(_cycA):
    """One row per cycle for the coordinator's own decision.

    The Global DSS states its verdict explicitly every cycle: it ranks the
    regions on operational priority, assigns the attention shares and says
    which regions it will only monitor. That statement was written to the
    chronicle but never shown anywhere.
    """
    _rows = []
    for _i, _c in enumerate(_cycA, start=1):
        _gl = _c.get("global_dss") or {}
        if not _gl:
            continue
        _rank = _gl.get("ranking") or []
        _att = set(_gl.get("attended") or [])
        _shr = _gl.get("shares") or {}
        _rows.append(dict(
            cycle=_i,
            step=int(_c.get("step", 0)),
            t_min=round(float(_c.get("t_min", 0.0)), 1),
            hotspot=_gl.get("hotspot") or "—",
            ranking=" > ".join(f"{_n} {float(_v):.2f}" for _n, _v in _rank)
            or "—",
            shares=" | ".join(f"{_n} {float(_v):.2f}"
                              for _n, _v in sorted(
                                  _shr.items(),
                                  key=lambda kv: -float(kv[1]))) or "—",
            monitored=", ".join(_n for _n in _shr if _n not in _att)
            or "none",
            statement=_gl.get("statement") or ""))
    _rows.reverse()
    return _rows


def _gate_marks(ad: dict) -> dict:
    """Per-gate verdict of one adaptation attempt.

    `+` passed, `-` stopped here, `·` never reached, `n/a` not applicable.
    Stages 1 and 2 do not go through the generative gate chain, so their row
    says n/a rather than pretending to a pass."""
    ORDER = ("G1", "G2", "G2b", "G3", "G4", "G5")
    out = {k: "n/a" for k in ORDER}
    if int(ad.get("tried") or 0) != 3:
        return out
    info = ad.get("info") or {}
    gates = info.get("gates") or {}
    det = str(gates.get("verdict") or ad.get("detail") or "")
    pkg = bool(info.get("package"))
    # The ENGINE's own verdict decides, not a recomputation here. Deriving
    # pass/fail from the logged j values used a different tolerance than the
    # gate itself, which produced rows marked "+" on every gate and still
    # rejected, contradicting the reason printed beside them.
    fail = None
    low = det.lower()
    # "admitted" wins: the verdict of an accepted package reads "G5 margin
    # CLEARED", and matching on the word margin alone marked it as a failure
    if bool(ad.get("accepted")) or low.startswith("admitted"):
        low = ""
    if "rejected" in low or "no improvement" in low or "failed" in low:
        for g in ("G5", "G4", "G3", "G2b", "G2", "G1"):
            if g.lower() in low:
                fail = g
                break
        if fail is None:
            fail = "G1"          # died before any named gate
    if fail is None:
        # admitted: everything that applies was cleared
        for k in ORDER:
            out[k] = "+"
        if not pkg:
            out["G2b"] = out["G5"] = "n/a"
        return out
    hit = False
    for k in ORDER:
        if k == fail:
            out[k] = "-"
            hit = True
        elif not hit:
            out[k] = "+"
        else:
            out[k] = "·"        # never reached
    if not pkg:
        # a plain rule never faces the vocabulary gates
        if out["G2b"] != "-":
            out["G2b"] = "n/a"
        if out["G5"] != "-":
            out["G5"] = "n/a"
    return out


def _adapt_target(ad: dict, recs) -> str:
    """Which rule or variable the stage acted on."""
    for _sec, r in recs or []:
        if _sec == "genai_rules":
            return str(r.get("name") or "")
        mt = r.get("modification_type")
        if mt == "rule_add":
            return str(((r.get("after") or {}).get("rule") or {}).get("name"))
        if mt == "consequent_update":
            return str(r.get("base_rule_id") or "")
        if mt in ("membership_shift", "term_insert"):
            return str(r.get("variable") or "")
    # nothing was stored, so the attempt was REJECTED: name what it aimed at
    info = ad.get("info") or {}
    _r = info.get("rule")
    if isinstance(_r, str) and _r and " " not in _r:
        return _r                       # stage 1 records the rule name
    cell = info.get("cell")
    if cell:
        return " AND ".join(f"{c[0]}={c[1]}" for c in cell[:2])
    det = str(ad.get("detail") or "")
    for tok in det.replace(",", " ").split():
        if tok[:1] in "RAG" and tok[1:2].isdigit():
            return tok
    return "—"


def _adapt_change(ad: dict, recs) -> str:
    """What it turned INTO: the point of the whole view."""
    for _sec, r in recs or []:
        mt = r.get("modification_type")
        aft, bef = r.get("after") or {}, r.get("before") or {}
        if mt == "consequent_update":
            o = dict((str(i), float(v))
                     for i, v in bef.get("consequents", []))
            n = dict((str(i), float(v))
                     for i, v in aft.get("consequents", []))
            d = ", ".join(f"{k} {o.get(k, 0.0):.2f}→{v:.2f}"
                          for k, v in n.items()
                          if abs(v - o.get(k, 0.0)) > 1e-9)
            return d or "no net change"
        if mt == "term_insert":
            new = sorted(set(aft.get("partition") or {})
                         - set(bef.get("partition") or {}))
            return f"term {', '.join(new)} inserted (catalog grew)"
        if mt == "membership_shift":
            return "membership boundary moved"
        if mt == "rule_add" or _sec == "genai_rules":
            sp = (aft.get("rule") or {}) if mt == "rule_add" else r
            ants = " AND ".join(f"{a[0]} is {a[1]}"
                                for a in sp.get("antecedents", []))
            cons = ", ".join(f"{i} {float(v):.2f}"
                             for i, v in sp.get("consequents", []))
            return f"NEW: IF {ants} THEN {cons}"
    # rejected: say what was TRIED, so the row still explains itself
    info = ad.get("info") or {}
    tr = info.get("trials") or []
    if tr:
        _t = tr[0]
        if _t.get("kind") == "consequent":
            return (f"tried {_t.get('rule')} consequents "
                    f"{float(_t.get('delta', 0)):+g}, kept={_t.get('kept')}")
        if _t.get("kind") == "membership":
            return (f"tried {_t.get('var')}.{_t.get('term')} "
                    f"{_t.get('move', 'boundary move')}, "
                    f"kept={_t.get('kept')}")
    _r = info.get("rule")
    if isinstance(_r, str) and " " in _r:
        return "proposed: " + _r[:90]
    if info.get("cell"):
        return "aimed at an already covered cell"
    return "—"


def _applied_orders(cyc: dict) -> str:
    """The orders that actually went out this cycle, per region."""
    out = []
    for name, r in (cyc.get("regions") or {}).items():
        u = r.get("orders_final") or r.get("orders") or r.get("u") or {}
        bits = ", ".join(f"{k.split('_')[0]} {float(v):.2f}"
                         for k, v in u.items() if float(v) > 0.02)
        if bits:
            out.append(f"{name}: {bits}")
    return " | ".join(out) or "none"


def _gstate_counts(path: str | None = None) -> dict:
    d = _read_gstate(path)
    return {k: len(d.get(k) or []) for k in
            ("evfis_rule_modifications", "genai_rules",
             "genai_concepts", "genai_interventions")}


def _ignition_warning(world, gx: int, gy: int, radius: int = 0):
    """Say so when an ignition lands where nothing can burn.

    A road, a rock outcrop or open water cannot carry fire, and roads ring
    every settlement: measured on a generated landscape, 38% of the ring of
    cells immediately around the towns is unburnable. Clicking there is a
    perfectly reasonable thing to do and the map used to answer with
    silence, which reads as a broken click rather than as a fuel break
    doing its job.
    """
    try:
        import numpy as _np
        fl0 = _np.asarray(world.fuel.fload0)
        eps = float(world.config.spread.eps_fuel)
        ny, nx = fl0.shape
        r = max(0, int(radius))
        y0, y1 = max(0, int(gy) - r), min(ny, int(gy) + r + 1)
        x0, x1 = max(0, int(gx) - r), min(nx, int(gx) + r + 1)
        patch = fl0[y0:y1, x0:x1]
        if patch.size and float(patch.max()) > eps:
            return None
        from disaster_phyengine.config import FUEL_MODELS
        _ft = int(_np.asarray(world.fuel.ftype)[int(gy), int(gx)])
        _nm = FUEL_MODELS[_ft].name if _ft in FUEL_MODELS else "this cover"
        return (f"Nothing to burn at ({int(gx)}, {int(gy)}): "
                f"{_nm.replace('_', ' ')} carries no fuel, so the ignition "
                "cannot take. Roads and bare ground ring most settlements "
                "and act as fuel breaks; place the ignition on vegetation "
                "or inside the built-up area itself.")
    except Exception:
        return None


def _reset_dss_state(drop_engine: bool = False) -> None:
    """A fire reset clears the DECISION state (gating priors, feature
    histories, per-run transients) but the engine SURVIVES: learned
    rules, membership moves and the controller value table are knowledge, not
    decisions, and persist across fires. drop_engine=True (map
    regeneration) discards the engine too."""
    for _k in list(st.session_state.keys()):
        if _k.startswith(("l3_gate_", "dss_featprev_")):
            del st.session_state[_k]
    st.session_state.pop("dss_decision_log", None)   # fresh fire, fresh log
    _eng_fr = st.session_state.get("dss_engine")
    if drop_engine or _eng_fr is None:
        for _k in ("dss_engine", "dss_engine_sig"):
            st.session_state.pop(_k, None)
    else:
        try:
            # ONE FIRE = ONE LOG DIRECTORY: rotate the run logger
            # BEFORE new_fire so the rules snapshot of the fresh fire
            # lands in the fresh directory; the meta travels along
            _lgo = getattr(_eng_fr, "run_logger", None)
            if _lgo is not None:
                import json as _json_rl
                import os as _os_lg
                _root_lg = _os_lg.path.dirname(_lgo.dir)
                import dss as _dss_lg
                _lgn = _dss_lg.RunLogger(
                    _root_lg, tag=getattr(_lgo, "tag", "run"))
                try:
                    with open(_os_lg.path.join(_lgo.dir,
                                               "meta.json")) as _fm:
                        _lgn.write_meta(_json_rl.load(_fm))
                except Exception:
                    pass
                _eng_fr.run_logger = _lgn
            _eng_fr.new_fire()
        except Exception:
            for _k in ("dss_engine", "dss_engine_sig"):
                st.session_state.pop(_k, None)


def _record_costs() -> None:
    st.session_state.cost_series.append(compute_costs(st.session_state.sim).to_dict())


def _model_ids() -> dict:
    """The pinned generative and review model identities, for meta.json.

    A campaign's claim "all runs used model X" is only verifiable if
    every run log carries the identity; the env overrides are read
    HERE, once, so a mid-campaign environment change is visible in
    the logs instead of silent."""
    import os as _os_mi
    try:
        from dss import genai as _gn_mi
        _g = _gn_mi.current_model()
    except Exception:
        _g = _os_mi.environ.get("DSS_GENAI_MODEL", "?")
    return dict(genai=_g,
                rca=_os_mi.environ.get("DSS_RCA_MODEL", "opus"))


def _fit_scale(nx) -> int:
    return int(max(4, min(16, 900 // max(nx, 1))))


_IV_COLOR = {"suppression_effort": "#2878ff",
             "resource_deployment": "#9aa0a6",
             "containment_line": "#96501e",
             "asset_protection": "#28dc5a",
             "evacuation": "#ff8c00",
             "public_warning": "#e6c400",
             "tactical_burn": "#ff5a1e",
             "water_drafting": "#3caaff",
             "retardant_drop": "#c85ac8"}


def _iv_bar(label: str, value: float, color: str) -> str:
    """One colored intensity bar (the intervention palette matches
    the map icons and the legend)."""
    v = max(0.0, min(1.0, float(value)))
    return (
        "<div style='margin:3px 0'>"
        f"<span style='display:inline-block;width:11px;height:11px;"
        f"background:{color};border-radius:2px;margin-right:6px'>"
        "</span>"
        f"<span style='font-size:0.86em'>{label}: "
        f"<b>{v:.2f}</b></span>"
        "<div style='background:#8882;border-radius:3px;height:7px;"
        "margin-top:2px'>"
        f"<div style='width:{v * 100:.0f}%;background:{color};"
        "height:7px;border-radius:3px'></div></div></div>")


XLSX_MIME = ("application/vnd.openxmlformats-officedocument"
             ".spreadsheetml.sheet")


def _xlsx_bytes(sheets: dict, meta: dict | None = None) -> bytes:
    """Rows to a formatted workbook, one sheet per entry.

    The tables were offered as TSV, which Excel opens as one column unless
    the reader knows to run the import wizard, and which loses every number
    to text. A workbook keeps the columns, keeps the numbers as numbers so
    they sort and chart, and can carry the related tables side by side
    instead of as three files to line up by hand.

    `sheets` maps a sheet name to a list of row dicts. `meta` becomes a
    leading sheet recording what the numbers describe, because a table
    detached from its run configuration is not evidence.
    """
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)
    head_font = Font(bold=True, color="FFFFFF")
    head_fill = PatternFill("solid", fgColor="44546A")

    if meta:
        ws = wb.create_sheet("Run")
        ws.append(["field", "value"])
        for c in (1, 2):
            ws.cell(row=1, column=c).font = head_font
            ws.cell(row=1, column=c).fill = head_fill
        for k, v in meta.items():
            ws.append([str(k), str(v)])
        ws.column_dimensions["A"].width = 26
        ws.column_dimensions["B"].width = 80
        ws.freeze_panes = "A2"

    for name, rows in sheets.items():
        ws = wb.create_sheet(str(name)[:31])
        rows = list(rows or [])
        if not rows:
            ws.append(["(nothing recorded yet)"])
            continue
        cols = list(rows[0].keys())
        ws.append(cols)
        for i in range(1, len(cols) + 1):
            c = ws.cell(row=1, column=i)
            c.font, c.fill = head_font, head_fill
            c.alignment = Alignment(vertical="center", wrap_text=True)
        for r in rows:
            ws.append([_xl_cell(r.get(k)) for k in cols])
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = (f"A1:{get_column_letter(len(cols))}"
                              f"{len(rows) + 1}")
        for i, k in enumerate(cols, start=1):
            # width from the CONTENT, capped: a reason column carries whole
            # sentences and would otherwise push everything off the screen
            w = max(len(str(k)),
                    *(len(str(r.get(k, ""))) for r in rows[:400]))
            ws.column_dimensions[get_column_letter(i)].width = \
                min(60, max(9, w + 2))
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _xl_cell(v):
    """Numbers stay numbers so the sheet can sort and chart them."""
    if isinstance(v, bool) or v is None:
        return "" if v is None else v
    if isinstance(v, (int, float)):
        return v
    return str(v)


_SWATCH_CACHE: dict = {}


def _legend_swatch(hexc: str, glyph: str, px: int = 16) -> str:
    """One legend swatch, drawn by the MAP's own code.

    These used to be hand-written CSS shapes living in this file while the
    map was drawn with PIL in viz.py. Two definitions of one symbol drift
    apart by construction, and they had: the reader was left matching a
    coloured blob in the legend against a different glyph on the map. The
    swatch is now the same function call the renderer makes, at a smaller
    radius, embedded as a PNG. There is one definition of every symbol.
    """
    key = (hexc, glyph, px)
    if key in _SWATCH_CACHE:
        return _SWATCH_CACHE[key]
    try:
        import base64 as _b64
        from disaster_phyengine.viz import legend_icon_png
        _h = hexc.lstrip("#")
        rgb = tuple(int(_h[k:k + 2], 16) for k in (0, 2, 4))
        png = legend_icon_png(glyph, rgb, px=px)
        html = ("<img alt='' style='width:%dpx;height:%dpx;flex:none;"
                "vertical-align:middle;image-rendering:auto' "
                "src='data:image/png;base64,%s'>"
                % (px, px, _b64.b64encode(png).decode()))
    except Exception:
        # a legend that cannot draw itself must not take the page with it
        html = ("<span style='width:%dpx;height:%dpx;display:inline-block;"
                "flex:none;background:%s;border:1px solid #555'></span>"
                % (px, px, hexc))
    _SWATCH_CACHE[key] = html
    return html


_LEGEND_ORDER = ["Land cover", "Assets",
                 "Fire", "Markers",
                 "Sensors (+ coverage fill)", "Resources",
                 "DSS orders (base)", "DSS orders (GenAI-generated)",
                 "Agents"]
# a separator line is drawn BEFORE each of these groups (block boundaries)
_LEGEND_BLOCK_START = {"Fire", "Sensors (+ coverage fill)",
                       "DSS orders (base)"}


def _legend_ordered(groups):
    return ([g for g in _LEGEND_ORDER if g in groups]
            + [g for g in groups if g not in _LEGEND_ORDER])


# Single source of truth for the map layer toggles. Both the Simulation
# page's Display panel and the Map editor's Layers box render THIS list, so
# the two stay identical and a tick made on either page holds on both (the
# state lives in one session_state key per layer).
_LAYER_DEFS = [
    ("ly_relief_v", "Relief", True, None),
    ("ly_fire_v", "Fire", True, None),
    ("ly_val_v", "Protection value", True,
     "Tints asset cells by the protection priority $V_{prio}$: pale pink "
     "= lower, deep purple = higher priority. Nothing shows until the map "
     "has buildings, critical facilities or population (Asset tool)."),
    ("ly_roads_v", "Roads", True, None),
    ("ly_grid_v", "Grid", True, None),
    ("ly_per_v", "Fire perimeter", True, None),
    ("ly_orders_v", "DSS orders (icons)", True, None),
    ("ly_alloc_v", "DSS allocation glow", True, None),
    ("ly_sens_v", "Sensors + coverage", True, None),
    ("ly_deps_v", "Resource depots + service areas", True, None),
    ("ly_agents_v", "Agent regions (Local DSS boundaries)", True, None),
]


def _render_layer_toggles(ncols: int = 4, key_prefix: str = "lyr"):
    """Draw the shared layer checkboxes. Rendered ONCE per run (the global
    Map layers bar at the top of every page); every page's render reads the
    resulting ly_*_v flags through _layer_flags()."""
    _cols = st.columns(ncols)
    for _i, (_k, _lab, _d, _help) in enumerate(_LAYER_DEFS):
        with _cols[_i % ncols]:
            st.session_state[_k] = st.checkbox(
                _lab, value=bool(st.session_state.get(_k, _d)),
                help=_help, key=f"{key_prefix}_{_k}")


def _pct_to_100(values):
    """Integer percentages of `values` that ALWAYS sum to 100 (largest
    remainder / Hamilton rounding), so a shares readout never shows 101%."""
    _s = float(sum(values))
    if _s <= 1e-12:
        return [0 for _ in values]
    _raw = [v / _s * 100.0 for v in values]
    _floor = [int(x // 1) for x in _raw]
    _rem = 100 - sum(_floor)
    # hand the leftover points to the largest fractional parts, in order
    _order = sorted(range(len(_raw)),
                    key=lambda i: (_raw[i] - _floor[i]), reverse=True)
    for _i in _order[:max(0, _rem)]:
        _floor[_i] += 1
    return _floor


def _layer_flags():
    """Base render flags shared by both pages, read from the toggles."""
    return dict(show_hillshade=bool(st.session_state.get("ly_relief_v", True)),
                show_fire=bool(st.session_state.get("ly_fire_v", True)),
                show_value=bool(st.session_state.get("ly_val_v", True)),
                show_roads=bool(st.session_state.get("ly_roads_v", True)),
                show_grid=bool(st.session_state.get("ly_grid_v", True)),
                show_perimeter=bool(st.session_state.get("ly_per_v", True)))


def legend_html(horizontal: bool = False, macros=None) -> str:
    if macros is None:
        macros = getattr(st.session_state.get("dss_engine"), "macros", None)
    if not macros:
        # No live engine yet, or one built before the store was read: the
        # generated interventions are still in the store and belong in the
        # legend. Reading only the engine made every macro created in an
        # earlier simulation disappear from the legend on restart, which
        # looked like the knowledge had been lost.
        macros = _all_macros(None) or None
    groups = {}
    for grp, lab, hexc, glyph in viz.legend_entries(macros=macros):
        groups.setdefault(grp, []).append((lab, hexc, glyph))
    _ordered = _legend_ordered(groups)
    if horizontal:
        # one line per CATEGORY, in fixed order, with a bolder separator
        # between the meta-blocks (cover/assets | fire/markers | sensors/
        # resources | DSS orders).
        _sep = ("<div style='border-top:2px solid #8888;"
                "margin:5px 0 3px'></div>")
        html = "<div style='font-size:0.8em;margin-top:2px'>"
        for _n, grp in enumerate(_ordered):
            if _n and grp in _LEGEND_BLOCK_START:
                html += _sep
            html += ("<div style='display:flex;flex-wrap:wrap;"
                     "gap:3px 12px;align-items:center;margin:2px 0;"
                     "padding:1px 0;border-bottom:1px solid #8882'>"
                     f"<span style='font-weight:600;min-width:110px'>"
                     f"{grp}</span>")
            for lab, hexc, glyph in groups[grp]:
                html += ("<span style='display:inline-flex;align-items:"
                         "center;gap:4px'>"
                         + _legend_swatch(hexc, glyph, px=14)
                         + f"<span>{lab}</span></span>")
            html += "</div>"
        return html + "</div>"
    html = "<div style='font-size:0.9em'>"
    for _n, grp in enumerate(_ordered):
        items = groups[grp]
        if _n and grp in _LEGEND_BLOCK_START:
            html += ("<div style='border-top:2px solid #8888;"
                     "margin:6px 0 3px'></div>")
        html += f"<div style='font-weight:600;margin:6px 0 2px'>{grp}</div>"
        for lab, hexc, glyph in items:
            html += ("<div style='display:flex;align-items:center;"
                     "gap:6px;margin:1px 0'>"
                     + _legend_swatch(hexc, glyph)
                     + f"<span>{lab}</span></div>")
    return html + "</div>"


# ---- engine freshness gate: a zombie server or stale module otherwise
# surfaces as confusing TypeErrors deep inside the pages ----
import disaster_phyengine as _dpe
import dss as _dss_pkg
_EXPECTED_ENGINE_BUILD = 51
_EXPECTED_DSS_BUILD = 91
if (getattr(_dpe, "ENGINE_BUILD", 0) != _EXPECTED_ENGINE_BUILD
        or getattr(_dss_pkg, "DSS_BUILD", 0) != _EXPECTED_DSS_BUILD):
    st.error(
        "**Old engine code is still running in this process.**\n\n"
        "The files on disk are up to date, but an older DisasterAware "
        "server is still alive and your browser is talking to it (check "
        "the port in the address bar vs. the one the launcher printed).\n\n"
        "Fix:\n"
        "1. Close **every** DisasterAware terminal window \u2014 or run "
        "`taskkill /F /IM python.exe` in a command prompt,\n"
        "2. start `run_dashboard.bat` again,\n"
        "3. open the **new** URL it prints (usually "
        "http://localhost:8501).")
    st.stop()

_ensure_state()
sim: Simulator = st.session_state.sim
world: World = st.session_state.world
cfg = world.config

# one-time migration: the elliptical kernel and ember spotting became
# default-ON (engine v0.2). Worlds created earlier carry the old flags in
# their stored config, so flip them once per session; afterwards the
# Parameters checkboxes stay authoritative.
if not st.session_state.get("realism_on_migrated"):
    cfg.spread.elliptical = True
    cfg.spread.spotting = True
    st.session_state["realism_on_migrated"] = True
# RUN LIMIT: 50 steps by default. It used to be lifted to 100,000 on first
# load, which made every run effectively unbounded and let a scenario run for
# minutes before anyone noticed. A short default is the honest starting point
# for an interactive experiment; raise it in Parameters when a run needs it.
if not st.session_state.get("maxsteps_default_50"):
    cfg.max_steps = 50
    st.session_state["maxsteps_default_50"] = True


# ------------------------------------------------------------ canvas parsing
def _clip(gx, gy):
    return (int(np.clip(gx, 0, cfg.nx - 1)), int(np.clip(gy, 0, cfg.ny - 1)))


def _path_points(obj):
    pts = []
    for cmd in obj.get("path", []):
        nums = [v for v in cmd[1:] if isinstance(v, (int, float))]
        for i in range(0, len(nums) - 1, 2):
            pts.append((nums[i], nums[i + 1]))
    return pts


def _do_point(gx, gy, kw):
    tool = kw["tool"]
    if tool == "Fuel":
        world.paint_disk(gx, gy, kw.get("brush", 0), FUEL_NAME_TO_ID[kw["ftype"]],
                         load=kw["load"], moisture=kw["moisture"])
    elif tool == "Firebreak":
        world.paint_disk(gx, gy, kw.get("brush", 0), kw["fbid"], load=0.0)
    elif tool == "Access":
        world.add_road_disk(gx, gy, kw.get("brush", 1))
    elif tool == "Asset":
        name = kw["aname"] or ASSET_LABELS.get(kw["akind"], "Asset")
        world.add_asset(Asset(name, kw["akind"], gx, gy, kw["aradius"],
                              kw["avalue"], kw["apop"]))
    elif tool == "Settlement":
        # A WHOLE TOWN, NOT A MARKER. The Asset tool drops one point; a
        # settlement is a painted block of built-up ground with a street
        # grid, its people spread across it and civic facilities around the
        # centre. It is built by the SAME function the generator uses, so a
        # hand-placed town and a generated one are the same kind of thing.
        from disaster_phyengine import terrain as _terr
        _terr.place_settlement(
            world, int(gx), int(gy), int(kw.get("spop", 5000)),
            building_scale=float(kw.get("sdens", 1.0)),
            rng=np.random.default_rng(int(kw.get("sseed", 0)) or None),
            main=bool(kw.get("smain", False)),
            name=(kw.get("sname") or None),
            label_index=1 + sum(1 for a in world.assets
                                if getattr(a, "kind", "") == "population"))
        world.rebuild_value_layers()
    elif tool == "Elevation":
        world.bump_terrain(gx, gy, kw.get("brush", 3),
                           kw.get("elev_delta", 40.0), recompute=False)


def _apply_edits(objects, scale, kw):
    tool = kw["tool"]

    def _g(px, py):
        return px / scale, py / scale

    n = 0
    for obj in objects:
        otype = obj.get("type")
        if otype == "rect":
            left, top = obj.get("left", 0), obj.get("top", 0)
            w = obj.get("width", 0) * obj.get("scaleX", 1)
            h = obj.get("height", 0) * obj.get("scaleY", 1)
            x0, y0 = _clip(*_g(left, top))
            x1, y1 = _clip(*_g(left + w, top + h))
            if tool == "Fuel":
                world.paint_rect(x0, y0, x1, y1, FUEL_NAME_TO_ID[kw["ftype"]],
                                 load=kw["load"], moisture=kw["moisture"])
            elif tool == "Firebreak":
                world.paint_rect(x0, y0, x1, y1, kw["fbid"], load=0.0)
            elif tool == "Access":
                world.add_road_rect(x0, y0, x1, y1)
            n += 1
        elif otype == "path":
            # walk the recorded points and fill the gap between each pair so a
            # fast free-hand stroke stays continuous instead of dropping dots
            seen = set()
            prev = None
            for px, py in _path_points(obj):
                cur = _g(px, py)
                if prev is None:
                    seg = [cur]
                else:
                    d = max(abs(cur[0] - prev[0]), abs(cur[1] - prev[1]))
                    m = max(1, int(np.ceil(d)))
                    seg = [(prev[0] + (cur[0] - prev[0]) * t / m,
                            prev[1] + (cur[1] - prev[1]) * t / m)
                           for t in range(1, m + 1)]
                for sx, sy in seg:
                    gxy = _clip(sx, sy)
                    if gxy in seen:
                        continue
                    seen.add(gxy)
                    _do_point(gxy[0], gxy[1], kw)
                prev = cur
            n += 1
        elif otype == "circle":
            left, top = obj.get("left", 0), obj.get("top", 0)
            rad = obj.get("radius", obj.get("width", 0) / 2)
            gx, gy = _clip(*_g(left + rad, top + rad))
            _do_point(gx, gy, kw)
            n += 1
    if kw.get("tool") == "Elevation" and n:
        world.recompute_slope_aspect()
    return n


def _wind_speed_label(ws: float) -> str:
    for lim, name in [(0.4, "calm"), (1.5, "light air"), (3.3, "light breeze"),
                      (5.4, "gentle breeze"), (7.9, "moderate breeze"),
                      (10.7, "fresh breeze"), (13.8, "strong breeze"),
                      (17.1, "near gale"), (20.7, "gale"), (24.4, "strong gale"),
                      (28.4, "storm"), (32.6, "violent storm")]:
        if ws <= lim:
            break
    else:
        name = "hurricane force"
    return f"{ws:.1f} m/s  \u2248 {ws*3.6:.0f} km/h  \u00b7  {name}"


def _fmt_sim_time(minutes: float) -> str:
    m = int(round(minutes))
    d, rem = divmod(m, 1440)
    h, mm = divmod(rem, 60)
    parts = []
    if d:
        parts.append(f"{d} d")
    if h:
        parts.append(f"{h} h")
    if mm or not parts:
        parts.append(f"{mm} min")
    return " ".join(parts)


def _wind_compass(cur_deg: float, key: str):
    """Clickable compass dial (cartesian scatter, so point clicks work).

    Returns the clicked direction in degrees (math convention: 0 = east,
    counter-clockwise) or None. The arrow points where the wind blows toward;
    the ring is labelled at 0/90/180/270 degrees."""
    import plotly.graph_objects as go
    opts = list(range(0, 360, 15))
    tho = np.radians(opts)
    near = min(opts, key=lambda o: min((o - cur_deg) % 360, (cur_deg - o) % 360))
    cols = ["#d62728" if o == near else "rgba(130,130,130,0.55)" for o in opts]
    fig = go.Figure()
    tt = np.linspace(0, 2 * np.pi, 91)
    fig.add_trace(go.Scatter(x=np.cos(tt), y=np.sin(tt), mode="lines",
                             line=dict(width=1, color="rgba(128,128,128,0.35)"),
                             hoverinfo="skip"))
    a = np.radians(cur_deg)
    fig.add_trace(go.Scatter(x=[0, 0.72 * np.cos(a)], y=[0, 0.72 * np.sin(a)],
                             mode="lines", line=dict(width=5, color="#1f77b4"),
                             hoverinfo="skip"))
    fig.add_trace(go.Scatter(x=[0.80 * np.cos(a)], y=[0.80 * np.sin(a)],
                             mode="markers",
                             marker=dict(symbol="arrow", size=16,
                                         color="#1f77b4",
                                         angle=(90 - cur_deg) % 360),
                             hoverinfo="skip"))
    fig.add_trace(go.Scatter(x=np.cos(tho), y=np.sin(tho), mode="markers",
                             marker=dict(size=10, color=cols),
                             customdata=opts,
                             hovertemplate="%{customdata}\u00b0"
                                           "<extra>click to set</extra>"))
    for ang, (tx, ty) in [(0, (1.24, 0)), (90, (0, 1.24)),
                          (180, (-1.26, 0)), (270, (0, -1.26))]:
        fig.add_annotation(x=tx, y=ty, text=f"{ang}\u00b0", showarrow=False,
                           font=dict(size=11))
    fig.update_layout(height=205, margin=dict(l=6, r=6, t=6, b=6),
                      showlegend=False, dragmode=False,
                      paper_bgcolor="rgba(0,0,0,0)",
                      plot_bgcolor="rgba(0,0,0,0)",
                      xaxis=dict(visible=False, range=[-1.4, 1.4],
                                 fixedrange=True),
                      yaxis=dict(visible=False, range=[-1.4, 1.4],
                                 fixedrange=True, scaleanchor="x"))
    ev = st.plotly_chart(fig, use_container_width=True,
                         config={"displayModeBar": False},
                         on_select="rerun", selection_mode="points", key=key)
    for pt in _selection_points(ev):
        cd = pt.get("customdata")
        if cd is not None:
            try:
                cd = cd[0] if isinstance(cd, (list, tuple)) else cd
                return float(cd) % 360.0
            except (TypeError, ValueError):
                pass
        x_, y_ = pt.get("x"), pt.get("y")
        if x_ is not None and y_ is not None and (x_ or y_):
            return float(np.degrees(np.arctan2(y_, x_))) % 360.0
    return None



def _clock_info():
    """(label, brightness) of the simulation clock for the map display."""
    import math as _m
    import datetime as _dt
    sm = float(getattr(cfg, "step_minutes", 1.0))
    t_min = sim.state.step * sm
    mode = st.session_state.get("wx_mode", "Manual")
    if mode == "Real case weather" and st.session_state.get("real_wx"):
        wx = st.session_state["real_wx"]
        ts = (_dt.datetime.fromisoformat(wx["start"])
              + _dt.timedelta(hours=float(wx["h0"]), minutes=t_min))
        hh = ts.hour + ts.minute / 60.0
        label = ts.strftime("%Y-%m-%d %H:%M UTC")
    elif mode == "Diurnal cycle":
        start = float(st.session_state.get("dr_start", 12.0))
        tot = start * 60.0 + t_min
        hh = (tot / 60.0) % 24.0
        label = f"Day {int(tot // 1440) + 1} \u00b7 " \
                f"{int(hh):02d}:{int((hh % 1) * 60):02d}"
    else:
        return f"t = {_fmt_sim_time(t_min)}", 1.0
    bright = 0.60 + 0.40 * max(0.0, _m.cos((hh - 13.0) / 24.0
                                           * 2.0 * _m.pi)) ** 0.8
    return label, bright


def _drive_weather(wobj, t_min: float) -> None:
    """Exogenous weather at simulated minute t_min, applied to wobj.

    Deterministic in t_min and the UI settings, so the counterfactual
    replay drives its CLONE through exactly the weather history the
    factual run saw (diurnal wave, EMC fuel-moisture drying, wind
    veer, shower windows). Without this the replay froze the weather
    at the rewind point and burned a different fire."""
    import math as _m
    mode = st.session_state.get("wx_mode", "Manual")
    if mode == "Real case weather" and st.session_state.get("real_wx"):
        wx = st.session_state["real_wx"]
        h = min(int(wx["h0"] + t_min // 60.0), len(wx["ws"]) - 1)
        wobj.meteo.wws[:] = float(wx["ws"][h])
        wobj.meteo.wwd[:] = _m.radians((270.0 - float(wx["wd"][h]))
                                       % 360.0)
        wobj.meteo.temp[:] = float(wx["t"][h])
        wobj.meteo.rh[:] = float(wx["rh"][h])
        if st.session_state.get("emc_on", True):
            from disaster_phyengine.fuel_moisture import (
                update_dead_fuel_moisture)
            update_dead_fuel_moisture(wobj)
    elif mode == "Diurnal cycle":
        start = float(st.session_state.get("dr_start", 12.0))
        h = (start + t_min / 60.0) % 24.0
        phase = _m.cos((h - 15.0) / 24.0 * 2.0 * _m.pi)  # peak 15:00
        td = float(st.session_state.get("dr_tday", 34.0))
        tn = float(st.session_state.get("dr_tnight", 22.0))
        rd = float(st.session_state.get("dr_rhday", 20.0))
        rn = float(st.session_state.get("dr_rhnight", 70.0))
        wobj.meteo.temp[:] = (td + tn) / 2.0 + (td - tn) / 2.0 * phase
        wobj.meteo.rh[:] = (rd + rn) / 2.0 - (rn - rd) / 2.0 * phase
        if st.session_state.get("emc_on", True):
            from disaster_phyengine.fuel_moisture import (
                update_dead_fuel_moisture)
            update_dead_fuel_moisture(wobj)
        # diurnal wind: calmer at night, stronger mid-afternoon,
        # direction veers; spatial pattern preserved by scaling a
        # captured base field rather than overwriting it
        _mv = st.session_state.get("map_version", 0)
        _bkey = f"_wind_base_{_mv}"
        if _bkey not in st.session_state:
            st.session_state[_bkey] = (wobj.meteo.wws.copy(),
                                       wobj.meteo.wwd.copy())
        _bws, _bwd = st.session_state[_bkey]
        _dayf = 0.55 + 0.55 * max(0.0, phase)
        _veer = 0.6 * _m.sin((h - 9.0) / 24.0 * 2.0 * _m.pi)
        wobj.meteo.wws[:] = np.clip(_bws * _dayf, 0.0, None)
        wobj.meteo.wwd[:] = _bwd + _veer
        wobj.meteo.gust[:] = wobj.meteo.wws * 1.4
        # optional shower window: only sets W_prec; the ENGINE wets
        # the fuel from precipitation every step and stops ember
        # spotting above 1 mm/h
        _mm = float(st.session_state.get("dr_rain_mm", 0.0))
        _rs = float(st.session_state.get("dr_rain_start", 18.0))
        _rd_ = float(st.session_state.get("dr_rain_dur", 3.0))
        _raining = _mm > 0 and ((h - _rs) % 24.0) < _rd_
        wobj.meteo.prec[:] = _mm if _raining else 0.0


def _step_sim(n: int = 1):
    """Advance the simulation, driving the exogenous weather first.

    With the diurnal cycle on, air temperature and humidity follow a daily
    wave (peak mid-afternoon) and dead fuel moisture tracks them through the
    EMC model. Nights are cool and humid, so the moisture can exceed the
    extinction threshold m_ext and the fire stalls or dies out on its own -
    the same mechanism that stops real fires overnight."""
    diag = None
    # INTEGRATION SUBSTEP CAP: a property of the SCENARIO, not of the button
    # that advances it. It used to be applied only while "Animate step by
    # step" was latched on, which made the same scenario come out differently
    # depending on how it was advanced. Applied uniformly here, so Step,
    # Step X, Run to end and Animate all integrate identically and a run
    # reproduces; set it in Parameters (0 = uncapped, full fidelity).
    try:
        _cap = int(st.session_state.get("sim_substep_cap", 8))
        sim._substep_cap = _cap if _cap > 0 else None
    except Exception:
        pass
    for _ in range(int(n)):
        t_min = sim.state.step * float(getattr(cfg, "step_minutes", 1.0))
        _drive_weather(world, t_min)
        _ov = None
        if (st.session_state.get("dss_apply")
                and st.session_state.get("dss_res_base") is not None
                and st.session_state.get("dss_res_base_v")
                == st.session_state.get("map_version")):
            import dss as _dss_step
            _net0 = st.session_state.get("dss_network")
            _obs0 = (_net0 if (st.session_state.get("dss_use_obs", True)
                               and _net0 is not None) else None)
            _sv0 = st.session_state.get
            _esig = (st.session_state.get("map_version"),
                     int(_sv0("dss_n", 1)),
                     float(_sv0("dss_cycle_min", 1.0)),
                     float(_sv0("dss_horizon_min", 30.0)),
                     float(_sv0("dss_jth", 0.35)),
                     float(_sv0("dss_eta", 0.60)),
                     # DERIVE the master adaptation flag HERE, at the point of
                     # use. It used to be written as a side effect of
                     # rendering the Layer 4 Decision panel, so after a wipe
                     # (which sets it False) it stayed False for as long as
                     # that panel was not open, and turning "evFIS active" on
                     # from anywhere else changed nothing: the run came out
                     # bit-for-bit identical to evFIS off.
                     bool(_sv0("dss_evfis_on", True)
                          or _sv0("dss_genai_on", True)),
                     bool(_sv0("dss_genai_on", True)),
                     bool(_sv0("dss_evfis_on", True)),
                     float(_sv0("dss_evfis_step", 0.05)),
                     float(_sv0("dss_ctrl_eps", 0.10)),
                     float(_sv0("dss_ctrl_lr", 0.05)),
                     float(_sv0("dss_attn_thr", 0.35)),
                     float(_sv0("dss_min_gain", 0.05)),
                     str(_sv0("dss_seed_profile", "minimal")),
                     bool(_sv0("dss_use_stage12", True)),
                     bool(_sv0("dss_use_stage3", True)),)
            _eng = st.session_state.get("dss_engine")
            if _eng is None or st.session_state.get(
                    "dss_engine_sig") != _esig:
                import os as _os_rl
                _lg = _dss_step.RunLogger(
                    _os_rl.path.join(_os_rl.path.dirname(
                        _os_rl.path.dirname(_os_rl.path.abspath(
                            __file__))), "logs"),
                    tag=f"m{st.session_state.map_version}")
                _eng = _dss_step.DecisionEngine(
                    _dss_step.partition_n(cfg.nx, cfg.ny, _esig[1]),
                    base_pool=st.session_state["dss_res_base"],
                    network=_obs0, j_threshold=_esig[4], eta=_esig[5],
                    cycle_min=_esig[2], horizon_min=_esig[3],
                    evfis_step=_esig[9], adapt_on=_esig[6],
                    genai_on=_esig[7], evfis_on=_esig[8],
                    ctrl_eps=_esig[10], ctrl_lr=_esig[11],
                    attention_thr=_esig[12], min_gain=_esig[13],
                    seed_profile=_esig[14],
                    use_evfis=_esig[15], use_genai=_esig[16],
                    # ONE source of truth. The legacy learned_rules.json is no
                    # longer passed: while both were live the engine reasoned
                    # from the generated-state store while the panels read the
                    # old file, so a view could report "nothing learned" with
                    # two dozen records sitting in the store.
                    state_path=(
                        st.session_state.get("dss_learned_override")
                        or _gstate_path()),
                    run_logger=_lg)
                # THE CONTROLLER REMEMBERS THE SCENE. Its value table is kept
                # in the store and restored whenever the engine is rebuilt on
                # the SAME map, so what it learns about which stage pays off
                # accumulates over a campaign of fires instead of restarting
                # from zero every run. A different map drops it.
                try:
                    _eng.bind_map(_map_key())
                except Exception:
                    pass
                try:
                    _lg.write_meta(dict(
                        map=dict(nx=cfg.nx, ny=cfg.ny,
                                 cell_m=cfg.cell_size_m,
                                 step_min=float(cfg.step_minutes)),
                        engine=dict(
                            regions=_esig[1], cycle_min=_esig[2],
                            horizon_min=_esig[3], j_th=_esig[4],
                            eta=_esig[5], adapt=_esig[6],
                            genai=_esig[7], evfis=_esig[8],
                            evfis_step=_esig[9],
                            ctrl_eps=_esig[10], ctrl_lr=_esig[11],
                            attn=_esig[12], min_gain=_esig[13],
                            seed_profile=_esig[14],
                            use_stage12=_esig[15],
                            use_stage3=_esig[16]),
                        weather=dict(
                            wx_mode=_sv0("wx_mode", "Manual"),
                            emc_on=bool(_sv0("emc_on", True)),
                            dr_start=float(_sv0("dr_start", 12.0)),
                            dr_tday=float(_sv0("dr_tday", 34.0)),
                            dr_tnight=float(_sv0("dr_tnight", 22.0)),
                            dr_rhday=float(_sv0("dr_rhday", 20.0)),
                            dr_rhnight=float(_sv0("dr_rhnight", 70.0)),
                            dr_rain_mm=float(_sv0("dr_rain_mm", 0.0)),
                            dr_rain_start=float(_sv0("dr_rain_start",
                                                     18.0)),
                            dr_rain_dur=float(_sv0("dr_rain_dur", 3.0))),
                        sensors=list(_sv0("dss_sensors", []) or []),
                        depots=list(_sv0("dss_res_items", []) or []),
                        models=_model_ids()))
                    # the snapshot must be the t=0 BASELINE even if
                    # the engine is (re)built mid-run: swap in the
                    # pristine fuel state around the dump
                    _fl_bk = world.fuel.fload
                    _fm_bk = world.fuel.fmoist
                    world.fuel.fload = world.fuel.fload0.copy()
                    world.fuel.fmoist = getattr(
                        sim, "_fmoist0", world.fuel.fmoist).copy()
                    try:
                        _lg.save_world(world)
                    finally:
                        world.fuel.fload = _fl_bk
                        world.fuel.fmoist = _fm_bk
                except Exception:
                    pass
                _eng_prev = st.session_state.get("dss_engine")
                if (_eng_prev is not None
                        and getattr(_eng_prev, "seed_profile", None)
                        == _eng.seed_profile
                        and st.session_state.get("dss_engine_map")
                        == st.session_state.get("map_version")):
                    # a SETTINGS change rebuilds the engine, but the
                    # learned knowledge survives: rules and the controller
                    # value table transplant (membership moves live in
                    # the global registry, which the new engine
                    # resets; the tuned consequents ride along in
                    # the rules)
                    _eng.rules = _eng_prev.rules
                    _eng.controller.q = _eng_prev.controller.q
                st.session_state["dss_engine"] = _eng
                st.session_state["dss_engine_sig"] = _esig
                st.session_state["dss_engine_map"] = \
                    st.session_state.get("map_version")
            _eng.network = _obs0
            _eng.base_pool = st.session_state["dss_res_base"]
            _dtm = float(getattr(cfg, "step_minutes", 1.0))
            _cycm = (_eng.cycle_min if _eng.cycle_min is not None
                     else _eng.cycle_steps * _dtm)
            _due = (_eng.last_override is None
                    or (sim.state.step - _eng.last_cycle_step)
                    >= max(1, int(round(_cycm / _dtm))))
            if _due:
                with st.spinner("DSS decision cycle: shadow "
                                "forecasts + adaptation..."):
                    _ov = _eng.maybe_decide(sim)
            else:
                _ov = _eng.maybe_decide(sim)
            if _due:
                # DECISION LOG: one concise line per cycle - which
                # interventions were ordered (S D C P E W with intensity),
                # whether the adaptation was accepted, and the forecast gain.
                try:
                    _CH = [("S", "suppression_effort"),
                           ("D", "resource_deployment"),
                           ("C", "containment_line"),
                           ("P", "asset_protection"),
                           ("E", "evacuation"), ("W", "public_warning")]
                    _acts = getattr(_eng, "last_actions", None)
                    _parts = []
                    if _acts and _acts.get("regions"):
                        for _ro in _acts["regions"]:
                            _u = _ro.get("u", {})
                            _iv = " ".join(
                                f"{_c}{_u.get(_k, 0):.1f}" for _c, _k in _CH
                                if _u.get(_k, 0) > 0.05)
                            if _iv:
                                _rn = _ro.get("name")
                                _bx = _ro.get("box")
                                _tag = (str(_rn) if _rn
                                        else (f"[{int(_bx[0])},{int(_bx[1])}]"
                                              if _bx else ""))
                                _parts.append(f"{_tag}: {_iv}" if _tag
                                              else _iv)
                    _tmin = sim.state.step * _dtm
                    _line = (f"t={_fmt_sim_time(_tmin)} — "
                             + (" | ".join(_parts) if _parts
                                else "no active orders"))
                    _cyc = (_eng.cycles[-1]
                            if getattr(_eng, "cycles", None) else None)
                    if _cyc:
                        _ad = _cyc.get("adaptation", {}) or {}
                        _acc = _ad.get("accepted")
                        _dj = _ad.get("dJ")
                        _line += "  · adapt " + ("✅" if _acc else "·")
                        if _dj is not None:
                            _line += f" ΔJ={float(_dj):+.2f}"
                        if getattr(_eng, "last_withheld", False):
                            _line += "  · NO-HARM withheld"
                    _dl = st.session_state.setdefault("dss_decision_log", [])
                    _dl.append(_line)
                    st.session_state["dss_decision_log"] = _dl[-80:]
                except Exception:
                    pass
        diag = sim.step(resource_override=_ov)
        _engL = st.session_state.get("dss_engine")
        if _engL is not None and _engL.run_logger is not None:
            try:
                _engL.run_logger.log_step(sim, compute_costs(sim),
                                          override=_ov)
            except Exception:
                pass
        _net = st.session_state.get("dss_network")
        if _net is not None and st.session_state.get("dss_use_obs", True):
            _net.update(sim, float(getattr(cfg, "step_minutes", 1.0)))
    return diag


# -------------------------------------------------------------------- sidebar
def _anim_toggle_cb():
    # ANY user interaction with the Animate toggle clears the deferred stop
    # flags: a pending stop from a previous finish / new map would otherwise
    # be popped a moment later and reset a FRESH ON back to OFF (the old
    # two-clicks-to-start bug). Callbacks run BEFORE the script body, so
    # clearing here makes one click start the animation.
    st.session_state.pop("anim_stop", None)
    if st.session_state.get("anim_on"):
        st.session_state.pop("runend_on", None)


def _runend_toggle_cb():
    st.session_state.pop("runend_stop", None)
    st.session_state.pop("anim_stop", None)


_PAGES = ["Simulation", "Step analysis", "Map editor",
          "Data layers", "Parameters",
 "GIS import", "Validation", "System Description"]
_PAGE_ICONS = {"Simulation": "\U0001F525", "Map editor": "✏️", "Data layers": "\U0001F5FA️",
               "Parameters": "⚙️", "Validation": "✅",
               "Step analysis": "\U0001F50E",
 "GIS import": "\U0001F30D", "System Description": "\U0001F4D8"}

def _rca_status_body():
    """Where the after-action review stands. Safe to call from anywhere.

    Reads the background job WITHOUT waiting on it and lifts a finished
    report into the session, so the result is picked up no matter which
    page happened to be open when the model finished.
    """
    from dss import rca as _r
    _d = st.session_state.get("rca_dir")
    if not _d:
        return
    if st.session_state.get("rca_report"):
        return                      # already collected, the panel shows it
    _j = _r.poll(_d)
    _state = _j.get("state")
    if _state == "running":
        _el = _r.elapsed_s(_d)
        st.info(f"Root cause analysis running on "
                f"**{_j.get('model') or 'opus'}** — {_el:.0f} s so far. "
                "Keep working; this refreshes itself and the report lands "
                "under Layer 4 · Analysis when it is ready.")
    elif _state == "error":
        st.error(f"Root cause analysis failed: {_j.get('error')}")
        st.session_state.pop("rca_dir", None)
    elif _state == "done":
        st.session_state["rca_report"] = _j.get("report")
        st.session_state["rca_recs"] = _j.get("recs") or {}
        if hasattr(st, "toast"):
            st.toast("Root cause analysis is ready", icon="✅")
        st.rerun()


# A FRAGMENT RERUNS ITSELF, NOT THE PAGE. That is what lets the review be
# waited on without freezing the simulation: every few seconds Streamlit
# re-executes just this block, sees whether the thread has finished, and
# leaves everything else alone. On builds without fragments the status is
# still correct, it simply updates on the next interaction.
_rca_live = (st.fragment(run_every=3.0)(_rca_status_body)
             if hasattr(st, "fragment") else _rca_status_body)


def _rca_sidebar_badge():
    """The same status, on EVERY page, so the review can be started and
    then forgotten about until it is done."""
    from dss import rca as _r
    _d = st.session_state.get("rca_dir")
    if not _d:
        return
    if st.session_state.get("rca_report"):
        st.sidebar.success("Root cause analysis ready — open "
                           "Layer 4 · Analysis")
        return
    if _r.poll(_d).get("state") == "running":
        st.sidebar.info(f"Root cause analysis running "
                        f"({_r.elapsed_s(_d):.0f} s)")


with st.sidebar:
    st.title("DisasterAware")
    st.caption("Enhanced Decision Support System for Wildfire "
               "Disaster Response and Management")
    import disaster_phyengine as _da
    _engine_ok = hasattr(Simulator, "rewind") and hasattr(sim, "rewindable_steps")
    if not _engine_ok:
        st.caption("\u26a0 engine outdated in memory")
    if not _engine_ok:
        st.error("The simulation engine changed on disk but the running "
                 "process still uses the old version (Python caches "
                 "imported modules). Close the app completely and start "
                 "it again (run_dashboard.bat). Until then rewind and "
                 "time-step scaling stay inactive.")

    # current page is read early because the run-to-end gate in the Simulation
    # card (rendered first, at the top of the panel) needs it
    if st.session_state.get("nav_page") not in _PAGES:
        st.session_state.pop("nav_page", None)
    page = st.session_state.get("nav_page", _PAGES[0])

    # --- simulation control panel (top of the left panel) ---
    if st.session_state.pop("anim_stop", False):
        st.session_state.anim_on = False
        st.session_state.runend_on = False
    if st.session_state.pop("runend_stop", False):
        st.session_state.runend_on = False
    with st.container(border=True):
        st.markdown("**Simulation**")
        # controls stacked one under another so they read top to bottom
        if st.button("Step", use_container_width=True,
                     help="Advance the fire by one time step."):
            _step_sim(); _record_costs(); st.rerun()
        _sc1, _sc0 = st.columns([0.6, 0.4])
        xsteps = int(_sc0.number_input(
            "X", 1, 1000, 10, key="step_x",
            label_visibility="collapsed",
            help="How many steps the 'Step X' button advances at "
                 "once."))
        if _sc1.button(f"Step {xsteps}", use_container_width=True,
                       help="Advance X steps at once (set X in the box "
                            "beside)."):
            _step_sim(xsteps); _record_costs(); st.rerun()
        st.toggle("Animate step by step", key="anim_on",
                  on_change=_anim_toggle_cb,
                  help="Advance automatically, one step per refresh, until "
                       "the fire is over.")
        st.toggle("Run to end", key="runend_on",
                  on_change=_runend_toggle_cb,
                  help="Latches ON and keeps running until the fire "
                       "is out or the step cap (max_steps) is "
                       "reached; press again to stop. The map "
                       "refreshes after every chunk.")
        if st.button("Reset fire", use_container_width=True,
                     help="Clear the fire and the cost series; the map and "
                          "all edits stay."):
            sim.reset(); st.session_state.cost_series = []
            st.session_state.pop("anim_quiet", None)
            _reset_dss_state()
            st.session_state.pop("dss_net_sig", None)
            st.rerun()
        _telap = float(getattr(
            sim, "t_elapsed_min",
            sim.state.step * float(getattr(cfg, "step_minutes", 30.0))))
        st.caption(f"Step {sim.state.step} \u00b7 "
                   f"t = {_fmt_sim_time(_telap)}")
        st.caption("active fire "
                   f"{int((sim.state.burning > 0.5).sum())} cells")
        cfg.max_steps = int(st.number_input(
            "Run limit — max_steps", 20, 1_000_000,
            int(cfg.max_steps), 10,
            help="Safety stop for 'Run to end' and the animation. The fire "
                 "normally stops on its own when it burns out."))
        avail = (sim.rewindable_steps
                 if hasattr(sim, "rewindable_steps") else [])
        lo = int(avail[0]) if avail else 0
        hi = int(sim.state.step)
        if avail and hi > lo:
            st.markdown("**Rewind** \u2014 go back, change conditions, replay")
            def _do_rewind(_k):
                if sim.rewind(int(_k)):
                    st.session_state.cost_series = [
                        r for r in st.session_state.cost_series
                        if r.get("step", 0) <= int(_k)]
                    # the DSS rolls back WITH the physics: standing
                    # orders, overlays, logs and gating priors after
                    # the rewind point disappear (learned rules stay,
                    # they are knowledge, not decisions)
                    _eng_rw = st.session_state.get("dss_engine")
                    if _eng_rw is not None:
                        try:
                            _eng_rw.rewind_to(int(_k))
                        except Exception:
                            pass
                    for _kk in list(st.session_state.keys()):
                        if _kk.startswith(("l3_gate_",
                                           "dss_featprev_")):
                            del st.session_state[_kk]
                    st.rerun()
                else:
                    st.warning("Snapshot for that step is no longer stored.")

            rsl = st.slider("Rewind to step", lo, hi, hi, 1, key="rw_sl",
                            label_visibility="collapsed",
                            help="Dragging the slider rewinds immediately.")
            if int(rsl) < hi:
                _do_rewind(int(rsl))   # slider rewinds instantly
            rc1, rc2 = st.columns([1.4, 1])
            rnum = rc1.number_input("k", lo, hi, hi, 1, key="rw_num",
                                    label_visibility="collapsed",
                                    help="Type the exact step, then press Go.")
            if rc2.button("Go", use_container_width=True,
                          disabled=int(rnum) >= hi,
                          help="Restore the full state at step k. The "
                               "history after it is discarded; replay with "
                               "new wind, moisture or resources."):
                _do_rewind(int(rnum))

    # --- DSS SETUP card: the three things that must exist before the DSS can
    # decide anything, each with a light that says whether it is done. The
    # same actions live deep inside Layer 1 and Layer 2; having them here
    # means the readiness of a run can be seen and fixed without hunting
    # through panels. ---
    with st.container(border=True):
        st.markdown("**DSS setup**")

        def _lamp(ok: bool) -> str:
            return "🟢" if ok else "🔴"
        _mv = st.session_state.get("map_version")
        _sens = list(st.session_state.get("dss_sensors", []) or [])
        # same ordering trap as the pool below: dss_network is built later in
        # page_simulation, so the lamp judges the PLACED SENSORS, which is
        # what the button actually produces and what the user can see
        _net_ok = bool(_sens)
        # judge the STAGED ITEMS, not the rasterized pool. The pool is built
        # further down in page_simulation, which runs AFTER this sidebar, so
        # right after the button the lamp read a pool that did not exist yet
        # and went red until some other click forced one more rerun.
        _res_ok = (bool(st.session_state.get("dss_res_items"))
                   and st.session_state.get("dss_res_base_v") == _mv)
        _dss_ok = bool(st.session_state.get("dss_apply")) and _res_ok

        _c1, _c2 = st.columns([1, 4])
        _c1.markdown(_lamp(_net_ok))
        if _c2.button("Suggest sensors", use_container_width=True,
                      key="side_sugg_net",
                      help="Layer 1: greedy maximum weighted coverage, the "
                           "same action as 'Suggest network' in the Layer 1 "
                           "panel. Red means the agents are blind."):
            try:
                # exactly what the Layer 1 button does, same argument and
                # same session keys, so the two stay in step
                _covS = int(st.session_state.get("dss_cov_target", 60))
                _pl, _why = _dss_pkg.suggest_network(
                    world, coverage_target=float(_covS))
                st.session_state["dss_sensors"] = _pl
                st.session_state["dss_suggest_why"] = _why
                # write the TARGET as well, not only the applied value: the
                # panel re-suggests whenever the two disagree, and leaving
                # the target unset let that happen the moment Layer 1 opened
                st.session_state["dss_cov_target"] = _covS
                st.session_state["dss_cov_applied"] = _covS
                st.session_state.pop("dss_net_sig", None)
                st.rerun()
            except Exception as _e_sn:
                st.error(f"Sensor suggestion failed: {_e_sn}")

        _c3, _c4 = st.columns([1, 4])
        _c3.markdown(_lamp(_res_ok))
        if _c4.button("Suggest resources", use_container_width=True,
                      key="side_sugg_res",
                      help="Layer 1: build the baseline suppression pool "
                           "from the map's fire stations, roads and "
                           "helibases. The same action as 'Suggest "
                           "resources' in the Layer 1 panel."):
            try:
                # SAME DEFAULT AS THE LAYER 1 SLIDER (50, not 70). Reading a
                # different default here staged a pool for a much higher
                # efficiency target, which is why the map filled with
                # overlapping helibases; opening Layer 1 then set the slider
                # to 50, that differed from the applied value, and the panel
                # silently re-suggested and "fixed" the map. The target is
                # WRITTEN too, so the slider and the applied value agree and
                # no re-suggestion is triggered behind the user's back.
                _effS = int(st.session_state.get("dss_eff_target", 50))
                _denS = float(st.session_state.get("dss_res_density", 1.0))
                _its, _rwhy = _dss_pkg.suggest_resource_items(
                    world, efficiency_target=_effS / 100.0,
                    density=_denS,
                    coverage=float(st.session_state.get("dss_res_cov",
                                                        100.0)) / 100.0)
                st.session_state["dss_res_items"] = _its
                st.session_state["dss_res_why"] = _rwhy
                st.session_state["dss_res_base_v"] = _mv
                st.session_state["dss_eff_target"] = _effS
                st.session_state["dss_res_density"] = _denS
                st.session_state["dss_eff_applied"] = _effS
                st.session_state["dss_dens_applied"] = _denS
                st.rerun()
            except Exception as _e_sr:
                st.error(f"Resource suggestion failed: {_e_sr}")

        # THE PLANNERS EXPLAIN THEMSELVES HERE TOO. Both suggesters
        # return their optimization trace; showing it only in the
        # Layer 1 panel made these two buttons look like magic.
        if st.session_state.get("dss_suggest_why"):
            with st.expander("Why these sensor positions"):
                for _ln_w in st.session_state["dss_suggest_why"]:
                    st.caption(_ln_w)
        if st.session_state.get("dss_res_why"):
            with st.expander("Why this resource pool"):
                for _ln_w in st.session_state["dss_res_why"]:
                    st.caption(_ln_w)

        _c5, _c6 = st.columns([1, 4])
        _c5.markdown(_lamp(_dss_ok))
        _lbl = ("Local DSS: ON" if st.session_state.get("dss_apply")
                else "Local DSS: OFF")
        if _c6.button(_lbl, use_container_width=True, key="side_dss_apply",
                      help="Layer 2 and up: let the local agents observe and "
                           "issue orders. Needs a resource pool, so the lamp "
                           "stays red until 'Suggest resources' has run."):
            st.session_state["dss_apply"] = not bool(
                st.session_state.get("dss_apply"))
            st.rerun()
        # how many local agents the map is split into. The same setting lives
        # in the Layer 1 panel; it is here too because it decides how many
        # DSS instances the run has, which is a setup question, not a detail
        _c7, _c8 = st.columns([1, 4])
        _c7.markdown("👥")
        _mirror("dss_n", "side_dss_n", 1)
        _c8.number_input(
            "Local DSS agents", 1, 12, step=1, key="side_dss_n",
            on_change=_adopt, args=("dss_n", "side_dss_n"),
            help="The map is split into exactly this many regions, one "
                 "local DSS each, covering every cell. The same setting "
                 "lives in Layer 1; the two always agree.")

        # Layer 4 settings at a glance: which model stage 3 would call and
        # whether the generative stage is switched on at all
        # SAME DEFAULTS AS THE PANEL. The Layer 4 toggles read
        # _sv("dss_evfis_on", True); reading them here without that default
        # returned None until that panel had been opened once, so a module
        # that was plainly switched on showed a red lamp.
        _gon = bool(st.session_state.get("dss_genai_on", True))
        _evon = bool(st.session_state.get("dss_evfis_on", True))
        _dact = bool(st.session_state.get("dss_apply"))
        try:
            _greach = _dss_pkg.genai_config().get("mode") == "cli"
        except Exception:
            _greach = False
        # a stage only RUNS when the DSS itself is running, so the master
        # switch is part of the lamp rather than a separate thing to notice
        _ev_run = _dact and _evon
        _ga_run = _dact and _gon and _greach

        _c11, _c12 = st.columns([1, 4])
        _c11.markdown(_lamp(_ev_run))
        _c12.markdown("**evFIS ①②**")
        # the lamp already says it is running; only a red lamp needs a reason
        if not _ev_run:
            _c12.caption("DSS active is off" if not _dact
                         else "evFIS active is off")
        _c9, _c10 = st.columns([1, 4])
        _c9.markdown(_lamp(_ga_run))
        _c10.markdown("**GenAI ③**")
        # the model is CHOSEN here, not only reported: it is the one stage ③
        # setting that changes what a run costs in wall time, and hunting for
        # it inside Layer 4 Settings every time was the wrong place for it
        _MODELS_SB = ["(plan default)", "opus", "sonnet", "haiku"]
        _cur_sb = str(st.session_state.get("genai_model_ui",
                                           "(plan default)"))
        if _cur_sb not in _MODELS_SB:
            _cur_sb = "(plan default)"
        _mirror("genai_model_ui", "side_genai_model", "(plan default)")
        _pick_sb = _c10.selectbox(
            "GenAI ③ model", _MODELS_SB, key="side_genai_model",
            on_change=_adopt, args=("genai_model_ui", "side_genai_model"),
            help="Which model stage ③ calls. Same setting as the one in "
                 "Layer 4 Settings; changing it here changes it there. "
                 "Measured on this transport the wait is roughly 3.4 s of "
                 "fixed startup plus 7 ms per output token, so the answer "
                 "LENGTH matters more than the family: opus and sonnet "
                 "honour the short-JSON instruction, haiku does not and "
                 "ends up slowest.")
        if _pick_sb != _cur_sb:
            if _pick_sb != "(plan default)":
                os.environ["DSS_GENAI_MODEL"] = _pick_sb
            else:
                os.environ.pop("DSS_GENAI_MODEL", None)
            # PROBE ON CHANGE. Picking a model without checking it leaves the
            # question the label is meant to answer open, and with "(plan
            # default)" the served model cannot be known any other way. One
            # call, a few seconds, and the lamp and the label are truthful.
            st.session_state["genai_probe_pending"] = True
            st.rerun()
        if st.session_state.pop("genai_probe_pending", False):
            with st.spinner(f"Testing GenAI on `{_pick_sb}`…"):
                try:
                    _prS = _dss_pkg.genai_probe()
                except Exception as _e_pb:
                    _prS = dict(ok=False, error=str(_e_pb))
            st.session_state["genai_last_probe"] = _prS
            if _prS.get("ok"):
                st.session_state["genai_served_model"] = str(
                    _prS.get("reported_model") or "")
                _mmS = dict(st.session_state.get("genai_model_ms", {}))
                _mmS[_pick_sb] = int(_prS.get("latency_ms") or 0)
                st.session_state["genai_model_ms"] = _mmS
        _prL = st.session_state.get("genai_last_probe") or {}
        # green lamp = running; say the resolved model, and only explain
        # when the lamp is red
        _c10.caption(
            _genai_model_label() if _ga_run else
            ("DSS active is off" if not _dact else
             ("GenAI active is off" if not _gon else
              "`claude` not reachable")))
        _t1, _t2 = st.columns([1, 4])
        _t1.markdown("🧪")
        if _t2.button("Test GenAI", use_container_width=True,
                      key="side_genai_test",
                      help="One real call over the same path stage ③ uses, "
                           "so the latency and the served model are the ones "
                           "the simulation will see."):
            st.session_state["genai_probe_pending"] = True
            st.rerun()
        if _prL:
            if _prL.get("ok"):
                _uL = _prL.get("usage") or {}
                st.caption(
                    f"✅ {_prL.get('finished_at', '')} · served "
                    f"`{_prL.get('reported_model', '?')}` · "
                    f"{_prL.get('latency_ms', 0)} ms · "
                    f"{_uL.get('output_tokens', '?')} output tokens")
            else:
                st.caption(f"❌ {str(_prL.get('error', ''))[:120]}")
    # --- page navigation (grouped card, below the Simulation card) ---
    _NAV_GROUPS = [["Simulation", "Step analysis"],
                   ["Map editor", "GIS import", "Data layers"],
                   ["Parameters", "Validation"],
                   ["System Description"]]
    # A radio, not a column of buttons. Every page was a full-width button,
    # which read as eight competing actions instead of one setting with eight
    # values, and it cost eight rows of height in a panel that is short.
    _PAGES_FLAT = [p for g in _NAV_GROUPS for p in g]
    with st.container(border=True):
        # the radio keeps its own state, so a button elsewhere that set
        # nav_page was overwritten by the radio's stale value on the very
        # next line and the page never changed
        _mirror("nav_page", "nav_radio", _PAGES_FLAT[0])
        _sel_pg = st.radio(
            "Page", _PAGES_FLAT,
            format_func=lambda p: f"{_PAGE_ICONS[p]}  {p}",
            key="nav_radio", on_change=_adopt,
            args=("nav_page", "nav_radio"), label_visibility="collapsed")
        if _sel_pg != page:
            st.session_state["nav_page"] = _sel_pg
            st.rerun()

    # THE REVIEW FOLLOWS YOU. It is started from one panel but takes
    # minutes, so the status belongs where every page can see it.
    try:
        _rca_sidebar_badge()
    except Exception:
        pass

    st.divider()


    if not HAS_CANVAS:
        st.warning("Install streamlit-drawable-canvas for mouse editing.")

    st.divider()
    import datetime as _dt_about
    _ver = getattr(_da, "__version__", "0.2.1")
    _year = _dt_about.date.today().year
    _today = _dt_about.date.today().strftime("%d %b %Y")
    with st.expander("About", expanded=False):
        st.markdown(
            f"**DisasterAware**  \n"
            f"Enhanced Decision Support System for Wildfire Disaster "
            f"Response and Management  \n\n"
            f"Version {_ver}  \n"
            f"Date: {_today}  \n"
            f"Author: Çağlar Akman  \n\n"
            f"© {_year} Çağlar Akman. All rights reserved.  \n"
            f"License: for academic and research use only.")

# the DisasterAware title / caption already live in the sidebar header and
# the About box, so the main content area starts straight at the content


# ============================================================== SIMULATION ===
def page_simulation():
    view_col, side_col = st.columns([2.9, 1.45], gap="medium")

    def _sv(key, default):
        return st.session_state.get(key, default)

    with side_col:
        # ---- SITUATIONAL AWARENESS BOARD (cockpit) ----
        # The live picture a jury reads at a glance: the fire, every
        # cost term, the capacity flow, each agent's current decision
        # and the coordinator's word. Fed by the SAME cycle record the
        # chronicle logs (engine.cycles[-1]), so the board can never
        # disagree with the log.
        _engSA = st.session_state.get("dss_engine")
        _cylSA = (getattr(_engSA, "cycles", None) or []) \
            if _engSA is not None else []
        _cySA = _cylSA[-1] if _cylSA else None
        with st.expander("DSS Dashboard — situational awareness",
                         expanded=True):
            if st.button("Cost settings & charts",
                         key="sa_cost_dlg", use_container_width=True,
                         help="The J_k equation, the operational "
                              "priority and protection weights, the "
                              "advanced thresholds and the per-term "
                              "cost charts, in a popup."):
                _cost_dialog()
            # THE SITUATION IS NOT THE DSS's TO REPORT. What is burning,
            # what it has cost, the wind and the fuel moisture are
            # properties of the SIMULATION; only the agent rows and the
            # coordinator's word need a decision cycle. The board used to
            # refuse to show any of it without the DSS, so a free-running
            # fire had no situational awareness at all.
            # ALWAYS THE LIVE SIMULATION, never the last decision cycle.
            # The board used to read engine.cycles[-1], which is a SNAPSHOT
            # taken when that cycle ran. Switch the DSS off, or let it fall
            # silent, and the chronicle stops while the fire goes on: the
            # board then froze at the moment of the last cycle and reported
            # assets 0.00 and population 0.00 for a town that had since
            # burned to the ground, while the Cost panel, which reads the
            # simulator, showed the losses correctly. Only the agent rows
            # and the coordinator's word belong to a cycle.
            _simSA = st.session_state.get("sim")
            _smS = _coS = _poS = {}
            if _simSA is not None:
                try:
                    _wSA = _simSA.world
                    _smS = dict(
                        burning=int((_simSA.state.burning > 0.5).sum()),
                        burned=int(_simSA.ever_burned.sum()),
                        wws_mean=float(_wSA.meteo.wws.mean()),
                        fmoist_mean=float(_wSA.fuel.fmoist.mean()))
                    _rSA = compute_costs(_simSA)
                    _coS = dict(j_total=_rSA.j_total,
                                j_physical=_rSA.j_physical,
                                j_burn=_rSA.j_burn,
                                j_asset=_rSA.j_asset,
                                j_pop=_rSA.j_pop,
                                j_resp=_rSA.j_resp,
                                j_delay=_rSA.j_delay)
                except Exception:
                    pass
            _bp = st.session_state.get("dss_res_base")
            if _bp is not None:
                try:
                    _poS = dict(rcap_total=float(np.asarray(_bp.rcap).sum()))
                except Exception:
                    pass
            if not _poS:
                _poS = (_cySA.get("pool") or {}) if _cySA else {}
            if True:
                # A GAUGE THAT EMPTIES, AND NEVER PASSES FULL. The meter
                # used to divide fielded capacity by staged capacity with
                # neither side weighted by availability, a ratio the engine
                # never forms, so it read 138% and told the reader nothing.
                #
                # Capacity here is a FLOW, how much force can act per
                # minute, not a stock that drains: the same pool is there
                # next step. Scarcity therefore shows up as DEMAND, what
                # the orders asked for, running past the BUDGET the
                # allocator has to spend. So the gauge is what is left of
                # the budget this step: it falls to zero exactly when the
                # response saturates, and when the orders want more than
                # exists the shortfall is named instead of being hidden
                # behind a number above 100%.
                _actS = getattr(_engSA, "last_actions", None) or {}
                _demS = _actS.get("demand")
                _budS = _actS.get("budget")
                _stagS = _rawS = 0.0
                try:
                    _bpS = st.session_state.get("dss_res_base")
                    if _bpS is not None:
                        _stagS = float(np.asarray(_bpS.rcap).sum())
                    _resS = getattr(st.session_state.get("sim"),
                                    "last_applied_resource", None)
                    if _resS is not None:
                        _rawS = float(np.asarray(_resS.rcap).sum())
                except Exception:
                    pass
                if not _stagS:
                    _stagS = float(_poS.get("rcap_total") or 0.0)
                _freeS = None
                _shortS = 0.0
                if _demS is not None and _budS:
                    _useS = float(_demS) / float(_budS)
                    _freeS = max(0.0, min(1.0, 1.0 - _useS))
                    _shortS = max(0.0, _useS - 1.0)
                mS1, mS2, mS3 = st.columns(3)
                mS1.metric("burning", int(_smS.get("burning", 0)))
                mS2.metric("burned", int(_smS.get("burned", 0)))
                mS3.metric(
                    "capacity free",
                    ("—" if _freeS is None else f"{100.0 * _freeS:.0f}%"),
                    delta=(None if _shortS <= 0.0
                           else f"short {100.0 * _shortS:.0f}%"),
                    delta_color="inverse",
                    help="How much of this step's allocation budget the "
                         "orders did NOT need. It falls toward zero as the "
                         "response commits, and it cannot pass 100%.\n\n"
                         "Capacity is a RATE, not a stock: the pool does "
                         "not empty, the same force is available next "
                         "minute. What runs short is the amount that can "
                         "act at once. At 0% every cell the orders asked "
                         "for cannot be funded, so the allocator pays for "
                         "the highest-value ones and the rest get nothing "
                         "— which is how an under-resourced fire escapes "
                         "in this model.\n\n"
                         "The red figure is the shortfall: how much more "
                         "capacity the orders wanted than exists. While it "
                         "shows, the fire is being fought with less than "
                         "it asked for.")
                if _shortS > 0.0:
                    st.markdown(
                        "<div style='font-size:0.84em;color:#c0392b'>"
                        "&#9888; response saturated — the orders want "
                        f"{100.0 * (1.0 + _shortS):.0f}% of the budget; "
                        "the lowest-value cells are going unfunded"
                        "</div>", unsafe_allow_html=True)
                # THE WHOLE COST, NOT PART OF IT. The delay term was never
                # shown and the physical outcome had nowhere to sit beside
                # the decision cost, so a reader could not check the total
                # against its parts nor tell a good FIRE from a cheap one.
                _barsS = [
                    ("J total (decision)", _coS.get("j_total"), "#8c8c8c"),
                    ("J phys (outcome)", _coS.get("j_physical"), "#5c5c5c"),
                    ("burned area", _coS.get("j_burn"), "#ff5a1e"),
                    ("assets", _coS.get("j_asset"), "#b42828"),
                    ("population", _coS.get("j_pop"), "#2878ff"),
                    ("response", _coS.get("j_resp"), "#a08cff"),
                    ("delay", _coS.get("j_delay"), "#c8a0ff")]
                st.markdown(
                    "".join(_iv_bar(_lb, float(_vb or 0.0), _cb)
                            for _lb, _vb, _cb in _barsS),
                    unsafe_allow_html=True)
                st.caption(
                    f"wind {float(_smS.get('wws_mean') or 0):.1f} m/s"
                    f" · fuel moisture "
                    f"{float(_smS.get('fmoist_mean') or 0):.2f}"
                    f" · staged pool {_stagS:.0f}"
                    f" · fielded {_rawS:.0f}"
                    + (f" · asked {float(_demS):.0f} of {float(_budS):.0f}"
                       if (_demS is not None and _budS) else ""))
                # the agents and the coordinator ARE the DSS, so these
                # stay empty until a decision cycle has run
                if not _cySA:
                    st.caption("No decision cycle yet — the agent rows and "
                               "the coordinator's ranking appear once the "
                               "DSS runs. The situation above is the "
                               "simulation's own.")
                _glS = (_cySA.get("global_dss") or {}) if _cySA else {}
                _hotS = _glS.get("hotspot")
                for _nS, _rS in ((_cySA.get("regions") or {})
                                 if _cySA else {}).items():
                    _uS = _rS.get("orders_final") or {}
                    _topS = sorted(
                        ((k, float(v)) for k, v in _uS.items()
                         if k != "_share" and float(v) > 0.05),
                        key=lambda t: -t[1])[:3]
                    _roleS = ("FOCUS" if _nS == _hotS
                              else ("attended" if _rS.get("attended")
                                    else "monitor"))
                    st.markdown(
                        "<div style='font-size:0.84em;margin:2px 0'>"
                        f"<b>{_nS}</b> · {_roleS} · share "
                        f"{float(_rS.get('coord_share') or 0):.2f} · "
                        f"Q {float(_rS.get('quality') or 0):.2f}"
                        + (" · <b style='color:#c00'>FAIL-SAFE</b>"
                           if _rS.get("failsafe") else "")
                        + "<br>"
                        + (", ".join(f"{k.split('_')[0]} {v:.2f}"
                                     for k, v in _topS)
                           or "no offensive orders")
                        + "</div>", unsafe_allow_html=True)
                if _glS.get("statement"):
                    st.caption(str(_glS.get("statement")))
                for _dS in (_glS.get("directives") or []):
                    st.markdown(
                        "<div style='font-size:0.82em;color:#c40'>"
                        f"\u25c6 {_dS}</div>",
                        unsafe_allow_html=True)
        # one panel at a time: no scrolling, DSS first
        # panels laid out in fixed rows (a button grid, so the grouping is
        # stable): row 1 environment/time/ignition, rows 2-3 the DSS layers,
        # then the display settings.
        _rows = [["Environment", "Time", "Ignition"],
                 ["Layer 1 \u00b7 Input", "Layer 2 \u00b7 Perception",
                  "Layer 3 \u00b7 Concepts"],
                 ["Layer 4 \u00b7 Decision", "Layer 4 Settings"],
                 ["Layer 4 Rules", "Layer 4 \u00b7 Analysis", "Layer 4 \u00b7 Logs"]]
        _panels = [p for _r in _rows for p in _r]
        # shadow state keeps the selection across a Step rerun.
        _pcur = st.session_state.get("sim_panel_v", _panels[0])
        if _pcur not in _panels:
            _pcur = _panels[0]
        for _r in _rows:
            _cols = st.columns(len(_r))
            for _ci, _pn in enumerate(_r):
                # show the FULL name incl. the descriptor, e.g.
                # "Layer 1 \u00b7 Input", not just "Layer 1"
                _lbl = _pn
                if _cols[_ci].button(
                        _lbl, key=f"panelbtn_{_pn}",
                        use_container_width=True,
                        type=("primary" if _pcur == _pn else "secondary")):
                    st.session_state["sim_panel_v"] = _pn
                    st.rerun()
        panel = _pcur
        import dss as _dss
        # the sensor network lives OUTSIDE the panels: the map overlay
        # and Layer 2 need it no matter which panel is open
        _slist = list(_sv("dss_sensors", []))
        use_obs = bool(_sv("dss_use_obs", True))
        # the resource pool is rasterized from its editable item rows
        # whenever the rows change (signature check), panel-independent
        _rit = st.session_state.get("dss_res_items")
        if _rit and (st.session_state.get("dss_res_base_v")
                     != st.session_state.get("map_version")):
            # SELF-HEAL: rows exist but the bookkeeping says another
            # map. If every row still fits this map, adopt the rows
            # (the user DID stage a pool; a stale version stamp must
            # not un-stage it); only rows that fall off the map void
            # the pool for real.
            if all((0 <= int(it.get("x", 0)) < cfg.nx
                    and 0 <= int(it.get("y", 0)) < cfg.ny)
                   for it in _rit if "x" in it):
                st.session_state["dss_res_base_v"] = \
                    st.session_state.get("map_version")
            else:
                st.session_state["dss_res_items"] = None
                st.session_state["dss_res_base"] = None
                _rit = None
        if _rit and (st.session_state.get("dss_res_base_v")
                     == st.session_state.get("map_version")):
            _rsig = tuple(sorted((it["kind"], it.get("x", -1),
                                  it.get("y", -1), it.get("cap", 0),
                                  it.get("radius", 0),
                                  it.get("avail", 1.0),
                                  it.get("t_disp", 10.0))
                                 for it in _rit))
            if st.session_state.get("dss_res_sig") != _rsig:
                _rl_new = _dss.build_resource_layer(world, _rit)
                st.session_state["dss_res_base"] = _rl_new
                st.session_state["dss_res_sig"] = _rsig
                # the response-cost term is charged as the FRACTION
                # of the staged pool committed (with the 1.2 surge),
                # so J_resp reads "how much of what exists is in the
                # field", not an absolute vs an arbitrary constant
                cfg.cost.capacity_reference = max(
                    100.0, 1.2 * float((_rl_new.rcap
                                        * _rl_new.ravail).sum()))
        elif st.session_state.get("dss_res_base") is not None \
                and not _rit:
            st.session_state["dss_res_base"] = None
        # draw BOTH ground depots and helibases as map markers (road
        # corridors have no point, so they are excluded). The label carries
        # the kind's short name so the renderer picks the helibase glyph
        # (it keys off "heli" in the label).
        st.session_state["dss_depots_draw"] = ([
            (int(it["x"]), int(it["y"]), int(it.get("radius", 4)),
             float(it.get("cap", 0.8)),
             f"D{_k + 1} "
             + _dss.RESOURCE_KINDS.get(it.get("kind"), {}).get(
                 "short", str(it.get("kind", ""))))
            for _k, it in enumerate(_rit or [])
            if it.get("kind") in ("depot", "helibase")] or None) \
            if _rit else None
        # (re)build the network when the map or the fleet changes
        _sig = (st.session_state.get("map_version"),
                tuple((d["kind"], d["x"], d["y"],
                       d.get("radius_m"), d.get("latency_min"))
                      for d in _slist))
        if st.session_state.get("dss_net_sig") != _sig:
            _net = _dss.SensorNetwork(
                [_dss.Sensor(d["kind"], d["x"], d["y"],
                             radius_m=d.get("radius_m"),
                             latency_min=d.get("latency_min"))
                 for d in _slist], cfg.ny, cfg.nx, cfg.cell_size_m)
            _net.update(sim, 0.0)
            st.session_state["dss_network"] = _net
            st.session_state["dss_net_sig"] = _sig
        _net = st.session_state.get("dss_network")
        st.session_state["dss_sensors_draw"] = [
            (d["x"], d["y"],
             (None if _dss.SENSOR_CATALOG[d["kind"]]["radius_m"] is None
              else max(1, int(round(_dss.SENSOR_CATALOG[d["kind"]]
                                    ["radius_m"] / cfg.cell_size_m)))),
             d["kind"], f"S{_i + 1} {d['kind']}")
            for _i, d in enumerate(_slist)] or None
        _obsnet = _net if (use_obs and _net is not None) else None

        if panel == "Layer 1 \u00b7 Input":
            st.caption("Layer 1 \u2014 input space: heterogeneous, "
                       "spatio-temporal sources. Sensing assets are "
                       "managed here; terrain, fuel, values, roads and "
                       "resources are known priors; weather drives "
                       "$U_{Meteo}$.")
            # ---- sensor network (partial observation) ----
            st.markdown("**Sensor network**")
            use_obs = st.checkbox(
                "Partial observation via sensors",
                value=bool(_sv("dss_use_obs", True)),
                help="ON: epistemic uncertainty enters the system "
                     "through partial observation \u2014 agents only know "
                     "what the sensors deliver, uncovered or stale cells "
                     "keep old values, and the per-cell confidence "
                     "$conf=\\min\\{\\gamma_{obs},\\gamma_{cov},"
                     "\\gamma_{fre},\\gamma_{rel}\\}$ decays "
                     "(observability, coverage, freshness "
                     "$e^{-\\lambda_{conf}\\Delta t}$, source "
                     "reliability). OFF: perfect observation (debug).")
            st.session_state["dss_use_obs"] = use_obs
            _slist = list(_sv("dss_sensors", []))
            _kinds = list(_dss.SENSOR_CATALOG)
            sa1, sa2, sa3 = st.columns([1.7, 0.8, 0.8])
            _kadd = sa1.selectbox(
                "Type", _kinds,
                format_func=lambda k: _dss.SENSOR_CATALOG[k]["label"],
                key="dss_sens_kind")
            _spec_add = _dss.SENSOR_CATALOG[_kadd]
            _ranged = _spec_add.get("radius_m") is not None
            _xadd = int(sa2.number_input("x", 0, cfg.nx - 1, cfg.nx // 2,
                                         key="dss_sens_x"))
            _yadd = int(sa3.number_input("y", 0, cfg.ny - 1, cfg.ny // 2,
                                         key="dss_sens_y"))
            sa5, sa6, sa7 = st.columns([1.0, 1.0, 0.8])
            # range only for ranged kinds (satellite is whole-map)
            _radd = None
            if _ranged:
                _radd = float(sa5.number_input(
                    "range (km)", 0.2, 50.0,
                    float(_spec_add["radius_m"]) / 1000.0, 0.1,
                    key="dss_sens_r",
                    help="Coverage radius. Realistic defaults per type; "
                         "override here."))
            else:
                sa5.caption("range: whole map")
            _ladd = float(sa6.number_input(
                "latency (min)", 0.0, 120.0,
                float(_spec_add["latency_min"]), 1.0, key="dss_sens_lat",
                help="Reporting delay: this sensor's data arrives this many "
                     "minutes late (satellite ~20, public reports ~15, "
                     "aerial ~2, in-situ ~0).") )
            sa7.markdown("<div style='height:1.75em'></div>",
                         unsafe_allow_html=True)
            if sa7.button("Add", use_container_width=True):
                _new = dict(kind=_kadd, x=_xadd, y=_yadd,
                            latency_min=float(_ladd))
                if _radd is not None:
                    _new["radius_m"] = float(_radd) * 1000.0
                _slist.append(_new)
                st.session_state["dss_sensors"] = _slist
                st.rerun()
            _covt = st.slider(
                "Coverage target (%)", 0, 100,
                int(_sv("dss_cov_target", 60)), 5,
                help="Suggest network keeps adding the best next "
                     "asset (families in rotation, greedy maximum "
                     "weighted coverage) until this fraction of "
                     "the risk-weighted map is covered. 0 = only "
                     "the satellite; 100 = blanket the risk.")
            st.session_state["dss_cov_target"] = int(_covt)
            if (st.session_state.get("dss_suggest_why")
                    and st.session_state.get("dss_cov_applied")
                    is not None
                    and int(_covt) != int(
                        st.session_state["dss_cov_applied"])):
                # the network came from the optimizer: a moved target
                # re-runs it live (adds or removes assets), exactly
                # like pressing 'Suggest network' again
                _pl, _why = _dss.suggest_network(
                    world, coverage_target=float(_covt))
                st.session_state["dss_sensors"] = _pl
                st.session_state["dss_suggest_why"] = _why
                st.session_state["dss_cov_applied"] = int(_covt)
                st.rerun()
            sb1, sb2 = st.columns(2)
            if sb1.button(
                    "Suggest network", use_container_width=True,
                    help="Optimized field deployment: greedy maximum "
                         "weighted coverage, i.e. place assets one at a "
                         "time to maximize $\\sum_{x,y} r(x,y)\\,"
                         "c(x,y)$ where the risk field "
                         "$r=0.45\\,\\hat R_{spread}+0.35\\,"
                         "\\hat V_{prio}+0.20\\,F_{load}$ and each "
                         "asset lifts the coverage $c$ of its footprint "
                         "by its sensing quality. Constraints: cameras "
                         "prefer high ground (line of sight), field "
                         "posts sit on the road network, public-report "
                         "sources are pinned to the settlements, one "
                         "satellite is always tasked, same-type assets "
                         "keep a footprint apart."):
                _pl, _why = _dss.suggest_network(
                    world,
                    coverage_target=float(_covt))
                st.session_state["dss_sensors"] = _pl
                st.session_state["dss_suggest_why"] = _why
                st.session_state["dss_cov_applied"] = int(_covt)
                st.rerun()
            if st.session_state.get("dss_suggest_why"):
                with st.expander("Why these positions (optimization "
                                 "trace)"):
                    for _ln in st.session_state["dss_suggest_why"]:
                        st.caption(_ln)
            if _slist and sb2.button("Clear sensors",
                                     use_container_width=True):
                st.session_state["dss_sensors"] = []
                st.rerun()
            if _slist:
                _net0 = st.session_state.get("dss_network")
                _stat = None
                if (_net0 is not None
                        and len(getattr(_net0, "sensors", [])) == len(_slist)
                        and hasattr(_net0, "status")):
                    _stat = _net0.status()
                for _i, _sd in enumerate(list(_slist)):
                    _c1, _c2, _c3 = st.columns([5, 0.7, 0.7])
                    _spec = _dss.SENSOR_CATALOG[_sd["kind"]]
                    _cc = _spec.get("color", (120, 200, 255))
                    _scol = f"#{_cc[0]:02x}{_cc[1]:02x}{_cc[2]:02x}"
                    _full = _spec["label"]
                    # effective range/latency = per-sensor override or default
                    _rm = (_sd.get("radius_m") if _sd.get("radius_m")
                           is not None else _spec.get("radius_m"))
                    _lat = (_sd.get("latency_min")
                            if _sd.get("latency_min") is not None
                            else _spec.get("latency_min", 0.0))
                    _rng = ("whole map" if _rm is None
                            else f"{_rm / 1000.0:.1f} km range")
                    _tx = (f"S{_i + 1} \u2014 {_full} @ "
                           f"({_sd['x']}, {_sd['y']}) \u00b7 {_rng} \u00b7 "
                           f"{_lat:.0f} min latency")
                    if _stat is not None:
                        _si = _stat[_i]
                        _lr = _si["last_report_min"]
                        _tx += (f" \u00b7 next pass "
                                f"{_si['next_pass_min']:.0f} min")
                        _tx += (f" \u00b7 data {_lr:.0f} min old"
                                if _lr is not None else
                                " \u00b7 no data yet")
                        if _si["in_transit"]:
                            _tx += (f" \u00b7 {_si['in_transit']} report "
                                    f"en route ({_si['latency_min']:.0f} "
                                    f"min latency)")
                    _c1.markdown(
                        f"<span style='color:{_scol}'>\u25cf</span> "
                        f"<small>{_tx}</small>",
                        unsafe_allow_html=True)
                    if _c2.button("\u270e", key=f"dss_sed_{_i}",
                                  help="Edit this sensor"):
                        st.session_state["dss_sens_edit"] = (
                            None if st.session_state.get("dss_sens_edit")
                            == _i else _i)
                        st.rerun()
                    if _c3.button("\u2716",
                                  key=(f"dss_srm_{_i}_{_sd['kind']}_"
                                       f"{_sd['x']}_{_sd['y']}"),
                                  help="Remove this sensor"):
                        _slist.pop(_i)
                        st.session_state["dss_sensors"] = _slist
                        st.session_state["dss_sens_edit"] = None
                        st.rerun()
                    if st.session_state.get("dss_sens_edit") == _i:
                        e1, e2, e3 = st.columns([1.7, 0.8, 0.8])
                        _nk = e1.selectbox(
                            "Type", _kinds,
                            index=_kinds.index(_sd["kind"]),
                            format_func=lambda k:
                                _dss.SENSOR_CATALOG[k]["label"],
                            key=f"dss_se_k_{_i}")
                        _nspec = _dss.SENSOR_CATALOG[_nk]
                        _nranged = _nspec.get("radius_m") is not None
                        _nx_ = int(e2.number_input(
                            "x", 0, cfg.nx - 1, int(_sd["x"]),
                            key=f"dss_se_x_{_i}"))
                        _ny_ = int(e3.number_input(
                            "y", 0, cfg.ny - 1, int(_sd["y"]),
                            key=f"dss_se_y_{_i}"))
                        e5, e6, e7 = st.columns([1.0, 1.0, 0.8])
                        # keep the sensor's override only when the type is
                        # unchanged; changing type falls back to the new
                        # type's realistic defaults (keyed by _nk so the
                        # field resets on a type switch)
                        _same = _sd["kind"] == _nk
                        _cr = (_sd.get("radius_m")
                               if (_same and _sd.get("radius_m") is not None)
                               else _nspec.get("radius_m"))
                        _nr = None
                        if _nranged:
                            _nr = float(e5.number_input(
                                "range (km)", 0.2, 50.0,
                                float(_cr) / 1000.0, 0.1,
                                key=f"dss_se_r_{_i}_{_nk}"))
                        else:
                            e5.caption("range: whole map")
                        _cl = (_sd.get("latency_min")
                               if (_same and _sd.get("latency_min")
                                   is not None)
                               else _nspec.get("latency_min", 0.0))
                        _nl = float(e6.number_input(
                            "latency (min)", 0.0, 120.0, float(_cl), 1.0,
                            key=f"dss_se_lat_{_i}_{_nk}"))
                        e7.markdown("<div style='height:1.75em'></div>",
                                    unsafe_allow_html=True)
                        if e7.button("Save", key=f"dss_se_s_{_i}",
                                     use_container_width=True):
                            _upd = dict(kind=_nk, x=_nx_, y=_ny_,
                                        latency_min=float(_nl))
                            if _nr is not None:
                                _upd["radius_m"] = float(_nr) * 1000.0
                            _slist[_i] = _upd
                            st.session_state["dss_sensors"] = _slist
                            st.session_state["dss_sens_edit"] = None
                            st.rerun()
            elif use_obs:
                st.warning("No sensors placed \u2014 the agents are "
                           "BLIND ($conf \\approx 0$): they keep "
                           "assuming no fire. Add sensors or use 'Suggest "
                           "network'.")
            st.divider()
            st.markdown("**Resource pool \u2014 $U_{Res}$ "
                        "($R_{cap}, R_{avail}, R_{eff}, R_{time}$)**")
            _ritems = st.session_state.get("dss_res_items")
            if (st.session_state.get("dss_res_base_v")
                    != st.session_state.map_version):
                _ritems = None      # map changed: the old pool is void
            _rt1, _rt2 = st.columns(2)
            st.session_state["dss_eff_target"] = int(_rt1.slider(
                "Target effectiveness (%)", 10, 90,
                int(_sv("dss_eff_target", 50)), 5,
                help="The planner stages the baseline pool (depots + "
                     "road corridor + one helibase), then keeps ADDING "
                     "aerial units on the worst risk-weighted reach "
                     "gaps until the expected intervention "
                     "effectiveness meets this target (up to 10 "
                     "additions). Every added unit stays an editable "
                     "row below."))
            st.session_state["dss_res_cov"] = float(st.slider(
                "Base coverage (% of candidate sites staged)", 10, 100,
                int(_sv("dss_res_cov", 100)), 5,
                help="How many of the candidate sites actually get a base. "
                     "Capacity alone is not scarcity: a thin depot in every "
                     "town still puts a response everywhere, while a "
                     "province with three brigades leaves ground with none "
                     "near it. The sites with the most to protect are kept "
                     "first."))
            st.session_state["dss_res_density"] = float(_rt2.slider(
                "Resource density (×R_cap)", 0.2, 2.0,
                float(_sv("dss_res_density", 1.0)), 0.1,
                help="Scales the staged capacity R_cap of every unit — how "
                     "much suppression sits on the map (the eta_cap = "
                     "R_cap/R_cap_max term). <1 = a sparse, "
                     "under-resourced pool that may FAIL to hold the fire; "
                     ">1 = a dense pool. Separate from effectiveness, which "
                     "closes reach gaps with aerial units."))
            if (st.session_state.get("dss_res_why") and _ritems
                    and st.session_state.get("dss_eff_applied")
                    is not None
                    and (int(st.session_state["dss_eff_target"])
                         != int(st.session_state["dss_eff_applied"])
                         or float(st.session_state["dss_res_density"])
                         != float(st.session_state.get("dss_dens_applied",
                                                       1.0))
                         # ITS OWN KEY. This started out as
                         # "dss_cov_applied", which is the SENSOR panel's
                         # coverage-target memory: the two panels then
                         # overwrote each other's number, each saw a change
                         # it had not made, each re-suggested and called
                         # st.rerun(), and the screen blinked in a loop
                         # while the suggested sensors were thrown away on
                         # every pass.
                         or float(st.session_state["dss_res_cov"])
                         != float(st.session_state.get("dss_rescov_applied",
                                                       100.0)))):
                _its, _rwhy = _dss.suggest_resource_items(
                    world, efficiency_target=float(
                        st.session_state["dss_eff_target"]) / 100.0,
                    density=float(st.session_state["dss_res_density"]),
                    coverage=float(st.session_state["dss_res_cov"]) / 100.0)
                st.session_state["dss_res_items"] = _its
                st.session_state["dss_res_why"] = _rwhy
                st.session_state["dss_res_base_v"] = \
                    st.session_state.map_version
                st.session_state["dss_eff_applied"] = int(
                    st.session_state["dss_eff_target"])
                st.session_state["dss_dens_applied"] = float(
                    st.session_state["dss_res_density"])
                st.session_state["dss_rescov_applied"] = float(
                    st.session_state["dss_res_cov"])
                st.rerun()
            rp1, rp2 = st.columns(2)
            if rp1.button(
                    "Suggest resources", use_container_width=True,
                    help="Builds the baseline pool as EDITABLE rows: a "
                         "ground depot at each FIRE STATION (capacity "
                         "$0.8\\,R_{cap}^{max}$, station radius; a town "
                         "base only if the map has no fire station) plus a "
                         "thin road-corridor capacity and helibase(s). "
                         "Hospitals, schools and offices are protected "
                         "values, not depots. "
                         "$R_{eff}$ comes from the terrain access "
                         "field, $R_{time}$ from the road-network "
                         "distance (dispatch at the nearest base, then "
                         "60 km/h on roads). The Layer 4 decisions "
                         "allocate THIS pool and cannot exceed 1.5x "
                         "the staged capacity anywhere."):
                _its, _rwhy = _dss.suggest_resource_items(
                    world, efficiency_target=float(
                        st.session_state["dss_eff_target"]) / 100.0,
                    coverage=float(st.session_state.get("dss_res_cov",
                                                        100.0)) / 100.0)
                st.session_state["dss_res_items"] = _its
                st.session_state["dss_res_why"] = _rwhy
                st.session_state["dss_res_base_v"] = \
                    st.session_state.map_version
                st.session_state["dss_eff_applied"] = int(
                    st.session_state["dss_eff_target"])
                st.rerun()
            if _ritems and rp2.button("Clear pool",
                                      use_container_width=True):
                st.session_state["dss_res_items"] = None
                st.session_state["dss_res_base"] = None
                st.rerun()
            if st.session_state.get("dss_res_why") and _ritems:
                for _ln in st.session_state["dss_res_why"]:
                    st.caption(_ln)
            if _ritems:
                _bpe = st.session_state.get("dss_res_base")
                if _bpe is None:   # first render after Suggest
                    _bpe = _dss.build_resource_layer(world, _ritems)
                _pe, _ped = _dss.pool_efficiency(world, _bpe)
                st.progress(min(1.0, _pe),
                            text=f"Expected intervention "
                                 f"effectiveness: {_pe:.0%}")
                st.caption(f"= reach {_ped['reach']:.0%} \u00d7 "
                           f"capacity {_ped['capacity']:.0%} \u00b7 "
                           f"aerial covers {_ped.get('air', 0.0):.0%} "
                           "of the risk \u2014 "
                           "risk-weighted: can the crews reach "
                           "the ground that matters, and is the "
                           "staged pool big enough for it? Add "
                           "depots near the risk or raise their "
                           "capacity to push this up.")
            if _ritems:
                for _j, _it in enumerate(list(_ritems)):
                    _c1, _c2, _c3 = st.columns([5, 0.7, 0.7])
                    _is_air = _it["kind"] == "helibase"
                    _typ = "AERIAL" if _is_air else "GROUND"
                    if _it["kind"] == "road_corridor":
                        _c1.caption(
                            f"{_j + 1}. [GROUND] road corridor \u00b7 "
                            f"Rcap {_it['cap']:.2f} \u00b7 "
                            f"Ravail {_it.get('avail', 1.0):.2f}")
                    else:
                        _unit = _dss.RESOURCE_KINDS.get(
                            _it["kind"], {}).get("short", _it["kind"])
                        _c1.caption(
                            f"{_j + 1}. [{_typ}] D{_j + 1} {_unit} @ "
                            f"({_it['x']}, {_it['y']}) \u00b7 Rcap "
                            f"{_it['cap']:.2f} \u00b7 Ravail "
                            f"{_it.get('avail', 1.0):.2f} \u00b7 "
                            f"dispatch {_it.get('t_disp', 10.0):.0f} min "
                            f"\u00b7 r {_it['radius']} \u00b7 "
                            f"{_it.get('label', '')}")
                    if _c2.button("\u270e", key=f"res_ed_{_j}",
                                  help="Edit this row"):
                        st.session_state["res_edit"] = (
                            None if st.session_state.get("res_edit")
                            == _j else _j)
                        st.rerun()
                    if _c3.button("\u2716", key=f"res_rm_{_j}",
                                  help="Remove this row"):
                        _ritems.pop(_j)
                        st.session_state["dss_res_items"] = _ritems
                        st.session_state["res_edit"] = None
                        st.rerun()
                    if st.session_state.get("res_edit") == _j:
                        if _it["kind"] == "road_corridor":
                            f1, f2, f3 = st.columns([1.0, 1.0, 0.8])
                            _nc = float(f1.number_input(
                                "$R_{cap}$", 0.0, 1.0, float(_it["cap"]),
                                0.05, key=f"res_e_c_{_j}"))
                            _na = float(f2.number_input(
                                "$R_{avail}$", 0.0, 1.0,
                                float(_it.get("avail", 1.0)), 0.05,
                                key=f"res_e_a_{_j}"))
                            f3.markdown("<div style='height:1.75em'>"
                                        "</div>", unsafe_allow_html=True)
                            if f3.button("Save", key=f"res_e_s_{_j}",
                                         use_container_width=True):
                                _it.update(cap=_nc, avail=_na)
                                st.session_state["res_edit"] = None
                                st.rerun()
                        else:
                            f1, f2, f3, f4 = st.columns(4)
                            _nx_ = int(f1.number_input(
                                "x", 0, cfg.nx - 1, int(_it["x"]),
                                key=f"res_e_x_{_j}"))
                            _ny_ = int(f2.number_input(
                                "y", 0, cfg.ny - 1, int(_it["y"]),
                                key=f"res_e_y_{_j}"))
                            _nc = float(f3.number_input(
                                "$R_{cap}$", 0.0, 1.0, float(_it["cap"]),
                                0.05, key=f"res_e_c_{_j}"))
                            # the generator sizes a helibase as
                            # max(6, min(nx, ny) // 8), which on a big map
                            # exceeds 20 and crashed this editor. The bound
                            # follows the map instead of a constant.
                            _rmaxE = max(20, int(min(cfg.nx, cfg.ny)))
                            _nr = int(f4.number_input(
                                "r (cells)", 1, _rmaxE,
                                int(min(_it["radius"], _rmaxE)),
                                key=f"res_e_r_{_j}"))
                            f5, f6, f7 = st.columns([1.0, 1.0, 0.8])
                            _na = float(f5.number_input(
                                "$R_{avail}$", 0.0, 1.0,
                                float(_it.get("avail", 1.0)), 0.05,
                                key=f"res_e_a_{_j}"))
                            _nt = float(f6.number_input(
                                "dispatch (min)", 0.0, 120.0,
                                float(_it.get("t_disp", 10.0)), 1.0,
                                key=f"res_e_t_{_j}"))
                            f7.markdown("<div style='height:1.75em'>"
                                        "</div>", unsafe_allow_html=True)
                            if f7.button("Save", key=f"res_e_s_{_j}",
                                         use_container_width=True):
                                _it.update(x=_nx_, y=_ny_, cap=_nc,
                                           radius=_nr, avail=_na,
                                           t_disp=_nt)
                                st.session_state["res_edit"] = None
                                st.rerun()
            else:
                st.caption("No pool staged yet \u2014 add a unit below, or "
                           "press Suggest resources.")
            # --- Add a resource unit: ALWAYS available (same logic as the
            # sensor add). Type comes from the canonical vocabulary
            # dss.RESOURCE_KINDS; only the user-addable kinds show. ---
            st.caption("Add a resource unit")
            _ritems = list(st.session_state.get("dss_res_items") or [])
            _addable = [k for k, v in _dss.RESOURCE_KINDS.items()
                        if v.get("addable")]
            a0, a1, a2, a3, a4, a5 = st.columns(
                [1.1, 0.8, 0.8, 0.8, 0.8, 0.8])
            _akind = a0.selectbox(
                "type", _addable,
                format_func=lambda k: _dss.RESOURCE_KINDS[k]["label"],
                key="res_a_kind",
                help="Ground depot = road-bound crews; Helibase = aerial, "
                     "map-wide reach.")
            _kd = _dss.RESOURCE_KINDS[_akind]
            _ax = int(a1.number_input("x", 0, cfg.nx - 1,
                                      cfg.nx // 2, key="res_a_x"))
            _ay = int(a2.number_input("y", 0, cfg.ny - 1,
                                      cfg.ny // 2, key="res_a_y"))
            _ac = float(a3.number_input("cap", 0.0, 1.0,
                                        float(_kd["cap"]), 0.05,
                                        key="res_a_c"))
            _ar = int(a4.number_input("r", 1, 40, int(_kd["radius"] or 4),
                                      key="res_a_r"))
            a5.markdown("<div style='height:1.75em'></div>",
                        unsafe_allow_html=True)
            if a5.button("Add", key="res_a_b", use_container_width=True):
                _ritems.append(dict(kind=_akind, x=_ax, y=_ay,
                                    cap=_ac, radius=_ar, avail=1.0,
                                    t_disp=float(_kd["t_disp"]),
                                    label="manual " + _kd["short"]))
                st.session_state["dss_res_items"] = _ritems
                st.rerun()

        elif panel == "Environment":
            st.caption("Environment \u2014 the exogenous drivers that force "
                       "the fire before any decision layer: weather "
                       "(U_Meteo) and, through the diurnal model, fuel "
                       "moisture.")
            st.markdown("**Weather / environment drivers \u2014 $U_{Meteo}$**")
            ws = st.slider(
                "$W_{ws}$ — wind speed (m/s)", 0.0, 30.0,
                float(world.meteo.wws.mean()), 0.5,
                help="Synoptic (large-scale) wind. The engine modulates it "
                     "per cell by the terrain: exposed ridges speed it up, "
                     "valleys shelter it "
                     "($\\mathrm{clip}(1+g_{tw}(2\\hat e-1),0.4,1.8)$) "
                     "and on steep ground the direction veers toward the "
                     "local valley axis. Spread saturates toward $w_0$ "
                     "(Parameters).")
            st.caption(_wind_speed_label(ws))
            cur_deg = float(np.degrees(world.meteo.wwd.mean())) % 360.0
            st.markdown("$W_{wd}$ — **wind direction**: click the "
                        "dial; the arrow shows where the wind blows "
                        "**toward**")
            sel = _wind_compass(cur_deg,
                                key=f"wind_dial_{st.session_state.map_version}")
            if sel is not None and min((sel - cur_deg) % 360,
                                       (cur_deg - sel) % 360) > 0.5:
                world.meteo.wwd[:] = np.radians(sel)
                st.rerun()
            st.caption(f"Current: {cur_deg:.0f}\u00b0 "
                       "(0\u00b0 = east, 90\u00b0 = north). Terrain bends "
                       "this synoptic direction toward valley axes on "
                       "steep ground.")
            emc = st.checkbox(
                "Auto fuel moisture from weather (EMC)", value=True,
                help="Equilibrium Moisture Content (Simard 1968): computes "
                     "dead surface fuel moisture from air temperature and "
                     "relative humidity every step. Hotter and drier air "
                     "gives drier fuel and a faster fire. When on, the "
                     "slider below is ignored.")
            st.session_state["emc_on"] = emc
            _modes = ["Manual", "Diurnal cycle", "Real case weather"]
            _mcur = st.session_state.get("wx_mode", "Manual")
            wxm = st.radio(
                "Weather source", _modes,
                index=_modes.index(_mcur) if _mcur in _modes else 0,
                help="Manual: the sliders above stay fixed. Diurnal "
                     "cycle: a synthetic day/night wave (afternoon "
                     "peak); humid nights push fuel moisture above "
                     "$m_{ext}$ and the fire can die out overnight. "
                     "Real case weather: the hourly ERA5 series already "
                     "downloaded for a validation case drives wind, "
                     "temperature and humidity with real timestamps.")
            st.session_state["wx_mode"] = wxm
            st.session_state["diurnal_on"] = (wxm == "Diurnal cycle")
            if wxm == "Real case weather":
                import glob as _gl
                import json as _js
                _root = os.path.dirname(os.path.dirname(
                    os.path.abspath(__file__)))
                _found = {}
                for _wf in _gl.glob(os.path.join(
                        _root, "validation", "cache", "*",
                        "weather_*.json")):
                    _cid = os.path.basename(os.path.dirname(_wf))
                    _found[_cid] = _wf
                if not _found:
                    st.warning("No cached case weather yet \u2014 run "
                               "a case once on the Validation page.")
                else:
                    _cids = sorted(_found)
                    _csel = st.selectbox("Case weather", _cids,
                                         key="wx_case")
                    _h0 = st.number_input(
                        "Start hour (UTC) in the series", 0, 23, 11, 1,
                        key="wx_h0",
                        help="Simulation step 0 maps to this hour of "
                             "the case's first day (fires typically "
                             "start late morning).")
                    try:
                        _av2 = _load_auto_validate()
                        _start = _av2.CASES.get(_csel, {}).get(
                            "start", "2021-01-01")
                    except Exception:
                        _start = "2021-01-01"
                    _js0 = _js.load(open(_found[_csel]))
                    st.session_state["real_wx"] = {
                        "start": _start, "h0": float(_h0),
                        "ws": _js0["wind_speed_10m"],
                        "wd": _js0["wind_direction_10m"],
                        "t": _js0["temperature_2m"],
                        "rh": _js0["relative_humidity_2m"]}
                    st.caption(f"\u2713 {_csel}: {_start} + "
                               f"{int(_h0):02d}:00 UTC \u2014 wind, "
                               "T and RH follow the real series; the "
                               "map clock shows the real date-time.")
            if wxm == "Diurnal cycle":
                rr1, rr2, rr3 = st.columns(3)
                st.session_state["dr_rain_mm"] = float(rr1.number_input(
                    "Rain (mm/h)", 0.0, 30.0,
                    float(st.session_state.get("dr_rain_mm", 0.0)), 1.0,
                    help="0 = no rain. Rain wets the fuel toward "
                         "extinction moisture and stops ember spotting "
                         "\u2014 fires die out under sustained rain."))
                st.session_state["dr_rain_start"] = float(rr2.number_input(
                    "Rain start (h)", 0.0, 23.0,
                    float(st.session_state.get("dr_rain_start", 18.0)),
                    1.0))
                st.session_state["dr_rain_dur"] = float(rr3.number_input(
                    "Duration (h)", 0.5, 24.0,
                    float(st.session_state.get("dr_rain_dur", 3.0)), 0.5))
                dc1, dc2 = st.columns(2)
                st.session_state["dr_tday"] = float(dc1.number_input(
                    "Day $T$ (\u00b0C)", -10.0, 50.0,
                    float(st.session_state.get("dr_tday", 34.0)), 1.0))
                st.session_state["dr_tnight"] = float(dc2.number_input(
                    "Night $T$ (\u00b0C)", -20.0, 40.0,
                    float(st.session_state.get("dr_tnight", 22.0)), 1.0))
                st.session_state["dr_rhday"] = float(dc1.number_input(
                    "Day RH (%)", 5.0, 100.0,
                    float(st.session_state.get("dr_rhday", 20.0)), 5.0))
                st.session_state["dr_rhnight"] = float(dc2.number_input(
                    "Night RH (%)", 5.0, 100.0,
                    float(st.session_state.get("dr_rhnight", 70.0)), 5.0))
            mo = st.slider(
                "$F_{moist}$ — fuel moisture, whole map (mass fraction)",
                0.0, 0.6,
                float(world.fuel.fmoist.mean()), 0.01, disabled=emc,
                help="Surface fuel moisture as a mass fraction, applied "
                     "uniformly to every cell. It is a daily weather "
                     "condition, which is why it lives here; spatially "
                     "varying moisture can be painted in the Map editor "
                     "(moving this slider overwrites painted values). "
                     "At the extinction moisture $m_{ext}$ of a fuel class "
                     "the spread stops entirely.")
            # apply weather immediately, each control independently, so an
            # unrelated change never overwrites a painted field
            if abs(ws - float(world.meteo.wws.mean())) > 1e-9:
                world.meteo.wws[:] = ws
            if emc:
                from disaster_phyengine.fuel_moisture import update_dead_fuel_moisture
                update_dead_fuel_moisture(world)
            else:
                _last_mo = st.session_state.get("last_mo")
                if _last_mo is not None and abs(mo - _last_mo) > 1e-9:
                    world.fuel.fmoist[:] = mo
                    # the user's setting is the new PRE-FIRE baseline:
                    # 'Reset fire' must return here, not to the value
                    # the map generator happened to produce
                    _simM = st.session_state.get("sim")
                    if _simM is not None:
                        _simM._fmoist0 = world.fuel.fmoist.copy()
                st.session_state["last_mo"] = mo

        elif panel == "Layer 3 \u00b7 Concepts":
            st.caption("Layer 3 \u2014 concept space: the features are "
                       "fuzzified on the five-term partition and "
                       "aggregated up the four-level hierarchy with "
                       "weights $\\omega$; every activation is gated "
                       "by the observation confidence and blended with "
                       "a persistence prior before any rule reads it.")
            # what the GenAI stage has grown ON TOP of the base hierarchy,
            # and whether it is ACTIVE in the engine right now
            _engC = st.session_state.get("dss_engine")
            _stg3 = bool(st.session_state.get("dss_use_stage3", True))
            # the hierarchy this panel renders. With a live engine it is the
            # engine's. WITHOUT one (fresh app start, nothing run yet) the
            # STORE is the ground truth, so what has already been learned is
            # visible immediately instead of only after the first run.
            _hierC = getattr(_engC, "hierarchy", None)
            _macC = dict(getattr(_engC, "macros", {}) or {})
            _fromStore = False
            if _hierC is None:
                _hierC = dict(_dss.HIERARCHY)
                if _stg3:
                    try:
                        _scv, _smv = _store_vocab()
                        for _cn0, _cd0 in _scv.items():
                            _hierC[_cn0] = (
                                int(_cd0.get("level", 2)),
                                [(a, float(b))
                                 for a, b in _cd0.get("inputs", [])])
                        _macC = {k: dict(composition=[
                            (a, float(b)) for a, b in
                            v.get("composition", [])])
                            for k, v in _smv.items()}
                        _fromStore = bool(_scv or _smv)
                    except Exception as _e_voc:
                        # never swallow this: a silent failure here reads as
                        # "nothing has been learned", which is a lie
                        st.warning("Could not read the learned store "
                                   f"({type(_e_voc).__name__}: {_e_voc}).")
            _newC3 = [c for c in _hierC if c not in _dss.HIERARCHY]
            _newM3 = list(_macC)
            if _newC3 or _newM3:
                _cbits = ", ".join(f"{c} (L{_hierC[c][0]})"
                                   for c in _newC3) or "none"
                st.success("GENERATED VOCABULARY "
                           + ("in the persistent store, and it will load "
                              "with the next engine"
                              if _fromStore else "active in this engine")
                           + " · concepts: " + _cbits
                           + " · macro interventions: "
                           + (", ".join(_newM3) or "none")
                           + " — marked \U0001F7E9 GenAI in the "
                           "tables below.")
            else:
                st.caption("No generated concept: "
                           + ("the store holds none yet."
                              if _stg3 else
                              "“Use stage ③ rules” is OFF, "
                              "so stored concepts are not loaded."))
            _regsC = _dss.partition_n(cfg.nx, cfg.ny, int(_sv("dss_n", 1)))
            _namesC = [r.name for r in _regsC] + ["All agents (table)"]
            # default to the all-agents table (the last entry)
            _iC = min(int(_sv("dss_sel_i", len(_namesC) - 1)),
                      len(_namesC) - 1)
            _selC = st.selectbox("Agent (region)", _namesC, index=_iC,
                                 key="l3_agent")
            if _selC == "All agents (table)":
                _pool3 = st.session_state.get("dss_res_base")
                _head = "| concept |" + "".join(
                    f" {_r.name} |" for _r in _regsC)
                _sep = "|---|" + "---|" * len(_regsC)
                _cols = {}
                # engine hierarchy when a run is live, otherwise the
                # STORE-backed one built above
                _eng_h = _hierC
                for _r in _regsC:
                    _f3 = _dss.ten_features(sim, _r, network=_obsnet,
                                            pool=_pool3)
                    _g3 = _dss.concept_gates(
                        _dss.feature_confidence(_obsnet, _r),
                        hierarchy=_eng_h)
                    _gt = st.session_state.get(f"l3_gate_{_r.name}")
                    if _gt is None:
                        _gt = _dss.GatedConcepts()
                        st.session_state[f"l3_gate_{_r.name}"] = _gt
                    _cols[_r.name] = _dss.crisp(_gt.gate(
                        _dss.infer_concepts(_f3, hierarchy=_eng_h),
                        _g3, step=int(sim.state.step)))
                _rows3 = [_head, _sep]
                # level order, so a GENERATED concept sits with the other
                # concepts of its own level instead of dangling at the end
                for _cn, (_l, _ins3) in sorted(
                        (_eng_h or _dss.HIERARCHY).items(),
                        key=lambda kv: (kv[1][0], kv[0])):
                    _lab = _dss.CONCEPT_LABEL.get(
                        _cn, _cn.replace("_", " ")
                        + " \U0001F7E9 GenAI")
                    if _cn in _dss.DECISION_CONCEPTS:
                        _lab = f"**{_lab}** \u2605"
                    _cells = " | ".join(f"{_cols[_r.name][_cn]:.2f}"
                                        for _r in _regsC)
                    _rows3.append(f"| L{_l} {_lab} | {_cells} |")
                st.markdown("\n".join(_rows3))
                st.caption("gated (effective) activations, crisp "
                           "readout \u00b7 \u2605 = the five decision "
                           "concepts")
            else:
                _regC = _regsC[_namesC.index(_selC)]
                st.session_state["dss_region"] = (*_regC.box, _regC.name)
                _fC = _dss.ten_features(
                    sim, _regC, network=_obsnet,
                    pool=st.session_state.get("dss_res_base"))
                _fcC = _dss.feature_confidence(_obsnet, _regC)
                _eng_hC = _hierC
                _gamC = _dss.concept_gates(_fcC, hierarchy=_eng_hC)
                _gkey = f"l3_gate_{_regC.name}"
                _gater = st.session_state.get(_gkey)
                if _gater is None:
                    _gater = _dss.GatedConcepts()
                    st.session_state[_gkey] = _gater
                _actC = _dss.infer_concepts(_fC, hierarchy=_eng_hC)
                _effC = _gater.gate(_actC, _gamC, step=int(sim.state.step))
                _crO = _dss.crisp(_actC)
                _crE = _dss.crisp(_effC)
                st.caption("per-concept gate $\\gamma$ = min of the "
                           "feature confidences feeding the concept \u00b7 "
                           f"persistence $\\rho={_dss.RHO_PERSIST}$ \u00b7 "
                           "observed \u2192 gated (effective)")
                _lvlname = {1: "Level 1 \u2014 base", 2: "Level 2",
                            3: "Level 3", 4: "Level 4 \u2014 coordination"}
                for _lvl in (1, 2, 3, 4):
                    st.markdown(f"**{_lvlname[_lvl]}**")
                    for _cn, (_l, _ins) in (_eng_hC
                                            or _dss.HIERARCHY).items():
                        if _l != _lvl:
                            continue
                        _dec = _cn in _dss.DECISION_CONCEPTS
                        _lab = _dss.CONCEPT_LABEL.get(
                            _cn, _cn.replace("_", " ")
                            + " \U0001F7E9 GenAI")
                        if _dec:
                            _lab = f"{_lab} \u2605"
                        st.progress(min(1.0, _crE[_cn]),
                                    text=f"{_lab}: {_crO[_cn]:.2f} "
                                         f"\u2192 {_crE[_cn]:.2f} "
                                         f"(\u03b3 {_gamC[_cn]:.2f})")
                st.caption("\u2605 = the five decision concepts the "
                           "intervention rules read.")
            with st.expander("Feature \u2192 concept mapping "
 " \u2014 explicit"):
                _fn = dict(_dss.FEATURE_NAME)
                _fn["access_road_status_inv"] = \
                    "access / road status (inverted)"
                _eng_hM = _hierC or _dss.HIERARCHY
                for _cn, (_l, _ins) in _eng_hM.items():
                    _parts = " + ".join(
                        f"{_fn.get(_src, _dss.CONCEPT_LABEL.get(_src, _src))}"
                        f" ({_w:.2f})" for _src, _w in _ins)
                    _clab = _dss.CONCEPT_LABEL.get(
                        _cn, _cn.replace("_", " ") + " \U0001F7E9")
                    st.caption(f"L{_l} **{_clab}** "
                               f"\u2190 {_parts}")
        elif panel == "Layer 4 \u00b7 Decision":
            st.caption("Layer 4 \u2014 decision space (evolution). "
                       "These knobs configure the staged adaptation; "
                       "the stages themselves arrive with the decision "
                       "layer.")
            st.markdown("**DSS mode**")
            # master switch: DSS active vs no DSS (fire runs free)
            _mirror("dss_apply", "l4_dss_apply", False)
            st.toggle(
                "DSS active", key="l4_dss_apply",
                on_change=_adopt, args=("dss_apply", "l4_dss_apply"),
                help="ON: the DSS decides and its orders enter the "
                     "simulation. OFF: no DSS — the fire runs free "
                     "(baseline for comparison).")
            _active = bool(st.session_state.get("dss_apply"))
            # (1) LOAD: which pre-learned rules are brought in from the store
            st.caption("Use pre-learned rules (loaded from the store, act "
                       "like seeds):")
            _u1, _u2 = st.columns(2)
            _mirror("dss_use_stage12", "l4_use12", True)
            _u1.toggle(
                "Use stage ①② rules", key="l4_use12",
                on_change=_adopt, args=("dss_use_stage12", "l4_use12"),
                disabled=not _active,
                help="Load the previously-learned stage 1/2 rules (tuned "
                     "seeds, A# resolution rules, term inserts) from the "
                     "store. They are used like seeds. This does NOT by "
                     "itself run live evFIS. OFF: only the seed base (the "
                     "rules stay in the store, not deleted).")
            _mirror("dss_use_stage3", "l4_use3", True)
            _u2.toggle(
                "Use stage ③ rules", key="l4_use3",
                on_change=_adopt, args=("dss_use_stage3", "l4_use3"),
                disabled=not _active,
                help="Load the previously-generated stage 3 rules + concepts "
                     "+ macro interventions from the store. This does NOT by "
                     "itself run live GenAI. OFF: the generated rules / "
                     "concepts / interventions are not used (kept in the "
                     "store).")
            # (2) MODULES: which adaptation stages RUN LIVE during the sim
            st.caption("Run modules live (keep adapting / generating during "
                       "the run):")
            _m1, _m2 = st.columns(2)
            _mirror("dss_evfis_on", "l4_evfis_on", True)
            _m1.toggle(
                "evFIS active", key="l4_evfis_on",
                on_change=_adopt, args=("dss_evfis_on", "l4_evfis_on"),
                disabled=not _active,
                help="Run stage 1/2 adaptation live: tune "
                     "memberships/consequents (①) and instantiate new A# "
                     "rules (②). OFF: no live tuning — the DSS runs fuzzy on "
                     "the loaded rules.")
            _mirror("dss_genai_on", "l4_genai_on", True)
            _m2.toggle(
                "GenAI active", key="l4_genai_on",
                on_change=_adopt, args=("dss_genai_on", "l4_genai_on"),
                disabled=not _active,
                help="Run stage 3 live: Claude proposes new G# rules and, "
                     "when needed, new concepts / interventions. OFF: no live "
                     "generation.")
            st.session_state["dss_genai_on2"] = \
                st.session_state["dss_genai_on"]
            st.session_state["dss_adapt_on"] = bool(
                st.session_state["dss_evfis_on"]
                or st.session_state["dss_genai_on"])
            _ev = st.session_state["dss_evfis_on"]
            _ge = st.session_state["dss_genai_on"]
            # (the resulting mode is shown in the DSS configuration line below)
            # seed rule profile the DSS starts from (SAME key as the Layer 4
            # Rules panel, so the two stay in sync)
            _sp_opts = {"40 rules (full)": "full",
                        "core (22 rules)": "core",
                        "minimal (5 rules)": "minimal"}
            _sp_cur = str(_sv("dss_seed_profile", "minimal"))
            _sp_idx = next((i for i, v in enumerate(_sp_opts.values())
                            if v == _sp_cur), 0)
            _sp_sel = st.selectbox("Seed rule base", list(_sp_opts),
                                   index=_sp_idx, key="l4_seed_profile",
                                   help="How much doctrine the DSS starts "
                                        "with: 40 (full), core 22, or minimal "
                                        "5 (learns the rest). This is the "
                                        "OPERATIONAL profile the DSS runs; the "
                                        "Layer 4 Rules tab selector is "
                                        "view-only.")
            if _sp_opts[_sp_sel] != _sp_cur:
                st.session_state["dss_seed_profile"] = _sp_opts[_sp_sel]
                st.rerun()
            # one-line summary of the resulting DSS configuration
            _prof_lbl = {"full": "40", "core": "22 core",
                         "minimal": "5 minimal"}.get(
                             str(_sv("dss_seed_profile", "minimal")), "?")
            if not _active:
                st.markdown("**DSS configuration: No DSS** — the fire runs "
                            "free.")
            else:
                _u12 = st.session_state["dss_use_stage12"]
                _u3 = st.session_state["dss_use_stage3"]
                _loaded = ([] + (["evFIS rules"] if _u12 else [])
                           + (["GenAI rules"] if _u3 else []))
                _live = ([] + (["evFIS"] if _ev else [])
                         + (["GenAI"] if _ge else []))
                _parts = [f"seed {_prof_lbl}"]
                if _loaded:
                    _parts.append("+ " + " + ".join(_loaded))
                _line = " · ".join(_parts)
                _line += (" · live: " + " + ".join(_live)) if _live \
                    else " · live: none (fuzzy)"
                st.markdown(f"**DSS configuration: {_line}**")
                # GROUND TRUTH: what the ACTUAL running engine loaded. If this
                # disagrees with the toggles above, the process is running a
                # stale engine (fully restart the dashboard).
                _engD = st.session_state.get("dss_engine")
                if _engD is not None:
                    # count ONLY rules genuinely loaded from the store this
                    # build (R41/R42 seed examples carry an evFIS note but are
                    # NOT store-loaded, so they must not count)
                    # the resolver TAGS every rule it brought back from the
                    # store, so the count is read off that tag. It used to
                    # match a note string the resolver no longer writes,
                    # which pinned this number at 0 no matter what loaded.
                    _nlrn = sum(1 for r in _engD.rules
                                if getattr(r, "from_store", None))
                    _n2 = sum(1 for r in _engD.rules
                              if getattr(r, "from_store", None) == "stage2")
                    _n3 = sum(1 for r in _engD.rules
                              if getattr(r, "from_store", None) == "stage3")
                    _ntun = int(getattr(_engD, "applied_mods", 0) or 0)
                    _uev = getattr(_engD, "use_evfis", "?")
                    _ugn = getattr(_engD, "use_genai", "?")
                    _mismatch = (_uev != _u12) or (_ugn != _u3)
                    st.caption(
                        f"engine now: use_evfis={_uev}, use_genai={_ugn}, "
                        f"evFIS_active={getattr(_engD, 'evfis_on', '?')}, "
                        f"GenAI_active={getattr(_engD, 'genai_on', '?')} · "
                        f"learned rules currently in the engine: {_nlrn}"
                        + (f" (stage ② {_n2}, stage ③ {_n3})" if _nlrn else "")
                        + f", tunings applied: {_ntun}"
                        + ("  ⚠ MISMATCH with the toggles — restart the "
                           "dashboard (stale engine in memory)"
                           if _mismatch else ""))
                    # spec: an unresolved dependency in the store is a
                    # visible warning, never a silent drop
                    _wD = list(getattr(_engD, "resolve_warnings", None)
                               or [])
                    if _wD:
                        # ONE BOX, ONE CAUSE. Five records skipped for the
                        # same reason is one finding, not five, so the record
                        # id is stripped and identical causes are counted.
                        import re as _re_wD
                        _gD = {}
                        for _w in _wD:
                            _k = _re_wD.sub(r"^\S+?_\d+:\s*", "", _w)
                            _gD[_k] = _gD.get(_k, 0) + 1
                        st.warning(
                            "Records that could not be brought in:\n"
                            + "\n".join(
                                f"- {_k}" + (f"  ({_c} records)"
                                             if _c > 1 else "")
                                for _k, _c in _gD.items()))
            if (st.session_state["dss_apply"]
                    and st.session_state.get("dss_res_base") is None):
                _rit_fix = st.session_state.get("dss_res_items")
                if _rit_fix:
                    # rows are staged but the raster is missing (a
                    # condition the bookkeeping missed): build it NOW
                    st.session_state["dss_res_base"] = \
                        _dss.build_resource_layer(world, _rit_fix)
                    st.session_state["dss_res_base_v"] = \
                        st.session_state.get("map_version")
                    st.caption("Resource pool rebuilt from the "
                               "staged rows.")
                else:
                    st.warning("No resource pool staged \u2014 go to "
                               "Layer 1 and press 'Suggest "
                               "resources'.")
            # (resetting / wiping learned adaptations lives on the Rules tab)
            st.divider()
            _regs4 = _dss.partition_n(cfg.nx, cfg.ny,
                                      int(_sv("dss_n", 1)))
            _names4 = [r.name for r in _regs4]
            _eng4r = st.session_state.get("dss_engine")

            # base + whatever stage 3 generated, from the store when no
            # engine is running, so the candidate view computes the same
            # concepts the DSS would
            _eng_hX = _all_hierarchy(_eng4r)

            def _cand_for(_regX):
                _fX = _dss.ten_features(sim, _regX, network=_obsnet)
                _gamX = _dss.concept_gates(
                    _dss.feature_confidence(_obsnet, _regX),
                    hierarchy=_eng_hX)
                _gX = st.session_state.get(f"l3_gate_{_regX.name}")
                if _gX is None:
                    _gX = _dss.GatedConcepts()
                    st.session_state[f"l3_gate_{_regX.name}"] = _gX
                _effX = _gX.gate(
                    _dss.infer_concepts(_fX, hierarchy=_eng_hX),
                    _gamX, step=int(sim.state.step))
                return _dss.evaluate_rules(
                    _effX, _fX,
                    _eng4r.rules if _eng4r is not None else None,
                    macros=getattr(_eng4r, "macros", None) or None)

            _glbT = getattr(st.session_state.get("dss_engine"),
                            "last_global", None)
            st.markdown("**Global DSS (coordinator)**")
            if _glbT:
                _rankG = _glbT.get("ranking", [])
                _sharesG = _glbT.get("shares", {})
                _thrG = _glbT.get("thresholds", {})
                _attG = set(_glbT.get("attended", []))
                _focusG = _glbT.get("hotspot") or (
                    _rankG[0][0] if _rankG else "-")
                st.caption(
                    "The coordinator ranks every region by operational "
                    "urgency and concentrates the response on the top one "
                    f"(**{_focusG}**). Attended regions get the full "
                    "offensive tempo and budget share; monitored regions "
                    "are only watched and face a tighter acceptance gate, so "
                    "an order there must prove higher quality to draw on the "
                    "shared resources. What it does to each region:")
                _hg = "| # | agent | priority | budget share | role | gate \u03b7 |"
                _sg = "|---|---|---|---|---|---|"
                _rowsG = [_hg, _sg]
                for _ig, (_ng, _pg) in enumerate(_rankG, 1):
                    _roleG = ("focus" if _ng == _focusG
                              else "attended" if _ng in _attG
                              else "monitored")
                    _shG = _sharesG.get(_ng)
                    _etG = _thrG.get(_ng)
                    _rowsG.append(
                        f"| {_ig} | {_ng} | {_pg:.2f} | "
                        + (f"{_shG:.2f}" if _shG is not None else "-")
                        + f" | {_roleG} | "
                        + (f"{_etG:.2f}" if _etG is not None else "-")
                        + " |")
                st.markdown("\n".join(_rowsG))
                st.caption("priority = operational urgency \u00b7 budget share = "
                           "fraction of the shared capacity and offensive "
                           "tempo the region gets \u00b7 role = focus / attended / "
                           "monitored \u00b7 gate \u03b7 = decision-quality bar to apply "
                           "an order there (higher = stricter). Logged every "
                           "cycle to global.csv and cycles.jsonl.")
            else:
                st.caption("No global decision yet \u2014 turn on 'DSS active' and "
                           "step. With a single agent the ranking is trivial "
                           "but still logged.")
            # the all-agents view is a dropdown choice (not a checkbox),
            # matching Layer 2 / Layer 3
            _names4d = _names4 + ["All agents (table)"]
            _i4 = min(int(_sv("dss_sel_i", len(_names4d) - 1)),
                      len(_names4d) - 1)
            _sel4 = st.selectbox("Agent (region)", _names4d,
                                 index=_i4, key="l4_agent")
            if _sel4 == "All agents (table)":
                _res4 = {r.name: _cand_for(r) for r in _regs4}
                _head4 = "| intervention |" + "".join(
                    f" {n} |" for n in _names4)
                _sep4 = "|---|" + "---|" * len(_names4)
                _rows4 = [_head4, _sep4]
                _macros4 = list(_all_macros(_eng4r))
                for _iv in _visible_interventions(_eng4r) + _macros4:
                    _cells4 = " | ".join(
                        f"{_res4[n][0].get(_iv, 0.0):.2f}" for n in _names4)
                    _isM4 = _iv in _macros4
                    _chip4 = ("<span style='color:"
                              f"{_IV_COLOR.get(_iv, '#c000ff')}'>"
                              "\u25a0</span> ")
                    _lbl4 = (_dss.INTERVENTION_LABEL.get(_iv)
                             or _iv.replace('_', ' ') + " (GenAI macro)")
                    _rows4.append(
                        f"| {_chip4}{_lbl4} | {_cells4} |")
                _rows4.append("| fired rules | " + " | ".join(
                    str(sum(1 for _r, _w in _res4[n][1] if _w > 0.01))
                    for n in _names4) + " |")
                st.markdown("\n".join(_rows4),
                            unsafe_allow_html=True)
                st.caption("Every Local DSS agent side by side: the "
                           "candidate order intensities its rule "
                           "base produces from its own region right "
                           "now, and how many rules fired. Pick a "
                           "single agent above for the detail with "
                           "the rule trace.")

            else:
                _reg4 = _regs4[_names4.index(_sel4)]
                _u4, _tr4 = _cand_for(_reg4)
                _mac4b = list(_all_macros(_eng4r))
                st.markdown("".join(
                    _iv_bar(_dss.INTERVENTION_LABEL.get(_iv)
                            or _iv.replace('_', ' ') + " (GenAI macro)",
                            _u4.get(_iv, 0.0),
                            _IV_COLOR.get(_iv, "#c000ff"))
                    for _iv in _visible_interventions(_eng4r)
                    + _mac4b),
                    unsafe_allow_html=True)
                _fired = [(r, w) for r, w in _tr4 if w > 0.01]
                with st.expander(f"Fired rules ({len(_fired)}) \u2014 "
                                 "traceability"):
                    if not _fired:
                        st.caption("No rule fires in this region "
                                   "right now.")
                    for _r, _w in sorted(_fired, key=lambda t: -t[1]):
                        st.caption(f"[{_w:.2f}] {_r.text()}")

        elif panel == "Layer 4 Settings":
            st.caption("Tuning coefficients for the Layer 4 decision loop: cadence, forecasting, the adaptation stages and the stage controller.")
            st.caption("Protection priority weights ($V_{prio}$) live on "
                       "the Cost panel, beside the loss weights: both "
                       "express the same operational value judgment.")
            # TICK-AWARE RECOMMENDATION: the decision cycle cannot be
            # finer than the simulation tick T (decisions land on step
            # boundaries), and the horizon is meaningful in TICKS, not
            # minutes: 1 tick per cycle and 15 ticks of lookahead
            # reproduce the approved (1, 30) pairing at T <= 2 and
            # scale with a coarser clock. When T changes, the two
            # fields are set to the recommendation; they stay freely
            # editable afterwards.
            _T_now = float(getattr(cfg, "step_minutes", 1.0))
            _cyc_rec = float(min(240.0, max(1.0, _T_now)))
            _hor_rec = float(min(480.0, max(30.0, 15.0 * _T_now)))
            if float(st.session_state.get("_dss_reco_T", -1.0)) \
                    != _T_now:
                st.session_state["_dss_reco_T"] = _T_now
                st.session_state["dss_cycle_min"] = _cyc_rec
                st.session_state["dss_horizon_min"] = _hor_rec
            st.caption(f"Tick T = {_T_now:g} min \u2192 recommended: "
                       f"decision cycle {_cyc_rec:g} min (1 tick), "
                       f"forecast horizon {_hor_rec:g} min "
                       f"(\u2248 {_hor_rec / _T_now:.0f} ticks). Set "
                       "on tick change; adjust freely below.")
            dc1, dc2, dc3 = st.columns(3)
            st.session_state["dss_cycle_min"] = float(dc1.number_input(
                "Decision cycle (min)", 1.0, 240.0,
                float(_sv("dss_cycle_min", 1.0)), 1.0,
                help="A decision is recomputed every this many SIM "
                     "minutes regardless of the step length; between "
                     "cycles the last allocation holds. NOTE: every "
                     "cycle boundary runs shadow forecasts (about "
                     "1-3 s wall time); with a very short cycle the "
                     "animation will stall on every frame. The 1 min "
                     "default reacts fastest; raise it if "
                     "Animate feels slow."))
            st.session_state["dss_horizon_min"] = float(dc2.number_input(
                "Forecast horizon (min)", 10.0, 480.0,
                float(_sv("dss_horizon_min", 30.0)), 5.0,
                help="Lookahead used to judge every candidate. With "
                     "the 1-min decision cycle the decision is "
                     "re-checked every minute but each check looks 30 "
                     "min ahead (adaptation trials and the no-harm "
                     "guard always re-check at >= 45 min). The "
                     "shadow run steps at a coarse 5-min tick (the "
                     "reference-step scaling keeps the physics "
                     "identical), so a 1-minute live tick still gets "
                     "a real horizon."))
            st.session_state["dss_min_gain"] = float(dc3.number_input(
                "Required forecast gain", 0.0, 0.5,
                float(_sv("dss_min_gain", 0.05)), 0.01,
                help="ADAPTIVE satisficing: a candidate must clear "
                     "$J_{TH}$ OR beat the no-action forecast by this "
                     "relative margin; otherwise the adaptation stages "
                     "engage. 0 restores the absolute-threshold-only "
                     "behaviour."))
            with st.expander("What these three do (and higher vs lower)"):
                st.markdown(
                    "**Decision cycle (min)** — how often the DSS recomputes "
                    "a decision. Between cycles the last allocation holds.\n"
                    "- Higher = reacts more slowly, but far cheaper "
                    "(forecasts run less often, animation is smoother).\n"
                    "- Lower = reacts fast, but every cycle runs the shadow "
                    "forecasts (heavy; a very short cycle stalls animation).\n\n"
                    "**Forecast horizon (min)** — how far ahead the DSS "
                    "simulates each candidate to judge it.\n"
                    "- Higher = more foresight (sees a far-off asset threat), "
                    "but more compute per cycle.\n"
                    "- Lower = short-sighted, cheaper; may miss slow threats.\n\n"
                    "**Required forecast gain** — the minimum improvement a "
                    "candidate must show over doing nothing to be accepted.\n"
                    "- Higher = conservative: only clearly-worth-it changes "
                    "pass, so fewer adaptations and less resource commitment.\n"
                    "- Lower (toward 0) = eager: small gains are accepted, so "
                    "more adaptation and more orders, at higher response cost.")
            # (the full seed-rule catalog lives on the Rules tab)
            st.markdown("**evFIS \u2014 evolving fuzzy rule base**")
            st.session_state["dss_jth"] = float(st.slider(
                "$J_{TH}$ \u2014 acceptable cost threshold", 0.0, 1.0,
                float(_sv("dss_jth", 0.35)), 0.01,
                help="A candidate decision is applied as-is when its "
                     "cost satisfies $J \\le J_{TH}$; otherwise the "
                     "adaptation stages engage."))
            st.session_state["dss_eta"] = float(st.slider(
                "$\\eta$ \u2014 decision quality gate", 0.0, 1.0,
                float(_sv("dss_eta", 0.60)), 0.01,
                help="Adapted rules apply only while the decision "
                     "quality keeps $Q \\ge \\eta$; below it the "
                     "graduated fail-safe attenuates toward the safe "
                     "baseline."))
            st.session_state["dss_evfis_step"] = float(st.slider(
                "Membership adaptation step", 0.0, 0.2,
                float(_sv("dss_evfis_step", 0.05)), 0.01,
                help="How far stage \u2460 may move the membership "
                     "parameters (a,b,c,d) and the consequents per "
                     "adaptation."))
            st.markdown("**GenAI \u2014 rule proposer (stage \u2462)**")
            # NOTE: the single on/off for stage 3 is the "GenAI proposals
            # (stage 3)" toggle above; it sets dss_genai_on. Everything here
            # is disabled when that toggle is OFF.
            _genai_on = bool(st.session_state.get("dss_genai_on", True))
            st.session_state["dss_genai_temp"] = float(st.slider(
                "Proposal temperature", 0.0, 1.0,
                float(_sv("dss_genai_temp", 0.3)), 0.05,
                disabled=not _genai_on))
            # model passed to `claude --model`; stage 3 runs ONLY on your
            # Claude subscription via Claude Code (`claude -p`), no API key.
            # Use Claude Code's own aliases (opus / sonnet / haiku): full
            # version strings are not always recognised and `claude` then
            # silently falls back to its default model.
            _MODELS = ["(plan default)", "opus", "sonnet", "haiku"]
            # SPEED MATTERS HERE, not only quality: stage 3 runs inside the
            # decision cycle, which runs inside the animation frame, so the
            # model's latency is a pause the operator sees on every call.
            # Measured on THIS transport, not the usual family ordering. The
            # wait is set by how much the model writes, and opus and sonnet
            # obey "minified JSON only" while haiku keeps writing prose
            # around it, which makes the nominally fastest model the slowest
            # one here.
            _MSPEED = {
                "(plan default)": "whatever your plan serves",
                "opus": "fastest here — answers in ~22 tokens",
                "sonnet": "fast here — answers in ~16 tokens",
                "haiku": "slowest here — ignores the brevity rule, ~290 tok",
            }
            _mtimes = dict(st.session_state.get("genai_model_ms", {}))

            def _mlabel(m):
                _t = _mtimes.get(m)
                return (f"{m} — {_MSPEED[m]}"
                        + (f" · measured {_t / 1000:.1f} s" if _t else ""))
            _cur_m = st.session_state.get("genai_model_ui", "(plan default)")
            if _cur_m not in _MODELS:
                _cur_m = "(plan default)"
            _mirror("genai_model_ui", "l4_genai_model", "(plan default)")
            _mdl = st.selectbox(
                "Model", _MODELS, key="l4_genai_model",
                on_change=_adopt, args=("genai_model_ui", "l4_genai_model"),
                format_func=_mlabel,
                disabled=not _genai_on,
                help="Which model `claude` uses for stage 3 (Claude Code "
                     "aliases). '(plan default)' lets your Pro/Max plan "
                     "decide. Note: your plan may downgrade `claude -p` to a "
                     "smaller model when limits are hit.")
            st.caption(
                "Measured on this transport the call costs about "
                "**3.4 s of fixed startup plus ~7 ms per output token**, so "
                "the wait is set by how MUCH the model writes rather than by "
                "which model writes it. That inverts the usual ordering "
                "here: opus and sonnet honour the 'minified JSON only' "
                "instruction and come back in ~3.3 s and ~3.7 s, while "
                "haiku keeps writing prose around the JSON (~290 tokens "
                "that are discarded unread) and takes ~5.3 s. The smallest "
                "model is currently the slowest one for this job."
                + f" Each stage ③ call blocks the simulation until it "
                  f"answers, capped at {_dss.genai_timeout():.0f} s "
                  "(`DSS_GENAI_TIMEOUT` raises it for batch runs)."
                + ("" if _mtimes else
                   " Press 'Test Claude connection' to record what each "
                   "model actually costs on this machine."))
            if _mdl and _mdl != "(plan default)":
                os.environ["DSS_GENAI_MODEL"] = _mdl.strip()
            else:
                os.environ.pop("DSS_GENAI_MODEL", None)
            # --- GenAI status: stage 3 runs only on the Claude subscription ---
            _gcfg = _dss.genai_config()
            if _gcfg.get("mode") == "cli":
                st.success(
                    "GenAI link: **Claude Code (your subscription)** \u2014 "
                    "stage \u2462 runs through the local `claude` command on "
                    "your Pro/Max plan. No API key.")
            else:
                st.warning(
                    "GenAI link: **inactive** \u2014 stage \u2462 is skipped "
                    "(stages \u2460\u2461 still run). Install Claude Code, then in "
                    "a terminal run `claude` and type `/login` on your Pro/Max "
                    "plan. Restart the dashboard afterwards.")
            if st.button("Test Claude connection",
                         help="Sends ONE real, live request over the "
                              "configured transport (API key or your "
                              "subscription via Claude Code) and shows "
                              "Claude's own reply \u2014 proof the generative "
                              "stage is wired to Claude, not mocked."):
                with st.spinner("Calling Claude\u2026"):
                    _pr = _dss.genai_probe()
                if _pr["ok"]:
                    # remember what THIS machine measured for THIS model, so
                    # the selector shows real latency instead of only the
                    # documented ordering
                    _mm = dict(st.session_state.get("genai_model_ms", {}))
                    _mm[_mdl] = int(_pr.get("latency_ms") or 0)
                    st.session_state["genai_model_ms"] = _mm
                    # remember what the plan actually SERVED, so the sidebar
                    # can name the model even when the request said
                    # "(plan default)"
                    st.session_state["genai_served_model"] = str(
                        _pr.get("reported_model") or "")
                    _u = _pr.get("usage") or {}
                    _req_m = _pr.get("model") or "(subscription default)"
                    _act_m = _pr.get("reported_model") or "?"
                    _ot = _u.get("output_tokens")
                    st.success(
                        f"Live reply from Claude at "
                        f"**{_pr.get('finished_at', '?')}** in "
                        f"{_pr['latency_ms']} ms "
                        f"(sent {_pr.get('started_at', '?')}, requested "
                        f"`{_req_m}`, served `{_act_m}`, tokens in/out "
                        f"{_u.get('input_tokens', '?')}/"
                        f"{_ot if _ot is not None else '?'}):")
                    if isinstance(_ot, int):
                        # the length of the answer is the part of the wait
                        # this project controls, so it is called out
                        st.caption(
                            f"About {3400 + 7 * _ot:.0f} ms of this is "
                            "explained by the fixed startup plus the answer "
                            f"length ({_ot} output tokens). "
                            + ("The model kept it short, which is what the "
                               "schema asks for." if _ot < 60 else
                               "The model wrote well past the JSON it was "
                               "asked for; every one of those tokens is "
                               "discarded unread and still costs about "
                               "7 ms."))
                    st.code(_pr["reply"] or "(empty)", language=None)
                    # a single `claude -p` run can bill several models: the
                    # one that answered plus a small helper doing internal
                    # bookkeeping. Show the breakdown so "served X" is
                    # verifiable instead of having to be trusted.
                    _mus = _pr.get("model_usage") or {}
                    if len(_mus) > 1:
                        st.caption(
                            "Models billed by this one call: "
                            + ", ".join(
                                f"`{k}` {v.get('out_tokens', 0)} tok / "
                                f"${v.get('cost', 0.0):.4f}"
                                for k, v in sorted(
                                    _mus.items(),
                                    key=lambda kv: -kv[1].get("cost", 0.0)))
                            + " — the one matching the call's usage block "
                              "served the request; the rest are Claude "
                              "Code's internal helpers, which can emit MORE "
                              "tokens than the answer itself.")
                    # the plan can serve a smaller model than requested
                    if (_req_m not in ("(subscription default)", "", None)
                            and _act_m
                            and _req_m.split("-")[0].lower()
                            not in _act_m.lower()):
                        st.warning(
                            f"You asked for `{_req_m}` but the Claude Code "
                            f"subscription served `{_act_m}`. With `claude -p` "
                            "on a Pro/Max plan the `--model` choice is a "
                            "preference, not a guarantee: the plan downgrades "
                            "to a smaller model when the higher tier is "
                            "rate-limited or is not part of your plan. To "
                            "force a model, run `claude` interactively and set "
                            "it there, or use a plan that includes it.")
                    with st.expander("Raw `claude` JSON \u2014 what the "
                                     "CLI actually reported"):
                        st.caption("Command: `"
                                   + str(_pr.get("cmd") or "claude ...")
                                   + "`")
                        st.code(_pr.get("raw") or "(none)",
                                language="json")
                    st.caption(
                        "This text was generated by Claude just now over "
                        f"{_pr['endpoint']} \u2014 the same path stage "
                        "\u2462 uses to propose rules.")
                else:
                    st.error(f"No live link: {_pr['error']}")
            st.markdown("**Stage controller \u2014 associative search**")
            st.session_state["dss_ctrl_eps"] = float(st.slider(
                "$\\epsilon$ \u2014 exploration", 0.0, 1.0,
                float(_sv("dss_ctrl_eps", 0.10)), 0.01,
                help="Probability of trying a non-greedy adaptation "
                     "stage; reward = realized cost reduction."))
            st.session_state["dss_ctrl_lr"] = float(st.slider(
                "Learning rate", 0.001, 0.5,
                float(_sv("dss_ctrl_lr", 0.05)), 0.001, format="%.3f"))
            with st.expander("How the learning works \u2014 controller \u00b7 "
                             "evFIS \u00b7 GenAI, and what persists"):
                st.markdown(
                    "**Stage controller (associative search)** is an "
                    "$\\epsilon$-greedy contextual bandit. State = "
                    "the cost-deficit bucket (low / mid / high), "
                    "actions = the enabled adaptation stages, reward "
                    "= the realized forecast improvement of the "
                    "chosen stage. It is NOT pre-trained: it learns "
                    "online during the run, every untried "
                    "(bucket, stage) pair is tried once before "
                    "greedy takes over, and the value table is written "
                    "to the run log (cycles.jsonl) every cycle.\n\n"
                    "**evFIS (stages \u2460\u2461)** modifies the "
                    "ENGINE's own copy of the rule base: consequent "
                    "nudges and membership-shoulder moves that "
                    "survive for the lifetime of this engine. The "
                    "thesis seed catalog in dss/rules.py is never "
                    "touched.\n\n"
                    "**GenAI (stage \u2462)** proposes a brand-new "
                    "rule from the live concept situation \u2014 "
                    "the proposal comes from Claude through Claude "
                    "Code on your subscription (see the status line "
                    "above); a reachable model is required, so "
                    "without `claude` logged in the stage is skipped "
                    "and logged, never faked. Admitted rules (gates "
                    "G1-G4) join the runtime base with a G prefix.\n\n"
                    "**Lifecycle**: everything learned lives in the "
                    "running engine and its logs. A rebuilt engine "
                    "(changed settings, new map) starts from the "
                    "thesis seed state again. The button below "
                    "resets by hand.")
        elif panel == "Layer 4 Rules":
            st.markdown("**Rule catalog \u2014 thesis seeds + "
                        "everything this run has learned**")
            _profs = {
                "40 rules (full seed set)": "full",
                "core \u2014 doctrine R1-R22 only": "core",
                "minimal \u2014 5 strongest seeds (one per "
                "intervention family), LEARN the rest by trial":
                "minimal"}
            # VIEW-ONLY on this tab: picking a profile here just shows that
            # profile's seeds for inspection; it does NOT change the running
            # DSS. The operational seed base is set under Layer 4 Decision.
            # this VIEW opens on minimal: the 5-rule base is the one worth
            # inspecting rule by rule, and it is where the generated records
            # are easiest to tell apart from the seeds. The operational
            # profile under Layer 4 Decision is untouched by this.
            _pcur_r = str(_sv("rules_view_profile", "minimal"))
            _pidx = [i for i, v in enumerate(_profs.values())
                     if v == _pcur_r]
            _psel = st.selectbox(
                "View seed profile (display only)", list(_profs),
                index=(_pidx[0] if _pidx else 0),
                help="Shows that profile's seed rules here for inspection "
                     "only. It does NOT change the running DSS: the "
                     "operational seed base is chosen under Layer 4 "
                     "Decision.")
            st.session_state["rules_view_profile"] = _profs[_psel]
            st.caption("Display only \u2014 the DSS runs the seed base set under "
                       "Layer 4 Decision "
                       f"(**{_sv('dss_seed_profile', 'minimal')}**).")
            import os as _os_st
            _shared_store = _shared_store_path()
            # the store the DSS ACTUALLY uses (a loaded run's rules.json wins),
            # so display + reset/wipe operate on the SAME file the engine reads
            _eng_store = st.session_state.get("dss_engine")
            _store_p = (getattr(_eng_store, "state_path", None)
                        or st.session_state.get("dss_learned_override")
                        or _shared_store)
            _gcnt = _gstate_counts(_store_p)
            st.caption(f"Active store the DSS reads: `{_os_st.path.basename(_os_st.path.dirname(_store_p))}/"
                       f"{_os_st.path.basename(_store_p)}`"
                       + (" (a loaded run's snapshot)"
                          if _store_p != _shared_store else " (shared)"))
            # the counters come from the SAME file the engine resolves from,
            # so a panel can no longer report an empty store while the engine
            # is reasoning with records
            _gd = _read_gstate(_store_p)
            _gflags = _gd.get("runtime_flags") or {}
            st.caption(
                "Generated records · evFIS modifications: "
                f"**{_gcnt['evfis_rule_modifications']}** · GenAI rules: "
                f"**{_gcnt['genai_rules']}** · concepts: "
                f"**{_gcnt['genai_concepts']}** · interventions: "
                f"**{_gcnt['genai_interventions']}**"
                + (f" · config `{_dss.config_id(_gflags)}`"
                   if _gflags else ""))
            _warns = list(getattr(_eng_store, "resolve_warnings", []) or [])
            if _warns:
                # collapse repeats: five records tuning the same missing rule
                # is ONE finding, printed five times only adds noise
                import re as _re_w
                _grp = {}
                for _w in _warns:
                    _k = _re_w.sub(r"^\S+?_\d+:\s*", "", _w)
                    _grp[_k] = _grp.get(_k, 0) + 1
                st.warning(
                    "Records that could not be brought in:\n"
                    + "\n".join(
                        f"- {'' if n == 1 else f'{n}× '}{k}"
                        for k, n in sorted(_grp.items(),
                                           key=lambda kv: -kv[1])[:6]))
            _keepN = int(st.number_input(
                "keep N strongest learned rules", 0, 500, 10, 1,
                key="rules_keepn",
                help="'Reset (keep the strongest)' prunes the learned store "
                     "to this many rules, ranked by strength (accumulated "
                     "fired weight). If you have fewer learned rules than N, "
                     "nothing is dropped, so LOWER this to actually prune the "
                     "weak / never-fired (strength 0) rules."))
            _rb1, _rb2 = st.columns(2)
            if _rb1.button("Reset (keep the strongest)",
                           key="rules_reset",
                           help="Keeps the strongest N learned rules (and the "
                                "membership tunes), drops the rest, clears the "
                                "controller value table, and prunes the "
                                "persistent store. Natural-selection reset."):
                from dss.adapt import reset_partitions as _rspR
                _rspR()
                # prune the GENERATED store: keep the N strongest rule-adding
                # records and every parameter modification, drop the rest.
                # A modification has no strength of its own, so only the
                # records that introduced a rule are ranked.
                _gs = _dss.GeneratedState.load(_store_p)
                _mods = _gs.records("evfis_rule_modifications")
                _adds = [r for r in _mods
                         if r.get("modification_type") == "rule_add"]
                _rest = [r for r in _mods
                         if r.get("modification_type") != "rule_add"]
                _adds.sort(key=lambda r: -float(
                    ((r.get("after") or {}).get("rule") or {})
                    .get("strength", 0.0)))
                _keep = _adds[:max(0, _keepN)]
                _gs.data["evfis_rule_modifications"] = sorted(
                    _rest + _keep, key=lambda r: int(r.get("seq", 0)))
                _grules = _gs.records("genai_rules")
                _grules.sort(key=lambda r: -float(r.get("strength", 0.0)))
                _gs.data["genai_rules"] = sorted(
                    _grules[:max(0, _keepN)],
                    key=lambda r: int(r.get("seq", 0)))
                _gs.save()
                _nkeep = len(_keep) + len(_gs.data["genai_rules"])
                # drop the engine so it REBUILDS from the pruned store with the
                # CURRENT use-stage toggles (a manual re-merge here ignored the
                # toggles and reloaded everything)
                _reset_dss_state(drop_engine=True)
                st.toast(f"Reset: kept the {_nkeep} strongest learned "
                         "rules, dropped the rest.")
                st.rerun()
            if _rb2.button("Wipe learned store",
                           key="rules_wipe",
                           help="Clears every generated record (evFIS "
                                "modifications, GenAI rules, concepts and "
                                "interventions) and returns to the pure "
                                "seed profile. A backup is kept beside the "
                                "store. The baseline rule sets and the six "
                                "base interventions are never deleted, only "
                                "returned to factory value."):
                from dss.adapt import reset_partitions as _rspW
                _rspW()
                try:
                    _wc = _dss.GeneratedState.load(_store_p).wipe()
                except Exception as _e_w:
                    _wc = {}
                    st.error(f"The store could not be wiped: {_e_w}")
                # spec: a wipe turns the PRODUCTION toggles off so the
                # clean state is not re-dirtied within seconds; the
                # use-stage (consumption) toggles and DSS active stay
                st.session_state["dss_evfis_on"] = False
                st.session_state["dss_genai_on"] = False
                st.session_state["dss_adapt_on"] = False
                # rebuild the engine from the (now empty) store
                _reset_dss_state(drop_engine=True)
                st.toast("Wiped "
                         + ", ".join(f"{v} {k.replace('_', ' ')}"
                                     for k, v in (_wc or {}).items() if v)
                         + " (backup kept); production toggles off, "
                           "pure seed profile."
                         if _wc else "Nothing to wipe; the store was empty.")
                st.rerun()
            _engR = st.session_state.get("dss_engine")
            import os as _os_rs
            _store_v = _store_p   # the SAME store the DSS actually reads
            if _engR is not None:
                _rlist = _engR.rules
            else:
                # no engine (fresh map / apply off): the DISPLAY must
                # still show the persistent store merged into the seed
                # profile, respecting the SAME use-stage toggles the engine
                # would, otherwise the learned rules LOOK wrong
                from dss.adapt import reset_partitions as _rspD
                from dss.fuzzy import REGISTRY as _REGD
                _gsD = _dss.GeneratedState.load(
                    _store_v, active_rule_set=str(
                        _sv("dss_seed_profile", "minimal")))
                _gsD.set_flags(
                    active_rule_set=str(_sv("dss_seed_profile", "minimal")),
                    use_stage12_rules=bool(_sv("dss_use_stage12", True)),
                    use_stage3_rules=bool(_sv("dss_use_stage3", True)))
                _rlist = _dss.resolve_active_set(
                    _gsD, _dss.make_runtime_rules, _REGD, _rspD).rules
                st.caption("No engine running yet \u2014 showing "
                           "the seed profile MERGED with the "
                           "persistent learned store. Enable "
                           "'Apply decisions' and step: the engine "
                           "starts from exactly this base.")
            _bornN, _tunedN = _dss.load_learned(
                _store_v,
                profile=str(_sv("dss_seed_profile", "minimal")))
            if _os_rs.path.exists(_store_v):
                import time as _t_rs
                _mt = _t_rs.strftime(
                    "%H:%M:%S", _t_rs.localtime(
                        _os_rs.path.getmtime(_store_v)))
                st.caption(f"Persistent store: `logs/learned_rules"
                           f".json` \u00b7 lineage of THIS profile: "
                           f"{len(_bornN)} born rules, "
                           f"{len(_tunedN)} tuned seeds \u00b7 last "
                           f"saved {_mt} \u00b7 saved EVERY decision "
                           "cycle; survives fires, engines and maps. "
                           "Each seed profile keeps its OWN lineage: "
                           "the selected profile's rules are the "
                           "base, evFIS/GenAI grow on top of it, and "
                           "nothing leaks between profiles.")
            else:
                st.caption("Persistent store: not created yet "
                           "(first accepted adaptation or decision "
                           "cycle writes logs/dss_generated_state.json).")

            # PROVENANCE BY VALUE, NOT BY WORDING. R41 and R42 ship in the
            # doctrine catalog as worked examples of what an adapted rule
            # looks like, so their notes say "evFIS" and "GenAI" even though
            # they are pristine baseline. Classifying on the note text made
            # them read as learned modifications that had survived a wipe.
            _pristineO = {_r.name: tuple((str(i), round(float(v), 4))
                                         for i, v in _r.consequents)
                          for _r in _dss.make_runtime_rules(
                              str(_sv("dss_seed_profile", "minimal")))}

            def _origin_of(_r):
                _p = _pristineO.get(_r.name)
                _cur = tuple((str(i), round(float(v), 4))
                             for i, v in _r.consequents)
                if _p is not None and _cur == _p:
                    # in the catalog and unchanged: baseline, whatever the
                    # note happens to describe
                    return 0, "\U0001F7E6 seed"
                if _r.name.startswith("G"):
                    return 3, "\U0001F7E9 GenAI (stage \u2462)"
                if _r.name.startswith("A"):
                    return 2, ("\U0001F7E7 resolution "
                               "(stage \u2461)")
                if "evFIS" in (_r.note or ""):
                    return 1, ("\U0001F7E8 seed, evFIS-tuned "
                               "(stage \u2460)")
                return 0, "\U0001F7E6 seed"

            # ---- SEED table: the never-deleted doctrine base for the chosen
            # profile (40 / core 22 / minimal). Shown as defined; adaptation
            # never edits this table.
            _seedsPure = _dss.make_runtime_rules(
                str(_sv("rules_view_profile",
                        _sv("dss_seed_profile", "minimal"))))
            st.markdown(f"**Seed rules (never deleted) — "
                        f"{len(_seedsPure)} in this profile (view)**")
            _tblR = [dict(
                name=_r.name,
                IF=" AND ".join(f"{v} is {t}"
                                for v, t in _r.antecedents),
                THEN=", ".join(f"{iv} {x:.2f}"
                               for iv, x in _r.consequents),
                active="yes" if _r.active else "no")
                for _r in _seedsPure]
            st.dataframe(_tblR, use_container_width=True, height=360)

            # a rule counts as LEARNED only if it differs from the pristine
            # seed baseline: a born A#/G# rule, or a seed R# whose consequents
            # were retuned at runtime. Seed-defined provenance examples
            # (R41/R42, whose notes mention evFIS/GenAI but which ship unchanged
            # in the catalog) stay in the seed table, not here.
            # classify against the OPERATIONAL profile (what the engine runs),
            # not the view profile, so the learned table is always correct
            _seedsOp = _dss.make_runtime_rules(
                str(_sv("dss_seed_profile", "minimal")))
            _seed_sig = {r.name: tuple((str(i), round(float(v), 4))
                                       for i, v in r.consequents)
                         for r in _seedsOp}

            def _is_learned(r):
                base = _seed_sig.get(r.name)
                if base is None:
                    return True
                return tuple((str(i), round(float(v), 4))
                             for i, v in r.consequents) != base

            # ---- MODIFIED / GENERATED table: stage ① tuned seeds, stage ②
            # (A#) and stage ③ GenAI (G#). Kept SEPARATE from the seeds. They
            # persist in the store and the DSS applies them like seeds even in
            # Fuzzy-only mode (evFIS / GenAI off), until wiped.
            # ---- VIEWER, NOT A MIRROR OF THE ACTIVE SET ----
            # This table used to be derived from the rules the engine is
            # currently reasoning with, so turning a consumption toggle off
            # emptied it and the accumulated knowledge looked lost. It reads
            # the STORE instead: everything ever produced is listed, and the
            # "in use" column says whether the current toggles admit it.
            _gdR = _read_gstate(_store_p)
            _active_names = {r.name for r in _rlist}
            _recs = []

            def _fmt_cons(_cs):
                return ", ".join(f"{i} {float(v):.2f}" for i, v in _cs)

            def _fmt_rule(_sp):
                return ("IF " + " AND ".join(f"{v} is {t}" for v, t
                                             in _sp.get("antecedents", []))
                        + " THEN " + _fmt_cons(_sp.get("consequents", [])))
            for _m in sorted(_gdR.get("evfis_rule_modifications") or [],
                             key=lambda r: int(r.get("seq", 0))):
                _mt = _m.get("modification_type")
                _aft = _m.get("after") or {}
                _bef = _m.get("before") or {}
                _stp = (_m.get("trigger") or {}).get("step")
                if _mt == "rule_add":
                    # stage 2 CREATES rules: it instantiates the antecedent
                    # cell of a situation no active rule answered
                    _sp = _aft.get("rule") or {}
                    _nm = _sp.get("name", _m.get("id"))
                    _recs.append(dict(
                        stage="2 — resolution", rule=_nm,
                        description=(
                            "NEW RULE. No active rule covered the situation, "
                            "so the cell was instantiated: " + _fmt_rule(_sp)
                            + ". The consequents are seeded from what the "
                            "dominant concepts demanded."),
                        strength=round(float(_sp.get("strength", 0.0)), 1),
                        step=_stp,
                        in_use=("yes" if _nm in _active_names
                                else "no — stage ①② consumption off"),
                        seq=int(_m.get("seq", 0))))
                elif _mt == "consequent_update":
                    _ob = dict((str(i), float(v))
                               for i, v in _bef.get("consequents", []))
                    _an = dict((str(i), float(v))
                               for i, v in _aft.get("consequents", []))
                    _delta = ", ".join(
                        f"{k} {_ob.get(k, 0.0):.2f} → {v:.2f}"
                        for k, v in _an.items()
                        if abs(v - _ob.get(k, 0.0)) > 1e-9) or "no net change"
                    _recs.append(dict(
                        stage="1 — evFIS tuning", rule=_m.get("base_rule_id",
                                                              "?"),
                        description=(
                            "MODIFICATION of an existing rule, not a new one. "
                            "Ordered intensities retuned: " + _delta
                            + ". Kept because the physical forecast improved."),
                        strength=0.0, step=_stp,
                        in_use=("yes" if bool(_sv("dss_use_stage12", True))
                                else "no — stage ①② consumption off"),
                        seq=int(_m.get("seq", 0))))
                else:
                    _vv = _m.get("variable", "?")
                    _new = sorted(set(_aft.get("partition") or {})
                                  - set(_bef.get("partition") or {}))
                    if _new:
                        _recs.append(dict(
                            stage="2 — resolution", rule=_vv,
                            description=(
                                "NEW TERM " + ", ".join(_new)
                                + f" inserted into the partition of {_vv}. "
                                "No existing label described the situation "
                                "well enough, so the variable gained "
                                "resolution and the antecedent catalog grew "
                                "with it."),
                            strength=0.0, step=_stp,
                            in_use=("yes" if bool(_sv("dss_use_stage12", True))
                                    else "no — stage ①② consumption off"),
                            seq=int(_m.get("seq", 0))))
                    else:
                        _recs.append(dict(
                            stage="1 — evFIS tuning", rule=_vv,
                            description=(
                                "MODIFICATION of the membership functions of "
                                f"{_vv}: the shared boundary between two "
                                "neighbouring terms was moved. Both move "
                                "together, so the partition still sums to one "
                                "everywhere."),
                            strength=0.0, step=_stp,
                            in_use=("yes" if bool(_sv("dss_use_stage12", True))
                                    else "no — stage ①② consumption off"),
                            seq=int(_m.get("seq", 0))))
            for _g in sorted(_gdR.get("genai_rules") or [],
                             key=lambda r: int(r.get("seq", 0))):
                _nm = _g.get("name", _g.get("id"))
                _ants = " AND ".join(
                    (f"{a[0]} is {a[1]}" if not isinstance(a, dict)
                     else f"{a.get('concept')} is {a.get('term')}")
                    for a in _g.get("antecedents", _g.get("antecedent", [])))
                _cons = ", ".join(
                    (f"{c[0]} {float(c[1]):.2f}" if not isinstance(c, dict)
                     else f"{c.get('channel')} {float(c.get('value', 0)):.2f}")
                    for c in _g.get("consequents", _g.get("consequent", [])))
                _dep = _g.get("depends_on_concepts") or []
                _recs.append(dict(
                    stage="3 — GenAI", rule=_nm,
                    description=(
                        "NEW RULE proposed by Claude and admitted through the "
                        f"gates: IF {_ants} THEN {_cons}."
                        + (" Depends on the generated concept(s) "
                           + ", ".join(_dep) + "." if _dep else "")),
                    strength=round(float(_g.get("strength", 0.0)), 1),
                    step=(_g.get("trigger") or {}).get("step"),
                    in_use=("yes" if _nm in _active_names
                            else "no — stage ③ consumption off"),
                    seq=int(_g.get("seq", 0))))
            st.markdown("**Generated knowledge in the store (stage ①②③) — "
                        f"{len(_recs)} record(s)**")
            st.caption(
                "Everything the adaptation stages have ever produced, "
                "whatever the toggles currently say. The store is the "
                "record; the toggles only decide what the DSS is allowed to "
                "reason with right now, which is what the 'in use' column "
                "reports. Records survive until 'Wipe learned store'.")
            _perr = list(getattr(_engR, "persist_errors", []) or [])
            if _perr:
                st.error(
                    f"{len(_perr)} accepted adaptation(s) could NOT be "
                    "written to the store, so they are missing from this "
                    "table even though the Analysis tab counted them:\n"
                    + "\n".join(f"- {e}" for e in _perr[:5]))
            if _recs:
                _cols_g = ["stage", "rule", "description", "step",
                           "in_use", "strength", "seq"]
                st.dataframe([{k: r.get(k) for k in _cols_g} for r in _recs],
                             use_container_width=True, height=280,
                             column_order=_cols_g)
                st.caption(
                    "Which stage does what: **stage ①** only MODIFIES what "
                    "already exists (ordered intensities, membership "
                    "boundaries). **Stage ②** creates rules and inserts terms "
                    "when the base is silent rather than wrong. **Stage ③** "
                    "creates rules and may also propose a new concept or a "
                    "composite intervention. So a new rule comes from ② or "
                    "③, never from ①.")
                st.caption(
                    "step = the simulation step the record was produced at. "
                    "seq = the global replay order: a wipe reverts records in "
                    "reverse seq, a restart replays them forward, which is "
                    "why it is a single sequence across all record types "
                    "rather than a per-type counter. strength = accumulated "
                    "fired weight of the decisions a rule took part in; "
                    "'Reset (keep the strongest)' ranks the rule-creating "
                    "records by it, so it stays blank for modifications.")
            else:
                st.caption("The store holds no generated record yet.")
            _learned = [_r for _r in _rlist if _is_learned(_r)]

            st.markdown("**Membership modifications (evFIS stage "
                        "\u2460 + inserted terms, stage \u2461) "
 "\u2014 registry vs**")
            from dss.fuzzy import default_partition as _defpR
            from dss.fuzzy import REGISTRY as _REGR
            _dpR = _defpR()
            _modsR = []
            for _var in _REGR.variables():
                for _term, _abcd in _REGR.get(_var).items():
                    _d0 = _dpR.get(_term)
                    if _d0 is None:
                        _modsR.append(dict(
                            variable=_var, term=_term,
                            default="(INSERTED \u2014 catalog "
                                    "grew)",
                            current=str(tuple(round(float(v), 3)
                                              for v in _abcd))))
                        continue
                    if tuple(np.round(np.asarray(_abcd, float), 4)) \
                            != tuple(np.round(np.asarray(_d0, float),
                                              4)):
                        _modsR.append(dict(
                            variable=_var, term=_term,
                            default=str(tuple(round(float(v), 3)
                                              for v in _d0)),
                            current=str(tuple(round(float(v), 3)
                                              for v in _abcd))))
            # the registry above is the ACTIVE partition, so with stage \u2460\u2461
            # consumption off it is back at factory value and shows nothing.
            # The stored record is what the viewer must report.
            _pmods = [_m for _m in (_gdR.get("evfis_rule_modifications") or [])
                      if _m.get("modification_type") in ("membership_shift",
                                                         "term_insert")]
            if _modsR:
                st.dataframe(_modsR, use_container_width=True)
                st.caption("Currently applied to the registry.")
            elif _pmods:
                st.caption(
                    f"{len(_pmods)} membership record(s) are in the store but "
                    "not applied right now, because \u201cUse stage \u2460\u2461 rules\u201d is "
                    "off. They are listed in the table above and return the "
                    "moment the toggle goes back on.")
            else:
                st.caption("Every membership still sits on its "
                           "default \u2014 stage \u2460 has not moved anything (yet).")

            # ---- INTERVENTION -> RULES: pick an intervention, see every rule
            # (seed + learned) that ORDERS it ----
            st.markdown("**Interventions — which rules order each**")

            def _ivlbl(iv):
                return (_dss.INTERVENTION_LABEL.get(iv)
                        or iv.replace("_", " ") + " (GenAI macro)")
            _ivall = (list(_dss.INTERVENTIONS)
                      + list(_all_macros(_engR)))
            _ivsel = st.selectbox("Intervention", _ivall,
                                  format_func=_ivlbl, key="rules_iv_sel")
            _iv_rules = sorted(
                ((_r, dict(_r.consequents).get(_ivsel, 0.0))
                 for _r in _rlist
                 if any(str(_i) == _ivsel for _i, _x in _r.consequents)),
                key=lambda t: -(t[1] or 0.0))
            # an INACTIVE catalog entry (R41/R42 are shipped as provenance
            # examples with active=False) cannot order anything, so it is
            # separated out instead of being listed as if it fires
            _iv_live = [(r, y) for r, y in _iv_rules if getattr(r, "active",
                                                               True)]
            _iv_off = [(r, y) for r, y in _iv_rules
                       if not getattr(r, "active", True)]
            if _iv_live:
                st.caption(f"{len(_iv_live)} rule(s) order "
                           f"**{_ivlbl(_ivsel)}** (strongest first):")
                for _r, _iy in _iv_live:
                    _ifs = " AND ".join(f"{v} is {t}"
                                        for v, t in _r.antecedents)
                    _orig = ("🟦 seed" if _origin_of(_r)[0] == 0
                             else _origin_of(_r)[1])
                    st.caption(f"• **{_r.name}** [{_orig}] · intensity "
                               f"{_iy:.2f} — IF {_ifs}")
            else:
                st.caption(f"No active rule orders {_ivlbl(_ivsel)} right now.")
            if _iv_off:
                st.caption("Inactive catalog entries naming it (they never "
                           "fire): "
                           + ", ".join(f"**{r.name}**" for r, _ in _iv_off))

            _nnew = sum(1 for _r in _rlist if _r.name[0] in "AG")
            _ntun = sum(1 for _r in _rlist if "evFIS" in (_r.note or ""))
            _fullR = _dss.make_runtime_rules("full")
            _cellsN = [set(_r.antecedents) for _r in _rlist
                       if _r.active]
            _hitR = sum(1 for _fr in _fullR if _fr.active and any(
                set(_fr.antecedents) & _cn for _cn in _cellsN))
            _totR = sum(1 for _fr in _fullR if _fr.active)
            st.caption(f"{len(_rlist)} rules \u00b7 adaptation-born: "
                       f"{_nnew} \u00b7 evFIS-tuned consequents: "
                       f"{_ntun} \u00b7 convergence toward the "
 f" doctrine: {_hitR}/{_totR} doctrine "
                       f"cells touched ({_hitR / max(_totR, 1):.0%})."
                       " Learned rules persist in logs/"
                       "dss_generated_state.json across fires, engines and "
                       "MAPS; strength = accumulated fired weight.")
            from dss.fuzzy import REGISTRY as _REGR
            _dcL = (list(getattr(_engR, "decision_concepts", []))
                    or list(_dss.DECISION_CONCEPTS))
            _catN = 1
            for _dcC in _dcL:
                _catN *= len(_REGR.get(_dcC))
            # the VIEWER lists the whole stored vocabulary; whether each entry
            # is admitted right now is a separate statement, not a filter
            _svc, _svm = _store_vocab(_store_p)
            _liveC = set(getattr(_engR, "hierarchy", {}) or {})
            _liveM = set(getattr(_engR, "macros", {}) or {})
            _newC = [(c + ("" if c in _liveC else " (stored, stage \u2462 off)"))
                     for c in _svc]
            _newM = [(m + ("" if m in _liveM else " (stored, stage \u2462 off)"))
                     for m in _svm]
            st.caption("LEARNED VOCABULARY \u00b7 concepts: "
                       + (", ".join(_newC) or "none yet")
                       + " \u00b7 macro interventions: "
                       + (", ".join(_newM) or "none yet")
                       + " \u2014 vocabulary packages (new object + "
                       "a rule using it, gates G2/G2b/G3/G4/G5) come "
                       "from the LIVE Claude proposer only. Enable it with "
                       "an API key or with Claude Code on your subscription.")
            st.caption(f"Linguistic catalog: {_catN:,} antecedent "
                       "cells over the five decision concepts "
                       "(5\u2075 = 3,125 at the seed partition; "
                       "stage \u2461 term insertions GROW it).")
        elif panel == "Layer 4 · Analysis":
            # ONE table, one row per decision cycle: what the DSS decided,
            # what the adaptation tried, which gate stopped it, what came out
            # of it, and whether the fail-safe let the orders through. The
            # funnel and per-stage summaries were removed: they counted
            # things without ever saying what happened at a given step.
            _engA = st.session_state.get("dss_engine")
            _cycA = list(getattr(_engA, "cycles", []) or [])
            _rsA = dict(getattr(_engA, "run_stats", {}) or {})
            if not _cycA:
                st.info("No decision cycle has run yet. Turn on 'Apply "
                        "decisions' and step the simulation.")
            else:
                _acc = sum(1 for c in _cycA
                           if (c.get("adaptation") or {}).get("accepted"))
                _rows = build_step_rows(_cycA, _rsA)
                _COLS_A = list(STEP_COLS)
                st.caption(
                    f"{len(_cycA)} decision cycle(s) · {_acc} adaptation(s) "
                    f"accepted · seed base "
                    f"**{_sv('dss_seed_profile', 'minimal')}**")
                st.markdown("**What happened at each step** "
                            "(latest cycle first)")
                # THE PANEL COLUMN IS NARROW. Twenty-one columns of
                # provenance do not fit beside the map, so the full table
                # lives on its own page; what stays here is a preview.
                if st.button("Open the full table as a page",
                             key="goto_steps_page",
                             use_container_width=True,
                             help="Opens 'Step analysis' in the navigation: "
                                  "the same table across the whole window."):
                    st.session_state["nav_page"] = "Step analysis"
                    st.rerun()
                _ob1, _ob2 = st.columns([1, 1])
                if hasattr(st, "dialog"):
                    @st.dialog("What happened at each step", width="large")
                    def _steps_window():
                        st.dataframe(_rows, use_container_width=True,
                                     height=620, column_order=_COLS_A,
                                     column_config=_STEP_COL_CONFIG)
                        st.caption(
                            f"{len(_rows)} cycle(s). Columns are explained "
                            "under the preview table on the panel.")
                        st.download_button(
                            "⬇ Download as Excel",
                            _xlsx_bytes({"Adaptation": [
                                {c: r.get(c) for c in _COLS_A}
                                for r in _rows]}),
                            file_name="layer4_steps.xlsx", mime=XLSX_MIME,
                            key="dl_steps_xlsx")
                    if _ob1.button("Open in a window", key="open_steps",
                                   use_container_width=True,
                                   help="Shows the full table with every "
                                        "column, wider than this panel "
                                        "allows."):
                        _steps_window()
                _prev = _ob2.checkbox("Preview here", value=True,
                                      key="prev_steps")
                if _prev:
                    st.dataframe(_rows, use_container_width=True, height=300,
                                 column_order=_COLS_A,
                                 column_config=_STEP_COL_CONFIG)
                st.markdown(_STEP_COL_HELP)

                with st.expander("What the gates G1 to G5 mean"):
                    st.markdown(_GATE_HELP)

                st.markdown(
                    "**$J_{total}$ — full decision cost "
                    "(burned area + assets + population + response + "
                    "delay)**")
                _js = _rsA.get("j_series") or []
                if _js:
                    import pandas as _pdA
                    st.line_chart(_pdA.DataFrame(
                        [dict(step=s, candidate=jc, no_action=j0,
                              satisficing_bound=b)
                         for s, jc, j0, b in _js]).set_index("step"),
                        height=200)
                    st.caption(
                        "$J_{total}$ of the DSS candidate against $J_{total}$ "
                        "of doing nothing, plus the satisficing bound the "
                        "candidate has to clear. This one INCLUDES the price "
                        "of acting.")
                _ps_ = _rsA.get("phys_series") or []
                if _ps_:
                    st.markdown(
                        "**$J_{phys}$ — physical cost only "
                        "(burned area + assets + population)**")
                    import pandas as _pdP
                    st.line_chart(_pdP.DataFrame(
                        [dict(step=s, candidate=pc, no_action=p0)
                         for s, pc, p0 in _ps_]).set_index("step"),
                        height=200)
                    _win = sum(1 for _s, _pc, _p0 in _ps_ if _pc < _p0 - 1e-9)
                    st.caption(
                        f"Total $J$ above, PHYSICAL cost below. The orders "
                        f"were physically better than no action in "
                        f"**{_win}** of {len(_ps_)} cycle(s). Total $J$ also "
                        "carries the price of acting, so it sits higher "
                        "whenever the fleet is fielded; the physical curve "
                        "is the verdict on the orders themselves.")
                # ---- ROOT CAUSE ANALYSIS: the after-action review ----
                st.divider()
                st.markdown("**Root cause analysis — after-action "
                            "review (Opus)**")
                st.caption("Compiles this run's evidence (trajectory, "
                           "orders, vetoes, geometry, gates) and asks "
                           "a strong model for the incident analyst's "
                           "verdict: what worked, what failed, what "
                           "would have been better. Apply feeds the "
                           "findings back: rules / interventions / "
                           "concepts into the store, settings to the "
                           "panel, sensor and depot advice onto the "
                           "staging list. Then reset the fire and "
                           "rerun to SEE the difference.")
                _lgA = getattr(_engA, "run_logger", None)
                from dss import rca as _rca
                _rc1, _rc2 = st.columns([1, 1])
                _rmodel = _rc1.selectbox(
                    "Review model", ["opus", "sonnet"], index=0,
                    key="rca_model",
                    help="Opus: the deepest review, 1-3 minutes. "
                         "Sonnet: a faster review, usually well "
                         "under a minute.")
                if _rc2.button("Root Cause Analysis", key="rca_go",
                               disabled=_lgA is None):
                    _saved = _rca.load_saved(_lgA.dir)
                    _ev = _rca.build_evidence(
                        _lgA.dir, _engA, st.session_state.get("sim"))
                    st.session_state["rca_dir"] = _lgA.dir
                    st.session_state.pop("rca_report", None)
                    st.session_state.pop("rca_recs", None)
                    _rca.start_async(_lgA.dir, _ev, model=_rmodel)
                    st.rerun()
                # NEVER BLOCK ON THE REVIEW. This used to sleep two seconds
                # and rerun in a loop, which froze the WHOLE script for the
                # one to three minutes the deep model takes: the fire could
                # not be stepped and no other page would open. Worse,
                # leaving the panel stopped the polling altogether, so a
                # review that finished while you were elsewhere was never
                # picked up. The status now lives in a fragment that reruns
                # itself without touching the rest of the page.
                _rca_live()
                _repS = st.session_state.get("rca_report")
                if _repS:
                    st.markdown(_repS)
                    _recsS = st.session_state.get("rca_recs") or {}
                    _rlist = list(_recsS.get("recommendations") or [])
                    _nrec = len(_rlist)
                    st.caption(f"{_nrec} machine-applicable "
                               "recommendation(s); the report is also "
                               "saved to the run folder as "
                               "root_cause_analysis.md. Tick the ones "
                               "you want; untick to reject.")
                    _sel = []
                    for _iR, _recR in enumerate(_rlist):
                        _tR = str(_recR.get("type", "?"))
                        if _tR == "setting":
                            _lbl = (f"setting {_recR.get('key')} → "
                                    f"{_recR.get('value')}")
                        elif _tR in ("sensor", "depot"):
                            _lbl = (f"{_tR} {_recR.get('kind')} @ "
                                    f"({_recR.get('x')},"
                                    f"{_recR.get('y')})")
                        elif _tR == "tune_rule":
                            _lbl = f"tune rule {_recR.get('name')}"
                        else:
                            _lbl = f"{_tR} {_recR.get('name', '')}"
                        _why = str(_recR.get("why", ""))[:120]
                        if st.checkbox(f"{_lbl} — {_why}",
                                       value=True,
                                       key=f"rca_sel_{_iR}"):
                            _sel.append(_recR)
                    if _nrec and st.button(
                            f"Apply selected ({len(_sel)})",
                            key="rca_apply", disabled=not _sel):
                        from dss import rca as _rca2
                        _ap, _sk, _sess, _sns, _dps = (
                            _rca2.apply_recommendations(
                                {"recommendations": _sel}, _engA,
                                sim=st.session_state.get("sim"),
                                run_dir=(_lgA.dir if _lgA is not None
                                         else None)))
                        for _k2, _v2 in _sess.items():
                            st.session_state[_k2] = _v2
                        # PERSIST what was applied: settings and
                        # staged infrastructure survive an app
                        # restart exactly like the learned store does
                        try:
                            import json as _js_ra
                            import os as _os_ra
                            _pra = _os_ra.path.join(
                                _os_ra.path.dirname(_os_ra.path.dirname(
                                    _os_ra.path.abspath(__file__))),
                                "logs", "rca_applied.json")
                            try:
                                _prev_ra = _js_ra.load(open(_pra,
                                                            encoding="utf-8"))
                            except Exception:
                                _prev_ra = {}
                            _prev_ra.setdefault("settings", {}).update(
                                _sess)
                            for _lst, _new in (("sensors", _sns),
                                               ("depots", _dps)):
                                _cur_ra = _prev_ra.setdefault(_lst, [])
                                for _r_ra in _new:
                                    _sig_ra = (_r_ra.get("kind"),
                                               _r_ra.get("x"),
                                               _r_ra.get("y"))
                                    if _sig_ra not in [
                                            (q.get("kind"), q.get("x"),
                                             q.get("y"))
                                            for q in _cur_ra]:
                                        _cur_ra.append(_r_ra)
                            _js_ra.dump(_prev_ra, open(
                                _pra, "w", encoding="utf-8"), indent=1)
                        except Exception:
                            pass
                        if _sns:
                            _cur = list(st.session_state.get(
                                "dss_sensors", []) or [])
                            st.session_state["dss_sensors"] =                                 _cur + _sns
                            st.session_state.pop("dss_net_sig", None)
                        if _dps:
                            _cur = list(st.session_state.get(
                                "dss_res_items", []) or [])
                            st.session_state["dss_res_items"] =                                 _cur + _dps
                            st.session_state["dss_res_base"] = None
                            st.session_state.pop("dss_res_sig", None)
                        for _m3 in _ap:
                            st.success("applied: " + _m3)
                        for _m3 in _sk:
                            st.warning("skipped: " + _m3)
                        # the decision record goes NEXT TO the report
                        # in the analyzed run's folder, so what was
                        # accepted and what was refused is part of
                        # the run's own log trail
                        try:
                            import json as _js_rd
                            import os as _os_rd
                            import time as _tm_rd
                            _drd = st.session_state.get("rca_dir")
                            if _drd:
                                _js_rd.dump(
                                    dict(applied=_ap, skipped=_sk,
                                         at=_tm_rd.strftime(
                                             "%Y-%m-%d %H:%M:%S")),
                                    open(_os_rd.path.join(
                                        _drd, "rca_decisions.json"),
                                        "w", encoding="utf-8"),
                                    indent=1, ensure_ascii=False)
                        except Exception:
                            pass
                        # READBACK: prove the advice landed, from the
                        # live staging lists, not from the intent
                        _vs = [f"{q.get('kind')}@({q.get('x')},"
                               f"{q.get('y')})"
                               for q in (st.session_state.get(
                                   "dss_sensors") or [])]
                        _vd = [f"{q.get('kind')}@({q.get('x')},"
                               f"{q.get('y')})"
                               for q in (st.session_state.get(
                                   "dss_res_items") or [])
                               if q.get("label") == "RCA advice"]
                        st.caption("Now staged on the map: sensors "
                                   + (", ".join(_vs) or "none")
                                   + " | RCA depots: "
                                   + (", ".join(_vd) or "none")
                                   + ". Decision record saved to the "
                                   "run folder (rca_decisions.json).")
                        st.info("Reset the fire and rerun to see the "
                                "difference the review made.")
        elif panel == "Layer 4 · Logs":
            _eng4 = st.session_state.get("dss_engine")
            with st.expander("Saved runs \u2014 load & replay"):
                import os as _os_rp
                import json as _js_rp
                import gzip as _gz_rp
                _logroot = _os_rp.path.join(_os_rp.path.dirname(
                    _os_rp.path.dirname(_os_rp.path.abspath(__file__))),
                    "logs")
                # a run is CURRENT-FORMAT if it replays (world.json.gz) AND its
                # cycles.jsonl carries the fields this build reads
                _REQ = ("step", "sim", "forecast", "regions", "adaptation")

                def _run_ok(_d):
                    _p = _os_rp.path.join(_logroot, _d)
                    if not _os_rp.path.exists(
                            _os_rp.path.join(_p, "world.json.gz")):
                        return False
                    try:
                        with open(_os_rp.path.join(_p,
                                                   "cycles.jsonl")) as _fh:
                            _j0 = _js_rp.loads(_fh.readline())
                        return all(_k in _j0 for _k in _REQ)
                    except Exception:
                        return False
                _alldirs = [d for d in (_os_rp.listdir(_logroot)
                            if _os_rp.path.isdir(_logroot) else [])
                            if _os_rp.path.isdir(
                                _os_rp.path.join(_logroot, d))]
                _runs = sorted([d for d in _alldirs if _run_ok(d)],
                               reverse=True)
                _bad = sorted([d for d in _alldirs if d not in _runs])
                if _bad:
                    st.caption(f"⚠ {len(_bad)} saved run(s) do NOT match "
                               "the current log format (old / partial); they "
                               "cannot be analysed or replayed.")
                    _cfmP = st.checkbox("confirm delete", key="logs_purge_ok")
                    if st.button(f"Delete {len(_bad)} incompatible run(s)",
                                 disabled=not _cfmP, key="logs_purge"):
                        import shutil as _sh_p
                        _np = 0
                        for _d in _bad:
                            try:
                                _sh_p.rmtree(_os_rp.path.join(_logroot, _d))
                                _np += 1
                            except Exception:
                                pass
                        st.toast(f"Deleted {_np} incompatible run(s).")
                        st.rerun()
                if not _runs:
                    st.caption("No replayable runs yet \u2014 every "
                               "run with 'Apply decisions' ON saves "
                               "itself under 03_Codes/logs "
                               "(world.json.gz + meta.json + "
                               "cycles.jsonl).")
                else:
                    _rsel_rp = st.selectbox("Run", _runs, key="rp_run")
                    if st.button("Load this run for replay",
                                 help="Rebuilds the exact map, "
                                      "sensors, resource pool, "
                                      "engine and weather settings "
                                      "of the saved run. Then press "
                                      "'Run to end': the engine and "
                                      "the rng are deterministic, so "
                                      "the run reproduces itself "
                                      "(when GenAI is off; live "
                                      "Claude proposals are not "
                                      "deterministic)."):
                        _rd = _os_rp.path.join(_logroot, _rsel_rp)
                        with _gz_rp.open(_os_rp.path.join(
                                _rd, "world.json.gz"), "rt") as _f:
                            _wnew = World.from_dict(_js_rp.load(_f))
                        try:
                            _meta = _js_rp.load(open(_os_rp.path.join(
                                _rd, "meta.json")))
                        except Exception:
                            _meta = {}
                        _new_simulator(_wnew)
                        # rebuild EXACTLY this run's rule base: seed profile
                        # (restored below) + this run's learned snapshot
                        _rr = _os_rp.path.join(_rd, "rules.json")
                        if _os_rp.path.exists(_rr):
                            st.session_state["dss_learned_override"] = _rr
                        _me = _meta.get("engine", {}) or {}
                        for _mk, _sk in (
                                ("regions", "dss_n"),
                                ("cycle_min", "dss_cycle_min"),
                                ("horizon_min", "dss_horizon_min"),
                                ("j_th", "dss_jth"),
                                ("eta", "dss_eta"),
                                ("genai", "dss_genai_on"),
                                ("evfis", "dss_evfis_on"),
                                ("evfis_step", "dss_evfis_step"),
                                ("ctrl_eps", "dss_ctrl_eps"),
                                ("ctrl_lr", "dss_ctrl_lr"),
                                ("attn", "dss_attn_thr"),
                                ("min_gain", "dss_min_gain"),
                                ("seed_profile", "dss_seed_profile"),
                                ("use_stage12", "dss_use_stage12"),
                                ("use_stage3", "dss_use_stage3")):
                            if _mk in _me:
                                st.session_state[_sk] = _me[_mk]
                        for _wk, _wv in (_meta.get("weather", {})
                                         or {}).items():
                            st.session_state[_wk] = _wv
                        if _meta.get("sensors"):
                            st.session_state["dss_sensors"] =                                 list(_meta["sensors"])
                        if _meta.get("depots"):
                            st.session_state["dss_res_items"] =                                 list(_meta["depots"])
                            st.session_state["dss_res_base_v"] =                                 st.session_state.map_version
                        st.session_state["dss_apply"] = True
                        st.toast(f"{_rsel_rp} loaded \u2014 press "
                                 "'Run to end' to replay.")
                        st.rerun()

                    # ---- auto-analysis of the SELECTED run (from its
                    # cycles.jsonl), so a loaded run can be studied without
                    # replaying it ----
                    _acf = _os_rp.path.join(_logroot, _rsel_rp,
                                            "cycles.jsonl")
                    _acys = []
                    try:
                        with open(_acf) as _afh:
                            for _aln in _afh:
                                if _aln.strip():
                                    _acys.append(_js_rp.loads(_aln))
                    except Exception:
                        _acys = []
                    if _acys:
                        _SYM = {"suppression_effort": "S",
                                "resource_deployment": "D",
                                "containment_line": "C",
                                "asset_protection": "P",
                                "evacuation": "E", "public_warning": "W"}
                        _ivU = {iv: 0 for iv in _SYM}
                        _stc = {1: 0, 2: 0, 3: 0}
                        _sto = {1: 0, 2: 0, 3: 0}
                        _foc = {}
                        for _c in _acys:
                            for _rg in (_c.get("regions") or {}).values():
                                for _iv, _vv in (_rg.get("orders_final")
                                                 or {}).items():
                                    if _iv in _ivU and _vv > 0.05:
                                        _ivU[_iv] += 1
                            _ad = _c.get("adaptation") or {}
                            _t = _ad.get("tried")
                            if _t in _stc:
                                _stc[_t] += 1
                                if _ad.get("accepted"):
                                    _sto[_t] += 1
                            _h = (_c.get("global_dss") or {}).get("hotspot")
                            if _h:
                                _foc[_h] = _foc.get(_h, 0) + 1
                        _lastc = _acys[-1]
                        _lsm = _lastc.get("sim", {}) or {}
                        st.markdown(f"**Analysis of {_rsel_rp} \u00b7 "
                                    f"{len(_acys)} cycles**")
                        st.caption(
                            f"final situation: {_lsm.get('burning', '?')} "
                            f"burning, {_lsm.get('burned', '?')} burned \u00b7 "
                            "last forecast J="
                            + str((_lastc.get('forecast') or {}).get(
                                'j_candidate', '?')))
                        st.caption("intervention use (cycles ordered): "
                                   + " \u00b7 ".join(f"{_SYM[_iv]}={_n}"
                                                for _iv, _n in _ivU.items()))
                        st.caption("adaptations tried/accepted: "
                                   + " \u00b7 ".join(
                                       f"stage{_s} {_stc[_s]}/{_sto[_s]}"
                                       for _s in (1, 2, 3)))
                        if _foc:
                            _tp = sorted(_foc.items(),
                                         key=lambda x: -x[1])[:3]
                            st.caption("most-focused regions: "
                                       + ", ".join(f"{_n}\u00d7{_c}"
                                                   for _n, _c in _tp))
            st.markdown("**Decision log \u2014 one cycle, the full story**")
            st.caption("Pick a cycle and a region (or All agents) ONCE; the "
                       "situation, the agents, the forecast, the adaptation "
                       "and the what-if below all read that selection.")
            if _eng4 is None or not _eng4.log.records:
                st.caption("No decisions logged yet \u2014 stage a "
                           "pool, toggle 'Apply decisions' and step "
                           "the simulation.")
            else:
                _cyc = _eng4.log.cycles()
                st.caption(f"{len(_cyc)} decision cycles \u00b7 "
                           f"runtime rules: {len(_eng4.rules)} "
                           f"(adaptation-born: "
                           f"{sum(1 for r in _eng4.rules if r.name[0] in 'AG')})"
                           f" \u00b7 controller values: " + ", ".join(
                               f"{k[0]}/s{k[1]}={v:.2f}"
                               for k, v in sorted(_eng4.controller.q.items())))
                # ONE cycle selector + ONE region selector for the whole
                # section (situation, agents, forecast, adaptation, what-if)
                _ksel = st.selectbox("Cycle (step)", list(reversed(_cyc)),
                                     key="dlog_k")
                _recs = _eng4.log.at(int(_ksel))
                _ropt = ["All agents (table)"] + [r.region for r in _recs]
                _rsel = st.selectbox("Region", _ropt, key="dlog_r")
                _cy = next((c for c in _eng4.cycles
                            if c["step"] == int(_ksel)), None)
                _rec = (_recs[0] if _rsel == "All agents (table)"
                        else next(r for r in _recs if r.region == _rsel))

                # ---- SITUATION + GLOBAL + FORECAST + ADAPTATION ----
                if _cy is not None:
                    _sm0 = _cy["sim"]
                    _cst = _cy.get("costs", {})
                    st.markdown(f"**Situation · k={_cy['step']} "
                                f"(t={_cy['t_min']:.0f} min)**")
                    st.caption(
                        f"{_sm0['burning']} cells burning · "
                        f"{_sm0['burned']} burned · wind "
                        f"{_sm0['wws_mean']:.1f} m/s · rain "
                        f"{_sm0['prec_mean']:.1f} mm/h · fuel moisture "
                        f"{_sm0['fmoist_mean']:.3f}"
                        + (f" · population-loss term {_cst.get('pop', 0):.3f}"
                           if "pop" in _cst else ""))
                    _gd0 = _cy.get("global_dss")
                    if _gd0:
                        st.caption("Global DSS · "
                                   + _gd0.get("statement", ""))
                    _fc0 = _cy["forecast"]
                    _gain0 = _fc0["j_noaction"] - _fc0["j_candidate"]
                    st.caption(
                        f"Forecast · with orders J="
                        f"{_fc0['j_candidate']:.3f} vs no-action J="
                        f"{_fc0['j_noaction']:.3f} · expected gain "
                        f"ΔJ={_gain0:+.3f} "
                        + ("(the decision helps)" if _gain0 > 1e-4
                           else "(no measurable gain at this horizon)")
                        + " · costs "
                        + " ".join(f"{k}={v:.3f}"
                                   for k, v in _cst.items()))
                    _ad0 = _cy["adaptation"]
                    st.caption(
                        "Adaptation · " + (
                            (("✅ " if _ad0["accepted"] else "❌ ")
                             + f"stage {_ad0['tried']}: {_ad0['detail']} "
                             f"(dJ {_ad0['dJ']:+.4f})")
                            if _ad0["tried"] else
                            "not attempted — the seed rule base "
                            "satisficed"))

                # ---- REGION VIEW ----
                if _rsel == "All agents (table)":
                    _hd = "| |" + "".join(f" {r.region} |" for r in _recs)
                    _rws = [_hd, "|---|" + "---|" * len(_recs)]
                    _macL = list(_all_macros(_eng4))
                    for _ivL in list(_dss.INTERVENTIONS) + _macL:
                        _chipL = ("<span style='color:"
                                  f"{_IV_COLOR.get(_ivL, '#c000ff')}'>"
                                  "■</span> ")
                        _lblL = (_dss.INTERVENTION_LABEL.get(_ivL)
                                 or _ivL.replace('_', ' ') + " (GenAI macro)")
                        _rws.append(
                            f"| {_chipL}{_lblL} | "
                            + " | ".join(
                                f"{r.intensities.get(_ivL, 0.0):.2f}"
                                for r in _recs) + " |")
                    _rws.append("| quality Q | " + " | ".join(
                        f"{r.quality:.2f}" for r in _recs) + " |")
                    _rws.append("| global share | " + " | ".join(
                        f"{getattr(r, 'coord_share', 1.0):.2f}"
                        for r in _recs) + " |")
                    _rws.append("| attended | " + " | ".join(
                        ("●" if getattr(r, "attended", True) else "–")
                        for r in _recs) + " |")
                    st.markdown("\n".join(_rws), unsafe_allow_html=True)
                elif _cy is not None and _rsel in _cy["regions"]:
                    _rg = _cy["regions"][_rsel]
                    st.caption("z: " + " · ".join(
                        f"{_dss.FEATURE_SYM[k].replace('_', '')}={v:.2f}"
                        for k, v in _rg["features"].items()))
                    st.caption("concepts (gated): " + " · ".join(
                        f"{k.replace('_', ' ')}={v:.2f}"
                        for k, v in _rg["concepts_effective"].items()))
                    st.caption("fired: " + (" ".join(
                        f"{n}[{w:.2f}]" for n, w in _rg["fired"]) or "none"))
                    st.caption(
                        "orders (rules→final): " + " · ".join(
                            f"{k.split('_')[0]} "
                            f"{_rg['orders_from_rules'].get(k, 0):.2f}"
                            f"→{v:.2f}"
                            for k, v in _rg["orders_final"].items()
                            # undiscovered actuators carry permanent
                            # zeros; only doctrine families and cells
                            # with a real order belong in the story
                            if k in _dss.rules.DOCTRINE_INTERVENTIONS
                            or float(v) > 0.02
                            or float(_rg["orders_from_rules"]
                                     .get(k, 0)) > 0.02)
                        + f" · Q={_rg['quality']:.2f}"
                        + (" · FAILSAFE" if _rg["failsafe"] else "")
                        + f" · share {_rg['coord_share']:.2f}")

                if _cy is not None:
                    import json as _js_c
                    st.download_button(
                        "Download this cycle (JSON, full detail)",
                        _js_c.dumps(_cy, indent=1).encode(),
                        file_name=f"cycle_k{_cy['step']}.json",
                        mime="application/json", key="chron_dl")
                st.divider()
                st.markdown("**Trace table — every decision cycle as a row**")
                st.caption("One row per cycle × region with the FULL story: "
                           "the six order intensities by full name "
                           "(suppression_effort, resource_deployment, "
                           "containment_line, asset_protection, evacuation, "
                           "public_warning), quality Q, global share, "
                           "attended, forecast J, the situation (burning / "
                           "burned / wind / moisture), the cost terms and the "
                           "adaptation outcome. The CSV and the JSON below "
                           "carry the same information.")
                _sm_l = float(getattr(cfg, "step_minutes", 1.0))
                _cyc_by = {c["step"]: c for c in _eng4.cycles}
                _cost_keys = sorted({k for c in _eng4.cycles
                                     for k in (c.get("costs") or {})})

                def _row(r):
                    _c = _cyc_by.get(r.step, {})
                    _sm = _c.get("sim", {}) or {}
                    _cs = _c.get("costs", {}) or {}
                    _ad = _c.get("adaptation", {}) or {}
                    _gd = _c.get("global_dss", {}) or {}
                    d = dict(
                        step=r.step, t_min=int(r.step * _sm_l),
                        region=r.region,
                        provenance=_eng4.log.stage_story(r),
                        fired=" ".join(f"{n}[{w:.2f}]"
                                       for n, w in r.fired[:4]),
                        suppression_effort=round(
                            r.intensities.get("suppression_effort", 0), 2),
                        resource_deployment=round(
                            r.intensities.get("resource_deployment", 0), 2),
                        containment_line=round(
                            r.intensities.get("containment_line", 0), 2),
                        asset_protection=round(
                            r.intensities.get("asset_protection", 0), 2),
                        evacuation=round(
                            r.intensities.get("evacuation", 0), 2),
                        public_warning=round(
                            r.intensities.get("public_warning", 0), 2),
                        Q=round(r.quality, 2),
                        share=round(float(getattr(r, "coord_share", 1.0)), 2),
                        attended=int(bool(getattr(r, "attended", True))),
                        J_fc=round(r.j_forecast, 3),
                        J_no=round(r.j_noaction, 3),
                        burning=_sm.get("burning"),
                        burned=_sm.get("burned"),
                        wind=round(float(_sm.get("wws_mean", 0)), 2),
                        moisture=round(float(_sm.get("fmoist_mean", 0)), 3),
                        stage=(_ad.get("tried") or 0),
                        stage_ok=int(bool(_ad.get("accepted"))),
                        global_stmt=(_gd.get("statement", "") if _gd else ""))
                    for _k in _cost_keys:
                        d[f"cost_{_k}"] = round(float(_cs.get(_k, 0)), 3)
                    return d
                _tbl = [_row(r) for r in _eng4.log.records]
                st.dataframe(_tbl, use_container_width=True, height=280)
                import json as _js_all
                _dlc1, _dlc2 = st.columns(2)
                # a semicolon CSV "that opens in Excel" still arrives as
                # text in every column and depends on the reader's locale
                # separator; a workbook does not
                _dlc1.download_button(
                    "⬇ Download trace (Excel)",
                    _xlsx_bytes({"Decision trace": _tbl}),
                    file_name="dss_decision_trace.xlsx", mime=XLSX_MIME,
                    key="trace_xlsx")
                _dlc2.download_button(
                    "Download ALL cycles (JSON, full detail)",
                    _js_all.dumps(_eng4.cycles, indent=1).encode(),
                    file_name="dss_all_cycles.json",
                    mime="application/json", key="all_json_dl")
                with st.expander("Global DSS history (all "
                                 "cycles)"):
                    _gcyc = [(c.get("step"), c.get("t_min"),
                              c.get("global_dss"))
                             for c in _eng4.cycles
                             if c.get("global_dss")]
                    if not _gcyc:
                        st.caption("No global decisions logged yet.")
                    for _gs, _gt, _gg in _gcyc[-40:]:
                        st.caption(f"k={_gs} (t={_gt:.0f} min): "
                                   + _gg.get("statement", ""))
                    st.caption("Full history: `global.csv` in the "
                               "run folder (one row per cycle: "
                               "hotspot, attended set, shares, "
                               "statement).")
                with st.expander("Adaptation history (all cycles)"):
                    _seen_h = set()
                    for _r_h in _eng4.log.records:
                        if _r_h.step in _seen_h:
                            continue
                        _seen_h.add(_r_h.step)
                        st.caption(f"k={_r_h.step}: "
                                   + _eng4.log.stage_story(_r_h))
                if _eng4.run_logger is not None:
                    st.caption("Persistent run log: "
                               f"`{_eng4.run_logger.dir}` "
                               "(steps.csv + decisions.jsonl \u2014 "
                               "every simulation step and every "
                               "decision cycle).")
                st.markdown("**What-if — how much did the DSS orders help?**")
                st.caption("Clones the run, withdraws some orders on the "
                           "CLONE, replays it, and compares. The live run is "
                           "untouched; the difference is due to the removed "
                           "orders alone.")
                _cfsc = st.radio(
                    "Remove which orders?", [
                        "all orders, from the start (what if the DSS never "
                        "acted)",
                        f"only the orders from cycle k={_rec.step} onward"],
                    index=0, key="cf_scope",
                    help="First: the clone rewinds to step 0 and replays the "
                         "whole fire with no DSS orders at all. Second: keeps "
                         "everything up to the selected cycle and withdraws "
                         "only the orders from that cycle onward.")
                if st.button("What if these orders were NOT taken? "
                             "(counterfactual replay)",
                             help="Clones the live simulation, rewinds "
                                  "the CLONE and replays history "
                                  "without the selected orders; the "
                                  "live run is untouched. Fuel "
                                  "moisture and the rng state travel "
                                  "with the snapshot, so the "
                                  "difference is attributable to the "
                                  "withdrawn orders alone."):
                    _cf_from = (0 if _cfsc.startswith("all orders")
                                else int(_rec.step))
                    _dtm_cf = float(getattr(cfg, "step_minutes", 1.0))
                    with st.spinner("Replaying without orders..."):
                        _cf, _rcf = _dss.counterfactual(
                            sim, _cf_from,
                            step_hook=lambda w2, k2:
                                _drive_weather(w2, k2 * _dtm_cf))
                    if _cf is None:
                        st.warning("No snapshot left at that step "
                                   "(memory-capped history).")
                    else:
                        from disaster_phyengine.costs import (
                            compute_costs as _cc4)
                        _ract = _cc4(sim)
                        _pa = float(getattr(_ract, "j_physical", 0.0))
                        _pc = float(getattr(_rcf, "j_physical", 0.0))
                        m1, m2, m3 = st.columns(3)
                        m1.metric("Jφ actual (with DSS)",
                                  f"{_pa:.3f}",
                                  help="Fair outcome metric: "
                                       "physical loss only, both "
                                       "sides (a no-order run pays "
                                       "no response cost by "
                                       "definition, so the decision "
                                       "cost would be biased).")
                        m2.metric("Jφ without orders",
                                  f"{_pc:.3f}",
                                  delta=f"{_pc - _pa:+.3f}")
                        m3.metric("Burned cells: actual vs without",
                                  f"{int(sim.ever_burned.sum())} / "
                                  f"{int(_cf.ever_burned.sum())}")
                        i1, i2, i3 = st.columns(3)
                        _sc4 = max(3, 380 // max(cfg.nx, 1))
                        i1.image(viz.render_pil(world, sim=sim,
                                                scale=_sc4),
                                 caption="actual (orders applied)")
                        i2.image(viz.render_pil(_cf.world, sim=_cf,
                                                scale=_sc4),
                                 caption="counterfactual (no orders)")
                        from PIL import Image as _PILImg
                        _ba = sim.ever_burned
                        _bc = _cf.ever_burned
                        _dif = np.zeros((cfg.ny, cfg.nx, 3),
                                        dtype=np.uint8)
                        _dif[...] = 34
                        _dif[_ba & _bc] = (90, 90, 90)
                        _dif[_bc & ~_ba] = (40, 200, 90)
                        _dif[_ba & ~_bc] = (230, 60, 60)
                        _im = _PILImg.fromarray(_dif).resize(
                            (cfg.nx * _sc4, cfg.ny * _sc4),
                            _PILImg.NEAREST)
                        i3.image(_im, caption="difference \u2014 "
                                 "green: cells the orders SAVED \u00b7 "
                                 "red: burned only WITH orders \u00b7 "
                                 "grey: burned in both")
                        st.caption("Reading the maps: black = burned "
                                   "OUT, orange = still burning now; a "
                                   "slowed fire can look 'more orange' "
                                   "while having burned far less. The "
                                   "difference map is the honest "
                                   "comparison.")
        elif panel == "Time":
            st.caption("Simulation time: what one step represents and "
                       "where the clock stands.")
            st.session_state["dr_start"] = float(st.number_input(
                "Start hour of day", 0.0, 23.0,
                float(st.session_state.get("dr_start", 12.0)), 1.0,
                help="What time of day step 0 corresponds to (drives "
                     "the map clock, night dimming and the diurnal "
                     "weather wave)."))
            st.markdown("**Time per step** \u2014 what one step "
                        "$k \\to k{+}1$ represents")
            tc1, tc2 = st.columns([1.1, 1.1])
            _sm = float(getattr(cfg, "step_minutes", 30.0))
            if _sm % 1440 == 0 and _sm >= 1440:
                _u0, _v0 = "days", _sm / 1440
            elif _sm % 60 == 0 and _sm >= 60:
                _u0, _v0 = "hours", _sm / 60
            else:
                _u0, _v0 = "minutes", _sm
            tval = tc1.number_input("Value", 1.0, 1440.0, float(_v0), 1.0,
                                    key="dt_val")
            tunit = tc2.selectbox("Unit", ["minutes", "hours", "days"],
                                  index=["minutes", "hours",
                                         "days"].index(_u0), key="dt_unit")
            cfg.step_minutes = tval * {"minutes": 1, "hours": 60,
                                       "days": 1440}[tunit]
            st.caption("Elapsed fire time so far: "
                       f"{_fmt_sim_time(getattr(sim, 't_elapsed_min', sim.state.step * cfg.step_minutes))}"
                       f" in {sim.state.step} steps. Changing the "
                       "step length applies from the NEXT step on and "
                       "the elapsed total is PRESERVED (the clock "
                       "integrates each executed step). The physical "
                       "speed in m/min stays the same (System "
 "Description, note 8).")
            # live head-fire speed so the m/min meaning is always visible
            from disaster_phyengine.spread import rate_of_spread as _rosf
            _ros = _rosf(world.fuel, world.topo, world.meteo, cfg.spread)
            _pos = _ros[_ros > 0]
            if _pos.size:
                _v = float(np.percentile(_pos, 95))
                _cellm = cfg.cell_size_m / 30.0
                _vm = _v * _cellm          # m/min at the reference
                st.caption(f"Head-fire potential right now: \u2248 "
                           f"{_vm:.0f} m/min ({_vm*0.06:.1f} km/h) on "
                           "the fastest cells \u2014 independent of "
                           "the step length; per step the front moves "
                           f"\u2248 {_vm*cfg.step_minutes/cfg.cell_size_m:.2f} "
                           "cells.")
            _tmin = sim.state.step * float(getattr(cfg,
                                                   "step_minutes", 1.0))
            st.caption(f"Sim clock: {(_tmin/60.0) % 24.0:04.1f} h of "
                       "day (fire behaviour peaks ~15:00, minimum "
                       "before dawn).")
        elif panel == "Ignition":
            ig_live = st.checkbox(
                "At current step", value=True,
                help="Place the ignition at the current step; untick to "
                     "schedule it for a future step.")
            ig_step = sim.state.step if ig_live else st.number_input(
                "Step", 0, 5000, 0, key="ig_step")
            ig_rad = st.number_input(
                "Radius (cells)", 0, 20, 1, key="ig_rad",
                help=f"0 = a single cell; r ignites a disk of r cells "
                     f"around the click. One cell = {cfg.cell_size_m:g} m "
                     "on the ground.")
            th = st.number_input(
                "$\\Theta_{ign}$ — ignition threshold", 0.005, 1.0,
                float(cfg.spread.theta_ign), 0.005, format="%.3f",
                help="Activation threshold on the ignition influence "
 "buildup $A_k$. Unit: "
                     "accumulated influence = \u215b cell-widths of "
                     "front travel. Default $0.125=1/8$, which makes a "
                     "wind-aligned front advance at exactly "
                     "$R_{spread}$ cells per step. Lower = easier "
                     "ignition of new cells.")
            if abs(th - float(cfg.spread.theta_ign)) > 1e-12:
                cfg.spread.theta_ign = th
            st.session_state["ig_live_v"] = bool(ig_live)
            st.session_state["ig_step_v"] = int(ig_step)
            st.session_state["ig_rad_v"] = int(ig_rad)
            # --- placed ignitions: list + remove any, any time ---
            st.markdown("**Placed ignitions**")
            if world.ignitions:
                for _igi, _ev in enumerate(list(world.ignitions)):
                    _ic1, _ic2 = st.columns([5, 0.8])
                    _ic1.caption(
                        f"{_igi + 1}. @ ({int(_ev.x)}, {int(_ev.y)}) \u00b7 "
                        f"step {int(_ev.step)} \u00b7 r {int(_ev.radius)}")
                    if _ic2.button("\u2716",
                                   key=f"ig_rm_{_igi}_{_ev.x}_{_ev.y}",
                                   help="Remove this ignition"):
                        world.ignitions.pop(_igi)
                        st.rerun()
                if st.button("Clear all ignitions",
                             use_container_width=True):
                    world.ignitions.clear()
                    st.rerun()
            else:
                st.caption("No ignitions placed. Click the map to add one.")

        elif panel == "Layer 2 \u00b7 Perception":
            st.caption("Layer 2 \u2014 perception: multi-source data "
                       "fusion + feature extraction. Each Local DSS "
                       "reads the fused observation restricted to its "
                       "region and extracts the ten features "
                       "$z_1..z_{10}$.")
            if _obsnet is None:
                st.warning("OBSERVATION BYPASS \u2014 'Partial "
                           "observation via sensors' is OFF: every "
                           "feature reads the TRUE simulation state "
                           "with confidence 1. This is the idealized "
                           "physics-test mode, not the thesis "
                           "pipeline. Turn it on in Layer 1 for the "
                           "sensor-fused feed.")
            elif not _sv("dss_sensors", []):
                st.warning("No sensors staged: the fire-driven "
                           "features ($z_1$ intensity, $z_4$ "
                           "proximity, $z_{10}$ urgency) read 0 with "
                           "ZERO confidence \u2014 the DSS is blind "
                           "to the fire. The values you still see "
                           "come from sensor-independent sources "
                           "(see below). Stage a network in Layer 1.")
            st.caption("Where each value comes from \u00b7 "
                       "$z_1, z_4, z_{10}$: SENSOR fusion (burning/"
                       "intensity channels) \u00b7 $z_5$: pre-fire "
                       "fuel map prior, aged \u00b7 $z_2$: spread "
                       "model over maps + weather \u00b7 $z_3$: "
                       "weather service \u00b7 $z_6, z_7, z_8$: GIS "
                       "(assets, access, roads) \u00b7 $z_9$: own "
                       "resource ledger (staged pool).")
            _mirror("dss_n", "l1_dss_n", 1)
            n_agents = int(st.number_input(
                "Number of local DSS agents", 1, 12, step=1,
                key="l1_dss_n",
                on_change=_adopt, args=("dss_n", "l1_dss_n"),
                help="The map is split into exactly this many regions "
                     "covering every cell (near-square blocks, Agent_1 "
                     "at the north-west)."))
            _nrec = max(1, min(12, round((cfg.nx * cfg.ny) / 2500.0)))
            st.caption(f"Suggested: {_nrec} agent(s) \u2014 one "
                       "Local DSS per ~50\u00d750-cell "
                       "responsibility block. More agents pay off "
                       "when regional situations diverge (big maps, "
                       "several fronts); a single agent suffices on "
                       "small maps.")
            if _nrec != n_agents and st.button(
                    f"Use suggested ({_nrec})",
                    key="dss_n_sugg"):
                st.session_state["dss_n"] = int(_nrec)
                st.rerun()
            _regs = _dss.partition_n(cfg.nx, cfg.ny, n_agents)
            show_all = st.checkbox(
                "Show all regions on the map",
                value=bool(_sv("dss_show_all", True)))
            st.session_state["dss_show_all"] = show_all
            st.session_state["dss_regions_all"] = (
                [(*r.box, r.name) for r in _regs] if show_all else None)
            _names = [r.name for r in _regs] + ["All agents (table)"]
            # default to the all-agents table (the last entry)
            _icur = min(int(_sv("dss_sel_i", len(_names) - 1)),
                        len(_names) - 1)
            _selA = st.selectbox("Agent", _names, index=_icur)
            st.session_state["dss_sel_i"] = _names.index(_selA)
            if _selA == "All agents (table)":
                st.session_state["dss_region"] = None
                _thr = float(st.slider(
                    "Attention threshold (urgency)", 0.0, 1.0,
                    float(_sv("dss_attn_thr", 0.35)), 0.05,
                    help="The coordinator ATTENDS a region when its "
                         "temporal urgency reaches this value (fire "
                         "inside or approaching); the rest are only "
                         "monitored. Attended regions burn hot on the "
                         "map, ignored ones stay dim."))
                st.session_state["dss_attn_thr"] = _thr
                _feats = {_r.name: _dss.ten_features(sim, _r,
                                                     network=_obsnet)
                          for _r in _regs}
                _att_flags = {n: f["temporal_urgency"] >= _thr
                              for n, f in _feats.items()}
                # full table: all 10 features as rows, one column per agent
                # (the attention mark travels in the column header)
                _short = {k: f"${_dss.FEATURE_SYM[k].replace('_10', '_{10}')}$ "
                             + _dss.FEATURE_NAME[k]
                          for k in _dss.FEATURE_ORDER}
                _head = "| feature |"
                _sep = "|---|"
                for _r in _regs:
                    _mark = "\u25cf" if _att_flags[_r.name] else "\u2013"
                    _head += f" {_r.name} {_mark} |"
                    _sep += "---|"
                _rowsmd = [_head, _sep]
                if _obsnet is not None:
                    _kcells = " | ".join(
                        f"{_obsnet.region_conf(_r):.2f}" for _r in _regs)
                    _rowsmd.append(f"| $conf=\\min\\gamma$ | {_kcells} |")
                for _k in _dss.FEATURE_ORDER:
                    _cells = " | ".join(f"{_feats[_r.name][_k]:.2f}"
                                        for _r in _regs)
                    _lab = _short[_k]
                    if _k == "temporal_urgency":
                        _lab = f"**{_lab}**"
                    _rowsmd.append(f"| {_lab} | {_cells} |")
                st.markdown("\n".join(_rowsmd))
                with st.expander("What each $z_i$ measures"):
                    for _k, _sym, _nm, _ms in _dss.FEATURE_META:
                        _symtex = _sym.replace("_10", "_{10}")
                        st.markdown(f"${_symtex}$ **{_nm}** \u2014 {_ms}")
                _natt = sum(1 for v in _att_flags.values() if v)
                st.caption(f"\u25cf attended: {_natt} \u00b7 monitored "
                           f"only: {len(_regs) - _natt} (urgency \u2265 "
                           f"{_thr:.2f} \u21d2 attended). Allocation "
                           "logic arrives in the next phase.")
                if show_all:
                    st.session_state["dss_regions_all"] = [
                        (*_r.box, _r.name, bool(_att_flags[_r.name]))
                        for _r in _regs]
            else:
                _reg = _regs[_names.index(_selA)]
                st.session_state["dss_region"] = (*_reg.box, _reg.name)
                st.caption(f"{_reg.name}: x {_reg.x0}\u2013{_reg.x1 - 1}, "
                           f"y {_reg.y0}\u2013{_reg.y1 - 1} "
                           f"({_reg.n_cells} cells) \u2014 highlighted "
                           "on the map.")
                if _obsnet is not None:
                    _cf = _obsnet.region_conf(_reg)
                    st.markdown("**How well the sensors see this region**")
                    st.progress(min(1.0, _cf),
                                text=f"observation confidence {_cf:.2f} "
                                     "(0 = blind, 1 = fully fresh)")
                    _CHN = {"burning": "burning", "fload": "fuel",
                            "intensity": "intensity", "tau": "burn age"}
                    _cc = _obsnet.region_conf_components(_reg)
                    st.caption(
                        "freshness of each observed quantity (0 to 1): "
                        + " \u00b7 ".join(
                            f"{_dss.CHANNEL_SYMBOL[_ch]} {_v:.2f}"
                            for _ch, _v in _cc.items()))
                    _ra = _obsnet.region_age(_reg)
                    st.caption(
                        "time since the newest reading: " + " \u00b7 ".join(
                            f"{_dss.CHANNEL_SYMBOL[_ch]} "
                            + ("never" if _v > 9e5 else f"{_v:.0f} min")
                            for _ch, _v in _ra.items()))
                    _syx = _reg.slices()
                    _tb = int((sim.state.burning[_syx] > 0.5).sum())
                    _ob = int((_obsnet.obs["burning"][_syx] > 0.5).sum())
                    st.caption(f"burning cells the sensors report: {_ob} "
                               f"(really burning: {_tb}). The DSS acts on "
                               "what it can see, not the hidden truth; the "
                               "gap shrinks as fresh reports arrive.")
                    with st.expander("What these numbers mean"):
                        st.markdown(
                            "**Observation confidence** is how fresh and "
                            "complete the sensor picture of this region is. "
                            "It follows the weakest of the four observed "
                            "quantities.\n\n"
                            "**B / F / I / \u03c4** are those four quantities: "
                            "**B** whether a cell is burning, **F** fuel "
                            "load, **I** fire intensity, **\u03c4** how long a "
                            "cell has been burning.\n\n"
                            "**Freshness** is 1 right after an observation and "
                            "decays as the reading ages (halves about every "
                            "90 minutes), reaching 0 when the region has not "
                            "been seen for a long time.\n\n"
                            "Sensors are shared; this agent only reads the "
                            "combined observation inside its own region "
                            "$\\Omega_i$.")
                _f = _dss.ten_features(sim, _reg, network=_obsnet)
                _zsub = {"z_1": "z\u2081", "z_2": "z\u2082",
                         "z_3": "z\u2083", "z_4": "z\u2084",
                         "z_5": "z\u2085", "z_6": "z\u2086",
                         "z_7": "z\u2087", "z_8": "z\u2088",
                         "z_9": "z\u2089", "z_10": "z\u2081\u2080"}
                # change arrows: compare with the features at the PREVIOUS
                # simulation step (not the previous rerun)
                _pk = f"dss_featprev_{_reg.name}"
                _prev_rec = st.session_state.get(_pk)
                _step_now = int(sim.state.step)
                if _prev_rec is None or _prev_rec[0] == _step_now:
                    _fprev = (_prev_rec[1] if _prev_rec else dict(_f))
                else:
                    _fprev = _prev_rec[1]
                _fcL = _dss.feature_confidence(_obsnet, _reg)
                for _k in _dss.FEATURE_ORDER:
                    _d = float(_f[_k]) - float(_fprev.get(_k, _f[_k]))
                    _arr = ("\u2191" if _d > 0.005
                            else "\u2193" if _d < -0.005 else "")
                    _dtx = f"  {_arr}{abs(_d):.2f}" if _arr else ""
                    _ctx = ("" if _fcL[_k] >= 0.999
                            else f" \u00b7 conf {_fcL[_k]:.2f}")
                    st.progress(min(1.0, float(_f[_k])),
                                text=f"{_zsub[_dss.FEATURE_SYM[_k]]} "
                                     f"{_dss.FEATURE_NAME[_k]}: "
                                     f"{_f[_k]:.2f}{_dtx}{_ctx}")
                if _prev_rec is None or _prev_rec[0] != _step_now:
                    st.session_state[_pk] = (_step_now, dict(_f))
                with st.expander("What each $z_i$ measures"):
                    for _k, _sym, _nm, _ms in _dss.FEATURE_META:
                        _symtex = _sym.replace("_10", "_{10}")
                        st.markdown(f"${_symtex}$ **{_nm}** \u2014 {_ms}")

        else:  # fallback (layer show / hide now lives in the top bar)
            st.caption("Use the **Map layers** bar at the top of the page "
                       "to show or hide layers.")

    # values needed by the map regardless of which panel is open
    flags = _layer_flags()
    ig_live = bool(_sv("ig_live_v", True))
    ig_step = sim.state.step if ig_live else int(_sv("ig_step_v", 0))
    ig_rad = int(_sv("ig_rad_v", 1))
    with view_col:
        _views = ["2D map", "3D terrain"]
        _vcur = st.session_state.get("sim_view_sel", "2D map")
        vmode = st.radio("View", _views,
                         index=_views.index(_vcur) if _vcur in _views else 0,
                         horizontal=True, label_visibility="collapsed")
        st.session_state["sim_view_sel"] = vmode
        scale = _fit_scale(cfg.nx)
        playing = st.session_state.get("anim_on", False)
        if vmode == "3D terrain":
            # 3D is a pure viewer by design: browsers do not deliver click
            # events on 3D charts. Place ignitions on the 2D map view.
            _clk3, _ = _clock_info()
            st.caption(f"\u23f1 {_clk3} \u00b7 drag to rotate, scroll to "
                       "zoom; the camera survives the steps. To place "
                       "ignitions, switch to the 2D map and tick 'Click map "
                       "to place ignition'.")
            with _map_card():
                fig = viz.fire_surface_figure(world, sim=sim)
                key3d = f"plot3d_{st.session_state.map_version}"
                st.plotly_chart(fig, use_container_width=True,
                                config={"scrollZoom": True}, key=key3d)
        else:
            place = st.checkbox(
                "Click map to place ignition", value=False, key="sim_place",
                help="On: the map switches to a click canvas (same "
                     "mechanism as the Map editor) and every click drops "
                     "an ignition marker at that cell. Off: the map is the "
                     "pan / zoom view like 3D \u2014 drag = pan, mouse "
                     "wheel = zoom, toolbar top-right.")
            _clk, _nf = _clock_info()
            _dreg = st.session_state.get("dss_region")
            _rb, _rl = (_dreg[:4], _dreg[4]) if _dreg else (None, None)
            # THE OVERLAY IS DERIVED HERE, not inherited from a panel. It
            # used to be written only inside the Layer 2 panel, so changing
            # the agent count anywhere else left the map showing the old
            # split until that particular panel happened to be open. The
            # Layer 2 version carries an extra attendance flag per region;
            # it is kept when it still matches the current count.
            _rall = None
            # NOT BEFORE THERE ARE ANY. The default agent count is one, so
            # the map drew a single region over the whole world labelled
            # "Agent_1" on a scenario where no DSS had been set up at all:
            # a box around everything, named after a decision-maker that
            # does not exist yet. The overlay says something only once the
            # world is actually split, or once an engine is running.
            _eng_r = st.session_state.get("dss_engine")
            if bool(_sv("ly_agents_v", True)) and (
                    int(_sv("dss_n", 1)) > 1 or _eng_r is not None):
                _nreg = int(_sv("dss_n", 1))
                _rall = st.session_state.get("dss_regions_all")
                if bool(_sv("dss_show_all", True)):
                    if not _rall or len(_rall) != _nreg:
                        _rall = [(*r.box, r.name) for r in
                                 _dss.partition_n(cfg.nx, cfg.ny, _nreg)]
                        st.session_state["dss_regions_all"] = _rall
                else:
                    _rall = None
            _sens = (st.session_state.get("dss_sensors_draw")
                     if bool(_sv("ly_sens_v", True)) else None)
            _deps = (st.session_state.get("dss_depots_draw")
                     if bool(_sv("ly_deps_v", True)) else None)
            if not bool(_sv("ly_agents_v", True)):
                _rb, _rl = None, None
            _eng_m = st.session_state.get("dss_engine")
            _alloc = None
            _acts = None
            # only draw DSS orders when a resource POOL actually exists AND
            # the DSS is applying. When the pool is cleared the engine keeps
            # its last override cached (the decision cycle no longer runs),
            # which used to leave stale blue orders frozen on the map, so we
            # also drop that cache here.
            _pool_live = st.session_state.get("dss_res_base") is not None
            if (st.session_state.get("dss_apply") and _pool_live
                    and _eng_m is not None
                    and _eng_m.last_override is not None):
                if bool(_sv("ly_alloc_v", True)):
                    _alloc = _eng_m.last_override.rcap
                if bool(_sv("ly_orders_v", True)):
                    _acts = _eng_m.last_actions
                    # tag the regions where a GenAI-generated (G#) rule fired
                    # so the map can flag those orders distinctly
                    if isinstance(_acts, dict):
                        _acts["genai_regions"] = getattr(
                            _eng_m, "last_genai_regions", set())
                        # the macro definitions travel with the orders, so a
                        # generated intervention can be drawn in its OWN
                        # colour instead of one anonymous magenta badge
                        _acts["macros"] = dict(getattr(_eng_m, "macros", {})
                                               or {})
            elif (_eng_m is not None and not _pool_live
                    and getattr(_eng_m, "last_override", None) is not None):
                _eng_m.last_override = None
                _eng_m.last_actions = None
            if playing:
                # fast image frames while animating (keeps the loop
                # responsive); pause to pan / zoom
                with _map_card():
                    st.image(viz.render_pil(world, sim=sim, scale=scale,
                                            show_labels=True,
                                            clock_text=_clk,
                                            night_factor=_nf, region_box=_rb,
                                            region_label=_rl, regions=_rall,
                                            sensors=_sens, depots=_deps,
                                            alloc=_alloc, actions=_acts,
                                            **flags),
                             use_container_width=True)
            elif place and HAS_CANVAS:
                # ignition placement works exactly like the Map editor:
                # a click canvas over the rendered map; each click drops a
                # marker that is applied as an ignition
                bg = viz.render_pil(world, sim=sim, scale=scale,

                                    show_labels=True, clock_text=_clk,
                                    night_factor=_nf, region_box=_rb,
                                    region_label=_rl, regions=_rall,
                                    sensors=_sens, depots=_deps,
                                    alloc=_alloc, actions=_acts,
                                    **flags)
                with _map_card():
                    res = st_canvas(stroke_width=2, stroke_color="#a200de",
                                    background_image=bg, update_streamlit=True,
                                    height=cfg.ny * scale, width=cfg.nx * scale,
                                    drawing_mode="point", display_toolbar=False,
                                    point_display_radius=max(3, scale // 2),
                                    key=(f"simc_{st.session_state.canvas_key}_"
                                         f"{scale}"))
                if st.session_state.get("ign_warn"):
                    st.warning(st.session_state.pop("ign_warn"))
                if world.ignitions and st.button(
                        "Remove last ignition",
                        help="Deletes the most recently placed ignition "
                             "marker."):
                    world.ignitions.pop()
                    st.session_state.canvas_key += 1
                    st.session_state.sim_applied = 0
                    st.rerun()
                objs = (res.json_data or {}).get("objects", []) if res else []
                new = objs[st.session_state.get("sim_applied", 0):]
                if new:
                    for o in new:
                        if o.get("type") == "circle":
                            rad = o.get("radius", 0)
                            gx, gy = _clip((o["left"] + rad) / scale,
                                           (o["top"] + rad) / scale)
                            _wign2 = _ignition_warning(world, gx, gy,
                                                       int(ig_rad))
                            if _wign2:
                                st.session_state["ign_warn"] = _wign2
                            world.add_ignition(gx, gy, step=ig_step,
                                               radius=int(ig_rad))
                    # the canvas is NOT remounted between clicks (a
                    # remount swallowed the next click, so placing a
                    # second ignition needed re-ticking the box);
                    # instead only the objects beyond the applied
                    # counter are consumed, exactly like the editor
                    st.session_state.sim_applied = len(objs)
                    st.rerun()
            else:
                # plotly map exactly like the 3D view: top-right modebar
                # (zoom / pan / reset), drag to pan, wheel to zoom; the view
                # survives the steps via the figure uirevision
                _hov = st.checkbox(
                    "Cell tooltip on hover", value=True, key="map_hover",
                    help="Point at a cell to read its fuel, its burn "
                         "state, the assets on it, the DSS orders that "
                         "landed there and which agent owns it. It sends "
                         "one line per cell to the browser, so turn it off "
                         "if the page feels heavy on a large map.")
                # the plotly raster layer smooths on scale; pixelated
                # keeps cell edges honest, and the supersampled source
                # keeps them from looking jagged
                st.markdown("<style>.js-plotly-plot image "
                            "{ image-rendering: pixelated; }</style>",
                            unsafe_allow_html=True)
                with _map_card():
                    st.plotly_chart(
                        viz.map_figure_2d(world, sim=sim, scale=scale,
                                          hover=bool(_hov), engine=_eng_m,
                                          clock_text=_clk, night_factor=_nf,
                                          region_box=_rb, region_label=_rl,
                                          regions=_rall, sensors=_sens,
                                          depots=_deps, alloc=_alloc,
                                          actions=_acts,
                                          uirevision=f"m{st.session_state.map_version}",
                                          **flags),
                        use_container_width=True,
                        key=f"plot2d_{st.session_state.map_version}",
                        config={"scrollZoom": True,
                                "modeBarButtonsToRemove": ["lasso2d", "select2d"]})
                if world.ignitions and st.button(
                        "Remove last ignition",
                        help="Deletes the most recently placed ignition "
                             "marker."):
                    world.ignitions.pop()
                    st.rerun()
        # (the per-order explanation lived here; it duplicated the legend
        # below, so it was removed to declutter the screen)
        _export_panel(world, sim, key="expsim")
        st.markdown(legend_html(horizontal=True), unsafe_allow_html=True)

    # LAYOUT: the cost section (and the decision log) live in the LEFT column,
    # BELOW the map + legend (row 2 of column 1); column 2 holds the tab
    # buttons + the selected panel.
    with view_col:
        # DECISION LOG as a table a person can read: one row per region per
        # cycle, saying which local DSS ordered what and why, with the
        # coordinator's own decision for that cycle beside it. The previous
        # version was a wall of abbreviated text lines.
        _engDL = st.session_state.get("dss_engine")
        _cycDL = list(getattr(_engDL, "cycles", []) or [])
        if _cycDL:
            with st.expander(f"Decision log — who decided what, cycle by "
                             f"cycle ({len(_cycDL)})", expanded=False):
                _last = st.number_input(
                    "Show the last N cycles", 1, 500,
                    min(20, len(_cycDL)), 5, key="dlog_n")
                _rows_dl = []
                for _c in _cycDL[-int(_last):]:
                    _g = _c.get("global_dss") or {}
                    _hot = _g.get("hotspot")
                    _shares = _g.get("shares") or {}
                    _att = set(_g.get("attended") or [])
                    _ad = _c.get("adaptation") or {}
                    for _rn, _r in (_c.get("regions") or {}).items():
                        _u = _r.get("orders_final") or {}
                        _ordered = ", ".join(
                            f"{_k.replace('_', ' ')} {float(_v):.2f}"
                            for _k, _v in _u.items() if float(_v) > 0.02)
                        _fired = ", ".join(
                            str(_f[0]) for _f in (_r.get("fired") or [])[:5])
                        _rows_dl.append(dict(
                            step=int(_c.get("step", 0)),
                            t_min=round(float(_c.get("t_min", 0.0)), 1),
                            local_dss=_rn,
                            role=("HOTSPOT" if _rn == _hot
                                  else ("engaged" if _rn in _att
                                        else "monitored")),
                            share=round(float(_shares.get(_rn, 1.0)), 2),
                            ordered=_ordered or "nothing",
                            fired_rules=_fired or "none",
                            quality=round(float(_r.get("quality", 0.0)), 2),
                            failsafe=("attenuated" if _r.get("failsafe")
                                      else "full strength"),
                            withheld=("WITHHELD"
                                      if _c.get("no_harm_withheld")
                                      else ""),
                            adaptation=(
                                f"stage {_ad.get('tried')} "
                                + ("accepted" if _ad.get("accepted")
                                   else "rejected")
                                if _ad.get("tried") else ""),
                            global_dss=str(_g.get("statement") or "")[:150]))
                # same reason as the step table: this sits in the narrow left
                # column, and the coordinator's statement alone needs width
                if hasattr(st, "dialog"):
                    @st.dialog("Decision log", width="large")
                    def _dlog_window():
                        st.dataframe(_rows_dl, use_container_width=True,
                                     height=620)
                        _cd = list(_rows_dl[0].keys()) if _rows_dl else []
                        st.download_button(
                            "⬇ Download as Excel",
                            _xlsx_bytes({"Decision log": [
                                {c: r.get(c) for c in _cd}
                                for r in _rows_dl]}),
                            file_name="decision_log.xlsx", mime=XLSX_MIME,
                            key="dl_declog_xlsx")
                    if st.button("Open in a window", key="open_dlog",
                                 use_container_width=True):
                        _dlog_window()
                st.dataframe(_rows_dl, use_container_width=True, height=300)
                st.caption(
                    "**local_dss** is the regional agent that produced the "
                    "orders. **role** is the coordinator's call for that "
                    "cycle: the hotspot leads, engaged regions act at full "
                    "tempo, monitored ones keep watching with their "
                    "offensive share cut to **share**. **failsafe** says "
                    "whether the quality gate attenuated the orders, "
                    "**withheld** whether the no-harm guard pulled the "
                    "offensive allocation for that whole cycle. "
                    "**global_dss** is the coordinator's statement in its "
                    "own words.")
        # the cost panel moved into a dialog opened from the DSS
        # Dashboard (see _cost_dialog): the jury view keeps the map
        # tall and the numbers one click away


def _selection_points(ev):
    sel = getattr(ev, "selection", None)
    if sel is None and isinstance(ev, dict):
        sel = ev.get("selection")
    if not sel:
        return []
    if isinstance(sel, dict):
        return sel.get("points", []) or []
    return getattr(sel, "points", []) or []


def _place_from_selection(ev, ig_step, ig_rad, scale=None):
    """Drop an ignition from a plotly click selection. De-duplicates so the same
    selection is not re-applied on every rerun."""
    pts = _selection_points(ev)
    if not pts:
        return
    px, py = pts[0].get("x", 0), pts[0].get("y", 0)
    if scale:
        px, py = px / scale, py / scale
    gx, gy = _clip(round(px), round(py))
    sig = (gx, gy)
    if st.session_state.get("last_sel_sig") == sig:
        return
    st.session_state["last_sel_sig"] = sig
    _wign = _ignition_warning(world, gx, gy, ig_rad)
    if _wign:
        st.warning(_wign)
    world.add_ignition(gx, gy, step=ig_step, radius=ig_rad)
    st.toast(f"Ignition at ({gx}, {gy})")
    st.rerun()


_J_TERMS = [
    ("Jᵇᵘʳⁿ", "j_burn", "#2e8b57", "burned area"),
    ("Jᵃˢˢᵉᵗ", "j_asset", "#d9822b", "asset loss"),
    ("Jᵖᵒᵖ", "j_pop", "#8e44ad", "population"),
    ("Jʳᵉˢᵖ", "j_resp", "#2c3e50", "response cost"),
    ("Jᵈᵉˡ", "j_delay", "#7f8c8d", "response delay"),
]


@st.dialog("Cost function — settings & charts", width="large")
def _cost_dialog():
    """The whole cost story in a popup: the J_k equation, the
    operational priority weights, the protection targeting weights,
    the advanced thresholds and the per-term time charts. Opened from
    the DSS Dashboard so the map page itself stays uncluttered."""
    _cost_panel()


def _cost_panel():
    import matplotlib.pyplot as plt

    # `_sv` is a helper defined inside page_simulation, and Python does not
    # hand a caller's locals to a module-level function. The custom-weight
    # branch below referenced it and would have raised NameError the moment
    # anyone picked the Custom cost profile.
    def _sv(key, default):
        return st.session_state.get(key, default)

    thr = float(sim.cfg.cost.acceptance_fraction)
    st.divider()
    st.subheader("Cost function $J_k$")
    st.latex(r"J_k=w_1 J_k^{burn}+w_2 J_k^{asset}+w_3 J_k^{pop}"
             r"+w_4 J_k^{resp}+w_5 J_k^{del}")
    st.caption("Normalized cost-plus-loss of the run so far (System "
 "Description): each term is divided by its scenario "
               "reference scale, so every term and the weighted total lie in "
               "$[0,1]$. The weights encode operational priority. The dashed "
               f"line marks the acceptance threshold ({thr:g} of the "
               "do-nothing cost).")

    # ---- operational priority: doctrine presets for the weights ----
    _wpre = {
        "Balanced (default)":            (0.294, 0.294, 0.294,
                                          0.059, 0.059),
        "Life first":                    (0.20, 0.20, 0.45,
                                          0.075, 0.075),
        "Assets & infrastructure first": (0.20, 0.45, 0.25,
                                          0.05, 0.05),
        "Environment first":             (0.45, 0.20, 0.25,
                                          0.05, 0.05),
        "Custom":                        None,
    }
    _wmode = st.selectbox(
        "Operational priority \u2014 how the five losses are "
        "weighted", list(_wpre), key="cost_wmode",
        help="A doctrine choice assigns consistent weights (they "
             "sum to 1). 'Life first' makes population exposure "
             "dominant, 'Assets first' the built environment, "
             "'Environment first' the burned area. Custom frees "
             "the five weights and renormalizes them to sum to 1. "
             "The choice steers the DSS: forecasts, satisficing, "
             "adaptation gates and the counterfactual all read "
             "this J. All cost settings live here, not in "
             "Parameters.")
    _cwm = sim.cfg.cost
    if _wpre[_wmode] is not None:
        (_cwm.w_burn, _cwm.w_asset, _cwm.w_pop,
         _cwm.w_resp, _cwm.w_delay) = _wpre[_wmode]
    else:
        st.caption("Enter any five values; they are renormalized so the "
                   "weights sum to exactly 1.")
        _mw = st.columns(5)
        _defs = [("burn", _cwm.w_burn), ("asset", _cwm.w_asset),
                 ("pop", _cwm.w_pop), ("resp", _cwm.w_resp),
                 ("delay", _cwm.w_delay)]
        _vals = []
        for _i_w, (_nm_w, _dv_w) in enumerate(_defs):
            _vals.append(float(_mw[_i_w].number_input(
                f"w_{_nm_w}", 0.0, 1.0,
                float(_sv(f"cost_w_{_nm_w}",
                          round(float(_dv_w), 3))),
                0.05, key=f"cost_w_{_nm_w}")))
        _vsm = float(sum(_vals))
        if _vsm <= 1e-9:
            st.error("All weights are zero \u2014 falling back to Balanced.")
            _vals = list(_wpre["Balanced (default)"])
            _vsm = 1.0
        # store the NORMALIZED weights so cfg.cost weights literally sum to 1
        _vals = [v / _vsm for v in _vals]
        (_cwm.w_burn, _cwm.w_asset, _cwm.w_pop,
         _cwm.w_resp, _cwm.w_delay) = _vals
    _wsm = (_cwm.w_burn + _cwm.w_asset + _cwm.w_pop
            + _cwm.w_resp + _cwm.w_delay)
    if _wsm <= 1e-9:
        st.error("All weights are zero \u2014 J would be "
                 "undefined; falling back to Balanced.")
        (_cwm.w_burn, _cwm.w_asset, _cwm.w_pop,
         _cwm.w_resp, _cwm.w_delay) = _wpre["Balanced (default)"]
        _wsm = 1.0
    _shr = _pct_to_100([_cwm.w_burn, _cwm.w_asset, _cwm.w_pop,
                        _cwm.w_resp, _cwm.w_delay])
    st.caption(
        "Normalized shares (sum to 100%): burn "
        f"{_shr[0]}% \u00b7 asset {_shr[1]}% \u00b7 population {_shr[2]}% \u00b7 "
        f"response {_shr[3]}% \u00b7 delay {_shr[4]}%")
    if (_cwm.w_pop / _wsm) < 0.15:
        st.warning("Life-safety share below 15% is unusual for an "
                   "operational doctrine \u2014 make sure this is "
                   "intended.")
    if ((_cwm.w_resp + _cwm.w_delay) / _wsm) > 0.4:
        st.warning("Response cost + delay above 40% makes the "
                   "optimizer reluctant to field resources at all "
                   "\u2014 the fire terms should dominate.")
    # ---- protection priority weights (V_prio) ----
    # The SAME value judgment as the loss weights above, applied one
    # step earlier: J decides WHAT a bad outcome is, V_prio decides
    # WHERE the protective effort goes on the map (asset-protection
    # targeting and the asset exposure feature both read it). The two
    # belong side by side so a doctrine choice is made in one place.
    with st.container(border=True):
        st.markdown("**Protection priority weights** $V_{prio}$")
        st.caption("How the protective effort is targeted across what "
                   "is at risk. Dimensionless, renormalized to sum "
                   "to 1. A decision input, not a fire-behaviour "
                   "parameter; steers asset protection placement and "
                   "the asset exposure feature.")
        _vw = sim.cfg.value_weights
        _va, _vb = st.columns(2)
        _vw.w_crit = _va.number_input(
            "$w_{crit}$ — critical facility", 0.0, 1.0,
            float(_vw.w_crit), 0.05, key="vw_crit")
        _vw.w_pop = _vb.number_input(
            "$w_{pop}$ — population", 0.0, 1.0,
            float(_vw.w_pop), 0.05, key="vw_pop")
        _vw.w_bld = _va.number_input(
            "$w_{bld}$ — building", 0.0, 1.0,
            float(_vw.w_bld), 0.05, key="vw_bld")
        _vw.w_evac = _vb.number_input(
            "$w_{evac}$ — evacuation", 0.0, 1.0,
            float(_vw.w_evac), 0.05, key="vw_evac")
        _vsum = _vw.w_crit + _vw.w_pop + _vw.w_bld + _vw.w_evac
        if _vsum > 1e-9:
            _vs = _pct_to_100([_vw.w_crit, _vw.w_pop, _vw.w_bld,
                               _vw.w_evac])
            st.caption(f"Normalized shares (sum 100%): critical "
                       f"{_vs[0]}% · population {_vs[1]}% · building "
                       f"{_vs[2]}% · evacuation {_vs[3]}%")
        else:
            st.warning("All weights are zero — set at least one "
                       "above 0.")
    # the cost model's non-weight safeguards + reference scales live here too
    # (moved out of Parameters, so all cost settings are in one place)
    with st.expander("Advanced cost settings (thresholds & reference scales)"):
        _ac1, _ac2 = st.columns(2)
        _cwm.acceptance_fraction = _ac1.number_input(
            "acceptance threshold (fraction of do-nothing)", 0.0, 1.0,
            float(_cwm.acceptance_fraction), 0.05, format="%.2f",
            key="adv_accept")
        _cwm.population_at_risk_fraction = _ac2.number_input(
            "population at risk fraction", 0.0, 1.0,
            float(_cwm.population_at_risk_fraction), 0.005, format="%.3f",
            key="adv_rho")
        _cwm.horizon_steps = _ac1.number_input(
            "scenario horizon (steps)", 1.0, 5000.0,
            float(_cwm.horizon_steps), 10.0, key="adv_horizon")
        _cwm.capacity_reference = _ac2.number_input(
            "total capacity (response-cost reference)", 0.0, 1e6,
            float(_cwm.capacity_reference), 10.0, key="adv_caps")
        _cwm.delay_reference = _ac1.number_input(
            "reference delay (response-delay reference)", 0.0, 1e5,
            float(_cwm.delay_reference), 5.0, key="adv_delayref")
    rep = compute_costs(sim)
    d = rep.to_dict()

    # physical impact
    _pref = float(getattr(rep, "population_reference", 0.0) or 0.0)
    _pev = float(getattr(rep, "population_evacuated", 0.0) or 0.0)
    m = st.columns(6)
    m[0].metric("Burned area (ha)", f"{rep.burned_area_ha:,.1f}")
    m[1].metric("Burned forest (ha)", f"{rep.burned_forest_ha:,.1f}")
    # THE DENOMINATOR BELONGS ON SCREEN. "exposed 4,417" and "evacuated
    # 1,426" were shown with nothing to read them against, so neither could
    # be judged and the two looked like they should add up to something.
    m[2].metric("Population at risk", f"{_pref:,.0f}",
                help="The population the fire started with. J_pop is "
                     "normalized by this, so it is the denominator for "
                     "both figures beside it.")
    m[3].metric("Population exposed",
                f"{rep.population_exposed:,.0f}"
                + (f"  ({100.0 * rep.population_exposed / _pref:.0f}%)"
                   if _pref > 0 else ""),
                help="People inside cells that have burned. A HEADCOUNT: "
                     "J_pop is not this number, it is the exposure "
                     "integrated over time (person-steps), so a short "
                     "exposure of many people and a long exposure of few "
                     "can cost the same.")
    m[4].metric("Evacuated (safe)",
                f"{_pev:,.0f}"
                + (f"  ({100.0 * _pev / _pref:.0f}%)" if _pref > 0 else ""),
                help="People moved out by evacuation orders; they leave "
                     "vpop, so they stop accumulating exposure from that "
                     "moment. It cannot undo the person-steps already "
                     "accrued before the order landed, which is why a late "
                     "evacuation shows a large evacuated count and still "
                     "leaves J_pop high. Cells in or beside active flame "
                     "empty at ~30%/min under an order, elsewhere ~5%/min, "
                     "and a public warning roughly doubles that tempo.")
    m[5].metric("Asset value lost",
                f"{rep.asset_value_lost:,.1f} / {rep.asset_value_total:,.1f}")

    with st.expander("How each term is computed \u2014 actual "
                     "value, unit, reference, normalization"):
        _cw = sim.cfg.cost
        _bref_c = float(getattr(_cw, "burn_reference_fraction", 0.02))
        _aref_ha = _bref_c * rep.burnable_cells * sim.cfg.cell_area_ha
        _H = float(_cw.horizon_steps)
        _wsum = (_cw.w_burn + _cw.w_asset + _cw.w_pop
                 + _cw.w_resp + _cw.w_delay)
        _rows = [
            ("J_burn", f"{rep.burned_area_ha:,.1f} ha burned "
             f"({rep.burned_forest_ha:,.1f} ha forest)",
             f"A_ref = {_aref_ha:,.0f} ha (5% of burnable; 'major "
             "fire')",
             "x/(1+x), x = A/A_ref \u2014 rational saturation: "
             "0.5 at the reference fire, gradient survives at ANY "
             "size (the old 1-exp form went numerically flat past "
             "3\u00d7 the reference and the optimizer lost its "
             "incentive on big fires)",
             rep.j_burn, _cw.w_burn),
            ("J_asset", f"{rep.asset_value_lost:,.2f} of "
             f"{rep.asset_value_total:,.2f} value units lost",
             "total building + critical-infrastructure value",
             "lost / total", rep.j_asset, _cw.w_asset),
            ("J_pop", f"{rep.population_person_steps:,.0f} "
             "person-steps in burning cells "
             f"(exposed now {rep.population_exposed:,.0f}, "
             f"evacuated {float(getattr(rep, 'population_evacuated', 0)):,.0f}"
             " and no longer accumulating)",
             f"population AT RISK "
             f"{float(getattr(rep, 'population_reference', 0)):,.0f} "
             f"\u00d7 H={_H:g} steps",
             "\u03a3 exposed / (pop_at_risk \u00d7 H)", rep.j_pop,
             _cw.w_pop),
            ("J_resp", f"{rep.committed_capacity:,.0f} capacity now "
             "fielded (effort integrates over time)",
             f"staged pool \u00d7 1.2 = {rep.available_capacity:,.0f}"
             f", \u00d7 H={_H:g}",
             "\u03a3 committed / (ref \u00d7 H) \u2014 an army "
             "held for hours costs more than a strike", rep.j_resp,
             _cw.w_resp),
            ("J_del", f"{rep.mean_response_delay:,.1f} min mean "
             "dispatch (capacity-weighted)",
             f"delay_reference = {_cw.delay_reference:g} min",
             "mean delay / reference", rep.j_delay, _cw.w_delay),
        ]
        for _nm, _act, _ref, _frm, _jv, _wv in _rows:
            st.markdown(
                f"**{_nm}** = {_jv:.3f} \u00b7 weight {_wv:g} "
                f"\u2192 contribution {(_wv * _jv / _wsum):.3f}  \n"
                f"actual: {_act}  \n"
                f"reference: {_ref}  \n"
                f"normalization: {_frm}")
        st.caption(
            "J_resp and J_del price the response, never forbid it: "
            "their weights are small (0.2 vs 1.0) and J_burn is "
            "deliberately expensive, so 'never intervene' loses "
            "whenever the fire is real. The no-harm guard compares "
            "candidates on the PHYSICAL terms only (burn + asset + "
            "pop).")

    # normalized J terms and the total, all in [0,1]
    cc = st.columns(len(_J_TERMS) + 2)
    for i, (lab, key, _c, sub) in enumerate(_J_TERMS):
        cc[i].metric(f"{lab} · {sub}", f"{d[key]:.3f}")
    cc[-2].metric("J · DECISION", f"{rep.j_total:.3f}",
                  help="The 5-term cost the DSS optimizes (response "
                       "cost and delay INCLUDED: a commander "
                       "economizes the fleet).")
    cc[-1].metric("Jφ · OUTCOME",
                  f"{getattr(rep, 'j_physical', 0.0):.3f}",
                  help="Physical loss only (burn + asset + pop). "
                       "THE fair metric for comparing runs: a no-DSS "
                       "run pays no response cost by definition, so "
                       "J decision would reward doing nothing. Every "
                       "with/without comparison reads THIS.")

    series = st.session_state.cost_series
    if len(series) > 1:
        steps = [r["step"] for r in series]
        st.markdown("##### $J$ terms over time — one chart per term")
        titles = {k: lab for lab, k, _c, _s in _J_TERMS}
        cells = st.columns(3)
        panels = [(lab, key, col, titles[key]) for lab, key, col, _ in _J_TERMS]
        panels.append(("J", "j_total", "#111111", "J (decision)"))
        panels.append(("Jφ", "j_physical", "#b03a2e",
                       "Jφ (outcome, fair)"))
        for i, (lab, key, col, ttl) in enumerate(panels):
            with cells[i % 3]:
                f, a = plt.subplots(figsize=(3.4, 2.0))
                a.plot(steps, [r[key] for r in series], color=col, lw=1.8)
                a.fill_between(steps, [r[key] for r in series],
                               color=col, alpha=0.15)
                if key == "j_total":
                    a.axhline(thr, color="#c0392b", lw=1.0, ls="--")
                a.set_title(ttl, fontsize=11)
                a.set_xlabel("step k", fontsize=8)
                a.set_ylim(-0.02, 1.02)
                a.tick_params(labelsize=7)
                a.grid(alpha=0.25)
                f.tight_layout()
                st.pyplot(f, use_container_width=True)
                plt.close(f)


# =============================================================== MAP EDITOR ==
#: what the manager may edit, and what it is called on screen
_ASSET_COLS = ["keep", "kind", "name", "x", "y", "radius", "value",
               "population"]


def _asset_manager(world) -> None:
    """List, rename, move and delete everything the map holds.

    A generated map arrives with a couple of dozen assets and no way to
    touch any of them: the only tool was to paint another one on top. So a
    hospital in the wrong place stayed there, a name could not be corrected,
    and a settlement could not be thinned out.

    Deleting or moving is not a matter of the list alone. add_asset WRITES
    into the value layers with np.maximum and nothing takes a written value
    back out, so the layers are rebuilt from the edited list; otherwise a
    deleted hospital goes on being worth protecting where it used to stand.
    """
    _as = list(getattr(world, "assets", []) or [])
    with st.expander(f"Assets on this map ({len(_as)})", expanded=False):
        if not _as:
            st.caption("None yet. Click the map to place one.")
            return
        _rows = [dict(keep=True, kind=str(getattr(a, "kind", "")),
                      name=str(getattr(a, "name", "")),
                      x=int(getattr(a, "x", 0)), y=int(getattr(a, "y", 0)),
                      radius=int(getattr(a, "radius", 0)),
                      value=float(getattr(a, "value", 0.0) or 0.0),
                      population=float(getattr(a, "population", 0.0) or 0.0))
                 for a in _as]
        st.caption("Untick **keep** to delete. Edit a name, a position, a "
                   "radius or a value in place. Nothing changes until you "
                   "press Apply.")
        _ed = st.data_editor(
            _rows, key="asset_mgr", use_container_width=True, height=280,
            column_order=_ASSET_COLS, hide_index=True,
            disabled=["kind"],
            column_config={
                "keep": st.column_config.CheckboxColumn("keep", width="small"),
                "kind": st.column_config.TextColumn("kind", width="small"),
                "name": st.column_config.TextColumn("name", width="medium"),
                "x": st.column_config.NumberColumn(
                    "x", min_value=0, max_value=int(world.config.nx) - 1,
                    step=1, width="small"),
                "y": st.column_config.NumberColumn(
                    "y", min_value=0, max_value=int(world.config.ny) - 1,
                    step=1, width="small"),
                "radius": st.column_config.NumberColumn(
                    "radius", min_value=0, max_value=60, step=1,
                    width="small"),
                "value": st.column_config.NumberColumn(
                    "value", min_value=0.0, max_value=1.0, step=0.05,
                    width="small"),
                "population": st.column_config.NumberColumn(
                    "population", min_value=0.0, step=100.0,
                    width="small")})
        c1, c2 = st.columns([1, 2])
        if c1.button("Apply asset edits", use_container_width=True,
                     type="primary", key="asset_mgr_apply"):
            from disaster_phyengine.world import Asset
            _new = []
            # the rows and the assets are 1:1 (the editor adds no rows), so
            # the settlement tag rides along instead of being dropped and
            # breaking every town on the map into loose markers
            for r, _src in zip(_ed, _as):
                if not bool(r.get("keep", True)):
                    continue
                _new.append(Asset(
                    str(r.get("name") or "Asset"),
                    str(r.get("kind") or "building"),
                    int(np.clip(int(r.get("x", 0)), 0,
                                int(world.config.nx) - 1)),
                    int(np.clip(int(r.get("y", 0)), 0,
                                int(world.config.ny) - 1)),
                    int(max(0, int(r.get("radius", 0)))),
                    value=float(r.get("value", 0.0) or 0.0),
                    population=float(r.get("population", 0.0) or 0.0),
                    group=str(getattr(_src, "group", "") or "")))
            _dropped = len(_as) - len(_new)
            world.assets = _new
            # THE LAYERS ARE DERIVED, so they are rebuilt rather than patched
            world.rebuild_value_layers()
            st.session_state.map_version = \
                st.session_state.get("map_version", 0) + 1
            _reset_dss_state(drop_engine=True)
            c2.success(f"{len(_new)} asset(s) kept"
                       + (f", {_dropped} deleted" if _dropped else "")
                       + " — value layers rebuilt.")
            st.rerun()
        c2.caption("Deleting or moving rebuilds the protection value from "
                   "the list, so nothing is left behind at the old place. "
                   "The DSS is reset because its regions and its learned "
                   "targets refer to what was there before.")


def _settlement_manager(world) -> None:
    """List, rename, move and delete whole settlements.

    A town is not a marker. It is a block of built-up fuel, a street grid,
    its residents and its civic facilities, and the asset list could only
    ever delete those one at a time: the block of urban ground stayed
    behind, so the map kept a town-shaped patch that still burned like a
    town and no longer cost anything when it did. Here a settlement is one
    row, and removing it takes the ground back to the cover around it.
    """
    from disaster_phyengine import terrain as _tr
    _sets = _tr.settlements(world)
    with st.expander(f"Settlements on this map ({len(_sets)})",
                     expanded=False):
        if not _sets:
            st.caption("None. Place one with the Settlement tool.")
            return
        _rows = [dict(keep=True, name=str(v["name"]), x=int(v["x"]),
                      y=int(v["y"]), population=float(v["population"]),
                      facilities=int(v["facilities"]), parts=int(v["parts"]))
                 for v in _sets.values()]
        st.caption("Untick **keep** to remove the town, its people, its "
                   "facilities and its built-up ground. Change x/y to move "
                   "it (the block is repainted at the new place). Rename in "
                   "place. Nothing changes until you press Apply.")
        _ed = st.data_editor(
            _rows, key="sett_mgr", use_container_width=True, height=210,
            column_order=["keep", "name", "x", "y", "population",
                          "facilities", "parts"],
            hide_index=True, disabled=["facilities", "parts"],
            column_config={
                "keep": st.column_config.CheckboxColumn("keep", width="small"),
                "name": st.column_config.TextColumn("name", width="medium"),
                "x": st.column_config.NumberColumn(
                    "x", min_value=0, max_value=int(world.config.nx) - 1,
                    step=1, width="small"),
                "y": st.column_config.NumberColumn(
                    "y", min_value=0, max_value=int(world.config.ny) - 1,
                    step=1, width="small"),
                "population": st.column_config.NumberColumn(
                    "people", min_value=0.0, step=100.0, width="small"),
                "facilities": st.column_config.NumberColumn(
                    "facilities", width="small"),
                "parts": st.column_config.NumberColumn(
                    "assets", width="small")})
        c1, c2 = st.columns([1, 2])
        if c1.button("Apply settlement edits", use_container_width=True,
                     type="primary", key="sett_mgr_apply"):
            _keys = list(_sets)
            _rm = _mv = _rn = 0
            for _k, r in zip(_keys, _ed):
                _old = _sets[_k]
                if not bool(r.get("keep", True)):
                    _tr.remove_settlement(world, _k)
                    _rm += 1
                    continue
                _nx = int(np.clip(int(r.get("x", _old["x"])), 0,
                                  int(world.config.nx) - 1))
                _ny = int(np.clip(int(r.get("y", _old["y"])), 0,
                                  int(world.config.ny) - 1))
                _nm = str(r.get("name") or _k)
                if _nx != int(_old["x"]) or _ny != int(_old["y"]):
                    # MOVING IS A REBUILD. The block, the streets and the
                    # facility ring are painted onto the terrain, so they
                    # have to be unpainted here and painted there.
                    _tr.move_settlement(world, _k, _nx, _ny, name=_nm)
                    _mv += 1
                elif _nm != _k:
                    for _a in world.assets:
                        if str(getattr(_a, "group", "")) == _k:
                            _a.name = str(_a.name).replace(_k, _nm)
                            _a.group = _nm
                    _rn += 1
            world.rebuild_value_layers()
            st.session_state.map_version = \
                st.session_state.get("map_version", 0) + 1
            _reset_dss_state(drop_engine=True)
            c2.success(f"{_rm} removed, {_mv} moved, {_rn} renamed.")
            st.rerun()
        c2.caption("Removing a settlement returns its ground to the cover "
                   "around it and rebuilds the protection value. The roads "
                   "stay: the road through a town is also the road past it.")


def _map_library(world) -> None:
    """Save the map under a name, open a saved one, pick the default.

    A generated landscape used to be a throwaway: the only way to keep one
    was to download a scenario file and upload it again next session, and
    the app always opened on the same procedural mountain map whatever had
    been built in the editor.
    """
    from disaster_phyengine import maplib as _ml
    try:
        _maps = _ml.list_maps()
    except Exception as _e:
        st.warning(f"Map library unavailable: {_e}")
        return
    _cur = st.session_state.get("_opened_map") or ""
    with st.expander(f"Map library ({len(_maps)} saved)"
                     + (f" \u2014 open: {_cur}" if _cur else ""),
                     expanded=False):
        c1, c2, c3 = st.columns([2, 2, 1])
        _nm = c1.text_input("Name", value=_cur, key="lib_name",
                            placeholder="e.g. Marmaris kiyisi")
        _nt = c2.text_input("Note (optional)", key="lib_note",
                            placeholder="what this map is for")
        c3.markdown("&nbsp;", unsafe_allow_html=True)
        if c3.button("Save map", use_container_width=True, type="primary",
                     key="lib_save"):
            if not _nm.strip():
                c3.error("Name?")
            else:
                _r = _ml.save_map(world, _nm, _nt)
                st.session_state["_opened_map"] = _r["name"]
                st.success(f"Saved **{_r['name']}** "
                           f"({_r['nx']}x{_r['ny']} at {_r['cell_m']:.0f} m, "
                           f"{_r['settlements']} settlement(s), "
                           f"{_r['assets']} asset(s)).")
                st.rerun()
        st.caption("Saving writes the WHOLE map: terrain, cover, roads, "
                   "assets, settlements and scheduled ignitions. Saving "
                   "under a name that already exists replaces it.")

        if not _maps:
            st.info("Nothing saved yet.")
            return
        st.divider()
        st.dataframe(
            [{"default": bool(m.get("default")), "name": m.get("name"),
              "size": f"{m.get('nx', '?')}x{m.get('ny', '?')}"
                      f" @ {float(m.get('cell_m', 0)):.0f} m",
              "extent": f"{m.get('km', '?')} km", "towns": m.get("settlements"),
              "assets": m.get("assets"), "saved": m.get("saved"),
              "note": m.get("note")} for m in _maps],
            use_container_width=True, hide_index=True)
        _names = [str(m.get("name")) for m in _maps]
        _pick = st.selectbox("Map", _names, key="lib_pick")
        d1, d2, d3 = st.columns(3)
        if d1.button("Open", use_container_width=True, key="lib_open"):
            _new_simulator(_ml.load_map(_pick))
            st.session_state["_opened_map"] = _pick
            st.rerun()
        _isdef = any(m.get("default") and str(m.get("name")) == _pick
                     for m in _maps)
        if d2.button("Unset default" if _isdef else "Make default",
                     use_container_width=True, key="lib_def",
                     help="The default map opens with the app."):
            _ml.set_default(None if _isdef else _pick)
            st.rerun()
        if d3.button("Delete", use_container_width=True, key="lib_del"):
            # TWO PRESSES, NOT ONE. A map can be an afternoon's editing and
            # the button sits next to Open.
            if st.session_state.get("_lib_del_armed") == _pick:
                _ml.delete_map(_pick)
                st.session_state.pop("_lib_del_armed", None)
                st.rerun()
            st.session_state["_lib_del_armed"] = _pick
        if st.session_state.get("_lib_del_armed") == _pick:
            st.warning(f"Press Delete again to remove **{_pick}** "
                       "from the library.")
        st.caption(f"Files: `{_ml.library_dir()}`")


def page_editor():
    _map_library(world)
    with st.expander("New map / scenario \u2014 generate, load, import",
                     expanded=False):
        st.caption("Procedural landscapes for experiments. For a real region "
                   "(e.g. an actual Turkish province) import its elevation and "
                   "fuel rasters on the GIS import page; everything else "
                   "(assets, roads, ignition) is then edited here.")
        src = st.radio("Source", ["Landscape type", "Built in scenario",
                                  "Blank grid"])
        nx = st.number_input("Resolution X (nx)", 20, 600, 200, 10)
        ny = st.number_input("Resolution Y (ny)", 20, 600, 200, 10)
        cell = st.number_input(
            "Cell size (m)", 1.0, 1000.0, 30.0, 1.0,
            help="30 m is the wildfire reference calibration: the "
                 "spread physics (m/min), travel times, station and "
                 "flight radii and the substep logic all reference "
                 "30 m and AUTO-SCALE to other sizes. Pick the cell "
                 "size for the DETAIL you want; pick nx/ny for the "
                 "AREA you want.")
        _ext_x = nx * cell / 1000.0
        _ext_y = ny * cell / 1000.0
        st.caption(f"Physical extent: **{_ext_x:.1f} \u00d7 "
                   f"{_ext_y:.1f} km** ({_ext_x * _ext_y * 100:.0f} "
                   "ha). Everything cell-size-dependent scales "
                   "automatically (ROS, dispatch/flight minutes, "
                   "service radii, burn reference = % of burnable "
                   "area). A smaller cell does NOT change the "
                   "physics \u2014 it only refines the grid of a "
                   "smaller or same area.")
        if cell < 15:
            st.info(f"{cell:.0f} m cells over {nx}\u00d7{ny} gives "
                    f"only {_ext_x:.1f}\u00d7{_ext_y:.1f} km \u2014 "
                    "a neighbourhood, not a landscape. For wildfire "
                    "scenarios 30 m cells are recommended; increase "
                    "nx/ny instead if you need a larger area.")
        if src == "Landscape type":
            _D2R = 0.017453292519943295

            st.markdown("**Terrain — $U_{Geo}$**")
            g1, g2 = st.columns(2)
            relief = g1.slider(
                "Relief: plain ↔ mountain (m)", 20, 1200, 450, 10,
                help="Peak-to-valley elevation. Low = flat plain, high = "
                     "steep mountains where fire runs faster uphill.")
            access = g2.slider(
                "Accessibility", 0.0, 1.0, 1.0, 0.05,
                help="Scales the per-cell access field $A(x,y)\\in[0,1]$: "
                     "how quickly ground crews can reach EACH cell, "
                     "including forest interior (terrain sets it from "
                     "slope, roads force $A=1$ along them). $A=1$ means "
                     "every cell is reachable at nominal speed; lower "
                     "values slow suppression and raise the response-cost "
                     "and delay terms. It is not about villages: it is the "
                     "crew-reachability of the ground itself.")

            st.markdown("**Weather — $U_{Meteo}$**")
            m1, m2, m3 = st.columns(3)
            wspd = m1.slider("Wind speed (m/s)", 0.0, 30.0, 8.0, 0.5)
            wdir = m2.slider(
                "Wind direction (°)", 0, 359, 0, 5,
                help="Direction the wind blows toward (0 = +x / east).")
            moist = m3.slider(
                "Fuel moisture", 0.02, 0.40, 0.08, 0.01,
                help="Higher moisture slows spread and lowers intensity.")
            prec = st.slider(
                "Precipitation (mm/h)", 0.0, 30.0, 0.0, 1.0,
                help="Rain falls in scattered showers. While it rains the "
                     "engine drives fuel moisture up toward $0.35 > m_{ext}$ "
                     "(so the fire stalls and dies under sustained rain) and "
                     "ember spotting stops above $1$ mm/h.")

            st.markdown("**Vegetation / fuel — $U_{Fuel}$**")
            f1, f2 = st.columns(2)
            forest = f1.slider(
                "Forest density", 0.0, 1.0, 0.45, 0.05,
                help="Fraction of land covered by forest fuel clusters.")
            water = f2.slider(
                "Water level", 0.0, 0.40, 0.06, 0.02,
                help="Fraction of the lowest cells turned into water "
                     "(non-flammable).")
            fw1, fw2 = st.columns(2)
            river = fw1.checkbox("River", value=False)
            coast = fw2.checkbox("Coast (sea on east edge)", value=False)

            st.markdown("**Settlements & values at risk — $U_{Val}$**")
            v1, v2 = st.columns(2)
            nvill = v1.slider(
                "Number of settlements", 0, 20, 6, 1,
                help="Villages plus a central town (the first one). "
                     "0 = wildland only, no people or structures. "
                     "Settlements are scattered with blue-noise spacing "
                     "on suitable land (flat, low, near water). Keep this "
                     "small (a handful) so towns stay spread out and their "
                     "labels do not overlap.")
            popv = v2.slider(
                "Total population", 0, 500000, 60000, 5000,
                help="TOTAL population of the whole map, split across the "
                     "settlements with a skewed share: the town takes the "
                     "largest part, villages get smaller shares; the parts "
                     "sum exactly to this value.")
            bscale = st.slider(
                "Building / critical facility density", 0.0, 2.0, 1.0, 0.1,
                help="Scales the footprint of buildings and critical "
                     "facilities AND how many a settlement carries. At 0 "
                     "the settlements are houses only.")
            farmv = st.slider(
                "Farmland density (0 = none)", 0.0, 2.0, 1.0, 0.1,
                key="gen_farm",
                help="Cultivated parcels on the workable ground near the "
                     "settlements: how often a workable block is sown and "
                     "how far from the town the fields reach. A worked "
                     "field carries about half the fine fuel of natural "
                     "grass and a little more moisture, so it slows a "
                     "front. 0 gives a pure wildland scenario.")

            # the seed sits NEXT TO the generate button: it is what one
            # changes between quick trials, so no scrolling in between
            _sgA, _sgB = st.columns([1, 1.6],
                                    vertical_alignment="bottom")
            seed = _sgA.number_input("Seed", 0, 99999, 42)
            if _sgB.button("Generate map", use_container_width=True,
                           type="primary"):
                _new_simulator(terrain.generate_landscape(
                    SimConfig(nx=int(nx), ny=int(ny), cell_size_m=float(cell)),
                    seed=int(seed), relief_m=float(relief),
                    forest_density=float(forest), base_moisture=float(moist),
                    water_level=float(water), wind_speed=float(wspd),
                    wind_dir_rad=float(wdir) * _D2R,
                    precipitation=float(prec),
                    with_assets=int(nvill) > 0, with_roads=True,
                    n_settlements=int(nvill),
                    population_per_settlement=int(popv),
                    building_scale=float(bscale), accessibility=float(access),
                    farmland=float(farmv),
                    coast=bool(coast), river=bool(river)))
                st.rerun()
        elif src == "Built in scenario":
            scen = st.selectbox("Scenario", list(scenarios.SCENARIOS.keys()))
            if st.button("Load scenario", use_container_width=True,
                         type="primary"):
                _new_simulator(scenarios.SCENARIOS[scen]()); st.rerun()
        else:
            dfuel = st.selectbox("Default fuel", FUEL_TYPES)
            if st.button("Create blank grid", use_container_width=True,
                         type="primary"):
                _new_simulator(World.blank(
                    SimConfig(nx=int(nx), ny=int(ny), cell_size_m=float(cell)),
                    default_fuel=dfuel)); st.rerun()
        st.divider()
        st.markdown("**Resize current map** \u2014 keeps everything on it, "
                    "resamples the grid")
        _keep_ext = st.checkbox(
            "Keep the physical extent (change the RESOLUTION)", value=True,
            key="res_keep",
            help="On: the cell size is divided by the same factor the grid "
                 "is multiplied by, so the same ground is sampled more "
                 "finely. Off: the cell size is kept and the map covers "
                 "more ground, which makes the fire travel further and the "
                 "service radii cover less of it.")
        rz1, rz2, rz3 = st.columns([1, 1, 1])
        rnx = rz1.number_input("nx", 20, 600, int(cfg.nx), 10, key="res_nx")
        rny = rz2.number_input("ny", 20, 600, int(cfg.ny), 10, key="res_ny")
        rz3.markdown("<div style='height:1.75em'></div>",
                     unsafe_allow_html=True)
        if rz3.button("Resize", use_container_width=True,
                      disabled=(int(rnx) == cfg.nx and int(rny) == cfg.ny)):
            _new_simulator(_resize_world(world, int(rnx), int(rny),
                                         keep_extent=bool(_keep_ext)))
            st.rerun()
        st.divider()
        up = st.file_uploader("Load scenario file", type=["json", "yaml", "yml"])
        if up is not None:
            tmp = os.path.join(os.path.dirname(__file__), "_upload_" + up.name)
            with open(tmp, "wb") as fh:
                fh.write(up.getbuffer())
            _new_simulator(io_utils.load_scenario(tmp)); st.rerun()

    tools_col, view_col, legend_col = st.columns([1.2, 3.0, 0.9])
    with tools_col:
        st.markdown("**Tool palette**")
        _tools = ["Fuel", "Firebreak", "Access", "Asset", "Settlement",
                  "Elevation"]
        _cur = st.session_state.get("tool", "Fuel")
        if _cur not in _tools:
            _cur = "Fuel"
        # a radio needs no rerun, so switching tools never interrupts the run
        # and never resets the other widgets (view mode, brushes, layers)
        tool = st.radio("Tool palette", _tools, index=_tools.index(_cur),
                        label_visibility="collapsed")
        st.session_state.tool = tool
        st.divider()
        kw = {"tool": tool}
        if tool == "Fuel":
            use_std = st.checkbox("Use standard fuel model", value=False,
                                  help="Anderson 13 / Scott and Burgan catalogue "
                                       "mapped to the internal classes.")
            if use_std:
                from disaster_phyengine import fuels_standard
                cat = st.selectbox("Catalogue", ["Anderson 13", "Scott & Burgan"])
                code = st.selectbox("Model", list(fuels_standard.catalog(cat).keys()),
                                    format_func=lambda k:
                                    f"{k} - {fuels_standard.catalog(cat)[k][0]}")
                fid, ld, mo = fuels_standard.resolve(cat, code)
                inv = {v: k for k, v in FUEL_NAME_TO_ID.items()}
                kw["ftype"] = inv[fid]; kw["load"] = ld; kw["moisture"] = mo
                st.caption(f"-> {kw['ftype']}, load {ld}, moisture {mo}")
            else:
                kw["ftype"] = st.selectbox("Fuel type", FUEL_TYPES, index=2)
                mdl = FUEL_MODELS[FUEL_NAME_TO_ID[kw["ftype"]]]
                st.caption(f"$r_{{base}}$={mdl.r_base}  $m_{{ext}}$={mdl.m_ext}  "
                           f"$b_{{base}}$={mdl.b_base}")
                kw["load"] = st.slider("$F_{load}$ \u2014 fuel load", 0.0, 1.0,
                                       1.0, 0.05)
                kw["moisture"] = st.slider("$F_{moist}$ \u2014 moisture", 0.0,
                                           0.6, 0.08, 0.01)
            shape = st.radio("Shape", ["Brush", "Rectangle", "Point"], horizontal=True)
        elif tool == "Firebreak":
            fbname = st.selectbox("Firebreak type", list(FIREBREAK_TYPES.keys()))
            kw["fbid"] = FIREBREAK_TYPES[fbname]
            st.caption("Non flammable barrier the fire cannot cross.")
            shape = st.radio("Shape", ["Brush", "Rectangle", "Point"], horizontal=True)
        elif tool == "Access":
            st.caption("Roads set accessibility so resources can reach; non flammable.")
            kw["brush"] = st.slider("Road width", 1, 8, 1)
            shape = st.radio("Shape", ["Brush", "Rectangle"], horizontal=True)
        elif tool == "Asset":
            _settlement_manager(world)
            _asset_manager(world)
            kw["akind"] = st.selectbox("Asset kind", ASSET_KINDS,
                                       format_func=lambda k: ASSET_LABELS[k])
            kw["aname"] = st.text_input("Name (blank = kind)", "")
            kw["aradius"] = st.number_input("Radius (cells)", 0, 40, 3)
            kw["avalue"] = st.slider("Value", 0.0, 1.0, 1.0, 0.05)
            kw["apop"] = st.number_input("Population", 0, 1_000_000, 0)
            shape = "Point"
        elif tool == "Settlement":
            _settlement_manager(world)
            st.caption("Click the map to build a whole settlement: built-up "
                       "ground with streets, its people spread across it, "
                       "and civic facilities around the centre. The same "
                       "builder the map generator uses.")
            kw["sname"] = st.text_input("Name (blank = auto)", "",
                                        key="set_name")
            kw["spop"] = st.number_input("Population", 0, 1_000_000, 5000,
                                         500, key="set_pop")
            kw["sdens"] = st.slider(
                "Building / critical facility density", 0.0, 2.0, 1.0, 0.1,
                key="set_dens",
                help="Scales the footprint AND how many civic facilities "
                     "the settlement carries. At 0 it is houses only.")
            kw["smain"] = st.checkbox(
                "Main town (gets the full set of facilities)", value=False,
                key="set_main")
            kw["sseed"] = st.number_input("Seed (0 = random)", 0, 10 ** 6, 0,
                                          1, key="set_seed")
            shape = "Point"
        else:  # Elevation
            st.caption("Raise or lower the ground. Uphill slope speeds the fire.")
            direction = st.radio("Action", ["Raise", "Lower"], horizontal=True)
            amt = st.slider("Amount (m)", 5, 200, 40, 5)
            kw["elev_delta"] = float(amt if direction == "Raise" else -amt)
            shape = st.radio("Shape", ["Brush", "Point"], horizontal=True)
        if shape == "Brush":
            kw["brush"] = st.slider("Brush size", 1, 12, 3, key="brushsz")
            drawing_mode = "freedraw"
        elif shape == "Rectangle":
            drawing_mode = "rect"
        else:
            if "brush" not in kw:
                kw["brush"] = st.number_input("Point radius", 0, 10, 1)
            drawing_mode = "point"
        live = st.toggle("Live paint", value=True)

    with legend_col:
        with st.expander("Legend", expanded=True):
            st.markdown(legend_html(), unsafe_allow_html=True)
        # layer show / hide lives in the shared Map layers bar at the top of
        # the page; here we only read the resulting flags for the render
        eflags = _layer_flags()
        if eflags["show_value"] and float(world.value.vbld.max()
                + world.value.vcrit.max() + world.value.vpop.max()) <= 0:
            st.caption("\u26a0 No assets on the map yet \u2014 the "
                       "overlay shows nothing. Use the Asset tool.")

    with view_col:
        _views = ["2D canvas", "2D pan / zoom", "3D terrain"]
        _vcur = st.session_state.get("editor_view_sel", "2D canvas")
        vmode = st.radio("Editor view", _views,
                         index=_views.index(_vcur) if _vcur in _views else 0,
                         horizontal=True, label_visibility="collapsed")
        st.session_state["editor_view_sel"] = vmode
        if vmode == "2D pan / zoom":
            # view-only inspection with the same top-right toolbar as the
            # 3D view: drag = pan, wheel = zoom. Editing happens on the
            # 2D canvas view.
            st.caption("Inspect view \u00b7 drag to pan, wheel to zoom, "
                       "toolbar top-right. Switch back to the 2D canvas "
                       "to edit.")
            with _map_card():
                st.plotly_chart(
                    viz.map_figure_2d(world, sim=sim,
                                      scale=_fit_scale(cfg.nx), **eflags),
                    use_container_width=True,
                    key=f"edpz_{st.session_state.map_version}",
                    config={"scrollZoom": True,
                            "modeBarButtonsToRemove": ["lasso2d", "select2d"]})
        elif vmode == "3D terrain":
            # 3D is a pure preview: all editing happens on the 2D canvas
            _e = world.topo.elev
            st.caption(f"3D preview \u00b7 elevation {_e.min():.0f}\u2013"
                       f"{_e.max():.0f} m \u00b7 drag to rotate, scroll to "
                       "zoom. All tools work on the 2D canvas view; switch "
                       "back there to edit.")
            with _map_card():
                fig3 = viz.fire_surface_figure(world, sim=sim)
                key3e = f"edit3d_{st.session_state.map_version}"
                st.plotly_chart(fig3, use_container_width=True,
                                config={"scrollZoom": True}, key=key3e)
        else:
            from io import BytesIO
            scale = _fit_scale(cfg.nx)
            bg = viz.render_pil(world, sim=sim, scale=scale, show_labels=True,
                                **eflags)
            # our own controls (the built in canvas toolbar is hidden below):
            # a real PNG download, an edit undo and a stroke clear
            t1, t2, t3 = st.columns(3)
            _buf = BytesIO(); bg.save(_buf, format="PNG")
            t1.download_button("Download PNG", _buf.getvalue(),
                               file_name="map.png", mime="image/png",
                               use_container_width=True)
            if t2.button("Undo edit", use_container_width=True,
                         disabled=not st.session_state.get("edit_undo"),
                         help="Reverts the last APPLIED map edit (fuel, "
                              "firebreak, road, asset, elevation). Up to 12 "
                              "steps back."):
                _restore_snapshot(); st.session_state.canvas_key += 1; st.rerun()
            # Clear strokes only matters when strokes are NOT applied live
            _export_panel(world, sim, key="exped")
            if not live and t3.button(
                    "Clear strokes", use_container_width=True,
                    help="Discards the drawn strokes WITHOUT applying them; "
                         "the map itself never changes."):
                st.session_state.canvas_key += 1; st.rerun()
            if HAS_CANVAS:
                # .get, not [tool]: adding the Settlement tool to the
                # palette without a colour here would have raised a
                # KeyError the moment it was selected.
                stroke = {"Fuel": "#1f7a1f", "Firebreak": "#3070b0",
                          "Access": "#b08020", "Asset": "#ffd000",
                          "Settlement": "#ff8c00",
                          "Elevation": "#7a5230"}.get(tool, "#a200de")
                sw = kw.get("brush", 2) * scale if drawing_mode == "freedraw" else 2
                flagsig = abs(hash(tuple(sorted(eflags.items())))) % 100000
                ckey = f"canvas_{st.session_state.canvas_key}_{scale}_{flagsig}"
                with _map_card():
                    result = st_canvas(fill_color="rgba(255,160,0,0.20)",
                                       stroke_width=int(sw), stroke_color=stroke,
                                       background_image=bg, update_streamlit=True,
                                       height=cfg.ny * scale, width=cfg.nx * scale,
                                       drawing_mode=drawing_mode,
                                       display_toolbar=False,
                                       point_display_radius=max(3, scale // 2),
                                       key=ckey)
                objs = (result.json_data or {}).get("objects", []) if result else []
                if live and objs:
                    _push_snapshot(); _apply_edits(objs, scale, kw)
                    _simP = st.session_state.get("sim")
                    if _simP is not None:
                        _simP._fmoist0 = world.fuel.fmoist.copy()
                    st.session_state.canvas_key += 1; st.rerun()
                if not live and st.button("Apply edits", type="primary",
                                          use_container_width=True):
                    _push_snapshot(); _apply_edits(objs, scale, kw)
                    _simP = st.session_state.get("sim")
                    if _simP is not None:
                        _simP._fmoist0 = world.fuel.fmoist.copy()
                    st.session_state.canvas_key += 1; st.rerun()
            else:
                with _map_card():
                    st.image(bg, use_container_width=True)

    st.divider()
    e1, e2 = st.columns(2)
    if e1.button("Clear all assets", disabled=not world.assets,
                 help="Removes every asset marker and its value layers "
                      "(buildings, critical facilities, population, "
                      "evacuation routes)."):
        world.assets.clear()
        from disaster_phyengine.layers import ValueLayer
        world.value = ValueLayer.empty(cfg.ny, cfg.nx); st.rerun()
    if e2.button("Clear ignitions", disabled=not world.ignitions,
                 help="Removes all placed/scheduled ignition markers; "
                      "cells already burning keep burning."):
        world.ignitions.clear(); st.rerun()



# ============================================================== DATA LAYERS ==
# Field registry mirroring the System Description: external inputs,
# the state vector and derived fields.
_FIELD_DEFS = [
 ("W_temp \u2014 air temperature", "U_Meteo", "temp",
     r"U_{Meteo,k}=[\,W_{temp},W_{rh},W_{ws},W_{wd},W_{gust},W_{prec}\,]^T",
     "\u00b0C"),
 ("W_rh \u2014 relative humidity", "U_Meteo", "rh",
     r"F_{moist}=EMC(W_{temp},W_{rh})\ \text{(optional mode)}", "%"),
 ("W_ws \u2014 wind speed", "U_Meteo", "wws",
     r"g_{wind}=1+a_w\tanh(W_{ws}/w_0)", "m/s"),
 ("W_wd \u2014 wind direction", "U_Meteo", "wwd",
     r"g_{dir}=\max\{0,\cos(W_{wd}-\theta)\}", "rad"),
 ("W_gust \u2014 wind gust", "U_Meteo", "gust",
     r"\text{exogenous stochastic forcing channel}", "m/s"),
 ("W_prec \u2014 precipitation", "U_Meteo", "prec",
     r"\text{moistening channel (moisture dynamics mode)}", "mm/h"),
 ("G_elev \u2014 elevation", "U_Geo", "elev",
     r"U_{Geo}=[\,G_{elev},G_{slope},G_{aspect},G_{access}\,]^T", "m"),
 ("G_slope \u2014 slope", "U_Geo", "slope",
     r"g_{slope}=1+a_s\tan(G_{slope})", "rad"),
 ("G_aspect \u2014 aspect", "U_Geo", "aspect",
     r"g_{aspect}=1+a_{asp}\cos(G_{aspect}-W_{wd})", "rad"),
 ("G_access \u2014 accessibility", "U_Geo", "access",
     r"\eta_{reach}=e^{-\beta_t R_{time}}\,G_{access}", "[0,1]"),
 ("F_type \u2014 fuel class", "U_Fuel", "ftype",
     r"U_{Fuel,k}=[\,F_{type},F_{load,0},F_{moist,k}\,]^T", "class id"),
 ("F_load,0 \u2014 initial fuel load", "U_Fuel", "fload0",
     r"F_{load,0}\ \text{initializes the state}\ F_{load,k}", "[0,1] norm."),
 ("F_moist \u2014 fuel moisture", "U_Fuel", "fmoist",
     r"g_{moist}=\max\{0,\,1-F_{moist}/m_{ext}\}", "mass fraction"),
 ("V_bld \u2014 building footprint", "U_Val", "vbld",
     r"J^{val}=c_{bld}\lambda_{loss}\textstyle\sum A_k V_{bld}", "[0,1]"),
 ("V_crit \u2014 critical facilities", "U_Val", "vcrit",
     r"J^{inf}=c_{crit}\lambda_{loss}\textstyle\sum A_k V_{crit}", "[0,1]"),
 ("V_pop \u2014 population density", "U_Val", "vpop",
     r"P^{exp}=a_{km^2}\textstyle\sum A_k V_{pop}", "person/km\u00b2"),
 ("V_evac \u2014 evacuation distance", "U_Val", "vevac",
     r"V_{evac}^{norm}=1-\text{minmax}(V_{evac})", "m"),
 ("V_prio \u2014 protection priority", "U_Val", "vprio",
     r"V_{prio}=w_{bld}V_{bld}+w_{crit}V_{crit}"
     r"+w_{pop}V_{pop}^{norm}+w_{evac}V_{evac}^{norm}", "[0,1]"),
 ("R_cap \u2014 suppression capacity", "U_DSS", "rcap",
     r"\eta_{cap}=R_{cap}/R_{cap,max}", "capacity/step"),
 ("R_avail \u2014 availability", "U_DSS", "ravail",
     r"\eta_{avail}=R_{avail}\in\{0,1\}", "{0,1}"),
 ("R_eff \u2014 efficiency", "U_DSS", "reff",
     r"\eta_{eff}=R_{eff}/(1+\gamma_I I_k)", "[0,1]"),
 ("R_time \u2014 travel time", "U_DSS", "rtime",
     r"\eta_{reach}=e^{-\beta_t R_{time}}\,G_{access}", "min"),
 ("B_k \u2014 burning status", "State s_k", "burning",
     r"B_{k+1}=\max\{B^{pers},B^{prop},I_{Ign}\cdot H\}", "{0,1}"),
 ("F_load,k \u2014 fuel load", "State s_k", "fload",
     r"F_{load,k+1}=\max\{0,F_{load,k}-B_kF_{burn,k}F_{load,k}-F_{red,k}\}",
     "[0,1] norm."),
 ("I_k \u2014 fire intensity", "State s_k", "intensity",
     r"I_{k+1}=B_{k+1}\tanh\big(\beta(\tilde F+\gamma_W\tilde W"
     r"+\gamma_S\tilde S)\big)", "[0,1]"),
 ("\u03c4_k \u2014 time since ignition", "State s_k", "tau",
     r"\tau_{k+1}=\tau_k+\Delta t\ \text{(while burning)}", "steps"),
 ("A_k \u2014 ignition buildup", "Derived", "buildup",
     r"A_{k+1}=(1-B_{k+1})\big[(1-\lambda)A_k+\Psi_k\big]", "influence"),
 ("R_spread,k \u2014 rate of spread", "Derived", "ros",
     r"R_{spread}=r_{base}\,g_{moist}\,g_{wind}\,g_{slope}\,g_{aspect}",
     "cells/step (= m/min at 30 m / 30 min)"),
 ("t_ign \u2014 time of first ignition", "Derived", "tign",
     r"t_{ign}(x,y)=\min\{k:\,B_k(x,y)=1\}", "step index"),
]


def _field_array(key):
    from disaster_phyengine import behavior
    if key == "vprio":
        return world.priority_field()
    if key == "ros":
        return behavior.rate_of_spread_field(world)
    if key == "burning":
        return sim.state.burning
    if key == "intensity":
        return sim.state.intensity
    if key == "tau":
        return sim.state.tau
    if key == "fload":
        return sim.state.fload
    if key == "buildup":
        return sim.ign_buildup
    src = {"temp": world.meteo.temp, "rh": world.meteo.rh,
           "wws": world.meteo.wws, "wwd": world.meteo.wwd,
           "gust": world.meteo.gust, "prec": world.meteo.prec,
           "elev": world.topo.elev, "slope": world.topo.slope,
           "aspect": world.topo.aspect, "access": world.topo.access,
           "ftype": world.fuel.ftype.astype(float),
           "fload0": world.fuel.fload0,
           "fmoist": world.fuel.fmoist,
           "vbld": world.value.vbld, "vcrit": world.value.vcrit,
           "vpop": world.value.vpop, "vevac": world.value.vevac,
           "rcap": world.resource.rcap, "ravail": world.resource.ravail,
           "reff": world.resource.reff, "rtime": world.resource.rtime}
    return src[key]


def page_layers():
    import matplotlib.pyplot as plt
    st.subheader("Data layers")
    st.caption("The exact fields of the System Description: the external "
 "inputs of, the state vector of and the fields "
               "derived from them. Names, symbols and units match the System "
               "Description page.")

    st.markdown("#### Terrain \u2014 $U_{Geo}$")
    tc1, tc2 = st.columns(2)
    with tc1:
        st.markdown("**2D relief**")
        with _map_card():
            st.image(viz.terrain_pil(world, scale=max(4, 600 // max(cfg.nx, 1))))
    with tc2:
        st.markdown("**3D surface** (drag to rotate, scroll to zoom)")
        try:
            with _map_card():
                st.plotly_chart(viz.fire_surface_figure(world, sim=sim),
                                use_container_width=True)
        except Exception as exc:
            st.info(f"3D needs plotly. {exc}")

    st.divider()
    st.markdown("#### Field viewer")
    names = [f"{grp}  \u00b7  {name}" for name, grp, *_ in _FIELD_DEFS]
    pick = st.selectbox("Field", names, label_visibility="collapsed")
    name, grp, key, latex, unit = _FIELD_DEFS[names.index(pick)]
    st.latex(latex)
    st.caption(f"Defined in System Description {grp.split('(')[-1].rstrip(')')} "
               f"\u00b7 unit: {unit}")
    if key == "tign":
        st.caption("Red = ignites first, blue = ignites later; grey = never "
                   "ignited in this run.")
        with _map_card():
            st.image(viz.ignition_time_pil(sim, scale=max(4, 700 // max(cfg.nx, 1))))
    else:
        field = _field_array(key)
        _u = unit
        if key in ("wwd", "slope", "aspect"):
            field = np.degrees(field)
            _u = "°"
        fig, ax = plt.subplots(figsize=(8, 5))
        im = ax.imshow(field, origin="upper", cmap="viridis")
        fig.colorbar(im, ax=ax, shrink=0.8)
        ax.set_title(f"{name}  [{_u}]")
        ax.set_xticks([]); ax.set_yticks([])
        st.pyplot(fig, use_container_width=True); plt.close(fig)

    st.divider()
    with st.expander("Operational diagnostics (viewer only \u2014 not part "
                     "of the model)"):
        st.caption("Standard fire-behaviour products computed FROM the state "
                   "for interpretation and reporting; they never influence "
                   "the simulation. Fireline intensity is Byram's "
                   r"$I_B = H\,w\,R$ (kW/m) and flame length "
                   r"$L = 0.0775\,I_B^{0.46}$ (m), the measures used by "
                   "operational tools such as FARSITE/FlamMap; crown fire "
                   "marks burning forest cells above the intensity "
                   "threshold of the Parameters page.")
        bkind = st.selectbox("Diagnostic",
                             ["Fireline intensity", "Flame length",
                              "Rate of spread (m/min)", "Crown fire"])
        bmap = {"Fireline intensity": "fireline_intensity",
                "Flame length": "flame_length",
                "Rate of spread (m/min)": "rate_of_spread",
                "Crown fire": "crown_fire"}
        with _map_card():
            st.image(viz.behavior_pil(sim, kind=bmap[bkind],
                                      scale=max(4, 700 // max(cfg.nx, 1))))


# ============================================================== PARAMETERS ===
def page_params():
    st.subheader("Model parameters")
    st.caption("These shape the fire behaviour model. Day to day conditions live "
               "in the Simulation tab. Hover the ? on any control. Every symbol "
 "is defined in the System Description page; defaults follow its "
 "tables.")
    rc1, rc2 = st.columns([1.2, 3])
    if rc1.button("Reset to defaults", use_container_width=True):
        cfg.spread = SpreadParams()
        cfg.suppression = SuppressionParams()
        cfg.intensity = IntensityParams()
        cfg.value_weights = ValueWeights()
        cfg.cost = CostParams()
        for _i, _m0 in _REFERENCE_FUELS.items():
            FUEL_MODELS[_i] = dataclasses.replace(_m0)
        st.rerun()
    rc2.caption("Restores every parameter below, including the fuel class table.")
    with st.expander("Optional realism modes"):
        st.caption("Literature calibrated speeds, influence buildup and "
                   "the flank/backing floor are always active. The two "
                   "mechanisms below are **on by default**; turning them "
                   "off gives the plain cosine kernel without embers. "
 "Equations:")
        a, b = st.columns(2)
        st.markdown("**Elliptical spread kernel** (Cell2Fire / FARSITE)")
        a, b = st.columns(2)
        cfg.spread.elliptical = a.checkbox(
            "Enable elliptical kernel", value=cfg.spread.elliptical,
            help="Replaces the cosine directional weight by a wind-elongated "
                 "ellipse: g_dir = (1-e)/(1-e cos\u0394), with eccentricity "
                 "from the length-to-breadth ratio LB = LB0 + LBw\u00b7W_ws.")
        cfg.spread.lb_ratio_base = b.number_input(
            "$LB_0$ — length/breadth ratio at no wind (–)", 1.0, 3.0,
            float(cfg.spread.lb_ratio_base), 0.05)
        cfg.spread.lb_ratio_wind = b.number_input(
            "$LB_w$ — extra length/breadth per m/s of wind", 0.0, 0.5,
            float(cfg.spread.lb_ratio_wind), 0.01)
        st.markdown("**Ember spotting** (stochastic ignition ahead of the front)")
        a, b = st.columns(2)
        cfg.spread.spotting = a.checkbox(
            "Enable spotting", value=cfg.spread.spotting,
            help="Each intense burning cell (I_k above the threshold) throws "
                 "an ember downwind with the given probability per substep; "
                 "the ember ignites its landing cell if fuel is present.")
        cfg.spread.spot_prob = b.number_input(
            "$p_{spot}$ — spot probability (per cell and substep)", 0.0, 1.0,
            float(cfg.spread.spot_prob), 0.01)
        cfg.spread.spot_distance = int(a.number_input(
            "$d_{spot}$ — spot distance (cells downwind)", 1, 40,
            int(cfg.spread.spot_distance)))
        cfg.spread.spot_intensity_min = b.number_input(
            "$I_{min}$ — minimum $I_k$ to spot (–)", 0.0, 1.0,
            float(cfg.spread.spot_intensity_min), 0.05)
        st.markdown("**Crown fire flag** (diagnostic only)")
        cfg.intensity.crown_fire_threshold = st.number_input(
            "$I_{crown}$ — crown-fire intensity threshold (–)", 0.0, 1.0,
            float(cfg.intensity.crown_fire_threshold), 0.05,
            help="Burning forest cells with I_k above this value are "
                 "reported as crown fire in the Data layers diagnostics; "
                 "it does not change the dynamics.")

    # 'Run limit' (max_steps) now lives in the sidebar Simulation card, and
    # 'Suppression effectiveness' is not exposed here: real scenes overwrite it
    # (build_real_world calibration) and synthetic scenes use the defaults.
    with st.expander("Integration", expanded=False):
        st.number_input(
            "Substep cap per step (0 = uncapped)", 0, 200,
            int(st.session_state.get("sim_substep_cap", 8)), 1,
            key="sim_substep_cap",
            help="A long step is integrated in internal substeps so the "
                 "front advances at most about one cell per substep; the "
                 "cap trades a little spread accuracy for speed. It applies "
                 "to EVERY way of advancing the run (Step, Step X, Run to "
                 "end, Animate), so a scenario reproduces regardless of "
                 "which control was used. 0 removes the cap and runs at "
                 "full fidelity, roughly three times slower.")
        st.caption("Uncapped is the reference integration; the default cap "
                   "of 8 is what the interactive runs have always used.")

    with st.expander("Fire spread and propagation", expanded=True):
        a, b = st.columns(2)
        cfg.spread.w0 = a.number_input("$w_0$ — wind saturation speed (m/s)", 0.5, 30.0,
                                       float(cfg.spread.w0), 0.5,
                                       help="Wind where amplification saturates.")
        cfg.spread.eps_fuel = b.number_input("$\\epsilon_{fuel}$ — extinction fuel threshold (norm. fuel)", 1e-6, 0.1,
                                             float(cfg.spread.eps_fuel), format="%.5f",
                                             help="Fuel below which a cell stops burning.")
        cfg.spread.theta_ign = a.number_input(
            "$\\Theta_{ign}$ — ignition threshold (\u215b cell-widths)",
            0.005, 1.0, float(cfg.spread.theta_ign), 0.005, format="%.3f",
            help="Activation threshold on the ignition influence buildup "
 "$A_k$. Default "
                 "$0.125=1/8$: the front then advances at exactly "
                 "$R_{spread}$ cells per step.")
        cfg.spread.aniso_wind_full = b.number_input(
            "$w_{aniso}$ — wind for fully directional spread (m/s)", 0.0, 30.0,
            float(cfg.spread.aniso_wind_full), 0.5,
            help="Below this wind speed the directional weight blends toward "
 "isotropic spread.")
        cfg.spread.back_frac = a.number_input(
            "$f_{back}$ — flank/backing fraction (–)", 0.0, 1.0,
            float(cfg.spread.back_frac), 0.01,
            help="Directional-weight floor: backing fires run at this "
                 "fraction of the head fire rate (literature ~0.05-0.15).")
        cfg.spread.buildup_leak = b.number_input(
            "$\\lambda$ — influence buildup leak (fraction/step)", 0.0, 1.0,
            float(cfg.spread.buildup_leak), 0.01,
            help="Per reference step decay of the ignition buildup A_k: "
                 "heating dissipates when the fire source disappears.")
        cfg.spread.slope_wind_equiv = a.number_input(
            "$k_{slope}$ \u2014 slope-equivalent wind (m/s per tan)", 0.0,
            30.0, float(getattr(cfg.spread, "slope_wind_equiv", 10.0)), 1.0,
            help="Upslope acts like an extra wind of "
                 "$k_{slope}\\tan(G_{slope})$ m/s blowing uphill "
                 "(FARSITE-style vector combination): fire climbs "
                 "mountainsides even against a light gradient wind.")
        cfg.spread.slope_gain_max = b.number_input(
            "$g_{slope}^{max}$ \u2014 slope factor cap (\u2013)", 1.0,
            10.0, float(getattr(cfg.spread, "slope_gain_max", 3.0)), 0.5,
            help="Saturation of the slope speed factor; prevents tan() "
                 "blow-up on real DEM cliffs.")
        cfg.spread.slope_clip_rad = a.number_input("$G_{slope}$ clip (rad)", 0.1, 1.5,
                                                   float(cfg.spread.slope_clip_rad), 0.05)
        cfg.spread.diagonal_distance_weighting = b.checkbox(
            "Diagonal distance weighting", value=cfg.spread.diagonal_distance_weighting)
    with st.expander("Fire intensity"):
        a, b = st.columns(2)
        cfg.intensity.beta = a.number_input(
            "$\\beta$ — global intensity gain (–)", 0.1, 3.0,
            float(cfg.intensity.beta), 0.1,
            help="Scales the $\\tanh$ argument of the intensity proxy "
 "; typical range 1–3.")
        cfg.intensity.gamma_w = b.number_input(
            "$\\gamma_W$ — wind weight (–)", 0.0, 1.0,
            float(cfg.intensity.gamma_w), 0.05,
            help="Contribution of normalized wind to the intensity proxy "
 "; typical 0–0.7.")
        cfg.intensity.gamma_s = a.number_input(
            "$\\gamma_S$ — slope weight (–)", 0.0, 1.0,
            float(cfg.intensity.gamma_s), 0.05,
            help="Contribution of normalized slope to the intensity proxy "
 "; typical 0–0.5.")
        cfg.intensity.wws_max = b.number_input(
            "$W_{ref}$ — wind normalization (m/s)", 1.0, 60.0,
            float(cfg.intensity.wws_max), 1.0,
            help="Wind speed treated as 'extreme' in the intensity "
                 "normalization $\\tilde W=\\min\\{1,W_{ws}/W_{ref}\\}$.")
        cfg.intensity.slope_max_rad = a.number_input(
            "$S_{max}$ — slope normalization (rad)", 0.1, 1.4,
            float(cfg.intensity.slope_max_rad), 0.05,
            help="Slope normalization S_max. Default 0.7854 (45\u00b0).")
        cfg.intensity.fload_max = b.number_input(
            "$F_{max}$ — fuel normalization (norm. units)", 0.1, 5.0,
            float(cfg.intensity.fload_max), 0.1,
            help="Fuel normalization F_max. Default 1.0.")
    # 'Protection priority weights' (V_prio) are a DSS valuation, not a fire
    # behaviour parameter, so they moved to Simulation -> DSS settings.
    with st.expander("Fuel classes"):
        st.caption("Per class spread and combustion parameters: r_base "
                   "(cells/step), m_ext (mass fraction), a_w / a_s / a_asp "
 "(dimensionless), b_base, "
                   "e (economic value per cell unit, used by the cost model). "
                   "Note: b_base is cached when a simulator is created; press "
                   "Reset fire (or make a new map) after editing it.")
        import pandas as _pd
        _rows = [{"id": _i, "fuel": _m.name, "r_base": _m.r_base,
                  "m_ext": _m.m_ext, "a_w": _m.a_w, "a_s": _m.a_s,
                  "a_asp": _m.a_asp, "b_base": _m.b_base,
                  "e": _m.economic_value}
                 for _i, _m in FUEL_MODELS.items()]
        _edit = st.data_editor(_pd.DataFrame(_rows), hide_index=True,
                               disabled=["id", "fuel"], key="fuel_editor")
        for _, _r in _edit.iterrows():
            _m = FUEL_MODELS[int(_r["id"])]
            _m.r_base = float(_r["r_base"]); _m.m_ext = float(_r["m_ext"])
            _m.a_w = float(_r["a_w"]); _m.a_s = float(_r["a_s"])
            _m.a_asp = float(_r["a_asp"]); _m.b_base = float(_r["b_base"])
            _m.economic_value = float(_r["e"])
        if st.button("Reset fuel classes to defaults"):
            for _i, _m0 in _REFERENCE_FUELS.items():
                FUEL_MODELS[_i] = dataclasses.replace(_m0)
            st.rerun()

    # The cost model (priority weights + reference scales + safeguards) is not
    # set here anymore: it lives under Simulation -> Layer 4 -> Operational
    # priority (weights, with a Custom option that sums to 1) and its Advanced
    # cost settings expander, so there is a single place for all cost tuning.


# ====================================================== SYSTEM DESCRIPTION ===
def page_system_description():
    """Full mathematical description of the model.

    The content lives in app/system_description.py to keep this file small."""
    from system_description import render
    render()


# ============================================================== GIS IMPORT ===
def page_gis():
    st.subheader("Real-world maps")
    st.markdown("#### Automatic \u2014 download a real area "
                "(same open data as the Validation page)")
    st.caption("Terrain: Copernicus GLO-30 DEM \u00b7 fuel: ESA WorldCover "
               "10 m \u2192 fuel classes. Needs rasterio + internet on the "
               "first download; areas are cached under validation/cache/ "
               "and load offline afterwards.")
    try:
        _av = _load_auto_validate()
    except Exception:
        _av = None
    if _av is not None:
        rw1, rw2 = st.columns([1.4, 1])
        _opts = ["Custom bounding box"] + sorted(_av.CASES)
        _sel = rw1.selectbox(
            "Area", _opts,
            format_func=lambda k: (_av.CASES[k]["label"] + " (case area)"
                                   if k in _av.CASES else k))
        rw_cell = float(rw2.number_input("Cell size (m)", 30.0, 300.0, 90.0,
                                         10.0, key="rw_cell"))
        rwa, rwb = st.columns(2)
        rw_assets = rwa.checkbox(
            "Auto-place buildings + population (from WorldCover built-up)",
            value=True, key="rw_assets",
            help="Turns the built-up land-cover class into protectable "
                 "building and population assets, with a nominal urban "
                 "density. Water and forest already come from the fuel map.")
        rw_roads = rwb.checkbox(
            "Download roads + facilities from OpenStreetMap",
            value=True, key="rw_roads",
            help="Overpass API: road network becomes access corridors "
                 "(non-flammable, access=1); hospitals, fire stations and "
                 "power sites become critical assets. Needs internet.")
        rw_fire = False
        _focus_bbox = None       # a crop of the case rectangle, if chosen
        _focus_suffix = ""       # cache-dir suffix that keeps the crop apart
        if _sel in _av.CASES:
            bbox = dict(_av.CASES[_sel])   # full case: has start/hours too
            cdir = _sel
            # is the satellite ground truth already cached for this case?
            _root_g = os.path.dirname(os.path.dirname(os.path.abspath(
                __file__)))
            _cc = os.path.join(_root_g, "validation", "cache", cdir)
            _rw_cached = (
                os.path.exists(os.path.join(
                    _cc, f"firms_{_av.CASES[_sel]['start']}.csv"))
                or (os.path.exists(os.path.join(_cc, "firms.csv"))
                    and sum(1 for _ in open(os.path.join(_cc, "firms.csv")))
                    > 1))
            rw_fire = st.checkbox(
                "Replay the real fire: set ignition + wind from FIRMS/ERA5",
                value=False, key="rw_fire",
                help="Places the ignition at the satellite first-detection "
                     "front and sets a uniform driving wind + fuel moisture "
                     "from ERA5 at the ignition hour.")
            if rw_fire:
                _fkin = st.text_input(
                    "NASA FIRMS MAP_KEY"
                    + (" (not needed — case is cached)"
                       if _rw_cached else ""),
                    type="password",
                    value=st.session_state.get(
                        "firms_key", os.environ.get("FIRMS_MAP_KEY", "")),
                    key="rw_firms_key",
                    help="Free key (1 minute): "
                         "https://firms.modaps.eosdis.nasa.gov/api/map_key/ "
                         "— only needed the FIRST time a case is fetched; "
                         "afterwards the satellite data is cached and the "
                         "replay needs no key.")
                if _fkin:
                    st.session_state["firms_key"] = _fkin
                if _rw_cached:
                    st.caption(f"✓ {_av.CASES[_sel]['label']}: FIRMS data "
                               "cached — no key required.")
            # ---- focus / crop: work on just the relevant piece of the case,
            # not the whole (often huge) case rectangle. Nothing is lost: the
            # terrain, fuel, assets, roads and fire are all rebuilt for the
            # cropped window, and the fire truth is read from the full case
            # cache and clipped to the crop (no key needed). ----
            import math as _math_g
            _focus = st.radio(
                "Focus area", ["Whole case area", "Crop to fire + margin",
                               "Manual sub-box"],
                key="rw_focus", horizontal=True,
                help="Crop the case down to the region of interest so the "
                     "simulation runs only on the relevant pixels. 'Crop to "
                     "fire' fits the box to the satellite fire footprint plus "
                     "a margin; 'Manual sub-box' lets you type a tighter box.")
            if _focus == "Crop to fire + margin":
                _mk = st.slider("Margin around the fire (km)", 0.5, 15.0,
                                3.0, 0.5, key="rw_crop_margin")
                try:
                    _cb = _av.firms_footprint_bbox(
                        dict(_av.CASES[_sel]), _cc,
                        st.session_state.get("firms_key", "") or None,
                        margin_km=float(_mk))
                except Exception as _ce:
                    _cb = None
                    st.warning(f"Cannot crop to the fire yet: {_ce}")
                if _cb:
                    _focus_bbox = _cb
                    _focus_suffix = f"_fire{_mk:g}"
                    _lat0 = 0.5 * (_cb["south"] + _cb["north"])
                    _wkm = ((_cb["east"] - _cb["west"]) * 111.32
                            * _math_g.cos(_math_g.radians(_lat0)))
                    _hkm = (_cb["north"] - _cb["south"]) * 110.54
                    st.caption(f"Crop {_wkm:.1f} × {_hkm:.1f} km around the "
                               "fire footprint (the full case area is larger).")
            elif _focus == "Manual sub-box":
                _fc = dict(_av.CASES[_sel])
                s1, s2, s3, s4 = st.columns(4)
                _wv = float(s1.number_input(
                    "West", _fc["west"], _fc["east"], _fc["west"], 0.01,
                    format="%.3f", key="rw_sb_w"))
                _sv2 = float(s2.number_input(
                    "South", _fc["south"], _fc["north"], _fc["south"], 0.01,
                    format="%.3f", key="rw_sb_s"))
                _ev = float(s3.number_input(
                    "East", _fc["west"], _fc["east"], _fc["east"], 0.01,
                    format="%.3f", key="rw_sb_e"))
                _nv = float(s4.number_input(
                    "North", _fc["south"], _fc["north"], _fc["north"], 0.01,
                    format="%.3f", key="rw_sb_n"))
                if _ev > _wv and _nv > _sv2:
                    _focus_bbox = dict(_fc)
                    _focus_bbox.update(west=_wv, south=_sv2, east=_ev,
                                       north=_nv)
                    _focus_suffix = (f"_sub{_wv:.2f}_{_sv2:.2f}_"
                                     f"{_ev:.2f}_{_nv:.2f}")
                else:
                    st.warning("Sub-box is inverted: need west < east and "
                               "south < north.")
        else:
            b1, b2, b3, b4 = st.columns(4)
            bbox = dict(
                west=float(b1.number_input("West (lon)", -180.0, 180.0,
                                           31.30, 0.05, key="bb_w")),
                south=float(b2.number_input("South (lat)", -85.0, 85.0,
                                            36.72, 0.05, key="bb_s")),
                east=float(b3.number_input("East (lon)", -180.0, 180.0,
                                           31.80, 0.05, key="bb_e")),
                north=float(b4.number_input("North (lat)", -85.0, 85.0,
                                            37.08, 0.05, key="bb_n")))
            cdir = (f"custom_{bbox['west']:.2f}_{bbox['south']:.2f}_"
                    f"{bbox['east']:.2f}_{bbox['north']:.2f}")
        # a chosen crop replaces the build bbox; its terrain/fuel/asset/road
        # downloads live in a crop-specific cache dir so they never clobber the
        # full-area cache, while the FIRMS/weather truth is still read from the
        # full case cache (no key needed).
        _build_bbox = _focus_bbox if _focus_bbox is not None else bbox
        _build_cdir = cdir + _focus_suffix
        _truth_cdir = None
        if _focus_bbox is not None and _sel in _av.CASES:
            _troot = os.path.dirname(os.path.dirname(os.path.abspath(
                __file__)))
            _truth_cdir = os.path.join(_troot, "validation", "cache", cdir)
        if st.button("Download and load into the simulator", type="primary"):
            if (_build_bbox["east"] <= _build_bbox["west"]
                    or _build_bbox["north"] <= _build_bbox["south"]):
                st.error("Bounding box is inverted: need west < east and "
                         "south < north.")
            else:
                try:
                    import rasterio  # noqa: F401
                    root = os.path.dirname(os.path.dirname(
                        os.path.abspath(__file__)))
                    cache = os.path.join(root, "validation", "cache",
                                         _build_cdir)
                    with st.status("Building the real-world map \u2026",
                                   expanded=True):
                        if _focus_bbox is not None:
                            st.write("focus: cropped to the region of interest")
                        st.write("terrain + fuel (cache-aware)")
                        if rw_assets:
                            st.write("buildings + population from WorldCover "
                                     "built-up")
                        if rw_roads:
                            st.write("roads + facilities from OpenStreetMap "
                                     "(Overpass)")
                        if rw_fire:
                            st.write("ignition + wind from FIRMS / ERA5")
                        _fkey = st.session_state.get("firms_key", "") or None
                        _src_bbox = (dict(_av.CASES[_sel])
                                     if (_focus_bbox is not None
                                         and _sel in _av.CASES) else None)
                        wnew = _av.build_real_world(
                            _build_bbox, rw_cell, cache,
                            add_assets=bool(rw_assets),
                            add_roads=bool(rw_roads),
                            add_fire=bool(rw_fire), firms_key=_fkey,
                            truth_cache_dir=_truth_cdir,
                            source_bbox=_src_bbox)
                        _nroad = (0 if wnew.roads is None
                                  else int(wnew.roads.sum()))
                        _nign = len(wnew.ignitions)
                        st.write(f"grid {wnew.config.nx} \u00d7 "
                                 f"{wnew.config.ny} at {rw_cell:g} m \u00b7 "
                                 f"{len(wnew.assets)} assets \u00b7 "
                                 f"{_nroad} road cells \u00b7 {_nign} ignition "
                                 "cells")
                    _new_simulator(wnew)
                    if rw_fire and _nign == 0:
                        st.warning(
                            "Real fire requested but no ignition was set: "
                            "check the FIRMS MAP_KEY on the Validation page "
                            "and internet access. Assets/roads still loaded.")
                    _ign_msg = (f"{_nign} ignition cells from the real fire. "
                                "Run the Simulation with the DSS."
                                if _nign else
                                "Add ignitions in the Map editor, then run "
                                "the Simulation with the DSS.")
                    st.success(
                        f"Loaded: {len(wnew.assets)} assets, {_nroad} road "
                        f"cells. {_ign_msg}")
                    st.rerun()
                except ImportError:
                    st.error("Needs rasterio: pip install rasterio, then "
                             "restart.")
                except Exception as exc:
                    st.error(f"Download/build failed: {exc}")
    st.divider()
    st.markdown("#### Manual \u2014 your own GeoTIFF rasters")
    st.caption("Load a real elevation and optional fuel raster. Needs rasterio.")
    g1, g2, g3 = st.columns(3)
    gnx = g1.number_input("grid nx", 10, 400, 120, key="gnx")
    gny = g2.number_input("grid ny", 10, 400, 80, key="gny")
    gcell = g3.number_input("cell size (m)", 1.0, 1000.0, 30.0, key="gcell")
    dem_file = st.file_uploader("Elevation raster (GeoTIFF)", type=["tif", "tiff"], key="dem")
    fuel_file = st.file_uploader("Fuel raster (GeoTIFF)", type=["tif", "tiff"], key="fuelr")
    if st.button("Import rasters", type="primary"):
        try:
            from disaster_phyengine import gis
            paths = {}
            for tag, f in [("dem", dem_file), ("fuel", fuel_file)]:
                if f is not None:
                    p = os.path.join(os.path.dirname(__file__), "_gis_" + f.name)
                    with open(p, "wb") as fh:
                        fh.write(f.getbuffer())
                    paths[tag] = p
            _new_simulator(gis.world_from_rasters(
                SimConfig(nx=int(gnx), ny=int(gny), cell_size_m=float(gcell)),
                dem_path=paths.get("dem"), fuel_path=paths.get("fuel")))
            st.success("Imported."); st.rerun()
        except ImportError as exc:
            st.error(str(exc))
        except Exception as exc:  # pragma: no cover
            st.error(f"Import failed: {exc}")


# ------------------------------------------------------------- page dispatch
# =============================================================== VALIDATION ==
def _load_auto_validate():
    import importlib.util
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "..", "validation", "auto_validate.py")
    spec = importlib.util.spec_from_file_location("auto_validate", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def page_validation():
    import types as _types
    st.subheader("Validation \u2014 hindcast against a real fire")
    st.markdown(
        "A **hindcast**: the simulator receives ONLY the real inputs of a "
        "documented fire \u2014 real terrain (Copernicus GLO-30 DEM), real "
        "fuel map (ESA WorldCover 10 m \u2192 fuel classes), real hourly "
        "weather at the fire (ERA5: wind drives the spread, humidity drives "
        "fuel moisture via EMC) and the real initial fire front (first NASA "
        "FIRMS overpass) \u2014 and runs **blind**, without seeing the "
        "answer. The simulated burn is then scored against what actually "
        "burned.")
    with st.expander("Method \u2014 what is the ground truth?"):
        st.markdown(
            "**Primary ground truth** is the observed fire footprint. Two "
            "levels:\n\n"
            "1. **NASA FIRMS satellite detections** (automatic here): every "
            "375 m VIIRS pixel seen actively burning inside the simulated "
            "time window, consolidated by a morphological closing. Real, "
            "timestamped, citable \u2014 but it undersamples the burned "
            "interior and the real fire was actively **suppressed** while "
            "the simulation runs free, so hit rate and front error are the "
            "primary scores against it.\n"
            "2. **Official EFFIS / Copernicus EMS perimeter** (upload "
            "below): the mapped polygon of the final burned area \u2014 "
            "the referee-grade truth that makes Dice, false alarm and area "
            "bias meaningful. Sources: forest-fire.emergency.copernicus.eu "
            "(burnt areas), emergency.copernicus.eu (activations), "
            "mtbs.gov (US).\n\n"
            "Inputs come from open archives: Copernicus GLO-30 DEM (AWS), "
            "ESA WorldCover 2021 (AWS), ERA5 hourly via open-meteo, FIRMS "
            "VIIRS archive. Everything is fetched automatically and cached "
            "per case.")
    with st.expander("Validation criteria \u2014 definitions"):
        st.markdown(
            "$A$ = simulated burn, $B$ = observed burn. A free-spread "
            "model (no suppression) is validated on FRONT PROPAGATION, not "
            "final area.\n\n"
            "| Criterion | Definition | Checks |\n"
            "|---|---|---|\n"
            r"| Coverage (POD) | $\lvert A\cap B\rvert\,/\,\lvert B\rvert$ | fraction of the observed burn reproduced (target $>0.7$) |"
            "\n"
            r"| Front position error | mean / p90 of $\overline{d}(\partial A,\partial B)$ (m) | how far the simulated front is from the observed front |"
            "\n"
            "| Arrival-time agreement | mean abs. difference of simulated "
            "vs FIRMS first-detection time (h) and Spearman $\\rho$ of the "
            "arrival order | the rate-of-spread (propagation) test |\n\n"
            "Area-overlap scores (Dice, Jaccard, false alarm, area bias) "
            "are omitted: a free run overpredicts a suppressed real fire. "
            "Upload an official EFFIS/EMS perimeter for referee-grade area "
            "metrics.")
    with st.expander("Referee protocol \u2014 how to report these results"):
        st.markdown(
            "1. **Calibration/validation split**: tune free parameters on "
            "ONE fire only, freeze them, hindcast the other fires blind and "
            "report those.\n"
            "2. **Multi-seed**: spotting is stochastic \u2014 report mean "
            "\u00b1 sd over \u2265 5 seeds.\n"
            "3. **Sensitivity**: \u00b125% on wind, moisture, "
            "$\\Theta_{ign}$; plus the wind-direction ensemble below.\n"
            "4. **Baselines**: compare against a no-wind isotropic run and "
            "published FARSITE/Cell2Fire scores.\n"
            "5. **Temporal check**: simulated arrival times ($t_{ign}$ "
            "layer) vs FIRMS detection times.\n"
            "6. Archive every run (`validation_runs/`) for "
            "reproducibility.")
    st.caption("First run downloads ~100\u2013200 MB of open data; "
               "everything is cached in `validation_cache/` afterwards.")

    av = None
    try:
        av = _load_auto_validate()
    except Exception as exc:
        st.error(f"Could not load the validation module: {exc}")
        return

    c1, c2, c3 = st.columns([1.3, 1.6, 0.8])
    case_id = c1.selectbox("Case", sorted(av.CASES),
                           format_func=lambda k: av.CASES[k]["label"],
                           help="Documented historical fire to hindcast.")
    _root0 = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    _ccache = os.path.join(_root0, "validation", "cache", case_id)
    _f_new = os.path.join(_ccache,
                          f"firms_{av.CASES[case_id]['start']}.csv")
    _f_old = os.path.join(_ccache, "firms.csv")
    cached_truth = (os.path.exists(_f_new)
                    or (os.path.exists(_f_old)
                        and sum(1 for _ in open(_f_old)) > 1))
    key = c2.text_input("NASA FIRMS MAP_KEY"
                        + (" (not needed \u2014 case is cached)"
                           if cached_truth else ""),
                        type="password",
                        value=st.session_state.get("firms_key",
                                                   os.environ.get(
                                                       "FIRMS_MAP_KEY", "")),
                        help="Free key (1 minute): "
                             "https://firms.modaps.eosdis.nasa.gov/api/"
                             "map_key/ \u2014 only needed the FIRST time a "
                             "case is run; afterwards the satellite data is "
                             "cached and the run starts directly.")
    st.session_state["firms_key"] = key
    if cached_truth:
        st.caption(f"\u2713 {av.CASES[case_id]['label']}: all data cached "
                   "\u2014 press Run, no key required.")
    seeds = int(c3.number_input(
        "Seeds", 1, 12, 5,
        help="Ember spotting is stochastic, so each criterion is reported "
             "as mean +/- sd over this many seeds. 5 is a good default "
             "(3 minimum for a spread, 8-10 for final thesis numbers)."))
    c4, c5, c6 = st.columns(3)
    cell = float(c4.number_input("Cell size (m)", 30.0, 200.0, 90.0, 10.0,
                                 help="90 m matches the satellite truth "
                                      "resolution and keeps the run fast."))
    hours = float(c5.number_input("Time window (h)", 1.0, 120.0,
                                  float(av.CASES[case_id]["hours"]), 1.0,
                                  help="The time window that bounds BOTH the "
                                       "observed footprint (only what was seen "
                                       "burning by this hour) AND the "
                                       "simulation. Lower it to compare an "
                                       "earlier phase; the observed (blue) "
                                       "area shrinks with it."))
    stepm = float(c6.number_input(
        "Step length (min)", 10.0, 60.0, 30.0, 5.0,
        help="Temporal resolution of the spread. 30 min matches the hourly "
             "ERA5 weather and keeps runs fast; 15 min is finer. Do not "
             "exceed 60 min (the weather is hourly)."))
    stop_mode = st.selectbox(
        "Stopping rule (when does the simulation end?)",
        ["Match observed area (recommended)",
         "Fixed hours"],
        help="Both use the time window above for the observed footprint. "
             "'Match observed area' stops as soon as the simulated burn "
             "reaches the observed area OR the time window elapses, "
             "whichever comes first, so the shapes are compared at a "
             "comparable extent (the suppression-driven size difference "
             "cancels). 'Fixed hours' always runs the whole window and lets "
             "the free burn overpredict.")
    # a wind-direction offset can be adopted from the ensemble table below
    # (the "use this wind" button) and applied to the next run; apply the
    # pending pick BEFORE the widgets are created so it takes effect
    st.session_state.setdefault("val_wind_offset", 0.0)
    st.session_state.setdefault("val_wens", False)
    if st.session_state.get("_val_adopt_offset") is not None:
        st.session_state["val_wind_offset"] = float(
            st.session_state.pop("_val_adopt_offset"))
        st.session_state["val_wens"] = False    # clean, ensemble-off rerun
    c7, c8 = st.columns([1, 2])
    wind_offset = float(c7.number_input(
        "Wind-direction offset (deg)", -180.0, 180.0, step=15.0,
        key="val_wind_offset",
        help="Rotate the reanalysis wind direction by this angle for THIS "
             "run. Run the ensemble first, then adopt the best-matching "
             "member below and run again (ensemble off) to lock that wind."))
    if abs(wind_offset) > 1e-6:
        c8.caption(f"This run rotates the reanalysis wind by "
                   f"{wind_offset:+.0f}\u00b0.")
    wens = st.checkbox(
        "Wind-direction uncertainty ensemble (8 extra runs)", key="val_wens",
        help="Gridded reanalysis winds miss local channeling and the "
             "fire's own convection \u2014 the dominant input uncertainty. "
             "This reruns the hindcast with the wind rotated over the full "
             "circle and reports the sensitivity and the best-matching "
             "member.")
    up = st.file_uploader(
        "Optional: official EFFIS/EMS perimeter raster (stronger ground "
        "truth than the FIRMS footprint)", type=["tif", "tiff", "npy"])

    if st.button("Run validation", type="primary",
                 disabled=(not key and up is None)):
        try:
            import rasterio  # noqa: F401
            import requests  # noqa: F401
        except ImportError as exc:
            st.error(f"Missing package: {exc}. Run "
                     "`pip install rasterio requests` in the app "
                     "environment and restart.")
            return
        case = dict(av.CASES[case_id]); case["hours"] = hours
        root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        args = _types.SimpleNamespace(cell=cell, step_minutes=stepm,
                                      seeds=seeds,
                                      cache=os.path.join(root,
                                                         "validation",
                                                         "cache", case_id),
                                      out=os.path.join(root, "validation",
                                                       "runs", "report"))
        os.makedirs(args.cache, exist_ok=True)
        import datetime as _dt2
        import json as _json
        _stamp = _dt2.datetime.now().strftime("%Y%m%d-%H%M%S")
        _rdir = os.path.join(root, "validation", "runs",
                             f"{case_id}_{_stamp}")
        os.makedirs(os.path.join(_rdir, "frames"), exist_ok=True)
        # migrate any pre-restructure cache locations once
        for _flat in (os.path.join(root, "validation_cache", case_id),
                      os.path.join(root, "validation_cache")):
            if not os.path.isdir(_flat):
                continue
            for _fn in os.listdir(_flat):
                _o, _n = os.path.join(_flat, _fn), os.path.join(args.cache,
                                                                _fn)
                if os.path.isfile(_o) and not os.path.exists(_n):
                    os.replace(_o, _n)
        try:
            with st.status("Downloading the real data \u2026",
                           expanded=True) as stat:
                st.write("1/4 terrain \u2014 Copernicus GLO-30 DEM")
                dem, dlons, dlats = av._download_dem(case, args.cache)
                st.write("2/4 fuel \u2014 ESA WorldCover 10 m")
                wc, wlons, wlats = av._download_worldcover(case, args.cache)
                nx, ny, lons, lats = av._grid(case, cell)
                first = None
                fmask = None
                _ign_cells = []
                # run the FIRMS step whenever there is EITHER a key (first
                # fetch) OR a cached truth file (offline replay, no key). The
                # cache read ignores the key value, so an empty box is fine.
                if key or cached_truth:
                    st.write("3/4 fire truth \u2014 NASA FIRMS detections"
                             + (" (cached)" if cached_truth and not key
                                else ""))
                    # download on the documented window (fetches all pts)
                    _case_ign = dict(case)
                    _case_ign["hours"] = float(av.CASES[case_id]["hours"])
                    pts = av._download_firms(_case_ign, key or "", args.cache)
                    # OBSERVED footprint bounded by the TIME WINDOW the user
                    # picked: the detections seen burning within [t0, t0 + T].
                    # _firms_mask_and_ignition windows the detections by
                    # case["hours"] and morphologically closes them, so the
                    # blue mask is a FILLED footprint that grows with T_max.
                    case["hours"] = float(hours)
                    fmask, first, _ign_cells = \
                        av._firms_mask_and_ignition(
                            case, pts, nx, ny, lons, lats, cell)
                    if first is not None:
                        # align the weather series with the ignition hour
                        case["t0_hour"] = first[4].hour
                st.write("4/4 weather \u2014 hourly ERA5 at the fire "
                         "(open-meteo)")
                if first is not None:
                    _iglat = case["north"] - (first[1] + 0.5) * cell / 110540.0
                    import math as _math
                    _iglon = case["west"] + (first[0] + 0.5) * cell / (
                        111320.0 * _math.cos(_math.radians(_iglat)))
                    weather = av._download_weather(case, args.cache,
                                                   lat=_iglat, lon=_iglon)
                else:
                    weather = av._download_weather(case, args.cache)
                if up is not None:
                    tmpp = os.path.join(args.cache, "_upl_" + up.name)
                    with open(tmpp, "wb") as fh:
                        fh.write(up.getbuffer())
                    from disaster_phyengine.gis import _read_resampled
                    obs = (np.load(tmpp) if tmpp.endswith(".npy") else
                           _read_resampled(tmpp, ny, nx, nearest=True)) > 0.5
                else:
                    obs = fmask
                if obs is None or first is None and not up:
                    st.error("No ground truth: give a FIRMS key or upload "
                             "a perimeter raster.")
                    return
                ign = _ign_cells if _ign_cells else (nx // 2, ny // 2)
                if first:
                    st.write(f"initial fire front: {len(_ign_cells)} cells "
                             f"from the first overpass {first[2]} "
                             f"{first[3]} UTC (the fire had already grown "
                             "before the satellite first saw it)")
                stat.update(label="Data ready \u2014 running the blind "
                                  "hindcast", state="running")
            bar = st.progress(0.0, text="simulating \u2026")
            # live event view: the fire grows on the real map while it runs
            demg0 = av._sample(dem, dlons, dlats, lons, lats)
            wcg0 = av._sample(wc, wlons, wlats, lons, lats).astype(int)
            ftype0 = np.zeros_like(wcg0)
            for code, (fid, _ld) in av.WORLDCOVER_TO_FUEL.items():
                ftype0[wcg0 == code] = (case.get("tree_fuel", 3)
                                        if code == 10 else fid)
            base0 = av._basemap(ftype0, demg0)
            live = st.empty()
            log_lines = []
            cell_ha = (cell * cell) / 10000.0

            def _cb(seed, k, n):
                bar.progress((seed + k / n) / seeds,
                             text=f"seed {seed + 1}/{seeds} \u00b7 step "
                                  f"{k}/{n}")

            def _frame(k, n, burned, ws_now):
                t_h = k * stepm / 60.0
                img = base0.copy()
                img[obs] = (58, 110, 220)
                img[burned & obs] = (46, 160, 67)
                img[burned & ~obs] = (200, 55, 44)
                _i0 = ign[0] if isinstance(ign, list) else ign
                gx, gy = int(_i0[0]), int(_i0[1])
                rr = max(2, nx // 150)
                img[max(0, gy - rr):gy + rr + 1,
                    max(0, gx - rr):gx + rr + 1] = (255, 235, 60)
                ha = float(burned.sum()) * cell_ha
                live.image(img, use_container_width=True,
                           caption=f"t = {t_h:.1f} h \u00b7 simulated burn "
                                   f"{ha:,.0f} ha \u00b7 wind {ws_now:.1f} "
                                   "m/s \u00b7 red/green = simulation, "
                                   "blue = observed (final), yellow = "
                                   "ignition")
                log_lines.append(f"t={t_h:5.1f} h  burned={ha:9,.0f} ha  "
                                 f"wind={ws_now:4.1f} m/s")
                try:
                    from PIL import Image as _Img
                    _Img.fromarray(img).save(os.path.join(
                        _rdir, "frames", f"t{t_h:06.1f}h.png"))
                except Exception:
                    pass

            _pts = locals().get("pts")
            _obs_arr = None
            if _pts and first is not None and fmask is not None:
                _obs_arr = av.firms_arrival_hours(case, _pts, nx, ny, cell,
                                                  fmask, first[4])
            # observed arrival window = time span of the observed detections
            # (last minus first). The arrival-time error is judged against
            # THIS real fire duration, not the sim's (possibly short) run.
            _obs_window = None
            if _obs_arr is not None and np.isfinite(_obs_arr).any():
                _obs_window = float(_obs_arr[np.isfinite(_obs_arr)].max())
            # STOPPING RULE: keep the suppression-free run from overgrowing
            # the suppressed real fire (see the selectbox)
            # Match observed area: stop when the sim reaches the (windowed)
            # observed area OR the time window elapses, whichever first.
            # Fixed hours: run the whole window (stop_area = None).
            _stop_area = None
            if stop_mode.startswith("Match observed area"):
                _stop_area = float(np.asarray(obs, dtype=bool).sum())
            n_total = int(round(case["hours"] * 60.0 / stepm))
            weather_run = weather
            if abs(wind_offset) > 1e-6:
                weather_run = dict(weather)
                weather_run["wind_direction_10m"] = [
                    (d + wind_offset) % 360.0
                    for d in weather["wind_direction_10m"]]
                st.write(f"wind-direction offset {wind_offset:+.0f}° "
                         "applied to this run")
            runs, shape = av.run_case(case, args, dem, (dlons, dlats),
                                      wc, (wlons, wlats), weather_run, obs,
                                      ign, progress_cb=_cb,
                                      frame_cb=_frame,
                                      frame_every=max(1, n_total // 24),
                                      obs_arrival=_obs_arr,
                                      stop_area=_stop_area)
            # representative wind for the on-map arrow: mean DOWNWIND
            # direction over the simulated window (where the wind pushes the
            # fire). ERA5 gives the "from" direction; downwind = fire heading.
            _wtoward = None
            try:
                _wd = weather_run["wind_direction_10m"]
                _h0w = int(case.get("t0_hour", 0))
                _hi = min(len(_wd), _h0w + max(1, int(round(case["hours"]))))
                _ang = [np.radians((270.0 - float(_wd[h])) % 360.0)
                        for h in range(_h0w, max(_h0w + 1, _hi))]
                _u = float(np.mean([np.cos(a) for a in _ang]))
                _v = float(np.mean([np.sin(a) for a in _ang]))
                _wtoward = float(np.degrees(np.arctan2(_v, _u)) % 360.0)
            except Exception:
                _wtoward = None
            ens = None
            if wens:
                bar.progress(0.0, text="wind ensemble \u2026")

                def _ecb(i, nmem, k, n):
                    bar.progress((i + k / n) / nmem,
                                 text=f"wind ensemble member {i + 1}/{nmem}")

                members, _ = av.run_wind_ensemble(
                    case, args, dem, (dlons, dlats), wc, (wlons, wlats),
                    weather, obs, ign,
                    offsets=[-135, -90, -45, 0, 45, 90, 135, 180],
                    progress_cb=_ecb, stop_area=_stop_area)
                ens = [{"offset": m["offset_deg"], **m["rep"]}
                       for m in members]
            bar.empty()
            live.empty()
            st.session_state["val_log"] = log_lines
            st.session_state["val_ens"] = ens
            # ---- persistent run archive (for later analysis) ----
            try:
                import numpy as _np2
                from PIL import Image as _Img
                import disaster_phyengine as _da2
                _keys = ["hit_rate", "mean_m", "p90_m",
                         "arrival_mae_h", "arrival_rho"]
                _summary = {k: {"mean": float(_np2.nanmean(
                                    [r[0].get(k, _np2.nan) for r in runs])),
                                "sd": float(_np2.nanstd(
                                    [r[0].get(k, _np2.nan) for r in runs]))}
                            for k in _keys}
                _json.dump({
                    "case_id": case_id, "case": {k: v for k, v in
                                                 case.items()},
                    "settings": {"cell_m": cell, "step_minutes": stepm,
                                 "hours": hours, "seeds": seeds,
                                 "wind_ensemble": bool(wens),
                                 "wind_offset_deg": wind_offset,
                                 "truth": ("perimeter" if up is not None
                                           else "firms")},
                    "engine_version": getattr(_da2, "__version__", "?"),
                    "initial_front_cells": (len(ign)
                                            if isinstance(ign, list) else 1),
                    "summary": _summary,
                    "runs": [r[0] for r in runs],
                    "wind_ensemble": ens,
                }, open(os.path.join(_rdir, "report.json"), "w"), indent=2)
                with open(os.path.join(_rdir, "run_log.txt"), "w") as _fh:
                    _fh.write("\n".join(log_lines))
                _best = max(runs, key=lambda r: r[0]["hit_rate"])[1]
                _img = base0.copy()
                _img[obs] = (58, 110, 220)
                _img[_best & obs] = (46, 160, 67)
                _img[_best & ~obs] = (200, 55, 44)
                _Img.fromarray(_img).save(os.path.join(_rdir,
                                                       "agreement.png"))
                st.success(f"Run archived: validation/runs/"
                           f"{case_id}_{_stamp}/ (report.json, run_log.txt, "
                           "agreement.png, frames/)")
            except Exception as _exc:
                st.warning(f"Could not archive the run: {_exc}")
            demg = av._sample(dem, dlons, dlats, lons, lats)
            wcg = av._sample(wc, wlons, wlats, lons, lats).astype(int)
            ftype = np.zeros_like(wcg)
            for code, (fid, _ld) in av.WORLDCOVER_TO_FUEL.items():
                ftype[wcg == code] = (case.get("tree_fuel", 3)
                                      if code == 10 else fid)
            st.session_state["val_result"] = dict(
                runs=[r[0] for r in runs],
                best=max(runs, key=lambda r: r[0]["hit_rate"])[1],
                obs=obs, shape=shape, case=case["label"], cell=cell,
                base=av._basemap(ftype, demg), ign=ign,
                wind_offset=wind_offset, obs_window_h=_obs_window,
                wind_toward_deg=_wtoward, stop_rule=stop_mode,
                seeds=seeds, cap_hours=hours,
                truth="perimeter" if up is not None else "firms")
        except Exception as exc:
            st.error(f"Validation failed: {exc}")
            return

    res = st.session_state.get("val_result")
    if res:
        st.markdown(f"#### Result \u2014 {res['case']}")
        import numpy as _np
        # exact settings that produced THIS result, so two runs can be
        # compared at a glance (only these change the outcome; the engine is
        # deterministic and a longer run is a superset of a shorter one)
        _woff = float(res.get("wind_offset", 0.0) or 0.0)
        _shr = res.get("runs", [{}])
        _sh0 = float(_np.mean([r.get("stop_hours", float("nan"))
                               for r in _shr])) if _shr else float("nan")
        _Bcells0 = int(_np.asarray(res.get("obs"), dtype=bool).sum()) \
            if res.get("obs") is not None else 0
        st.info(
            f"**Run settings** \u2014 cell **{float(res.get('cell', 0)):.0f} m** \u00b7 "
            f"wind offset **{_woff:+.0f}\u00b0** \u00b7 stopping "
            f"**{res.get('stop_rule', '?')}** \u00b7 stopped at "
            f"**{_sh0:.1f} h** (cap {float(res.get('cap_hours', 0)):.0f} h) \u00b7 "
            f"seeds **{int(res.get('seeds', 0))}** \u00b7 |B| **{_Bcells0}** cells. "
            "If two runs differ, one of THESE changed (most often the wind "
            "offset, which can linger from an earlier ensemble-adopt). The "
            "max-hours cap alone cannot move or reshape the burn; the engine "
            "is deterministic and a longer run is a superset of a shorter "
            "one.")

        def _cm(k):
            vs = [r.get(k, float("nan")) for r in res["runs"]]
            vs = [v for v in vs if v == v]
            return (float(_np.mean(vs)) if vs else float("nan"),
                    float(_np.std(vs)) if vs else float("nan"))

        def _fmtv(k, p):
            mv, sv = _cm(k)
            return "n/a" if mv != mv else f"{mv:.{p}f} \u00b1 {sv:.{p}f}"
        _sh = _cm("stop_hours")[0]
        if _sh == _sh:
            st.caption(f"Simulation stopped at ~{_sh:.1f} h per the stopping "
                       "rule, so the fronts are scored at a comparable "
                       "extent and the run always terminates.")
        # normalization context: fire equivalent radius and effective time
        _obsm = _np.asarray(res["obs"], dtype=bool)
        _cellm = float(res.get("cell", 90.0))
        _Bc = float(_obsm.sum())
        _R = ((_Bc * _cellm * _cellm / 3.141592653589793) ** 0.5
              if _Bc > 0 else 1.0)
        # arrival error is judged against the OBSERVED fire duration (span of
        # the satellite detections), not the sim's own run length; this is a
        # per-case, physically meaningful reference for every case
        _T = res.get("obs_window_h")
        if not (_T and _T == _T and _T > 0):
            _T = _sh if (_sh == _sh and _sh > 0) else 24.0

        def _clip01(x):
            return max(0.0, min(1.0, float(x)))

        def _assess(key):
            mv, _sv = _cm(key)
            if mv != mv:
                return ("n/a", None, "n/a", "")
            if key == "hit_rate":
                pct = 100 * _clip01(mv); tgt = "\u2265 0.70"
                vd = ("\u2713 pass" if mv >= 0.7 else
                      "~ partial" if mv >= 0.4 else "\u2717 fail")
            elif key == "mean_m":
                pct = 100 * _clip01(1 - mv / _R); tgt = f"\u2264 {0.25*_R:.0f} m"
                vd = ("\u2713 pass" if mv <= 0.25*_R else
                      "~ partial" if mv <= 0.5*_R else "\u2717 fail")
            elif key == "p90_m":
                pct = 100 * _clip01(1 - mv / _R); tgt = f"\u2264 {0.5*_R:.0f} m"
                vd = ("\u2713 pass" if mv <= 0.5*_R else
                      "~ partial" if mv <= _R else "\u2717 fail")
            elif key == "arrival_mae_h":
                pct = 100 * _clip01(1 - mv / _T); tgt = f"\u2264 {0.25*_T:.1f} h"
                vd = ("\u2713 pass" if mv <= 0.25*_T else
                      "~ partial" if mv <= 0.5*_T else "\u2717 fail")
            elif key == "arrival_rho":
                # the observed detection times are satellite-overpass gated;
                # with few distinct time levels the rank correlation is
                # dominated by ties and is not meaningful, so mark it n/a and
                # keep it OUT of the overall score
                _lv = _cm("arrival_obs_levels")[0]
                if _lv != _lv or _lv < 5:
                    return ("n/a", None, "n/a (overpass-limited)",
                            "\u2265 0.60")
                pct = 100 * _clip01(mv); tgt = "\u2265 0.60"
                vd = ("\u2713 pass" if mv >= 0.6 else
                      "~ partial" if mv >= 0.3 else "\u2717 fail")
            elif key == "arrival_n":
                pct = 100 * _clip01(mv / 50.0); tgt = "\u2265 30 cells"
                vd = ("\u2713 pass" if mv >= 30 else
                      "~ partial" if mv >= 10 else "\u2717 fail")
            else:
                pct = 0.0; tgt = ""; vd = "n/a"
            return (f"{pct:.0f}%", pct, vd, tgt)
        _CR = [
            ("hit_rate", "Coverage (POD)", 3, "|A\u2229B| / |B|",
             "fraction of the observed burn reproduced"),
            ("mean_m", "Front error, mean (m)", 0, "mean d(\u2202A,\u2202B)",
             "average front-to-front distance"),
            ("p90_m", "Front error, p90 (m)", 0, "P90 d(\u2202A,\u2202B)",
             "90th-percentile front distance"),
            ("arrival_mae_h", "Arrival MAE (h)", 2, "mean |t_sim \u2212 t_obs|",
             "arrival-time error over shared cells"),
        ]
        _left, _right, _pcts = [], [], []
        for _key, _lab, _pr, _eqn, _meas in _CR:
            _score_s, _pct, _vd, _tgt = _assess(_key)
            if _pct is not None:
                _pcts.append(_pct)
            _left.append({"Criterion": _lab,
                          "Value (mean \u00b1 sd)": _fmtv(_key, _pr),
                          "Score": _score_s, "Verdict": _vd})
            _right.append({"Criterion": _lab, "Equation": _eqn,
                           "What it measures": _meas, "Target": _tgt})
        _overall = (sum(_pcts) / len(_pcts)) if _pcts else float("nan")
 # accumulate this case's Table 5.5 row so an export can hold all four
        # scenarios together (keyed by case, so a re-run updates its own row)
        _acell = float(res.get("cell", 0.0)) ** 2
        st.session_state.setdefault("val_rows", {})
        st.session_state["val_rows"][str(res.get("case", ""))] = {
            "POD": _cm("hit_rate")[0], "mean_m": _cm("mean_m")[0],
            "p90_m": _cm("p90_m")[0], "MAE": _cm("arrival_mae_h")[0],
            "R": _R, "T": _T, "A_cell": _acell,
            "cell": float(res.get("cell", 0.0))}
        st.markdown("**Success metrics** \u2014 normalized score = agreement with "
                    "the target (100% is ideal), pass / partial / fail per "
                    "criterion")
        _lc, _rc = st.columns([1.15, 1])
        _lc.table(_left)
        _rc.table(_right)
        if _overall == _overall:
            _ov = ("SUCCESS" if _overall >= 70 else
                   "PARTIAL" if _overall >= 45 else "WEAK")
            st.markdown(f"**Overall validation score: {_overall:.0f}% "
                        f"({_ov})** \u2014 mean of the four normalized criteria.")
        st.caption("Normalization: coverage and \u03c1 are already 0..1; the front "
                   f"errors are scored against the fire equivalent radius "
                   f"R = {_R:.0f} m (\u221a(|B|\u00b7cell\u00b2/\u03c0)); the arrival MAE against "
                   f"the OBSERVED fire duration T = {_T:.1f} h (span of the "
                   "satellite detections), so the timing error is judged "
                   "relative to how long the real fire actually spread. "
                   "Suppression-free model: validated on front propagation, "
                   "not final area (Dice/Jaccard/FAR/area-bias omitted; "
                   "upload an EFFIS/EMS perimeter for area metrics).")
        nx, ny = res["shape"]
        img = (res["base"].copy() if res.get("base") is not None
               else _np.zeros((ny, nx, 3), dtype=_np.uint8) + 24)
        best, obs = res["best"], res["obs"]
        img[best & obs] = (46, 160, 67)
        img[best & ~obs] = (200, 55, 44)
        img[~best & obs] = (58, 110, 220)
        if res.get("ign") is not None:
            _i0 = (res["ign"][0] if isinstance(res["ign"], list)
                   else res["ign"])
            gx, gy = int(_i0[0]), int(_i0[1])
            rr = max(2, nx // 150)
            img[max(0, gy - rr):gy + rr + 1,
                max(0, gx - rr):gx + rr + 1] = (255, 235, 60)
        # keep this case's agreement map so the export can embed the PNG of
        # EVERY scenario run so far, not just the current one
        st.session_state.setdefault("val_imgs", {})
        st.session_state["val_imgs"][str(res.get("case", ""))] = img.copy()
        # suggested wind-direction offset to steer the sim toward the
        # OBSERVED spread: rotate the wind by (observed heading - simulated
        # heading), both measured from the ignition to the burn centroid
        _sugg = None
        try:
            _obsm2 = _np.asarray(res["obs"], dtype=bool)
            _simm2 = _np.asarray(res["best"], dtype=bool)
            _i0s = (res["ign"][0] if isinstance(res.get("ign"), list)
                    else res.get("ign"))
            if _i0s is not None and _obsm2.any() and _simm2.any():
                _gx0, _gy0 = float(_i0s[0]), float(_i0s[1])
                _oy, _ox = _np.where(_obsm2)
                _sy, _sx = _np.where(_simm2)

                def _brg2(cx, cy):
                    _e = cx - _gx0
                    _n = -(cy - _gy0)          # image row grows south
                    return _np.degrees(_np.arctan2(_e, _n)) % 360.0
                _ob = _brg2(_ox.mean(), _oy.mean())
                _sb = _brg2(_sx.mean(), _sy.mean())
                # residual heading gap of THIS run, plus the offset already
                # applied, gives the ABSOLUTE offset to set (so it converges
                # instead of oscillating around the leftover gap)
                _applied = float(res.get("wind_offset", 0.0) or 0.0)
                _gap = ((_ob - _sb + 180.0) % 360.0) - 180.0
                _sugg = ((_applied + _gap + 180.0) % 360.0) - 180.0
        except Exception:
            _sugg = None
            _gap = 0.0
        # wind direction (mean over the run) shown as a compact line ABOVE
        # the map, so the map itself stays clean
        _wt = res.get("wind_toward_deg")
        if _wt is not None and _wt == _wt:
            _brg = (90.0 - float(_wt)) % 360.0
            _pts8 = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]
            _gly8 = ["↑", "↗", "→", "↘", "↓", "↙", "←", "↖"]
            _bi = int((_brg + 22.5) // 45) % 8
            _msg = (f"**Mean wind over the run:** {_gly8[_bi]} toward "
                    f"**{_pts8[_bi]}** (bearing {_brg:.0f}°, downwind). The "
                    "spread also feels upslope, so the burn direction can "
                    "differ from the raw wind.")
            if _sugg is not None and abs(_gap) >= 8:
                _msg += (f"  \n**Set the wind-direction offset to about "
                         f"{_sugg:+.0f}°** to steer the sim toward the "
                         f"observed spread (current residual gap "
                         f"{_gap:+.0f}°). It may take one more pass since "
                         "slope also steers the fire; the wind ensemble is "
                         "the reliable optimizer.")
            elif _sugg is not None:
                _msg += ("  \nHeading matches the observed spread "
                         "(residual gap under 8°).")
            st.markdown(_msg)
        # interactive agreement map: its OWN zoom/pan (scroll + toolbar),
        # independent of the browser zoom, on a fixed-height canvas
        try:
            import plotly.graph_objects as _go
            _figm = _go.Figure(_go.Image(z=img))
            _figm.update_xaxes(visible=False,
                               scaleanchor="y", constrain="domain")
            _figm.update_yaxes(visible=False)
            _figm.update_layout(height=640, dragmode="pan",
                                margin=dict(l=0, r=0, t=6, b=0))
            with _map_card():
                st.plotly_chart(_figm, use_container_width=True,
                                config={"scrollZoom": True,
                                        "displayModeBar": True,
                                        "displaylogo": False})
        except Exception:
            with _map_card():
                st.image(img, use_container_width=True)
        st.caption("Agreement map on the real terrain \u2014 green: correctly "
                   "predicted burn, red: simulated only, blue: observed only, "
                   "yellow: ignition. Scroll or use the toolbar to zoom and "
                   "pan; double-click resets.")
        # --- Export: a Word report (Table 5.5 + computation) + map PNG ---
        st.markdown("**Export for the thesis** \u2014 a Word report (Table 5.5 for "
                    "all cases run so far, plus the R / A_cell / % / target "
                    "computation for this case) and the agreement map as PNG")
        _accum = st.session_state.get("val_rows", {})
        _ec1, _ec2 = st.columns([2, 1])
        _ec1.caption("Table 5.5 currently holds: "
                     + (", ".join(_accum.keys()) if _accum else "(none yet)")
                     + f"  ({len(_accum)}/4 scenarios). Run each case once, "
                     "then export; the table and the per-case maps accumulate.")
        if _ec2.button("Clear accumulated table"):
            # wipe EVERYTHING: the accumulated Table 5.5 rows + per-case maps,
            # and the current run so nothing lingers in the view or the export
            for _k in ("val_rows", "val_imgs", "val_result", "val_log",
                       "val_ens", "val_wens", "val_adopt_pick"):
                st.session_state.pop(_k, None)
            st.rerun()
        if st.button("Export to Word + PNG", type="primary"):
            try:
                _dbytes, _pbytes, _odir = _make_validation_report(
                    res, _left, _right, _overall, _R, _T, img,
                    st.session_state.get("val_rows", {}),
                    st.session_state.get("val_imgs", {}))
                st.success(f"Saved to {os.path.relpath(_odir)}")
                st.download_button("\u2b07 Download Word report (.docx)", _dbytes,
                                   file_name="validation_report.docx")
                if _pbytes:
                    st.download_button("\u2b07 Download agreement map (.png)",
                                       _pbytes, file_name="agreement_map.png")
            except Exception as _exc:
                st.error(f"Export failed: {_exc}")
        if st.session_state.get("val_log"):
            with st.expander("Run log \u2014 what happened, step by step"):
                st.code("\n".join(st.session_state["val_log"]),
                        language=None)
        _ens = st.session_state.get("val_ens")
        if _ens:
            st.markdown("**Wind-direction ensemble** - same fire, wind "
                        "rotated; ranked by coverage (POD):")
            _best_e = max(_ens, key=lambda x: x["hit_rate"])
            rows = ["| rotation | Coverage (POD) | Front err (km) |",
                    "|---|---|---|"]
            for m in sorted(_ens, key=lambda m: m["offset"]):
                mark = " **\u2190 best**" if m is _best_e else ""
                rows.append(f"| {m['offset']:+.0f}\u00b0 | "
                            f"{m['hit_rate']:.3f} | "
                            f"{m['mean_m']/1000:.1f}{mark} |")
            st.markdown("\n".join(rows))
            _zero = [m for m in _ens if m["offset"] == 0]
            _z = _zero[0]["hit_rate"] if _zero else float("nan")
            st.caption(f"Best member at {_best_e['offset']:+.0f}\u00b0 "
                       f"(coverage {_best_e['hit_rate']:.2f} vs {_z:.2f} "
                       "with the raw reanalysis wind). A large gain from "
                       "rotation means the case is limited by the INPUT "
                       "wind (terrain channeling, fire convection), not by "
                       "the spread model; report it as input sensitivity.")
            # --- adopt a wind offset and rerun without the ensemble ---
            _offs = sorted({m["offset"] for m in _ens})
            _dmap = {m["offset"]: m["hit_rate"] for m in _ens}
            _ca, _cb = st.columns([1.5, 1])
            _pick = _ca.selectbox(
                "Adopt a wind offset for a clean (ensemble-off) rerun",
                _offs, index=_offs.index(_best_e["offset"]),
                format_func=lambda o: f"{o:+.0f}\u00b0  (POD {_dmap[o]:.3f})",
                key="val_adopt_pick")
            _cb.write("")
            if _cb.button("Use this wind \u2192 rerun", type="primary"):
                st.session_state["_val_adopt_offset"] = float(_pick)
                st.rerun()
            st.caption("Adopting sets the wind-direction offset field above "
                       "and turns the ensemble off. Then press "
                       "**Run validation** to score that single wind against "
                       "the truth.")


def _make_validation_report(res, left, right, overall, R, T, img,
                            all_rows=None, all_imgs=None):
    """Build a Word report (Table 5.5 row + detailed R/%/target computation +
    the agreement map of EVERY case run so far) and a PNG of the map, always
    in the same format. Returns (docx_bytes, png_bytes, out_dir)."""
    import datetime as _dt
    import numpy as _np2
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image as _PILImg

    def _agg(k):
        vs = [r.get(k, float("nan")) for r in res["runs"]]
        vs = [v for v in vs if v == v]
        return (float(_np2.mean(vs)) if vs else float("nan"),
                float(_np2.std(vs)) if vs else float("nan"))

    def _f(v, p):
        return "n/a" if (v != v) else f"{v:.{p}f}"

    stamp = _dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    _slug = "".join(c for c in str(res.get("case", "case"))
                    if c.isalnum() or c in "-_")[:40]
    out_dir = os.path.join(root, "validation", "runs",
                           f"export_{_slug}_{stamp}")
    os.makedirs(out_dir, exist_ok=True)

    png_path = os.path.join(out_dir, "agreement_map.png")
    try:
        _PILImg.fromarray(img).save(png_path)
    except Exception:
        png_path = None

    d = Document()
    d.styles["Normal"].font.name = "Times New Roman"
    d.styles["Normal"].font.size = Pt(11)
    d.add_heading(f"Simulation Core Validation: {res.get('case', '')}", 1)

    _woff = float(res.get("wind_offset", 0.0) or 0.0)
    _sh = _agg("stop_hours")[0]
    _Bc = int(_np2.asarray(res.get("obs"), dtype=bool).sum()) \
        if res.get("obs") is not None else 0
    d.add_paragraph(
        f"Run settings: cell {float(res.get('cell', 0)):.0f} m, time window "
        f"{float(res.get('cap_hours', 0)):.0f} h, wind offset {_woff:+.0f} deg, "
        f"stopping {res.get('stop_rule', '?')}, stopped at {_sh:.1f} h, seeds "
        f"{int(res.get('seeds', 0))}, |B| {_Bc} cells, R {R:.0f} m, T {T:.1f} h "
        f"(exported {stamp}).")

    # ---- Table 5.5 (all cases run so far, paste into the thesis table) ----
    d.add_heading("Table 5.5 (paste into the thesis)", 2)
    # accumulate the current case into the row set so a single run still works
    _acell_now = float(res.get("cell", 0.0)) ** 2
    rows = dict(all_rows or {})
    rows.setdefault(str(res.get("case", "")), {
        "POD": _agg("hit_rate")[0], "mean_m": _agg("mean_m")[0],
        "p90_m": _agg("p90_m")[0], "MAE": _agg("arrival_mae_h")[0],
        "R": R, "T": T, "A_cell": _acell_now,
        "cell": float(res.get("cell", 0.0))})
    _order = list(rows.keys())
    t5 = d.add_table(rows=1 + len(_order) + 1, cols=6); t5.style = "Table Grid"
    for i, x in enumerate(["Real Sce. (Table 5.4)", "cell (m)", "POD (-)",
                           "d_Front mean (m)", "d_90 (m)", "MAE (h)"]):
        c = t5.rows[0].cells[i]; c.text = ""
        c.paragraphs[0].add_run(x).bold = True
    for ri, ck in enumerate(_order, start=1):
        rd = rows[ck]
        vals = [ck, _f(rd.get("cell", float("nan")), 0),
                _f(rd.get("POD", float("nan")), 3),
                _f(rd.get("mean_m", float("nan")), 0),
                _f(rd.get("p90_m", float("nan")), 0),
                _f(rd.get("MAE", float("nan")), 2)]
        for i, x in enumerate(vals):
            t5.rows[ri].cells[i].text = str(x)
    sc = t5.rows[-1].cells
    sc[0].text = ""; sc[0].paragraphs[0].add_run("Success Criteria").bold = True
    for i, x in enumerate(["-", ">= 0.70", "<= 0.25 R", "<= 0.50 R",
                           "<= 0.25 T"], start=1):
        sc[i].text = x
    d.add_paragraph("The front and arrival targets are relative to the "
                    "per-case R and T listed below, since these differ by "
                    "case.").italic = True

    # ---- per-case normalization constants ----
    d.add_heading("Per-case normalization (R, T, A_cell)", 2)
    nt = d.add_table(rows=1 + len(_order), cols=6); nt.style = "Table Grid"
    for i, x in enumerate(["Real Sce.", "cell (m)", "A_cell (m^2)", "R (m)",
                           "0.25 R / 0.50 R (m)", "T (h)  /  0.25 T (h)"]):
        c = nt.rows[0].cells[i]; c.text = ""
        c.paragraphs[0].add_run(x).bold = True
    for ri, ck in enumerate(_order, start=1):
        rd = rows[ck]; _Rc = rd.get("R", 0.0); _Tc = rd.get("T", 0.0)
        _ac = rd.get("A_cell", 0.0); _cl = rd.get("cell", 0.0)
        vals = [ck, f"{_cl:.0f}", f"{_ac:.0f}", f"{_Rc:.0f}",
                f"{0.25*_Rc:.0f} / {0.50*_Rc:.0f}",
                f"{_Tc:.1f} / {0.25*_Tc:.1f}"]
        for i, x in enumerate(vals):
            nt.rows[ri].cells[i].text = str(x)

    # ---- detailed computation for EVERY case run so far ----
    def _clip01(v):
        return 0.0 if v < 0 else 1.0 if v > 1 else v

    def _detail_rows(rd):
        """(Criterion, Value, Normalization, Score, Target, Verdict) list for
        one case, computed from its stored POD / front / MAE and R / T."""
        _Rc = float(rd.get("R", 0.0)) or float("nan")
        _Tc = float(rd.get("T", 0.0)) or float("nan")
        pod = rd.get("POD", float("nan"))
        fm = rd.get("mean_m", float("nan"))
        f9 = rd.get("p90_m", float("nan"))
        mae = rd.get("MAE", float("nan"))
        out = []
        # Coverage (POD)
        _s = 100 * _clip01(pod) if pod == pod else float("nan")
        _v = ("pass" if pod >= 0.70 else "partial" if pod >= 0.50 else "fail") \
            if pod == pod else "n/a"
        out.append(("Coverage (POD)", _f(pod, 3), "value (already 0..1)",
                    _f(_s, 0) + "%" if _s == _s else "n/a", ">= 0.70", _v))
        # Front mean
        _s = 100 * _clip01(1 - fm / _Rc) if (fm == fm and _Rc == _Rc) \
            else float("nan")
        _v = ("pass" if fm <= 0.25 * _Rc else "partial" if fm <= 0.50 * _Rc
              else "fail") if (fm == fm and _Rc == _Rc) else "n/a"
        out.append(("Front error, mean (m)", _f(fm, 0), "1 - value / R",
                    _f(_s, 0) + "%" if _s == _s else "n/a",
                    f"<= 0.25 R = {0.25*_Rc:.0f} m" if _Rc == _Rc else "n/a",
                    _v))
        # Front p90
        _s = 100 * _clip01(1 - f9 / _Rc) if (f9 == f9 and _Rc == _Rc) \
            else float("nan")
        _v = ("pass" if f9 <= 0.50 * _Rc else "partial" if f9 <= _Rc
              else "fail") if (f9 == f9 and _Rc == _Rc) else "n/a"
        out.append(("Front error, p90 (m)", _f(f9, 0), "1 - value / R",
                    _f(_s, 0) + "%" if _s == _s else "n/a",
                    f"<= 0.50 R = {0.50*_Rc:.0f} m" if _Rc == _Rc else "n/a",
                    _v))
        # Arrival MAE
        _s = 100 * _clip01(1 - mae / _Tc) if (mae == mae and _Tc == _Tc) \
            else float("nan")
        _v = ("pass" if mae <= 0.25 * _Tc else "partial" if mae <= 0.50 * _Tc
              else "fail") if (mae == mae and _Tc == _Tc) else "n/a"
        out.append(("Arrival MAE (h)", _f(mae, 2), "1 - value / T",
                    _f(_s, 0) + "%" if _s == _s else "n/a",
                    f"<= 0.25 T = {0.25*_Tc:.1f} h" if _Tc == _Tc else "n/a",
                    _v))
        return out

    for ck in _order:
        rd = rows[ck]
        _Rc = float(rd.get("R", 0.0)); _Tc = float(rd.get("T", 0.0))
        _ac = float(rd.get("A_cell", 0.0)); _cl = float(rd.get("cell", 0.0))
        d.add_heading(f"Detailed computation: {ck}", 2)
        p = d.add_paragraph()
        p.add_run(
            "A_cell = cell area = (cell size)^2 = "
            f"{_ac:.0f} m^2 (cell = {_cl:.0f} m). "
            f"R = sqrt(|B| * A_cell / pi) = {_Rc:.0f} m is the equivalent "
            "radius of the observed burned area. "
            f"T = {_Tc:.1f} h is the observed fire duration (span of the "
            "satellite detections). Coverage is already in [0, 1]; front "
            "errors are scored as 1 - value/R; the arrival error as "
            "1 - value/T.").italic = True
        drows = _detail_rows(rd)
        dt = d.add_table(rows=1 + len(drows), cols=6); dt.style = "Table Grid"
        for i, x in enumerate(["Criterion", "Value", "Normalization",
                               "Score", "Target", "Verdict"]):
            c = dt.rows[0].cells[i]; c.text = ""
            c.paragraphs[0].add_run(x).bold = True
        _scores = []
        for ri, row6 in enumerate(drows, start=1):
            for i, x in enumerate(row6):
                dt.rows[ri].cells[i].text = str(x)
            try:
                _scores.append(float(str(row6[3]).rstrip("%")))
            except ValueError:
                pass
        if _scores:
            d.add_paragraph(
                f"Overall validation score: {sum(_scores)/len(_scores):.0f}% "
                "(mean of the four normalized criteria).")

    # ---- basis of the acceptance targets ----
    d.add_heading("Basis of the acceptance targets", 2)
    bt = d.add_table(rows=1, cols=3); bt.style = "Table Grid"
    for i, x in enumerate(["Criterion", "Target", "Basis"]):
        c = bt.rows[0].cells[i]; c.text = ""
        c.paragraphs[0].add_run(x).bold = True
    _basis = [
        ("Coverage (POD)", ">= 0.70",
         "Literature level for good agreement between simulated and observed "
         "fire footprints: FARSITE reached a Sorensen coefficient of 0.70 with "
         "improved fuel inputs (Price et al., 2022); calibrated cell-based "
         "models reach 0.7-0.9 and uncalibrated semi-empirical models 0.5-0.7 "
         "(Pais et al., 2021). The POD/Sorensen index set follows Filippi et "
         "al. (2014)."),
        ("Front error, mean (m)", f"<= 0.25 R = {0.25*R:.0f} m",
         "Acceptance threshold adopted in this study: the mean front lies "
         "within a quarter of the fire equivalent radius R, so the criterion "
         "is scale aware. Front position error in metres itself is a standard "
         "FARSITE-style output (Finney, 1998)."),
        ("Front error, p90 (m)", f"<= 0.50 R = {0.50*R:.0f} m",
         "Acceptance threshold adopted in this study: the worst-matching "
         "tenth of the front lies within half of R."),
        ("Arrival MAE (h)", f"<= 0.25 T = {0.25*T:.1f} h",
         "Acceptance threshold adopted in this study: the arrival-time error "
         "is within a quarter of the observed fire duration T (scale aware)."),
    ]
    for cr, tg, bs in _basis:
        cells = bt.add_row().cells
        cells[0].text = cr; cells[1].text = tg; cells[2].text = bs

    d.add_heading("References", 3)
    for _r in [
        "Filippi, J.-B., Mallet, V., & Nader, B. (2014). Representation and "
        "evaluation of wildfire propagation simulations. International Journal "
        "of Wildland Fire, 23(1), 46-57.",
        "Finney, M. A. (1998). FARSITE: Fire Area Simulator, model development "
        "and evaluation. USDA Forest Service, RMRS-RP-4.",
        "Pais, C., Carrasco, J., Martell, D. L., Weintraub, A., & Woodruff, "
        "D. L. (2021). Cell2Fire: a cell-based forest fire growth model to "
        "support strategic landscape management planning. Frontiers in Forests "
        "and Global Change, 4, 692706.",
        "Price, S., et al. (2022). Modeling of fire spread in sagebrush "
        "steppe using FARSITE: an approach to improving input data and "
        "simulation accuracy. Fire Ecology, 18, 22.",
    ]:
        _p = d.add_paragraph(_r)
        _p.paragraph_format.left_indent = Pt(18)
        _p.paragraph_format.first_line_indent = Pt(-18)

    # ---- figures: the agreement map of EVERY case run so far ----
    # build the image set: the accumulated per-case maps, and always the
    # current one (so a single run still exports its map).
    _imgs = dict(all_imgs or {})
    _imgs.setdefault(str(res.get("case", "")), img)
    d.add_heading("Agreement maps", 2)
    d.add_paragraph(
        "Green: correctly predicted burn, Red: simulated only, Blue: "
        "observed only, Yellow: ignition.").italic = True
    _fno = 0
    for _ck, _im in _imgs.items():
        if _im is None:
            continue
        _fno += 1
        _cslug = "".join(c for c in str(_ck) if c.isalnum() or c in "-_")[:40]
        _fp = os.path.join(out_dir, f"agreement_map_{_cslug or _fno}.png")
        try:
            _PILImg.fromarray(_np2.asarray(_im)).save(_fp)
        except Exception:
            continue
        d.add_heading(str(_ck), 3)
        d.add_picture(_fp, width=Inches(5.6))
        cap = d.add_paragraph()
        cap.add_run(f"Figure 5.{_fno} Agreement map for {_ck}.").italic = True

    docx_path = os.path.join(out_dir, "validation_report.docx")
    d.save(docx_path)
    with open(docx_path, "rb") as fh:
        docx_bytes = fh.read()
    png_bytes = open(png_path, "rb").read() if png_path else b""
    return docx_bytes, png_bytes, out_dir


# =========================================================== STEP ANALYSIS ===
def page_steps():
    """The per-step adaptation table, across the whole window.

    The same rows the Layer 4 Analysis panel previews, but this is a page, so
    all twenty-one columns get real width instead of the narrow side panel."""
    st.subheader("Step analysis — what happened at each decision cycle")
    _eng = st.session_state.get("dss_engine")
    _cyc = list(getattr(_eng, "cycles", []) or [])
    _rs = dict(getattr(_eng, "run_stats", {}) or {})
    if not _cyc:
        st.info("No decision cycle has run yet. Go to Simulation, turn on "
                "'Apply decisions' and step the fire.")
        return
    # the tabs are built AFTER the guard, so an empty run shows the message
    # instead of three empty tables with the message stranded underneath
    _tabA, _tabB, _tabC = st.tabs(
        ["Adaptation, one row per cycle",
         "Agents, one row per agent per cycle",
         "Global DSS, the coordination decision"])
    _rows = build_step_rows(_cyc, _rs)
    _acc = sum(1 for c in _cyc
               if (c.get("adaptation") or {}).get("accepted"))
    _gd = _read_gstate()

    with _tabA:
        st.caption(
            f"{len(_cyc)} decision cycle(s) · {_acc} adaptation(s) accepted · "
            f"seed base "
            f"**{st.session_state.get('dss_seed_profile', 'minimal')}**"
            f" · config `{_dss_pkg.config_id(_gd.get('runtime_flags') or {})}`")
        st.caption("The adaptation runs on ONE region per cycle, the "
                   "coordinator's hotspot, so this table has one row per "
                   "cycle. What each agent decided is on the next tab.")
        _c1, _c2, _c3 = st.columns([2, 2, 3])
        _only = _c1.checkbox("Only cycles where a stage was tried",
                             value=False, key="steps_only_tried")
        _acc_only = _c2.checkbox("Only accepted", value=False,
                                 key="steps_only_acc")
        _view = list(_rows)
        if _only:
            _view = [r for r in _view
                     if r.get("stage", "").strip()[:1] != "0"]
        if _acc_only:
            _view = [r for r in _view if r.get("verdict") == "ACCEPTED"]
        _c3.caption(f"showing {len(_view)} of {len(_rows)} cycle(s), "
                    f"latest first")
        st.dataframe(_view, use_container_width=True, height=560,
                     column_order=STEP_COLS, column_config=_STEP_COL_CONFIG)
        # ONE WORKBOOK, ALL THREE TABLES. They describe the same cycles
        # from three angles, so shipping them as three files left the
        # reader lining up cycle numbers by hand.
        st.download_button(
            "⬇ Download all three tables (Excel)",
            _xlsx_bytes({"Adaptation": [{c: r.get(c) for c in STEP_COLS}
                                        for r in _view],
                         "Agents": [{c: r.get(c) for c in AGENT_COLS}
                                    for r in build_agent_rows(_cyc)],
                         "Global DSS": [{c: r.get(c) for c in GLOBAL_COLS}
                                        for r in build_global_rows(_cyc)]},
                        meta={"cycles": len(_cyc),
                              "adaptations accepted": _acc,
                              "seed base": st.session_state.get(
                                  "dss_seed_profile", "minimal"),
                              "configuration": _dss_pkg.config_id(
                                  _gd.get("runtime_flags") or {}),
                              "rows shown (Adaptation)": len(_view)}),
            file_name="layer4_step_analysis.xlsx", mime=XLSX_MIME,
            key="steps_xlsx")
        with st.expander("What the columns mean", expanded=False):
            st.markdown(_STEP_COL_HELP)
        with st.expander("What the gates G1 to G5 mean", expanded=False):
            st.markdown(_GATE_HELP)

    with _tabB:
        _arows = build_agent_rows(_cyc)
        _names = sorted({r["agent"] for r in _arows})
        st.caption(
            f"{len(_arows)} agent-cycle(s) across {len(_names)} local DSS "
            f"agent(s). Every agent decides in EVERY cycle; the coordinator "
            f"then scales the offensive tempo by the attention share, which "
            f"is the difference between orders_from_rules and orders_final.")
        _a1, _a2 = st.columns([3, 2])
        _pick = _a1.multiselect("Agents", _names, default=_names,
                                key="steps_agent_pick")
        _fs_only = _a2.checkbox("Only cycles the fail-safe touched",
                                value=False, key="steps_agent_fs")
        _av = [r for r in _arows if r["agent"] in _pick]
        if _fs_only:
            _av = [r for r in _av if r.get("failsafe") != "ok"]
        _a2.caption(f"showing {len(_av)} of {len(_arows)} row(s)")
        st.dataframe(_av, use_container_width=True, height=560,
                     column_order=AGENT_COLS,
                     column_config=_STEP_COL_CONFIG)
        st.download_button(
            "⬇ Download this table (Excel)",
            _xlsx_bytes({"Agents": [{c: r.get(c) for c in AGENT_COLS}
                                    for r in _av]}),
            file_name="layer4_agents.xlsx", mime=XLSX_MIME,
            key="steps_agents_xlsx")
        st.markdown(_AGENT_COL_HELP)

    with _tabC:
        _grows = build_global_rows(_cyc)
        st.caption(
            "One row per cycle: how the Global DSS ranked the regions, what "
            "share each got, and which ones it decided to only monitor. "
            "The share scales the offensive tempo of that agent and steers "
            "the budget concentration in the allocator.")
        st.dataframe(_grows, use_container_width=True, height=560,
                     column_order=GLOBAL_COLS,
                     column_config=_STEP_COL_CONFIG)
        st.download_button(
            "⬇ Download this table (Excel)",
            _xlsx_bytes({"Global DSS": [{c: r.get(c) for c in GLOBAL_COLS}
                                        for r in _grows]}),
            file_name="layer4_global.xlsx", mime=XLSX_MIME,
            key="steps_global_xlsx")


PAGES = {"Simulation": page_simulation, "Map editor": page_editor,
         "Data layers": page_layers, "Parameters": page_params,
         "GIS import": page_gis, "Validation": page_validation,
         "Step analysis": page_steps,
 "System Description": page_system_description}

# Map layers bar: pinned at the top of the content area, but only on the pages
# that actually draw a map (Simulation, Map editor). It is the single
# interactive source for the ly_*_v flags; those pages only read them.
if page in ("Simulation", "Map editor"):
    with st.expander("🗺️  Map layers — show / hide", expanded=False):
        _render_layer_toggles(ncols=4, key_prefix="lyrtop")

PAGES[page]()

# run to end: a CHUNK of steps per rerun, AFTER the page has rendered,
# so the map shows every intermediate state instead of a blank flicker
# (the old version stepped inside the sidebar and reran before the map
# was ever drawn). The chunk is time-budgeted: about one second of
# stepping per frame keeps the motion visible without crawling.
if page == "Simulation" and st.session_state.get("runend_on"):
    _t0 = time.time()
    _limit = int(cfg.max_steps)
    _done = None
    while sim.state.step < _limit and time.time() - _t0 < 1.0:
        _d = _step_sim(); _record_costs()
        if (_d.n_burning == 0 and sim.state.step > 1
                and sim.ever_burned.any()
                and not any(ev.step >= sim.state.step
                            for ev in world.ignitions)):
            _done = "fire is out"
            break
    if sim.state.step >= _limit:
        _done = f"step cap {_limit} reached"
    if _done:
        st.session_state["runend_stop"] = True
        st.toast(f"Run to end stopped at step "
                 f"{sim.state.step}: {_done}")
    st.rerun()

# animate: advance one step per rerun while playing (only on the Simulation
# page, otherwise the rerun loop would trap every other tab)
if page == "Simulation" and st.session_state.get("anim_on", False):
    _pending = any(ev.step >= sim.state.step for ev in world.ignitions)
    # is_quiescent() is a SNAPSHOT: no cell above the burning threshold in
    # THIS frame. A fire dips below it all the time without being out (the
    # ignition still ramping up, a front briefly knocked down by suppression,
    # the overnight moisture stall). Treating one such frame as "the fire is
    # over" is what stopped the animation on the first click and again in the
    # middle of a run. So: no pending ignition, and quiet for several
    # consecutive frames, matching what Run to end already required.
    _QUIET_NEED = 3
    if sim.is_quiescent():
        _quiet = int(st.session_state.get("anim_quiet", 0)) + 1
    else:
        _quiet = 0
    st.session_state["anim_quiet"] = _quiet
    _finished = (_quiet >= _QUIET_NEED and not _pending
                 and sim.state.step > 1 and bool(sim.ever_burned.any()))
    if not _finished and sim.state.step < cfg.max_steps:
        # no artificial pause: a frame already costs physics + DSS + render
        # (well over 100 ms on a normal map), so the sleep only added dead
        # time to every frame without making the animation any smoother
        _step_sim(); _record_costs(); st.rerun()
    else:
        # cannot write a widget key after the toggle rendered; flag it for
        # the next run (handled just before the toggle in the sidebar)
        if sim.state.step >= cfg.max_steps:
            st.toast(f"Animation stopped at the max_steps cap "
                     f"({cfg.max_steps}); raise it in Parameters "
                     "> Run limit.")
        else:
            st.toast("Animation stopped: the fire is out.")
        st.session_state["anim_stop"] = True
        st.session_state.pop("anim_quiet", None)
        st.rerun()
